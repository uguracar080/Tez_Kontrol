# ============================================================
# OKÜ LEE - Tez Yazım Kuralları
# ------------------------------------------------------------
# Bu program Dr. Öğr. Üyesi Uğur ACAR tarafından geliştirilmiştir,2025.
# E-posta:uguracar@hotmail.com, uguracar@osmaniye.edu.tr adresinden ulaşabilirsiniz

import sys
import yaml
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import time


# ============================================================
# 0. Logger ayarı (hem ekrana hem log.txt dosyasına yazacak)
# ============================================================

import logging
import io
# Windows'ta emoji karakterleri için UTF-8 encoding ayarı
_stdout_handler = logging.StreamHandler(io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace'))
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        #logging.FileHandler("log.txt", mode="w", encoding="utf-8"),  # log.txt dosyasına yaz
        _stdout_handler  # ekrana da yaz (UTF-8)
    ]
)
logger = logging.getLogger(__name__)

# ============================================================
# DEBUG klasörü ayarı (tüm debug çıktıları Debug/ içine)
# ============================================================
BASE_DIR = Path(__file__).resolve().parent
DEBUG_DIR = BASE_DIR / "Debug"
DEBUG_DIR.mkdir(parents=True, exist_ok=True)

def dbg_path(filename: str) -> str:
    return str(DEBUG_DIR / filename)


# ============================================================
# 1. run_check
# ============================================================
def format_location_by_page(global_line: int, lines_per_page: int = 40) -> str:
    """
    Global satır numarasını (1, 2, 3, ...) tahmini sayfa/satır bilgisine çevirir.
    lines_per_page: bir sayfada yaklaşık kaç satır/paragraf olduğunu varsayıyoruz.
    Örnek çıktı: 'Sayfa 3, satır 7'
    """
    if global_line < 1:
        global_line = 1

    page = (global_line - 1) // lines_per_page + 1
    line_on_page = (global_line - 1) % lines_per_page + 1

    return f"Sayfa {page}, satır {line_on_page}"

def run_check(doc, paragraphs, check, rules_data):
    """
    Belgedeki paragrafları (paragraphs) ve doc nesnesini verilen kurallara (check) göre kontrol eder.
    Her kural için tabloya şu formatta döner:
        (paragraf_indexi, durum(True/False), kural_adı, açıklama)
    """
    global memo
    
    section = check.get("section", "")
    
    results = []  # Çıktılar burada toplanır 

    # ------------------------------------------------------
    # Yardımcı fonksiyonlar
    # ------------------------------------------------------
    def is_blank_para_text(txt: str) -> bool:
        """Sadece boşluk/tab içeren paragrafları da boş kabul et."""
        return (not txt) or (txt.strip() == "") or ("".join(txt.split()) == "")

    def para_font_size_pt(para):
        """Run yoksa stil punto değerini oku; yoksa None döner."""
        for run in para.runs:
            if run.font.size:
                return run.font.size.pt
        if para.style and para.style.font and para.style.font.size:
            return para.style.font.size.pt
        return None

    def para_font_name(para):
        """Run yoksa stil font adını oku; yoksa None döner."""
        for run in para.runs:
            if run.font.name:
                return run.font.name
        if para.style and para.style.font and para.style.font.name:
            return para.style.font.name
        return None
    
    from docx.oxml.ns import qn

    def has_page_break(paragraph):
        """
        Bir paragrafın içinde page break (sayfa sonu) var mı kontrol eder.
        """
        for r in paragraph.runs:
            el = r._element
            for br in el.findall(".//w:br", el.nsmap):
                if br.get(qn("w:type")) == "page":
                    return True
        return False

    # ======================================================
    # PARAGRAF SOL GİRİNTİ (Left Indent) HESAPLAMA
    # ------------------------------------------------------
    # Bu fonksiyon, paragrafın gerçek sol girintisini (cm)
    # hem paragraph_format hem de style_format üzerinden okur.
    # ======================================================
    def effective_left_indent_cm(p):
        try:
            # Öncelikle paragrafın kendi ayarını oku
            if p.paragraph_format.left_indent is not None:
                return p.paragraph_format.left_indent.cm
            # Yoksa stil üzerinden dene
            elif p.style and p.style.paragraph_format.left_indent is not None:
                return p.style.paragraph_format.left_indent.cm
            else:
                return 0.0
        except Exception:
            return 0.0

    # ======================================================
    # ASILI GİRİNTİ (Hanging Indent) HESAPLAMA
    # ------------------------------------------------------
    # Bu fonksiyon, paragrafın asılı (hanging) girintisini cm cinsinden döndürür.
    # ======================================================
    def effective_hanging_indent_cm(p):
        try:
            if p.paragraph_format.first_line_indent is not None:
                val = p.paragraph_format.first_line_indent.cm
            elif p.style and p.style.paragraph_format.first_line_indent is not None:
                val = p.style.paragraph_format.first_line_indent.cm
            else:
                val = 0.0
            # Asılı girintiler negatif olarak geçer, örn. -0.63 cm
            return float(val)
        except Exception:
            return 0.0

    # ======================================================
    # PARAGRAF BİÇİMİ – ETKİN (EFFECTIVE) DEĞER OKUMA
    # ======================================================

    def _resolve_from_styles_parfmt(para, attr_name):
        """
        paragraph_format değeri doğrudan boşsa, stil zincirinden okur.
        attr_name: 'line_spacing', 'space_before', 'space_after', 'first_line_indent', 'left_indent'
        """
        val = getattr(para.paragraph_format, attr_name, None)
        if val is not None:
            return val
        s = para.style
        while s is not None:
            pf = getattr(s, "paragraph_format", None)
            if pf is not None:
                v = getattr(pf, attr_name, None)
                if v is not None:
                    return v
            s = getattr(s, "base_style", None)
        return None


    def effective_line_spacing(para, default=1.0):
        ls = _resolve_from_styles_parfmt(para, "line_spacing")
        if ls is None:
            return float(default)
        if hasattr(ls, "pt"):
            return round(ls.pt, 1)
        try:
            return round(float(ls), 2)
        except Exception:
            return default


    def effective_space_pt(para, which: str) -> float:
        attr = "space_before" if which == "before" else "space_after"
        length = _resolve_from_styles_parfmt(para, attr)
        if length is None:
            return 0.0
        return round(getattr(length, "pt", 0.0), 1)


    def effective_first_line_indent_cm(para) -> float:
        ind = _resolve_from_styles_parfmt(para, "first_line_indent")
        try:
            return 0.0 if ind is None else round(float(ind.cm), 2)
        except Exception:
            return 0.0


    def effective_left_indent_cm(para) -> float:
        ind = _resolve_from_styles_parfmt(para, "left_indent")
        try:
            return 0.0 if ind is None else round(float(ind.cm), 2)
        except Exception:
            return 0.0


    def effective_hanging_indent_cm(para) -> float:
        ind = _resolve_from_styles_parfmt(para, "first_line_indent")
        try:
            return 0.0 if ind is None else round(float(ind.cm), 2)
        except Exception:
            return 0.0


    # ======================================================
    # Yardımcı Fonksiyonlar – Etkin (Effective) Biçim Özellikleri
    # ======================================================

    def effective_alignment(para):
        """Paragrafın etkin hizalamasını (stil kalıtımı dahil) döndürür."""
        a = para.paragraph_format.alignment
        if a is not None:
            return a
        s = para.style
        while s is not None:
            pf = getattr(s, "paragraph_format", None)
            if pf and pf.alignment is not None:
                return pf.alignment
            s = getattr(s, "base_style", None)
        return None

    def effective_bold(para):
        """Paragrafın kalın (bold) olup olmadığını run + stil zincirine göre belirler."""
        any_text = False
        any_bold_true = False
        any_bold_specified = False

        for r in para.runs:
            if not r.text:
                continue
            any_text = True
            if r.font and r.font.bold is not None:
                any_bold_specified = True
                if r.font.bold:
                    any_bold_true = True
            rs = getattr(r, "style", None)
            if rs and rs.font and rs.font.bold:
                any_bold_true = True

        if any_bold_true:
            return True
        if any_bold_specified:
            return False

        # Run'larda belirtilmemişse stil zincirine bak
        s = para.style
        while s is not None:
            if s.font and s.font.bold is not None:
                return bool(s.font.bold)
            s = getattr(s, "base_style", None)
        return False

    def effective_font_name(para):
        """Paragrafın etkin font adını döndürür (run + stil zinciri)."""
        for r in para.runs:
            if r.font and r.font.name:
                return r.font.name
            rs = getattr(r, "style", None)
            if rs and rs.font and rs.font.name:
                return rs.font.name
        s = para.style
        while s is not None:
            if s.font and s.font.name:
                return s.font.name
            s = getattr(s, "base_style", None)
        return None

    def effective_font_size_pt(para):
        """Paragrafın etkin punto değerini döndürür (run + stil zinciri)."""
        for r in para.runs:
            if r.font and r.font.size:
                try:
                    return float(r.font.size.pt)
                except Exception:
                    pass
            rs = getattr(r, "style", None)
            if rs and rs.font and rs.font.size:
                try:
                    return float(rs.font.size.pt)
                except Exception:
                    pass
        s = para.style
        while s is not None:
            if s.font and s.font.size:
                try:
                    return float(s.font.size.pt)
                except Exception:
                    pass
            s = getattr(s, "base_style", None)
        return None

    def effective_italic(para):
        """
        Paragrafın etkin (effective) italik durumunu döndürür.
        Mantık: run -> run style -> paragraf style -> base_style zinciri.
        Not: python-docx'te italic=None "devralıyorum" demektir, False demek değildir.
        """

        any_text = False
        any_italic_true = False
        any_italic_specified = False  # run düzeyinde italic açıkça set edilmiş mi?

        # 1) Run'larda italik açıkça set edilmiş mi bak
        for r in para.runs:
            if not r.text or not r.text.strip():
                continue

            any_text = True

            # Run font italic (True / False / None)
            if r.font and r.font.italic is not None:
                any_italic_specified = True
                if r.font.italic:
                    any_italic_true = True

            # Run'ın kendi style'ı üzerinden italik (varsa)
            rs = getattr(r, "style", None)
            if rs and rs.font and rs.font.italic:
                any_italic_true = True

        # Herhangi bir run italik ise -> italik kabul et
        if any_italic_true:
            return True

        # Run'larda italic açıkça set edilmiş ama hiç True yoksa -> italik değildir
        # (Örn: tüm run'larda italic=False)
        if any_italic_specified:
            return False

        # 2) Run'larda bilgi yoksa paragrafın style zincirine bak
        s = para.style
        while s is not None:
            if s.font and s.font.italic is not None:
                return bool(s.font.italic)
            s = getattr(s, "base_style", None)

        # 3) Hiçbir yerden gelmiyorsa italik değildir
        return False


    # ======================================================
    # 0. Başlık kontrolü
    # ======================================================
    if check["check"] == "heading":
        expected_text = check.get("text", "").lower()
        rule_title = f"Başlık: {expected_text.upper()}"
        for i, para in enumerate(paragraphs):
            if para.text.strip().lower() == expected_text:
                results.append((i, True, rule_title, ""))  # uygun
                break
        else:
            results.append((0, False, rule_title, f"Başlık bulunamadı: {expected_text}"))


    # ======================================================
    # 1-2. GENEL MARJİN KONTROLÜ
    # ======================================================
    elif check["check"] == "margins":
        orientation = check.get("orientation", "portrait")
        expected_top = check.get("top")
        expected_bottom = check.get("bottom")
        expected_left = check.get("left")
        expected_right = check.get("right")

        orientation_map = {"portrait": "Dikey", "landscape": "Yatay"}
        orientation_tr = orientation_map.get(orientation.lower(), orientation)

        rule_title = (
            f"Sayfa Marjinleri ({orientation_tr})\n"
            f"Üst:{expected_top}cm, Alt:{expected_bottom}cm, "
            f"Sol:{expected_left}cm, Sağ:{expected_right}cm"
        )

        errors = []
        for section in doc.sections:
            # None kontrolü - page_width veya page_height None ise portrait varsay
            if section.page_width is None or section.page_height is None:
                is_landscape = False
            else:
                is_landscape = section.page_width > section.page_height
            actual_orientation = "landscape" if is_landscape else "portrait"
            if actual_orientation != orientation:
                continue

            top = round(section.top_margin.cm, 1)
            bottom = round(section.bottom_margin.cm, 1)
            left = round(section.left_margin.cm, 1)
            right = round(section.right_margin.cm, 1)

            if top != expected_top:
                errors.append(f"Üst kenar {top} cm, beklenen {expected_top} cm")
            if bottom != expected_bottom:
                errors.append(f"Alt kenar {bottom} cm, beklenen {expected_bottom} cm")
            if left != expected_left:
                errors.append(f"Sol kenar {left} cm, beklenen {expected_left} cm")
            if right != expected_right:
                errors.append(f"Sağ kenar {right} cm, beklenen {expected_right} cm")

        if errors:
            results.append((0, False, rule_title, "; ".join(errors)))
        else:
            results.append((0, True, rule_title, ""))

    # ======================================================
    # 3. TÜM METİN FONT ADI KONTROLÜ (istisnalı, paragraf bazlı)
    # ======================================================
    elif check["check"] == "font_name_all":

        expected_name = check.get("expected", "Times New Roman")
        allowed_exceptions = check.get("exceptions", [])  # YAML’den gelen istisnalar
        preview_word_count = int(check.get("preview_words", 6))  # kaç kelime yazdırılsın

        rule_title = (
            f"Tüm Metin Yazı Tipi\n"
            f"Beklenen: {expected_name} "
            f"(istisnalar: {', '.join(allowed_exceptions) if allowed_exceptions else 'yok'})"
        )

        def preview_text(text, n_words=6):
            words = (text or "").strip().split()
            if not words:
                return "(boş paragraf)"
            return " ".join(words[:n_words]) + ("..." if len(words) > n_words else "")

        errors = []

        for para in paragraphs:
            para_text = para.text.strip()
            if not para_text:
                continue

            # Paragraf içindeki run’ları tara
            bad_fonts = set()
            for run in para.runs:
                fn = run.font.name
                if fn and fn != expected_name and fn not in allowed_exceptions:
                    bad_fonts.add(fn)

            if bad_fonts:
                preview = preview_text(para_text, preview_word_count)
                fonts_str = ", ".join(sorted(bad_fonts))
                errors.append(
                    f"'{preview}': yazı tipi {fonts_str} olamaz"
                )

        if errors:
            results.append((0, False, rule_title, "; ".join(errors)))
        else:
            results.append((0, True, rule_title, "Tüm metin yazı tipi kurallara uygundur."))

    # ======================================================
    # 4. KAĞIT BOYUTU KONTROLÜ
    # ======================================================
    elif check["check"] == "paper_size":
        expected = check.get("expected", "A4")
        rule_title = f"Kağıt Boyutu\nBeklenen: {expected}"

        page_width = round(doc.sections[0].page_width.cm, 1)
        page_height = round(doc.sections[0].page_height.cm, 1)

        if abs(page_width - 21.0) <= 0.1 and abs(page_height - 29.7) <= 0.1:
            results.append((0, True, rule_title, ""))
        else:
            results.append((0, False, rule_title,
                            f"Bulunan: {page_width} × {page_height} cm, Beklenen: 21.0 × 29.7 cm"))


    # ======================================================
    # 5. SAYFA NUMARALARI (WORD SECTION / PAGE NUMBERING) KONTROLÜ
    # - Paragraf → sayfa eşlemesi yerine Word'ün section "pgNumType" ayarını okur.
    # - Kritik düzeltme #1:
    #   Roman section'ı "ilk lowerRoman/start=1" diye seçmek yanlış olabiliyor.
    #   Çünkü bazı section'larda pgNumType var ama PAGE alanı yok (numara görünmüyor).
    #   Bu yüzden roman section = (lowerRoman,start=1) + (PAGE alanı gerçekten var) olan ilk section.
    # - Kritik düzeltme #2:
    #   Footer/Header "Link to Previous" ile miras alınıyorsa python-docx ilgili section’da boş görünebilir.
    #   Bu yüzden linked zincirini geriye doğru çözerek gerçek footer/header içeriğinde PAGE arıyoruz.
    # ======================================================
    elif check["check"] == "page_numbers" and check.get("enabled", True):
        import os
        import re

        # ---------------- YAML parametreleri ----------------
        main_arabic_from = int(check.get("main_arabic_from", 1))

        preliminaries_as_roman = bool(check.get("preliminaries_as_roman", True))
        prelim_roman_start = int(check.get("prelim_roman_start", 1))
        prelim_roman_fmt = str(check.get("prelim_roman_fmt", "lowerRoman"))

        main_arabic_fmt = str(check.get("main_arabic_fmt", "decimal"))

        # ✅ YAML: sayfa numarası font/punto kontrolü açık mı?
        check_pn_font = bool(check.get("check_page_number_font", True))

        # ✅ (YAML’de zaten vardı) Sayfa numarası yazı tipi kontrolü
        # Örn: font.name: Times New Roman, font.size_pt: 12
        expected_pn_font_name = (((check.get("font") or {}) .get("name")) or "").strip() if check_pn_font else ""
        expected_pn_font_size = (check.get("font") or {}).get("size_pt", None) if check_pn_font else None
        expected_pn_font_size = float(expected_pn_font_size) if expected_pn_font_size is not None else None

        # Debug
        debug_mode = bool(check.get("debug", False))
        debug_file_name = str(check.get("debug_file", "debug_page_numbers.txt")).strip() or "debug_page_numbers.txt"
        DEBUG_F = dbg_path(debug_file_name)

        rule_title = (
            "Sayfa Numaraları (Word Section/Page Numbering)\n"
            f"Beklenti: Ön sayfalar {'küçük roma' if preliminaries_as_roman else 'roman yok'} (i ile), "
            f"Tez metni arap rakamı ({main_arabic_from} ile)"
        )

        # ---------------- Yardımcılar ----------------
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        def _attr(el, local: str):
            """w:fmt / w:start gibi attribute'ları güvenli oku."""
            if el is None:
                return None
            return el.get(f"{{{W_NS}}}{local}")

        def _pgnumtype_of_section(sec):
            """
            Section içindeki <w:pgNumType>:
            - fmt   : decimal, lowerRoman, upperRoman, ... (bazı belgelerde hiç yazılmayabilir)
            - start : 1, 2, ...
            """
            try:
                sectPr = sec._sectPr
                pg = sectPr.xpath("./w:pgNumType")
                if not pg:
                    return {"fmt": None, "start": None}
                pg = pg[0]
                fmt = _attr(pg, "fmt")
                start = _attr(pg, "start")
                try:
                    start_i = int(start) if start is not None else None
                except Exception:
                    start_i = None
                return {"fmt": fmt, "start": start_i}
            except Exception:
                return {"fmt": None, "start": None}

        def _has_page_field_in_element(el) -> bool:
            """
            Header/Footer XML içinde PAGE alanı var mı?
            - fldSimple @w:instr
            - instrText içinde "PAGE" (NUMPAGES hariç)
            """
            try:
                # fldSimple: <w:fldSimple w:instr=" PAGE \* MERGEFORMAT ">
                fld = el.xpath(".//w:fldSimple")
                for f in fld:
                    instr = f.get(f"{{{W_NS}}}instr") or ""
                    up = instr.upper()
                    if "PAGE" in up and "NUMPAGES" not in up:
                        return True

                # instrText: <w:instrText>PAGE \* MERGEFORMAT</w:instrText>
                instrs = el.xpath(".//w:instrText")
                for it in instrs:
                    t = (it.text or "").upper()
                    if "PAGE" in t and "NUMPAGES" not in t:
                        return True

                return False
            except Exception:
                return False

        def _resolve_linked_section_container(sections, idx: int, which: str):
            """
            which ∈ {"footer","header","first_page_footer","first_page_header"}
            Eğer ilgili container linked_to_previous ise geriye doğru gidip ilk linkli olmayanı bulur.
            """
            j = idx
            while j > 0:
                try:
                    container = getattr(sections[j], which)
                    if getattr(container, "is_linked_to_previous", False):
                        j -= 1
                        continue
                    return container
                except Exception:
                    return None
            try:
                return getattr(sections[0], which)
            except Exception:
                return None

        # ✅ Sayfa numarası fontunu yakalamak için: PAGE field’ın geçtiği paragrafta görünen metne en yakın run’ı bul
        def _extract_page_number_run_style(container, doc_obj=None) -> dict | None:
            """
            Dönen: {"font_name": str|None, "font_size_pt": float|None}

            Güçlendirilmiş fallback sırası:
            1) Target run rPr (w:rFonts + w:sz / w:szCs)
            2) Paragraf rPr (w:pPr/w:rPr)
            3) Run karakter stili (w:rStyle) -> styles.xml
            4) Paragraf stili (w:pStyle) -> styles.xml
            5) Belge varsayılanı docDefaults -> styles.xml
            6) Paragraftaki diğer run’larda explicit sz/Fonts var mı?
            """
            try:
                el = container._element
                if not _has_page_field_in_element(el):
                    return None

                # Word font size: w:sz değeri "half-point" (örn 24 => 12pt)
                def _sz_to_pt(sz_val: str | None):
                    try:
                        if sz_val is None:
                            return None
                        v = float(sz_val)
                        return v / 2.0
                    except Exception:
                        return None

                def _read_rpr_font(rpr):
                    """
                    rpr: <w:rPr> element
                    Dönüş: (font_name, font_size_pt)
                    """
                    if rpr is None:
                        return (None, None)

                    font_name = None
                    rfonts = rpr.xpath("./w:rFonts")
                    if rfonts:
                        rf = rfonts[0]
                        font_name = (
                            rf.get(f"{{{W_NS}}}ascii")
                            or rf.get(f"{{{W_NS}}}hAnsi")
                            or rf.get(f"{{{W_NS}}}cs")
                            or None
                        )

                    # size: önce w:sz, yoksa w:szCs
                    size_pt = None
                    sz = rpr.xpath("./w:sz")
                    if sz:
                        size_pt = _sz_to_pt(sz[0].get(f"{{{W_NS}}}val"))
                    if size_pt is None:
                        szcs = rpr.xpath("./w:szCs")
                        if szcs:
                            size_pt = _sz_to_pt(szcs[0].get(f"{{{W_NS}}}val"))

                    return (font_name, size_pt)

                def _get_styles_root():
                    """
                    doc_obj varsa styles.xml root elementini döndür.
                    python-docx'te doc.part.styles.element lxml root'tur.
                    """
                    try:
                        if doc_obj is None:
                            return None
                        return doc_obj.part.styles.element
                    except Exception:
                        return None

                def _style_rpr_from_styles(styles_root, style_id: str):
                    """
                    styles.xml içinde styleId -> w:rPr döndür (yoksa None)
                    """
                    if styles_root is None or not style_id:
                        return None
                    try:
                        nodes = styles_root.xpath(f".//w:style[@w:styleId='{style_id}']/w:rPr")
                        return nodes[0] if nodes else None
                    except Exception:
                        return None

                def _pstyle_rpr_from_styles(styles_root, pstyle_id: str):
                    return _style_rpr_from_styles(styles_root, pstyle_id)

                def _docdefaults_rpr(styles_root):
                    if styles_root is None:
                        return None
                    try:
                        nodes = styles_root.xpath(".//w:docDefaults/w:rPrDefault/w:rPr")
                        return nodes[0] if nodes else None
                    except Exception:
                        return None

                digit_re = re.compile(r"^\s*\d+\s*$")
                roman_re = re.compile(r"^\s*[ivxlcdm]+\s*$", re.IGNORECASE)

                fld_simples = el.xpath(
                    ".//w:fldSimple[contains(translate(@w:instr,'abcdefghijklmnopqrstuvwxyz','ABCDEFGHIJKLMNOPQRSTUVWXYZ'),'PAGE')]"
                )
                instr_texts = el.xpath(
                    ".//w:instrText[contains(translate(text(),'abcdefghijklmnopqrstuvwxyz','ABCDEFGHIJKLMNOPQRSTUVWXYZ'),'PAGE')]"
                )

                candidates_p = []
                for node in fld_simples + instr_texts:
                    try:
                        p_anc = node.xpath("ancestor::w:p[1]")
                        if p_anc:
                            candidates_p.append(p_anc[0])
                    except Exception:
                        continue

                if not candidates_p:
                    candidates_p = el.xpath(".//w:p")

                txbx_ps = el.xpath(".//w:txbxContent//w:p")
                for p in txbx_ps:
                    candidates_p.append(p)

                uniq = []
                seen_ids = set()
                for p in candidates_p:
                    pid = id(p)
                    if pid not in seen_ids:
                        seen_ids.add(pid)
                        uniq.append(p)
                candidates_p = uniq

                styles_root = _get_styles_root()

                for p in candidates_p:
                    ts = p.xpath(".//w:r/w:t")

                    target_r = None
                    for t in ts:
                        txt = (t.text or "").strip()
                        if not txt:
                            continue
                        if digit_re.match(txt) or roman_re.match(txt):
                            r = t.xpath("ancestor::w:r[1]")
                            if r:
                                target_r = r[0]
                                break

                    if target_r is None:
                        r0 = p.xpath(".//w:r[1]")
                        target_r = r0[0] if r0 else None
                    if target_r is None:
                        continue

                    rpr = target_r.xpath("./w:rPr")
                    rpr = rpr[0] if rpr else None
                    font_name, size_pt = _read_rpr_font(rpr)

                    rstyle_id = None
                    try:
                        if rpr is not None:
                            rs = rpr.xpath("./w:rStyle")
                            if rs:
                                rstyle_id = rs[0].get(f"{{{W_NS}}}val") or None
                    except Exception:
                        rstyle_id = None

                    if font_name is None or size_pt is None:
                        try:
                            p_rpr = p.xpath("./w:pPr/w:rPr")
                            p_rpr = p_rpr[0] if p_rpr else None
                            fn2, sz2 = _read_rpr_font(p_rpr)
                            if font_name is None and fn2:
                                font_name = fn2
                            if size_pt is None and sz2 is not None:
                                size_pt = sz2
                        except Exception:
                            pass

                    pstyle_id = None
                    try:
                        ps = p.xpath("./w:pPr/w:pStyle")
                        if ps:
                            pstyle_id = ps[0].get(f"{{{W_NS}}}val") or None
                    except Exception:
                        pstyle_id = None

                    if styles_root is not None:
                        if (font_name is None or size_pt is None) and rstyle_id:
                            srpr = _style_rpr_from_styles(styles_root, rstyle_id)
                            fn3, sz3 = _read_rpr_font(srpr)
                            if font_name is None and fn3:
                                font_name = fn3
                            if size_pt is None and sz3 is not None:
                                size_pt = sz3

                        if (font_name is None or size_pt is None) and pstyle_id:
                            prpr = _pstyle_rpr_from_styles(styles_root, pstyle_id)
                            fn4, sz4 = _read_rpr_font(prpr)
                            if font_name is None and fn4:
                                font_name = fn4
                            if size_pt is None and sz4 is not None:
                                size_pt = sz4

                        if font_name is None or size_pt is None:
                            drpr = _docdefaults_rpr(styles_root)
                            fn5, sz5 = _read_rpr_font(drpr)
                            if font_name is None and fn5:
                                font_name = fn5
                            if size_pt is None and sz5 is not None:
                                size_pt = sz5

                    if size_pt is None:
                        try:
                            any_sz = p.xpath(".//w:rPr/w:sz[1]")
                            if any_sz:
                                size_pt = _sz_to_pt(any_sz[0].get(f"{{{W_NS}}}val"))
                            if size_pt is None:
                                any_szcs = p.xpath(".//w:rPr/w:szCs[1]")
                                if any_szcs:
                                    size_pt = _sz_to_pt(any_szcs[0].get(f"{{{W_NS}}}val"))
                        except Exception:
                            pass

                    if font_name is None:
                        try:
                            any_rf = p.xpath(".//w:rPr/w:rFonts[1]")
                            if any_rf:
                                rf = any_rf[0]
                                font_name = (
                                    rf.get(f"{{{W_NS}}}ascii")
                                    or rf.get(f"{{{W_NS}}}hAnsi")
                                    or rf.get(f"{{{W_NS}}}cs")
                                    or None
                                )
                        except Exception:
                            pass

                    if font_name is None and size_pt is None:
                        continue

                    return {"font_name": font_name, "font_size_pt": size_pt}

                return None

            except Exception:
                return None

        def _font_mismatch_msg(sec_index: int, where: str, pn_style: dict | None) -> str | None:
            if not (expected_pn_font_name or expected_pn_font_size is not None):
                return None

            if pn_style is None:
                return (
                    f"{where} (section {sec_index}) için sayfa numarası yazı tipi/punto tespit edilemedi "
                    f"(Word stil mirası/field yapısı nedeniyle). Debug açıkken pn_style satırına bak."
                )

            seen_name = (pn_style.get("font_name") or "").strip()
            seen_size = pn_style.get("font_size_pt")

            problems = []

            if expected_pn_font_name:
                if (not seen_name) or (seen_name.lower() != expected_pn_font_name.lower()):
                    problems.append(
                        f"yazı tipi '{expected_pn_font_name}' bekleniyordu, görülen '{seen_name or 'None'}'"
                    )

            if expected_pn_font_size is not None:
                if seen_size is None:
                    problems.append(
                        f"punto {expected_pn_font_size:g} bekleniyordu ancak punto değeri XML'den tespit edilemedi "
                        f"(muhtemelen stil mirası; run üzerinde w:sz yok)"
                    )
                else:
                    if abs(float(seen_size) - float(expected_pn_font_size)) > 0.01:
                        problems.append(
                            f"punto {expected_pn_font_size:g} bekleniyordu, görülen {float(seen_size):g}"
                        )

            if not problems:
                return None

            return (
                f"{where} sayfa numarası yazı tipi hatalı: " + ", ".join(problems) +
                f" (section {sec_index})."
            )

        def _section_has_visible_page_number(sections, idx: int) -> bool:
            """
            Bu section’da sayfa numarası gerçekten GÖRÜNÜR mü?
            - footer/header (linked zinciri çözülerek) içinde PAGE alanı aranır.
            - first_page_* (different first page) için de aynı kontrol yapılır.
            """
            for which in ("footer", "header", "first_page_footer", "first_page_header"):
                cont = _resolve_linked_section_container(sections, idx, which)
                if cont is None:
                    continue
                try:
                    if _has_page_field_in_element(cont._element):
                        return True
                except Exception:
                    continue
            return False

        # ---------------- Debug dosyası ----------------
        dbg = None
        if debug_mode:
            new_file = not os.path.exists(DEBUG_F)
            dbg = open(DEBUG_F, "a", encoding="utf-8")
            if new_file:
                dbg.write("[DEBUG] Sayfa Numaraları (Section/Page Numbering)\n")
                dbg.write("===========================================\n\n")
            else:
                dbg.write("\n-------------------------------------------\n")
                dbg.write("Yeni kontrol çalıştırması başlatıldı.\n\n")

        def dlog(msg: str):
            if dbg:
                dbg.write(msg.rstrip() + "\n")

        # ---------------- Bölüm bilgilerini topla ----------------
        sections = list(doc.sections)
        sections_info = []
        for i, sec in enumerate(sections):
            pg = _pgnumtype_of_section(sec)
            has_page = _section_has_visible_page_number(sections, i)

            pn_style = None
            if has_page and (expected_pn_font_name or expected_pn_font_size is not None):
                for which in ("footer", "first_page_footer", "header", "first_page_header"):
                    cont = _resolve_linked_section_container(sections, i, which)
                    if cont is None:
                        continue
                    pn_style = _extract_page_number_run_style(cont, doc_obj=doc)
                    if pn_style:
                        break

            info = {
                "index": i,
                "fmt": pg.get("fmt"),
                "start": pg.get("start"),
                "has_page_field": bool(has_page),
                "pn_style": pn_style,
            }
            sections_info.append(info)

            dlog(f"[SECTION {i}] fmt={info['fmt']}, start={info['start']}, has_page_field={info['has_page_field']}, pn_style={info['pn_style']}")

        memo["page_numbering_sections"] = sections_info

        # ---------------- Beklenen section düzenini kontrol et ----------------
        errors = []

        # 1) ROMAN section'ı bul (senin mevcut mantığını koruyoruz)
        roman_idx = None
        if preliminaries_as_roman:
            for s in sections_info:
                if (s.get("fmt") == prelim_roman_fmt) and (s.get("start") == prelim_roman_start) and s.get("has_page_field"):
                    roman_idx = s["index"]
                    break
            if roman_idx is None:
                for s in sections_info:
                    if (s.get("fmt") == prelim_roman_fmt) and (s.get("start") == prelim_roman_start):
                        roman_idx = s["index"]
                        break
            if roman_idx is None:
                errors.append(
                    f"Ön sayfalar için Roma numaralandırma bölümü bulunamadı "
                    f"(beklenen: fmt='{prelim_roman_fmt}', start={prelim_roman_start})."
                )

        # 2) ANA METİN section'ı bul (senin mevcut mantığını koruyoruz)
        main_idx = None
        search_from = (roman_idx + 1) if (preliminaries_as_roman and roman_idx is not None) else 0

        def _is_decimal_like(fmt):
            # Word’de fmt=None çoğu zaman decimal kabul edilir
            return (fmt == main_arabic_fmt) or (fmt is None and main_arabic_fmt == "decimal")

        for s in sections_info[search_from:]:
            if _is_decimal_like(s.get("fmt")) and s.get("has_page_field"):
                main_idx = s["index"]
                break

        if main_idx is None:
            errors.append("Tez metni için (decimal) sayfa numarası GÖRÜNÜR olan bir section bulunamadı.")

        # 3) Start kontrolü (anlaşılır mesaj)
        if main_idx is not None:
            s = sections_info[main_idx]
            fmt_seen = s.get("fmt")
            start_seen = s.get("start")

            if start_seen != main_arabic_from:
                fmt_txt = "Arap rakamı (decimal)" if _is_decimal_like(fmt_seen) else (fmt_seen or "None")
                errors.append(
                    "Tez metni sayfa numarası başlangıcı beklenenden farklı.\n"
                    f"- Beklenen başlangıç: {main_arabic_from}\n"
                    f"- Görülen başlangıç: {start_seen}\n"
                    f"- Biçim (fmt): {fmt_txt}\n"
                    f"- Tez metni kabul edilen bölüm: section {main_idx}"
                )

        # ✅ 3B) Tez metni başladıktan sonra (section main_idx), ilerideki section’larda numara RESTART/GERİYE gitmiş mi?
        # - Render yapmadan tek tek sayfa numarası doğrulanamaz; ama section pgNumType@start ile "Start at" hatası yakalanır.
        if main_idx is not None:
            last_explicit_start = sections_info[main_idx].get("start")
            # last_explicit_start None ise yine de referans olarak main_arabic_from kullan
            last_explicit_start = int(last_explicit_start) if last_explicit_start is not None else int(main_arabic_from)

            for s in sections_info[main_idx + 1:]:
                if not s.get("has_page_field"):
                    continue
                if not _is_decimal_like(s.get("fmt")):
                    continue

                st = s.get("start")
                if st is None:
                    # "Continue from previous section" → normal
                    continue

                try:
                    st_i = int(st)
                except Exception:
                    continue

                # 1) Yeniden 1'den (veya main_arabic_from'dan) başlatma: tipik section break hatası
                if st_i <= int(main_arabic_from):
                    errors.append(
                        "Tez metni içinde sayfa numarası ileride bir section’da yeniden başlatılmış görünüyor.\n"
                        f"- Beklenti: previous section’dan devam (start=None)\n"
                        f"- Görülen: start={st_i}\n"
                        f"- Olası sebep: yatay/dikey bölüm geçişinde 'Continue from previous section' yerine 'Start at' seçilmiş\n"
                        f"- Problemli bölüm: section {s['index']}"
                    )
                    last_explicit_start = st_i
                    continue

                # 2) Geriye gitme (explicit start küçülüyorsa)
                if st_i < last_explicit_start:
                    errors.append(
                        "Tez metni içinde sayfa numarası sırası bozulmuş (geri gitmiş) görünüyor.\n"
                        f"- Önceki explicit başlangıç: {last_explicit_start}\n"
                        f"- Sonraki explicit başlangıç: {st_i}\n"
                        f"- Problemli bölüm: section {s['index']}"
                    )

                last_explicit_start = st_i

        # ✅ 4) FONT kontrolü (senin mevcut mantığını koruyoruz)
        def _font_mismatch_msg(sec_index: int, where: str, pn_style: dict | None) -> str | None:
            if not (expected_pn_font_name or expected_pn_font_size is not None):
                return None
            if pn_style is None:
                return (
                    f"{where} (section {sec_index}) için sayfa numarası yazı tipi tespit edilemedi "
                    f"(Word stil mirası/field yapısı nedeniyle). Debug açıkken pn_style satırına bak."
                )

            seen_name = (pn_style.get("font_name") or "").strip()
            seen_size = pn_style.get("font_size_pt")

            problems = []
            if expected_pn_font_name:
                if (not seen_name) or (seen_name.lower() != expected_pn_font_name.lower()):
                    problems.append(f"yazı tipi '{expected_pn_font_name}' bekleniyordu, görülen '{seen_name or 'None'}'")
            if expected_pn_font_size is not None:
                if seen_size is None or abs(float(seen_size) - float(expected_pn_font_size)) > 0.01:
                    problems.append(f"punto {expected_pn_font_size:g} bekleniyordu, görülen '{seen_size if seen_size is not None else 'None'}'")

            if not problems:
                return None

            return (
                f"{where} sayfa numarası yazı tipi hatalı: " + ", ".join(problems) +
                f" (section {sec_index})."
            )

        if preliminaries_as_roman and roman_idx is not None:
            msg = _font_mismatch_msg(roman_idx, "Ön sayfalar (Roma)", sections_info[roman_idx].get("pn_style"))
            if msg:
                errors.append(msg)

        if main_idx is not None:
            msg = _font_mismatch_msg(main_idx, "Tez metni (Arap rakamı)", sections_info[main_idx].get("pn_style"))
            if msg:
                errors.append(msg)

        if dbg:
            dbg.close()

        if errors:
            return [(0, False, rule_title, "\n- " + "\n- ".join(errors))]
        else:
            return [(0, True, rule_title, "")]


    # ===============================================================================================================#
    # ===============================================================================================================#    
    # ===============================================================================================================#
    # ===============================================================================================================#
    # İÇ KAPAK SAYFASI - "T.C." BAŞLIĞI KONTROLÜ
    # ======================================================
    elif check["check"] == "inner_cover_heading_tc":
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 16))
        expected_bold = check.get("bold", True)
        expected_align = check.get("alignment", "center").lower()
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        expected_text = check.get("expected_text", "T.C.")
        markers = check.get("markers", ["T.C."])
        debug_mode = check.get("debug", False)

        rule_title = check.get(
            "description",
            f"İç Kapak - 'T.C.' Başlığı: {expected_name}, {int(expected_size)} punto, "
            f"kalın={expected_bold}, hizalama={expected_align}, "
            f"satır aralığı={expected_spacing}, önce={int(expected_before)}, sonra={int(expected_after)}"
        )

        import re, os

        # ===========================================================
        # 1️⃣ DEBUG dosyası (tek dosya, append)
        # ===========================================================
        debug_file = None
        if debug_mode:
            debug_path = dbg_path("debug_inner_cover.txt")
            new_file = not os.path.exists(debug_path)
            debug_file = open(debug_path, "a", encoding="utf-8")
            if new_file:
                debug_file.write("[DEBUG] İç Kapak Sayfası Kontrolleri\n")
                debug_file.write("===========================================\n\n")
            else:
                debug_file.write("\n-------------------------------------------\n")
                debug_file.write("Yeni kontrol çalıştırması başlatıldı.\n\n")

        # ===========================================================
        # 2️⃣ "T.C." başlığını bulma
        # ===========================================================
        # Regex güvenliği için escape
        marker_patterns = [re.escape(m.strip()) for m in markers]
        pattern = r"^\s*(" + "|".join(marker_patterns) + r")\s*$"

        start_idx = None
        for i, p in enumerate(paragraphs):
            if re.match(pattern, p.text.strip(), re.IGNORECASE):
                start_idx = i
                break

        if start_idx is None:
            results.append((0, False, rule_title, "'T.C.' başlığı bulunamadı"))
            if debug_file:
                debug_file.write("❌ 'T.C.' başlığı bulunamadı.\n")
                debug_file.close()
            return results

        # ===========================================================
        # 3️⃣ Paragraf bilgileri
        # ===========================================================
        p = paragraphs[start_idx]

        if debug_file:
            debug_file.write(f"✅ Bulundu: paragraf index = {start_idx}\n")
            debug_file.write(f"Metin içeriği: '{p.text.strip()}'\n\n")

        errors = []

        # 3A) Metin içeriği kontrolü (bulunsa bile tam değilse ihlal)
        found_text = p.text.strip()
        if found_text.upper() != expected_text.strip().upper():
            errors.append(f"Metin '{found_text}' yerine '{expected_text}' olmalı")

        # ===========================================================
        # 4️⃣ Yazı tipi (font name) kontrolü
        # ===========================================================
        fn = None
        for r in p.runs:
            if r.font and r.font.name:
                fn = r.font.name
                break
        # Stil devralımı kontrolü
        style_font = getattr(p.style, "font", None)
        style_name = style_font.name if style_font and style_font.name else None

        actual_font = fn or style_name
        if actual_font and actual_font != expected_name:
            errors.append(f"Yazı tipi {actual_font} yerine {expected_name} olmalı")

        # ===========================================================
        # 5️⃣ Punto (font size) kontrolü
        # ===========================================================
        fs = None
        for r in p.runs:
            if r.font and r.font.size:
                fs = r.font.size.pt
                break
        style_size = style_font.size.pt if style_font and style_font.size else None
        actual_size = fs or style_size
        if actual_size and round(float(actual_size), 1) != round(expected_size, 1):
            errors.append(f"Punto {actual_size} yerine {expected_size} olmalı")

        # ===========================================================
        # 6️⃣ Kalınlık (bold) kontrolü (run veya stil üzerinden)
        # ===========================================================
        has_bold_run = any(r.bold for r in p.runs if r.text.strip())
        style_bold = style_font.bold if style_font else None
        if expected_bold:
            if not (has_bold_run or style_bold):
                errors.append("'T.C.' kalın değil")

        # ===========================================================
        # 7️⃣ Hizalama kontrolü (paragraf veya stil üzerinden)
        # ===========================================================
        align_map = {
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }
        expected_align_enum = align_map.get(expected_align, WD_PARAGRAPH_ALIGNMENT.CENTER)
        style_align = getattr(p.style.paragraph_format, "alignment", None)

        # Ortalı kabul durumu: doğrudan CENTER veya stil CENTER
        if p.alignment not in [expected_align_enum, None] and style_align != expected_align_enum:
            errors.append(f"Hizalama '{p.alignment}' yerine '{expected_align}' olmalı")

        # ===========================================================
        # 8️⃣ Satır aralığı ve paragraf boşlukları
        # ===========================================================
        pf = p.paragraph_format
        ls = pf.line_spacing or getattr(p.style.paragraph_format, "line_spacing", None)
        sb = (pf.space_before.pt if pf.space_before else None) or (
            getattr(p.style.paragraph_format, "space_before", None).pt
            if getattr(p.style.paragraph_format, "space_before", None)
            else 0.0
        )
        sa = (pf.space_after.pt if pf.space_after else None) or (
            getattr(p.style.paragraph_format, "space_after", None).pt
            if getattr(p.style.paragraph_format, "space_after", None)
            else 0.0
        )

        if ls is not None and round(float(ls), 2) != expected_spacing:
            errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
        if round(sb, 1) != round(expected_before, 1):
            errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
        if round(sa, 1) != round(expected_after, 1):
            errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

        # ===========================================================
        # 9️⃣ DEBUG bilgilerini yaz
        # ===========================================================
        if debug_file:
            debug_file.write("Biçimsel Özellikler:\n")
            debug_file.write(f"  Yazı tipi: {actual_font}\n")
            debug_file.write(f"  Punto: {actual_size}\n")
            debug_file.write(f"  Kalın (run/stil): {has_bold_run or style_bold}\n")
            debug_file.write(f"  Hizalama (run/stil): {p.alignment or style_align}\n")
            debug_file.write(f"  Satır aralığı: {ls}\n")
            debug_file.write(f"  Önce: {sb} pt, Sonra: {sa} pt\n")

        # ===========================================================
        # 🔟 Sonuç değerlendirmesi
        # ===========================================================
        if errors:
            results.append((start_idx, False, rule_title, "; ".join(errors)))
            if debug_file:
                debug_file.write("\n❌ Hatalar:\n")
                for err in errors:
                    debug_file.write(f" - {err}\n")
        else:
            results.append((start_idx, True, rule_title, ""))
            if debug_file:
                debug_file.write("\n✅ Tüm kontroller başarıyla geçti.\n")

        # ===========================================================
        # 11️⃣ Bulunan satırı hafızaya kaydet (memo)
        # ===========================================================
        memo["inner_cover_tc_index"] = start_idx

        if debug_file:
            debug_file.write("\n-------------------------------------------\n")
            debug_file.write("Kontrol tamamlandı.\n")
            debug_file.close()
        return results
    # ======================================================
    # İÇ KAPAK SAYFASI - "OSMANİYE KORKUT ATA ÜNİVERSİTESİ" KONTROLÜ
    # ======================================================
    elif check["check"] == "inner_cover_university_name":
        expected_text = "OSMANİYE KORKUT ATA ÜNİVERSİTESİ"
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 16))
        expected_bold = check.get("bold", True)
        expected_align = check.get("alignment", "center").lower()
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        debug_mode = check.get("debug", False)

        rule_title = check.get(
            "description",
            f"İç Kapak - Üniversite Adı: {expected_name}, {int(expected_size)} punto, "
            f"kalın={expected_bold}, hizalama={expected_align}, "
            f"satır aralığı={expected_spacing}, önce={int(expected_before)}, sonra={int(expected_after)}"
        )

        import os

        # ===========================================================
        # 1️⃣ DEBUG dosyası (tek dosya, append)
        # ===========================================================
        debug_file = None
        if debug_mode:
            debug_path = dbg_path("debug_inner_cover.txt")
            new_file = not os.path.exists(debug_path)
            debug_file = open(debug_path, "a", encoding="utf-8")
            if new_file:
                debug_file.write("[DEBUG] İç Kapak - Üniversite Adı Kontrolleri\n")
                debug_file.write("===========================================\n\n")
            else:
                debug_file.write("\n-------------------------------------------\n")
                debug_file.write("Yeni kontrol çalıştırması başlatıldı.\n\n")

        # ===========================================================
        # 2️⃣ "T.C." başlığı bulunmuş olmalı
        # ===========================================================
        tc_idx = memo.get("inner_cover_tc_index", None)
        if tc_idx is None or tc_idx + 1 >= len(paragraphs):
            results.append((0, False, rule_title, "'T.C.' başlığı bulunamadı veya alt satır eksik"))
            if debug_file:
                debug_file.write("❌ 'T.C.' başlığı bulunamadı veya alt satır mevcut değil.\n")
                debug_file.close()
            return results

        # ===========================================================
        # 3️⃣ 'T.C.' sonrası: boş satırları geç, ilk dolu satırı üniversite adı kabul et (DİNAMİK OFFSET)
        # ===========================================================
        i = tc_idx + 1
        while i < len(paragraphs) and not paragraphs[i].text.strip():
            i += 1

        if i >= len(paragraphs):
            results.append((0, False, rule_title, "Üniversite adı satırı bulunamadı (T.C. sonrası dolu satır yok)"))
            if debug_file:
                debug_file.write("❌ Üniversite adı satırı bulunamadı (T.C. sonrası dolu satır yok).\n")
                debug_file.close()
            return results

        p = paragraphs[i]
        if debug_file:
            debug_file.write(f"✅ Üniversite satırı index = {i}\n")
            debug_file.write(f"Metin içeriği: '{p.text.strip()}'\n\n")

        errors = []

        # ===========================================================
        # 4️⃣ Metin kontrolü
        # ===========================================================
        if p.text.strip().upper() != expected_text:
            errors.append(f"Metin '{p.text.strip()}' yerine '{expected_text}' olmalı")

        # ===========================================================
        # 5️⃣ Stil ve biçim kontrolleri (stilden devralım dahil)
        # ===========================================================
        # --- Font Adı ---
        fn = None
        for r in p.runs:
            if r.font and r.font.name:
                fn = r.font.name
                break
        style_font = getattr(p.style, "font", None)
        style_name = style_font.name if style_font and style_font.name else None
        actual_font = fn or style_name
        if actual_font and actual_font != expected_name:
            errors.append(f"Yazı tipi {actual_font} yerine {expected_name} olmalı")

        # --- Punto ---
        fs = None
        for r in p.runs:
            if r.font and r.font.size:
                fs = r.font.size.pt
                break
        style_size = style_font.size.pt if style_font and style_font.size else None
        actual_size = fs or style_size
        if actual_size and round(float(actual_size), 1) != round(expected_size, 1):
            errors.append(f"Punto {actual_size} yerine {expected_size} olmalı")

        # --- Kalınlık (run veya stil üzerinden) ---
        has_bold_run = any(r.bold for r in p.runs if r.text.strip())
        style_bold = style_font.bold if style_font else None
        if expected_bold:
            if not (has_bold_run or style_bold):
                errors.append("Üniversite adı kalın değil")

        # --- Hizalama (paragraf veya stil üzerinden) ---
        align_map = {
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }
        expected_align_enum = align_map.get(expected_align, WD_PARAGRAPH_ALIGNMENT.CENTER)
        style_align = getattr(p.style.paragraph_format, "alignment", None)

        if p.alignment not in [expected_align_enum, None] and style_align != expected_align_enum:
            errors.append(f"Hizalama '{p.alignment}' yerine '{expected_align}' olmalı")

        # --- Satır aralığı ve boşluklar ---
        pf = p.paragraph_format
        ls = pf.line_spacing or getattr(p.style.paragraph_format, "line_spacing", None)
        sb = (pf.space_before.pt if pf.space_before else None) or (
            getattr(p.style.paragraph_format, "space_before", None).pt
            if getattr(p.style.paragraph_format, "space_before", None)
            else 0.0
        )
        sa = (pf.space_after.pt if pf.space_after else None) or (
            getattr(p.style.paragraph_format, "space_after", None).pt
            if getattr(p.style.paragraph_format, "space_after", None)
            else 0.0
        )

        if ls is not None and round(float(ls), 2) != expected_spacing:
            errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
        if round(sb, 1) != round(expected_before, 1):
            errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
        if round(sa, 1) != round(expected_after, 1):
            errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

        # ===========================================================
        # 6️⃣ DEBUG yazımı
        # ===========================================================
        if debug_file:
            debug_file.write("Biçimsel Özellikler:\n")
            debug_file.write(f"  Yazı tipi: {actual_font}\n")
            debug_file.write(f"  Punto: {actual_size}\n")
            debug_file.write(f"  Kalın (run/stil): {has_bold_run or style_bold}\n")
            debug_file.write(f"  Hizalama (run/stil): {p.alignment or style_align}\n")
            debug_file.write(f"  Satır aralığı: {ls}\n")
            debug_file.write(f"  Önce: {sb} pt, Sonra: {sa} pt\n")

        # ===========================================================
        # 🔟 Sonuç değerlendirmesi
        # ===========================================================
        if errors:
            results.append((i, False, rule_title, "; ".join(errors)))
            if debug_file:
                debug_file.write("\n❌ Hatalar:\n")
                for err in errors:
                    debug_file.write(f" - {err}\n")
        else:
            results.append((i, True, rule_title, ""))
            if debug_file:
                debug_file.write("\n✅ Tüm kontroller başarıyla geçti.\n")

        if debug_file:
            debug_file.write("\n-------------------------------------------\n")
            debug_file.write("Kontrol tamamlandı.\n")
            debug_file.close()

        # 11️⃣ Bulunan satırı hafızaya kaydet (memo)
        memo["inner_cover_university_index"] = i
        return results
    # ======================================================
    # İÇ KAPAK SAYFASI - "LİSANSÜSTÜ EĞİTİM ENSTİTÜSÜ" KONTROLÜ
    # ======================================================
    elif check["check"] == "inner_cover_institute_name":
        expected_text = "LİSANSÜSTÜ EĞİTİM ENSTİTÜSÜ"
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 16))
        expected_bold = check.get("bold", True)
        expected_align = check.get("alignment", "center").lower()
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        debug_mode = check.get("debug", False)

        rule_title = check.get(
            "description",
            f"İç Kapak - Enstitü Adı: {expected_name}, {int(expected_size)} punto, "
            f"kalın={expected_bold}, hizalama={expected_align}, "
            f"satır aralığı={expected_spacing}, önce={int(expected_before)}, sonra={int(expected_after)}"
        )

        import os

        # ===========================================================
        # 1️⃣ DEBUG dosyası (tek dosya, append)
        # ===========================================================
        debug_file = None
        if debug_mode:
            debug_path = dbg_path("debug_inner_cover.txt")
            new_file = not os.path.exists(debug_path)
            debug_file = open(debug_path, "a", encoding="utf-8")
            if new_file:
                debug_file.write("[DEBUG] İç Kapak - Enstitü Adı Kontrolleri\n")
                debug_file.write("===========================================\n\n")
            else:
                debug_file.write("\n-------------------------------------------\n")
                debug_file.write("Yeni kontrol çalıştırması başlatıldı (Lisansüstü Eğitim Enstitüsü).\n\n")

        # ===========================================================
        # 2️⃣ Üniversite adı bulunmuş olmalı
        # ===========================================================
        uni_idx = memo.get("inner_cover_university_index", None)
        if uni_idx is None or uni_idx + 1 >= len(paragraphs):
            results.append((0, False, rule_title, "Üniversite adı satırı bulunamadı veya alt satır eksik"))
            if debug_file:
                debug_file.write("❌ Üniversite adı bulunamadı veya alt satır mevcut değil.\n")
                debug_file.close()
            return results

        # ===========================================================
        # 3️⃣ Üniversite sonrası: boş satırları geç, ilk dolu satırı enstitü adı kabul et (DİNAMİK OFFSET)
        # ===========================================================
        i = uni_idx + 1
        while i < len(paragraphs) and not paragraphs[i].text.strip():
            i += 1

        if i >= len(paragraphs):
            results.append((0, False, rule_title, "Enstitü adı satırı bulunamadı (Üniversite sonrası dolu satır yok)"))
            if debug_file:
                debug_file.write("❌ Enstitü adı satırı bulunamadı (Üniversite sonrası dolu satır yok).\n")
                debug_file.close()
            return results

        p = paragraphs[i]
        if debug_file:
            debug_file.write(f"✅ Enstitü satırı index = {i}\n")
            debug_file.write(f"Metin içeriği: '{p.text.strip()}'\n\n")

        errors = []

        # ===========================================================
        # 4️⃣ Metin kontrolü
        # ===========================================================
        if p.text.strip().upper() != expected_text:
            errors.append(f"Metin '{p.text.strip()}' yerine '{expected_text}' olmalı")

        # ===========================================================
        # 5️⃣ Stil + biçim kontrolleri (stilden devralım dahil)
        # ===========================================================
        fn = None
        for r in p.runs:
            if r.font and r.font.name:
                fn = r.font.name
                break

        style_font = getattr(p.style, "font", None)
        style_name = style_font.name if style_font and style_font.name else None
        actual_font = fn or style_name
        if actual_font and actual_font != expected_name:
            errors.append(f"Yazı tipi {actual_font} yerine {expected_name} olmalı")

        fs = None
        for r in p.runs:
            if r.font and r.font.size:
                fs = r.font.size.pt
                break
        style_size = style_font.size.pt if style_font and style_font.size else None
        actual_size = fs or style_size
        if actual_size and round(float(actual_size), 1) != round(expected_size, 1):
            errors.append(f"Punto {actual_size} yerine {expected_size} olmalı")

        has_bold_run = any(r.bold for r in p.runs if r.text.strip())
        style_bold = style_font.bold if style_font else None
        if expected_bold:
            if not (has_bold_run or style_bold):
                errors.append("Enstitü adı kalın değil")

        align_map = {
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }
        expected_align_enum = align_map.get(expected_align, WD_PARAGRAPH_ALIGNMENT.CENTER)
        style_align = getattr(p.style.paragraph_format, "alignment", None)

        if p.alignment not in [expected_align_enum, None] and style_align != expected_align_enum:
            errors.append(f"Hizalama '{p.alignment}' yerine '{expected_align}' olmalı")

        pf = p.paragraph_format
        ls = pf.line_spacing or getattr(p.style.paragraph_format, "line_spacing", None)
        sb = (pf.space_before.pt if pf.space_before else None) or (
            getattr(p.style.paragraph_format, "space_before", None).pt
            if getattr(p.style.paragraph_format, "space_before", None)
            else 0.0
        )
        sa = (pf.space_after.pt if pf.space_after else None) or (
            getattr(p.style.paragraph_format, "space_after", None).pt
            if getattr(p.style.paragraph_format, "space_after", None)
            else 0.0
        )

        if ls is not None and round(float(ls), 2) != expected_spacing:
            errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
        if round(sb, 1) != round(expected_before, 1):
            errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
        if round(sa, 1) != round(expected_after, 1):
            errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

        # ===========================================================
        # 6️⃣ DEBUG bilgileri
        # ===========================================================
        if debug_file:
            debug_file.write("Biçimsel Özellikler:\n")
            debug_file.write(f"  Yazı tipi: {actual_font}\n")
            debug_file.write(f"  Punto: {actual_size}\n")
            debug_file.write(f"  Kalın (run/stil): {has_bold_run or style_bold}\n")
            debug_file.write(f"  Hizalama (run/stil): {p.alignment or style_align}\n")
            debug_file.write(f"  Satır aralığı: {ls}\n")
            debug_file.write(f"  Önce: {sb} pt, Sonra: {sa} pt\n")

        # ===========================================================
        # 7️⃣ Sonuç değerlendirmesi
        # ===========================================================
        if errors:
            results.append((i, False, rule_title, "; ".join(errors)))
            if debug_file:
                debug_file.write("\n❌ Hatalar:\n")
                for err in errors:
                    debug_file.write(f" - {err}\n")
        else:
            results.append((i, True, rule_title, ""))
            if debug_file:
                debug_file.write("\n✅ Tüm kontroller başarıyla geçti.\n")

        # ===========================================================
        # 8️⃣ Bulunan satırı hafızaya kaydet (memo)
        # ===========================================================
        memo["inner_cover_institute_index"] = i

        # ===========================================================
        # 9️⃣ Dosya kapanışı ve return
        # ===========================================================
        if debug_file:
            debug_file.write("\n-------------------------------------------\n")
            debug_file.write("Kontrol tamamlandı.\n")
            debug_file.close()

        return results

    # ======================================================
    # İÇ KAPAK SAYFASI - "ANA BİLİM DALI" SATIRI KONTROLÜ
    # ======================================================
    elif check["check"] == "inner_cover_department_name":
        expected_suffix = "ANA BİLİM DALI"
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 16))
        expected_bold = check.get("bold", True)
        expected_align = check.get("alignment", "center").lower()
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        debug_mode = check.get("debug", False)

        rule_title = check.get(
            "description",
            f"İç Kapak - Ana Bilim Dalı Satırı: {expected_name}, {int(expected_size)} punto, "
            f"kalın={expected_bold}, hizalama={expected_align}, "
            f"satır aralığı={expected_spacing}, önce={int(expected_before)}, sonra={int(expected_after)}, "
            f"ve metin '{expected_suffix}' ile bitmeli"
        )

        import os, re

        # ===========================================================
        # 1️⃣ DEBUG dosyası (tek dosya, append)
        # ===========================================================
        debug_file = None
        if debug_mode:
            debug_path = dbg_path("debug_inner_cover.txt")
            new_file = not os.path.exists(debug_path)
            debug_file = open(debug_path, "a", encoding="utf-8")
            if new_file:
                debug_file.write("[DEBUG] İç Kapak - Ana Bilim Dalı Kontrolleri\n")
                debug_file.write("===========================================\n\n")
            else:
                debug_file.write("\n-------------------------------------------\n")
                debug_file.write("Yeni kontrol çalıştırması başlatıldı (Ana Bilim Dalı).\n\n")

        # ===========================================================
        # 2️⃣ Enstitü satırı bulunmuş olmalı
        # ===========================================================
        inst_idx = memo.get("inner_cover_institute_index", None)
        if inst_idx is None or inst_idx + 1 >= len(paragraphs):
            results.append((0, False, rule_title, "Enstitü adı satırı bulunamadı veya alt satır eksik"))
            if debug_file:
                debug_file.write("❌ Enstitü adı bulunamadı veya alt satır mevcut değil.\n")
                debug_file.close()
            return results

        # ===========================================================
        # 3️⃣ Enstitü sonrası: boş satırları geç, ilk dolu satırı Ana Bilim Dalı kabul et (DİNAMİK OFFSET)
        # ===========================================================
        i = inst_idx + 1
        while i < len(paragraphs) and not paragraphs[i].text.strip():
            i += 1

        if i >= len(paragraphs):
            results.append((0, False, rule_title, "Ana Bilim Dalı satırı bulunamadı (Enstitü sonrası dolu satır yok)"))
            if debug_file:
                debug_file.write("❌ Ana Bilim Dalı satırı bulunamadı (Enstitü sonrası dolu satır yok).\n")
                debug_file.close()
            return results

        p = paragraphs[i]
        text = p.text.strip().upper()
        if debug_file:
            debug_file.write(f"✅ Ana Bilim Dalı satırı index = {i}\n")
            debug_file.write(f"Metin içeriği: '{text}'\n\n")

        errors = []

        # ===========================================================
        # 4️⃣ Metin kontrolü (sonu ANA BİLİM DALI ile bitmeli)
        # ===========================================================
        if not text.endswith(expected_suffix):
            errors.append(f"Metin '{expected_suffix}' ifadesiyle bitmiyor")

        # ===========================================================
        # 5️⃣ Biçimsel kontroller (stil devralımı dahil)
        # ===========================================================
        # --- Yazı tipi ---
        fn = None
        for r in p.runs:
            if r.font and r.font.name:
                fn = r.font.name
                break

        style_font = getattr(p.style, "font", None)
        style_name = style_font.name if style_font and style_font.name else None
        actual_font = fn or style_name
        if actual_font and actual_font != expected_name:
            errors.append(f"Yazı tipi {actual_font} yerine {expected_name} olmalı")

        # --- Punto ---
        fs = None
        for r in p.runs:
            if r.font and r.font.size:
                fs = r.font.size.pt
                break
        style_size = style_font.size.pt if style_font and style_font.size else None
        actual_size = fs or style_size
        if actual_size and round(float(actual_size), 1) != round(expected_size, 1):
            errors.append(f"Punto {actual_size} yerine {expected_size} olmalı")

        # --- Kalınlık ---
        has_bold_run = any(r.bold for r in p.runs if r.text.strip())
        style_bold = style_font.bold if style_font else None
        if expected_bold and not (has_bold_run or style_bold):
            errors.append("Ana Bilim Dalı satırı kalın değil")

        # --- Hizalama ---
        align_map = {
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }
        expected_align_enum = align_map.get(expected_align, WD_PARAGRAPH_ALIGNMENT.CENTER)
        style_align = getattr(p.style.paragraph_format, "alignment", None)

        if p.alignment not in [expected_align_enum, None] and style_align != expected_align_enum:
            errors.append(f"Hizalama '{p.alignment}' yerine '{expected_align}' olmalı")

        # --- Satır aralığı ve boşluklar ---
        pf = p.paragraph_format
        ls = pf.line_spacing or getattr(p.style.paragraph_format, "line_spacing", None)
        sb = (pf.space_before.pt if pf.space_before else None) or (
            getattr(p.style.paragraph_format, "space_before", None).pt
            if getattr(p.style.paragraph_format, "space_before", None)
            else 0.0
        )
        sa = (pf.space_after.pt if pf.space_after else None) or (
            getattr(p.style.paragraph_format, "space_after", None).pt
            if getattr(p.style.paragraph_format, "space_after", None)
            else 0.0
        )

        if ls is not None and round(float(ls), 2) != expected_spacing:
            errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
        if round(sb, 1) != round(expected_before, 1):
            errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
        if round(sa, 1) != round(expected_after, 1):
            errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

        # ===========================================================
        # 6️⃣ DEBUG bilgileri
        # ===========================================================
        if debug_file:
            debug_file.write("Biçimsel Özellikler:\n")
            debug_file.write(f"  Yazı tipi: {actual_font}\n")
            debug_file.write(f"  Punto: {actual_size}\n")
            debug_file.write(f"  Kalın (run/stil): {has_bold_run or style_bold}\n")
            debug_file.write(f"  Hizalama (run/stil): {p.alignment or style_align}\n")
            debug_file.write(f"  Satır aralığı: {ls}\n")
            debug_file.write(f"  Önce: {sb} pt, Sonra: {sa} pt\n")

        # ===========================================================
        # 7️⃣ Sonuç değerlendirmesi
        # ===========================================================
        if errors:
            results.append((i, False, rule_title, "; ".join(errors)))
            if debug_file:
                debug_file.write("\n❌ Hatalar:\n")
                for err in errors:
                    debug_file.write(f" - {err}\n")
        else:
            results.append((i, True, rule_title, ""))
            if debug_file:
                debug_file.write("\n✅ Tüm kontroller başarıyla geçti.\n")

        # ===========================================================
        # 8️⃣ Bulunan satırı hafızaya kaydet (memo)
        # ===========================================================
        memo["inner_cover_department_index"] = i

        if debug_file:
            debug_file.write("\n-------------------------------------------\n")
            debug_file.write("Kontrol tamamlandı.\n")
            debug_file.close()

        return results

    # ======================================================
    # İÇ KAPAK SAYFASI - "ANA BİLİM DALI" SONRASI BOŞ SATIR SAYISI VE BİÇİM KONTROLÜ
    # ======================================================
    elif check["check"] == "inner_cover_spacing_after_department":
        min_blank = int(check.get("min_blank_lines", 5))
        max_blank = int(check.get("max_blank_lines", 6))
        expected_font = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 16))
        # expected_bold = check.get("bold", True)   # ❌ BOLD KONTROLÜ KALDIRILDI
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        debug_mode = check.get("debug", False)

        rule_title = check.get(
            "description",
            f"Ana Bilim Dalı satırından sonra {min_blank}–{max_blank} satır boşluk olmalı. "
            f"Her boş satır {expected_font}, {expected_size} pt, "
            f"satır aralığı={expected_spacing}, önce={expected_before}, sonra={expected_after} olmalı."
        )

        import os

        # ===========================================================
        # 1️⃣ DEBUG dosyası (tek dosya, append)
        # ===========================================================
        debug_file = None
        if debug_mode:
            debug_path = dbg_path("debug_inner_cover.txt")
            new_file = not os.path.exists(debug_path)
            debug_file = open(debug_path, "a", encoding="utf-8")
            if new_file:
                debug_file.write("[DEBUG] İç Kapak - Ana Bilim Dalı Sonrası Boşluk ve Biçim Kontrolleri\n")
                debug_file.write("===========================================\n\n")
            else:
                debug_file.write("\n-------------------------------------------\n")
                debug_file.write("Yeni kontrol çalıştırması başlatıldı (Boşluk + Biçim).\n\n")

        # ===========================================================
        # 2️⃣ Ana Bilim Dalı satırı bulunmuş olmalı
        # ===========================================================
        dep_idx = memo.get("inner_cover_department_index", None)
        if dep_idx is None or dep_idx + min_blank >= len(paragraphs):
            results.append((0, False, rule_title, "Ana Bilim Dalı satırı bulunamadı veya sonrası eksik"))
            if debug_file:
                debug_file.write("❌ Ana Bilim Dalı satırı bulunamadı veya sonrası eksik.\n")
                debug_file.close()
            return results

        # ===========================================================
        # 3️⃣ Boş satır sayısı ve biçimsel özellikleri kontrol et
        # ===========================================================
        blank_count = 0
        errors = []

        for j in range(1, max_blank + 2):
            idx = dep_idx + j
            if idx >= len(paragraphs):
                break
            p = paragraphs[idx]
            text = p.text.strip()

            # Metin varsa boşluk biter
            if text:
                break

            blank_count += 1

            # Biçimsel kontroller
            pf = p.paragraph_format
            ls = pf.line_spacing or getattr(p.style.paragraph_format, "line_spacing", None)

            # space_before/after (sadece pf varsa, yoksa 0.0)
            sb = (pf.space_before.pt if pf.space_before else 0.0)
            sa = (pf.space_after.pt if pf.space_after else 0.0)

            # -------------------------------
            # ✅ font adı / punto okuma:
            # run → XML(pPr/rPr) → stil
            # -------------------------------
            fn = None
            fs = None

            # 1) Run'dan oku
            for r in p.runs:
                if r.font:
                    if not fn and r.font.name:
                        fn = r.font.name
                    if fs is None and r.font.size:
                        fs = r.font.size.pt

            # 2) Run yoksa / run'da yoksa: XML pPr/rPr
            if fs is None or fn is None:
                try:
                    if fs is None:
                        sz_vals = p._p.xpath("./w:pPr/w:rPr/w:sz/@w:val")
                        if sz_vals:
                            fs = float(int(sz_vals[0])) / 2.0

                    if fn is None:
                        font_vals = p._p.xpath("./w:pPr/w:rPr/w:rFonts/@w:ascii")
                        if not font_vals:
                            font_vals = p._p.xpath("./w:pPr/w:rPr/w:rFonts/@w:hAnsi")
                        if font_vals:
                            fn = font_vals[0]
                except Exception:
                    pass

            # 3) Stil fallback
            style_font = getattr(p.style, "font", None)
            style_name = style_font.name if style_font and style_font.name else None
            style_size = style_font.size.pt if style_font and style_font.size else None

            actual_font = fn or style_name
            actual_size = fs if fs is not None else style_size

            # Hatalar
            if actual_font and actual_font != expected_font:
                errors.append(f"{j}. boş satırda yazı tipi {actual_font} yerine {expected_font} olmalı")
            if actual_size is not None and round(float(actual_size), 1) != round(expected_size, 1):
                errors.append(f"{j}. boş satırda punto {actual_size} yerine {expected_size} olmalı")
            if ls is not None and round(float(ls), 2) != expected_spacing:
                errors.append(f"{j}. boş satırın satır aralığı {ls} yerine {expected_spacing} olmalı")
            if round(sb, 1) != round(expected_before, 1):
                errors.append(f"{j}. boş satırın paragraf öncesi {sb} yerine {expected_before} olmalı")
            if round(sa, 1) != round(expected_after, 1):
                errors.append(f"{j}. boş satırın paragraf sonrası {sa} yerine {expected_after} olmalı")

            # DEBUG bilgisi
            if debug_file:
                debug_file.write(
                    f"  {j}. satır: boş ✔️ | Font={actual_font}, Size={actual_size}, "
                    f"LS={ls}, Before={sb}, After={sa}\n"
                )

        # ===========================================================
        # 4️⃣ Boş satır sayısı değerlendirmesi
        # ===========================================================
        if blank_count < min_blank or blank_count > max_blank:
            errors.append(f"{blank_count} boş satır var; {min_blank}–{max_blank} arası olmalı")
        if debug_file:
            debug_file.write(f"\nToplam boş satır sayısı = {blank_count}\n\n")

        # ===========================================================
        # 5️⃣ Sonuç değerlendirmesi
        # ===========================================================
        if errors:
            results.append((dep_idx + 1, False, rule_title, "; ".join(errors)))
            if debug_file:
                debug_file.write("❌ Hatalar:\n")
                for e in errors:
                    debug_file.write(f" - {e}\n")
        else:
            results.append((dep_idx + 1, True, rule_title, ""))
            if debug_file:
                debug_file.write("✅ Boşluk ve biçim kontrolleri başarıyla geçti.\n")

        # ===========================================================
        # 6️⃣ DEBUG dosyasını kapat
        # ===========================================================
        if debug_file:
            debug_file.write("-------------------------------------------\n")
            debug_file.write("Kontrol tamamlandı.\n")
            debug_file.close()

        return results

    # ======================================================
    # İÇ KAPAK SAYFASI - "TEZİN ADI" KONTROLÜ
    # ======================================================
    elif check["check"] == "inner_cover_thesis_title":
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 16))
        expected_bold = check.get("bold", True)
        expected_align = check.get("alignment", "center").lower()
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        forbid_italic = bool(check.get("forbid_italic", False))
        debug_mode = check.get("debug", False)

        rule_title = check.get(
            "description",
            f"Tez başlığı 1–2 satır olabilir. Her satır Times New Roman, "
            f"{int(expected_size)} punto, kalın={expected_bold}, ortalı, "
            f"satır aralığı={expected_spacing}, önce={int(expected_before)}, sonra={int(expected_after)} olmalı."
        )

        import os

        # ===========================================================
        # 1️⃣ DEBUG dosyası (tek dosya, append)
        # ===========================================================
        debug_file = None
        if debug_mode:
            debug_path = dbg_path("debug_inner_cover.txt")
            new_file = not os.path.exists(debug_path)
            debug_file = open(debug_path, "a", encoding="utf-8")
            if new_file:
                debug_file.write("[DEBUG] İç Kapak - Tez Başlığı Kontrolleri\n")
                debug_file.write("===========================================\n\n")
            else:
                debug_file.write("\n-------------------------------------------\n")
                debug_file.write("Yeni kontrol çalıştırması başlatıldı (Tez Başlığı).\n\n")

        # ===========================================================
        # 2️⃣ Ana Bilim Dalı satırı bulunmuş olmalı
        # ===========================================================
        dep_idx = memo.get("inner_cover_department_index", None)
        if dep_idx is None:
            results.append((0, False, rule_title, "Ana Bilim Dalı satırı bulunamadı"))
            if debug_file:
                debug_file.write("❌ Ana Bilim Dalı satırı bulunamadı.\n")
                debug_file.close()
            return results

        # ===========================================================
        # 3️⃣ Ana Bilim Dalı sonrası: önce boş satırları geç, sonra ilk dolu satırdan 1–2 başlık satırı al (DİNAMİK OFFSET)
        # ===========================================================
        i = dep_idx + 1
        while i < len(paragraphs) and not paragraphs[i].text.strip():
            i += 1

        title_indices = []
        # İlk başlık satırı
        if i < len(paragraphs):
            title_indices.append(i)
            
        # Sonraki satırları da kontrol et (boş olana kadar veya max 5 satıra kadar)
        curr = i + 1
        while curr < len(paragraphs) and paragraphs[curr].text.strip() and len(title_indices) < 6:
            title_indices.append(curr)
            curr += 1

        if not title_indices:
            results.append((dep_idx, False, rule_title, "Tez başlığı bulunamadı"))
            if debug_file:
                debug_file.write("❌ Tez başlığı bulunamadı.\n")
                debug_file.close()
            return results
            
        # Başlığı Memo'ya kaydet (tam metin)
        full_title_text = " ".join([paragraphs[idx].text.strip() for idx in title_indices])
        memo["thesis_title"] = full_title_text
        # Son satırın index'ini kaydet (diğer kurallar için)
        memo["inner_cover_title_index"] = title_indices[-1]

        errors = []

        # ===========================================================
        # 4️⃣ Başlık satırlarının biçim kontrolleri
        # ===========================================================
        for idx in title_indices:
            p = paragraphs[idx]
            text = p.text.strip()
            if debug_file:
                debug_file.write(f"✅ Tez başlığı satırı (index={idx}): {text[:80]}\n")

            if len(text) < 5:
                errors.append("Tez başlığı çok kısa görünüyor")
            if text.upper() != text:
                errors.append("Tez başlığı tamamen büyük harf olmalı")

            # Font adı
            fn = None
            for r in p.runs:
                if r.font and r.font.name:
                    fn = r.font.name
                    break
            style_font = getattr(p.style, "font", None)
            style_name = style_font.name if style_font and style_font.name else None
            actual_font = fn or style_name
            if actual_font and actual_font != expected_name:
                errors.append(f"Yazı tipi {actual_font} yerine {expected_name} olmalı")

            # Punto
            fs = None
            for r in p.runs:
                if r.font and r.font.size:
                    fs = r.font.size.pt
                    break
            style_size = style_font.size.pt if style_font and style_font.size else None
            actual_size = fs or style_size
            if actual_size and round(float(actual_size), 1) != round(expected_size, 1):
                errors.append(f"Punto {actual_size} yerine {expected_size} olmalı")

            # İtalik yasaklanmışsa kontrol et
            if forbid_italic:
                has_italic_run = any(r.italic for r in p.runs if r.text.strip())
                style_italic = style_font.italic if style_font else None
                if has_italic_run or style_italic:
                    errors.append("Tez başlığında italik yazı kullanılamaz")

            # Kalınlık
            has_bold_run = any(r.bold for r in p.runs if r.text.strip())
            style_bold = style_font.bold if style_font else None
            if expected_bold and not (has_bold_run or style_bold):
                errors.append("Tez başlığı kalın değil")

            # Hizalama
            align_map = {
                "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            }
            expected_align_enum = align_map.get(expected_align, WD_PARAGRAPH_ALIGNMENT.CENTER)
            style_align = getattr(p.style.paragraph_format, "alignment", None)
            if p.alignment not in [expected_align_enum, None] and style_align != expected_align_enum:
                errors.append("Tez başlığı ortalanmamış")

            # Satır aralığı ve boşluklar
            pf = p.paragraph_format
            ls = pf.line_spacing or getattr(p.style.paragraph_format, "line_spacing", None)
            sb = (pf.space_before.pt if pf.space_before else 0.0)
            sa = (pf.space_after.pt if pf.space_after else 0.0)
            if ls is not None and round(float(ls), 2) != expected_spacing:
                errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
            if round(sb, 1) != round(expected_before, 1):
                errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
            if round(sa, 1) != round(expected_after, 1):
                errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

        # ===========================================================
        # 5️⃣ Sonuç değerlendirmesi
        # ===========================================================
        if errors:
            results.append((title_indices[0], False, rule_title, "; ".join(errors)))
            if debug_file:
                debug_file.write("\n❌ Hatalar:\n")
                for e in errors:
                    debug_file.write(f" - {e}\n")
        else:
            results.append((title_indices[0], True, rule_title, ""))
            if debug_file:
                debug_file.write("\n✅ Tüm kontroller başarıyla geçti.\n")

        # ===========================================================
        # 6️⃣ Memo ve dosya kapatma
        # ===========================================================
        memo["thesis_title"] = text
        memo["inner_cover_title_index"] = title_indices[-1]

        if debug_file:
            debug_file.write("\n-------------------------------------------\n")
            debug_file.write("Kontrol tamamlandı.\n")
            debug_file.close()

        return results

    # ======================================================
    # İÇ KAPAK SAYFASI - TEZ BAŞLIĞINDAN SONRAKİ BOŞ SATIR SAYISI VE BİÇİM KONTROLÜ
    # ======================================================
    elif check["check"] == "inner_cover_spacing_after_title":
        min_blank = int(check.get("min_blank_lines", 4))
        max_blank = int(check.get("max_blank_lines", 5))
        expected_font = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 16))
        # expected_bold = check.get("bold", True)   # ❌ BOLD KONTROLÜ KALDIRILDI
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        debug_mode = check.get("debug", False)

        rule_title = check.get(
            "description",
            f"Tez başlığından sonra {min_blank}–{max_blank} satır boşluk olmalı. "
            f"Her boş satır {expected_font} {int(expected_size)} punto, "
            f"satır aralığı={expected_spacing}, önce={expected_before}, sonra={expected_after} olmalı."
        )

        import os

        # ===========================================================
        # 1️⃣ DEBUG dosyası (tek dosya, append)
        # ===========================================================
        debug_file = None
        if debug_mode:
            debug_path = dbg_path("debug_inner_cover.txt")
            new_file = not os.path.exists(debug_path)
            debug_file = open(debug_path, "a", encoding="utf-8")
            if new_file:
                debug_file.write("[DEBUG] İç Kapak - Tez Başlığı Sonrası Boşluk Kontrolleri\n")
                debug_file.write("===========================================\n\n")
            else:
                debug_file.write("\n-------------------------------------------\n")
                debug_file.write("Yeni kontrol çalıştırması başlatıldı (Tez Başlığı Sonrası Boşluklar).\n\n")

        # ===========================================================
        # 2️⃣ Tez başlığı satırı bulunmuş olmalı
        # ===========================================================
        title_idx = memo.get("inner_cover_title_index", None)
        if title_idx is None:
            results.append((0, False, rule_title, "Tez başlığı satırı bulunamadı"))
            if debug_file:
                debug_file.write("❌ Tez başlığı satırı bulunamadı.\n")
                debug_file.close()
            return results

        # ===========================================================
        # 3️⃣ Boş satır sayısı ve biçim kontrolü
        # ===========================================================
        blank_count = 0
        errors = []

        for j in range(title_idx + 1, len(paragraphs)):
            p = paragraphs[j]
            if not p.text.strip():
                blank_count += 1
            else:
                break

        if blank_count < min_blank or blank_count > max_blank:
            errors.append(f"{blank_count} boş satır var; {min_blank}–{max_blank} arası olmalı")

        # ===========================================================
        # 4️⃣ Her boş satırın biçimsel özellikleri ✅ (GÜNCELLENDİ)
        # ===========================================================
        for k in range(1, blank_count + 1):
            p_blank = paragraphs[title_idx + k]
            pf = p_blank.paragraph_format

            # line_spacing (pf/stil)
            ls = pf.line_spacing or getattr(p_blank.style.paragraph_format, "line_spacing", None)

            # space_before/after (sadece pf varsa, yoksa 0.0)
            sb = (pf.space_before.pt if pf.space_before else 0.0)
            sa = (pf.space_after.pt if pf.space_after else 0.0)

            # font adı / punto: run → XML(pPr/rPr) → stil
            fn = None
            fs = None

            for r in p_blank.runs:
                if r.font:
                    if not fn and r.font.name:
                        fn = r.font.name
                    if fs is None and r.font.size:
                        fs = r.font.size.pt

            if fs is None or fn is None:
                try:
                    if fs is None:
                        sz_vals = p_blank._p.xpath("./w:pPr/w:rPr/w:sz/@w:val")
                        if sz_vals:
                            fs = float(int(sz_vals[0])) / 2.0

                    if fn is None:
                        font_vals = p_blank._p.xpath("./w:pPr/w:rPr/w:rFonts/@w:ascii")
                        if not font_vals:
                            font_vals = p_blank._p.xpath("./w:pPr/w:rPr/w:rFonts/@w:hAnsi")
                        if font_vals:
                            fn = font_vals[0]
                except Exception:
                    pass

            style_font = getattr(p_blank.style, "font", None)
            style_name = style_font.name if style_font and style_font.name else None
            style_size = style_font.size.pt if style_font and style_font.size else None

            actual_font = fn or style_name
            actual_size = fs if fs is not None else style_size

            if actual_font and actual_font != expected_font:
                errors.append(f"{k}. boş satırda yazı tipi {actual_font} yerine {expected_font} olmalı")
            if actual_size is not None and round(float(actual_size), 1) != round(expected_size, 1):
                errors.append(f"{k}. boş satırda punto {actual_size} yerine {expected_size} olmalı")
            if ls is not None and round(float(ls), 2) != expected_spacing:
                errors.append(f"{k}. boş satırda satır aralığı {ls} yerine {expected_spacing} olmalı")
            if round(sb, 1) != round(expected_before, 1):
                errors.append(f"{k}. boş satırda paragraf öncesi {sb} yerine {expected_before} olmalı")
            if round(sa, 1) != round(expected_after, 1):
                errors.append(f"{k}. boş satırda paragraf sonrası {sa} yerine {expected_after} olmalı")

            if debug_file:
                debug_file.write(
                    f"  {k}. boş satır: Font={actual_font}, Size={actual_size}, "
                    f"LS={ls}, Before={sb}, After={sa}\n"
                )

        # ===========================================================
        # 5️⃣ Sonuç değerlendirmesi
        # ===========================================================
        if errors:
            results.append((title_idx + 1, False, rule_title, "; ".join(errors)))
            if debug_file:
                debug_file.write("\n❌ Hatalar:\n")
                for e in errors:
                    debug_file.write(f" - {e}\n")
        else:
            results.append((title_idx + 1, True, rule_title, ""))
            if debug_file:
                debug_file.write("\n✅ Boşluk kontrolleri başarıyla geçti.\n")

        # ===========================================================
        # 6️⃣ Dosya kapatma
        # ===========================================================
        if debug_file:
            debug_file.write("\n-------------------------------------------\n")
            debug_file.write("Kontrol tamamlandı.\n")
            debug_file.close()

        return results

    # ======================================================
    # İÇ KAPAK SAYFASI - TEZ BAŞLIĞI SONRASI ÖĞRENCİ ADI SATIRI KONTROLÜ
    # ======================================================
    elif check["check"] == "inner_cover_student_name":
        expected_font = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 14))
        expected_bold = check.get("bold", True)
        expected_align = check.get("alignment", "center").lower()
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        must_be_uppercase = check.get("must_be_uppercase", False)
        debug_mode = check.get("debug", False)

        rule_title = check.get(
            "description",
            f"Öğrenci adı satırı: Times New Roman, {expected_size} pt, kalın={expected_bold}, "
            f"ortalı, satır aralığı {expected_spacing}, önce {expected_before}, sonra {expected_after}, "
            f"{'tamamı büyük harf olmalı' if must_be_uppercase else 'büyük/küçük harf serbest'}."
        )

        import os, re

        # ===========================================================
        # 1️⃣ DEBUG dosyası
        # ===========================================================
        debug_file = None
        if debug_mode:
            debug_path = dbg_path("debug_inner_cover.txt")
            new_file = not os.path.exists(debug_path)
            debug_file = open(debug_path, "a", encoding="utf-8")
            if new_file:
                debug_file.write("[DEBUG] İç Kapak - Öğrenci Adı Kontrolleri\n")
                debug_file.write("===========================================\n\n")
            else:
                debug_file.write("\n-------------------------------------------\n")
                debug_file.write("Yeni kontrol çalıştırması başlatıldı (Öğrenci Adı).\n\n")

        # ===========================================================
        # 2️⃣ Tez başlığı son satırı bulunmuş olmalı
        # ===========================================================
        title_idx = memo.get("inner_cover_title_index", None)
        if title_idx is None:
            results.append((0, False, rule_title, "Tez başlığı son satırı bulunamadı"))
            if debug_file:
                debug_file.write("❌ Tez başlığı son satırı bulunamadı.\n")
                debug_file.close()
            return results

        # ===========================================================
        # 3️⃣ Boşlukları geç, sonraki dolu satırı al (öğrenci adı)
        # ===========================================================
        i = title_idx + 1
        while i < len(paragraphs) and not paragraphs[i].text.strip():
            i += 1

        if i >= len(paragraphs):
            results.append((title_idx, False, rule_title, "Öğrenci adı satırı bulunamadı"))
            if debug_file:
                debug_file.write("❌ Öğrenci adı satırı bulunamadı.\n")
                debug_file.close()
            return results

        p = paragraphs[i]

        # ===========================================================
        # 🔹 Gelişmiş metin birleştirici (alan kodları dahil)
        # ===========================================================
        def full_text_with_fields(para):
            texts = []
            for r in para.runs:
                t = r.text or ""
                # Word alan kodlarını (instrText) da dahil et
                if hasattr(r._element, "xpath"):
                    fld_texts = r._element.xpath(".//w:instrText")
                    for fld in fld_texts:
                        if fld.text:
                            t += fld.text
                texts.append(t)
            return "".join(texts).strip()

        text = full_text_with_fields(p)

        errors = []
        if not text:
            errors.append("Öğrenci adı satırı boş")

        # ===========================================================
        # 4️⃣ Büyük harf kontrolü (yaml parametresine göre)
        # ===========================================================
        if must_be_uppercase and text.upper() != text:
            errors.append("Öğrenci adı tamamen büyük harf olmalı")

        # ===========================================================
        # 5️⃣ Biçim kontrolleri
        # ===========================================================
        fn = None
        for r in p.runs:
            if r.font and r.font.name:
                fn = r.font.name
                break
        style_font = getattr(p.style, "font", None)
        style_name = style_font.name if style_font and style_font.name else None
        actual_font = fn or style_name
        if actual_font and actual_font != expected_font:
            errors.append(f"Yazı tipi {actual_font} yerine {expected_font} olmalı")

        fs = None
        for r in p.runs:
            if r.font and r.font.size:
                fs = r.font.size.pt
                break
        style_size = style_font.size.pt if style_font and style_font.size else None
        actual_size = fs or style_size
        if actual_size and round(float(actual_size), 1) != round(expected_size, 1):
            errors.append(f"Punto {actual_size} yerine {expected_size} olmalı")

        has_bold_run = any(r.bold for r in p.runs if r.text.strip())
        style_bold = style_font.bold if style_font else None
        if expected_bold and not (has_bold_run or style_bold):
            errors.append("Öğrenci adı kalın değil")

        align_map = {
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }
        expected_align_enum = align_map.get(expected_align, WD_PARAGRAPH_ALIGNMENT.CENTER)
        style_align = getattr(p.style.paragraph_format, "alignment", None)
        if p.alignment not in [expected_align_enum, None] and style_align != expected_align_enum:
            errors.append("Öğrenci adı ortalanmamış")

        pf = p.paragraph_format
        ls = pf.line_spacing or getattr(p.style.paragraph_format, "line_spacing", None)
        sb = (pf.space_before.pt if pf.space_before else 0.0)
        sa = (pf.space_after.pt if pf.space_after else 0.0)
        if ls is not None and round(float(ls), 2) != expected_spacing:
            errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
        if round(sb, 1) != round(expected_before, 1):
            errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
        if round(sa, 1) != round(expected_after, 1):
            errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

        # ===========================================================
        # 6️⃣ Memo'ya öğrenci adını kaydet (HATALI OLSA BİLE)
        # ===========================================================
        if text:
            memo["student_name"] = text
            if debug_file:
                debug_file.write(f"✅ Öğrenci adı '{text}' olarak kaydedildi.\n")

        # ===========================================================
        # 7️⃣ Sonuç değerlendirmesi
        # ===========================================================
        if errors:
            results.append((i, False, rule_title, "; ".join(errors)))
            if debug_file:
                debug_file.write("\n❌ Hatalar:\n")
                for e in errors:
                    debug_file.write(f" - {e}\n")
        else:
            results.append((i, True, rule_title, ""))
            if debug_file:
                debug_file.write("\n✅ Öğrenci adı kontrolleri başarıyla geçti.\n")

        if debug_file:
            debug_file.write("\n-------------------------------------------\n")
            debug_file.write("Kontrol tamamlandı.\n")
            debug_file.close()

        memo["inner_cover_student_index"] = i

        return results

    # ======================================================
    # İÇ KAPAK SAYFASI - ÖĞRENCİ ADI SONRASI BOŞLUK KONTROLÜ
    # ======================================================
    elif check["check"] == "inner_cover_post_student_spacing":
        import os

        expected_blank_min = 7
        expected_blank_max = 7
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 14))
        # expected_bold = check.get("bold", True)   # ❌ BOLD KONTROLÜ KALDIRILDI
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        debug_mode = check.get("debug", False)

        rule_title = check.get(
            "description",
            f"Öğrenci adı satırından sonra {expected_blank_min}-{expected_blank_max} satır boşluk bulunmalı "
            f"({expected_name}, {int(expected_size)} punto, {expected_spacing} satır aralığı, "
            f"önce={expected_before}, sonra={expected_after})"
        )

        # -------------------------------------------------------
        # 1️⃣ DEBUG dosyası (append)
        # -------------------------------------------------------
        debug_file = None
        if debug_mode:
            debug_path = dbg_path("debug_inner_cover.txt")
            new_file = not os.path.exists(debug_path)
            debug_file = open(debug_path, "a", encoding="utf-8")
            if new_file:
                debug_file.write("[DEBUG] Öğrenci Adı Sonrası Boşluk Kontrolü\n")
                debug_file.write("===========================================\n\n")
            else:
                debug_file.write("\n-------------------------------------------\n")
                debug_file.write("Yeni kontrol çalıştırması başlatıldı (Öğrenci adı sonrası boşluk).\n\n")

        # -------------------------------------------------------
        # 2️⃣ Öğrenci satırını MEMO’dan oku
        # -------------------------------------------------------
        student_name = memo.get("student_name", "").strip()
        student_idx = memo.get("inner_cover_student_index", None)

        if debug_file:
            debug_file.write(f"ℹ️ Memo'dan alınan öğrenci adı: '{student_name}'\n")
            debug_file.write(f"ℹ️ Memo'dan alınan öğrenci index: {student_idx}\n")

        if student_idx is None or student_idx >= len(paragraphs):
            results.append((0, False, rule_title, "Öğrenci adı satırı index bilgisi memo’da bulunamadı."))
            if debug_file:
                debug_file.write("❌ Öğrenci adı satırı index bilgisi memo’da yok.\n")
                debug_file.close()
            return results

        # -------------------------------------------------------
        # 3️⃣ Öğrenci adından sonraki boş satırları say
        # -------------------------------------------------------
        blank_count = 0
        for p in paragraphs[student_idx + 1:]:
            if not p.text.strip():
                blank_count += 1
            else:
                break

        # -------------------------------------------------------
        # 3B) Boş satırların biçim kontrolü ✅ (GÜNCELLENDİ - BOLD KALDIRILDI)
        #     - font size: run → XML(pPr/rPr) → stil
        #     - font name: run → XML(pPr/rPr) → stil
        # -------------------------------------------------------
        errors = []

        for k in range(1, blank_count + 1):
            p_blank = paragraphs[student_idx + k]
            pf = p_blank.paragraph_format

            # line_spacing (pf/stil)
            ls = pf.line_spacing or getattr(p_blank.style.paragraph_format, "line_spacing", None)

            # space_before/after (sadece pf varsa, yoksa 0.0)
            sb = (pf.space_before.pt if pf.space_before else 0.0)
            sa = (pf.space_after.pt if pf.space_after else 0.0)

            fn = None
            fs = None

            # 1) Run'dan oku
            for r in p_blank.runs:
                if r.font:
                    if not fn and r.font.name:
                        fn = r.font.name
                    if fs is None and r.font.size:
                        fs = r.font.size.pt

            # 2) XML pPr/rPr
            if fs is None or fn is None:
                try:
                    if fs is None:
                        sz_vals = p_blank._p.xpath("./w:pPr/w:rPr/w:sz/@w:val")
                        if sz_vals:
                            fs = float(int(sz_vals[0])) / 2.0

                    if fn is None:
                        font_vals = p_blank._p.xpath("./w:pPr/w:rPr/w:rFonts/@w:ascii")
                        if not font_vals:
                            font_vals = p_blank._p.xpath("./w:pPr/w:rPr/w:rFonts/@w:hAnsi")
                        if font_vals:
                            fn = font_vals[0]
                except Exception:
                    pass

            # 3) Stil fallback
            style_font = getattr(p_blank.style, "font", None)
            style_name = style_font.name if style_font and style_font.name else None
            style_size = style_font.size.pt if style_font and style_font.size else None

            actual_font = fn or style_name
            actual_size = fs if fs is not None else style_size

            if actual_font and actual_font != expected_name:
                errors.append(f"{k}. boş satırda yazı tipi {actual_font} yerine {expected_name} olmalı")
            if actual_size is not None and round(float(actual_size), 1) != round(expected_size, 1):
                errors.append(f"{k}. boş satırda punto {actual_size} yerine {expected_size} olmalı")
            if ls is not None and round(float(ls), 2) != expected_spacing:
                errors.append(f"{k}. boş satırda satır aralığı {ls} yerine {expected_spacing} olmalı")
            if round(sb, 1) != round(expected_before, 1):
                errors.append(f"{k}. boş satırda paragraf öncesi {sb} yerine {expected_before} olmalı")
            if round(sa, 1) != round(expected_after, 1):
                errors.append(f"{k}. boş satırda paragraf sonrası {sa} yerine {expected_after} olmalı")

            if debug_file:
                debug_file.write(
                    f"  {k}. boş satır: Font={actual_font}, Size={actual_size}, "
                    f"LS={ls}, Before={sb}, After={sa}\n"
                )

        # -------------------------------------------------------
        # 4️⃣ Bilgi amaçlı biçim özellikleri
        # -------------------------------------------------------
        p = paragraphs[student_idx]
        fn = None
        for r in p.runs:
            if r.font and r.font.name:
                fn = r.font.name
                break

        style_font = getattr(p.style, "font", None)
        style_name = style_font.name if style_font and style_font.name else None
        actual_font = fn or style_name

        fs = None
        for r in p.runs:
            if r.font and r.font.size:
                fs = r.font.size.pt
                break
        style_size = style_font.size.pt if style_font and style_font.size else None
        actual_size = fs or style_size

        pf = p.paragraph_format
        ls = pf.line_spacing or getattr(p.style.paragraph_format, "line_spacing", None)
        sb = pf.space_before.pt if pf.space_before else 0.0
        sa = pf.space_after.pt if pf.space_after else 0.0

        if debug_file:
            debug_file.write(f"🎯 Öğrenci adı: '{student_name}' (index={student_idx})\n")
            debug_file.write(f"Boş satır sayısı: {blank_count}\n\n")
            debug_file.write("Biçimsel Özellikler (bilgi amaçlı):\n")
            debug_file.write(f"  Yazı tipi: {actual_font}\n")
            debug_file.write(f"  Punto: {actual_size}\n")
            debug_file.write(f"  Satır aralığı: {ls}\n")
            debug_file.write(f"  Önce: {sb} pt, Sonra: {sa} pt\n")

        # -------------------------------------------------------
        # 5️⃣ Sonuç değerlendirmesi
        # -------------------------------------------------------
        if blank_count < expected_blank_min or blank_count > expected_blank_max:
            msg = f"{blank_count} satır boşluk var, {expected_blank_min}-{expected_blank_max} satır olmalı."
            results.append((student_idx, False, rule_title, msg + (("; " + "; ".join(errors)) if errors else "")))
            if debug_file:
                debug_file.write(f"\n❌ {msg}\n")
                if errors:
                    debug_file.write("❌ Boş satır biçim hataları:\n")
                    for e in errors:
                        debug_file.write(f" - {e}\n")
        else:
            if errors:
                results.append((student_idx, False, rule_title, "; ".join(errors)))
                if debug_file:
                    debug_file.write("\n❌ Boş satır biçim hataları:\n")
                    for e in errors:
                        debug_file.write(f" - {e}\n")
            else:
                results.append((student_idx, True, rule_title, ""))
                if debug_file:
                    debug_file.write("\n✅ Boşluk sayısı ve biçimi uygun.\n")

        if debug_file:
            debug_file.write("-------------------------------------------\n")
            debug_file.write("Kontrol tamamlandı.\n")
            debug_file.close()

        return results

    # ======================================================
    # İÇ KAPAK SAYFASI - TEZ TÜRÜ (YÜKSEK LİSANS / DOKTORA)
    # ======================================================
    elif check["check"] == "inner_cover_thesis_type":
        import os, re

        valid_texts = ["YÜKSEK LİSANS TEZİ", "DOKTORA TEZİ"]
        expected_font = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 16))
        expected_bold = check.get("bold", True)
        expected_align = check.get("alignment", "center").lower()
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        must_be_uppercase = check.get("must_be_uppercase", True)
        debug_mode = check.get("debug", False)

        rule_title = check.get(
            "description",
            f"Tez türü satırı ('YÜKSEK LİSANS TEZİ' veya 'DOKTORA TEZİ'), "
            f"{expected_font}, {expected_size} pt, kalın={expected_bold}, "
            f"ortalı, {expected_spacing} satır aralığı, önce {expected_before}, sonra {expected_after}."
        )

        # ===========================================================
        # 1️⃣ DEBUG dosyası
        # ===========================================================
        debug_file = None
        if debug_mode:
            debug_path = dbg_path("debug_inner_cover.txt")
            new_file = not os.path.exists(debug_path)
            debug_file = open(debug_path, "a", encoding="utf-8")
            if new_file:
                debug_file.write("[DEBUG] İç Kapak - Tez Türü Kontrolleri\n")
                debug_file.write("===========================================\n\n")
            else:
                debug_file.write("\n-------------------------------------------\n")
                debug_file.write("Yeni kontrol çalıştırması başlatıldı (Tez Türü).\n\n")

        # ===========================================================
        # 2️⃣ Öğrenci adı bulunmuş olmalı
        # ===========================================================
        student_idx = memo.get("inner_cover_student_index", None)
        if student_idx is None or student_idx + 1 >= len(paragraphs):
            results.append((0, False, rule_title, "Öğrenci adı satırı bulunamadı veya alt satır eksik"))
            if debug_file:
                debug_file.write("❌ Öğrenci adı bulunamadı veya alt satır mevcut değil.\n")
                debug_file.close()
            return results

        # ===========================================================
        # 3️⃣ Öğrenci adından sonra: boşları geç, ilk dolu satır (tez türü)  ✅ DİNAMİK OFFSET
        # ===========================================================
        i = student_idx + 1
        while i < len(paragraphs) and not paragraphs[i].text.strip():
            i += 1

        if i >= len(paragraphs):
            results.append((student_idx, False, rule_title, "Tez türü satırı bulunamadı"))
            if debug_file:
                debug_file.write("❌ Tez türü satırı bulunamadı.\n")
                debug_file.close()
            return results

        p = paragraphs[i]
        raw_text = p.text.strip()
        text = raw_text.upper()

        # ===========================================================
        # 4️⃣ Metin kontrolü
        # ===========================================================
        errors = []
        if text not in valid_texts:
            errors.append(f"Metin '{text}' yerine {' veya '.join(valid_texts)} olmalı")

        if must_be_uppercase and raw_text.upper() != raw_text:
            errors.append("Tüm metin büyük harf olmalı")

        # ===========================================================
        # 5️⃣ Biçim kontrolleri
        # ===========================================================
        fn = None
        for r in p.runs:
            if r.font and r.font.name:
                fn = r.font.name
                break
        style_font = getattr(p.style, "font", None)
        style_name = style_font.name if style_font and style_font.name else None
        actual_font = fn or style_name
        if actual_font and actual_font != expected_font:
            errors.append(f"Yazı tipi {actual_font} yerine {expected_font} olmalı")

        fs = None
        for r in p.runs:
            if r.font and r.font.size:
                fs = r.font.size.pt
                break
        style_size = style_font.size.pt if style_font and style_font.size else None
        actual_size = fs or style_size
        if actual_size and round(float(actual_size), 1) != round(expected_size, 1):
            errors.append(f"Punto {actual_size} yerine {expected_size} olmalı")

        has_bold_run = any(r.bold for r in p.runs if r.text.strip())
        style_bold = style_font.bold if style_font else None
        if expected_bold and not (has_bold_run or style_bold):
            errors.append("Metin kalın değil")

        align_map = {
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }
        expected_align_enum = align_map.get(expected_align, WD_PARAGRAPH_ALIGNMENT.CENTER)
        style_align = getattr(p.style.paragraph_format, "alignment", None)
        if p.alignment not in [expected_align_enum, None] and style_align != expected_align_enum:
            errors.append("Metin ortalanmamış")

        pf = p.paragraph_format
        ls = pf.line_spacing or getattr(p.style.paragraph_format, "line_spacing", None)
        sb = pf.space_before.pt if pf.space_before else 0.0
        sa = pf.space_after.pt if pf.space_after else 0.0
        if ls is not None and round(float(ls), 2) != expected_spacing:
            errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
        if round(sb, 1) != round(expected_before, 1):
            errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
        if round(sa, 1) != round(expected_after, 1):
            errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

        # ===========================================================
        # 6️⃣ DEBUG bilgileri
        # ===========================================================
        if debug_file:
            debug_file.write(f"Metin içeriği: {raw_text}\n")
            debug_file.write(f"  Yazı tipi: {actual_font}\n")
            debug_file.write(f"  Punto: {actual_size}\n")
            debug_file.write(f"  Kalın (run/stil): {has_bold_run or style_bold}\n")
            debug_file.write(f"  Hizalama (run/stil): {p.alignment or style_align}\n")
            debug_file.write(f"  Satır aralığı: {ls}\n")
            debug_file.write(f"  Önce: {sb} pt, Sonra: {sa} pt\n")

        # ===========================================================
        # 7️⃣ Sonuç değerlendirmesi
        # ===========================================================
        if errors:
            results.append((i, False, rule_title, "; ".join(errors)))
            if debug_file:
                debug_file.write("\n❌ Hatalar:\n")
                for err in errors:
                    debug_file.write(f" - {err}\n")
        else:
            results.append((i, True, rule_title, ""))
            if debug_file:
                debug_file.write("\n✅ Tüm kontroller başarıyla geçti.\n")

        # ===========================================================
        # 8️⃣ Memo kaydı  ✅ (KOPUKLUK GİDERİLDİ: index + text her durumda yazılır)
        # ===========================================================
        memo["inner_cover_thesis_type_index"] = i
        memo["inner_cover_thesis_type_text"] = raw_text

        if debug_file:
            debug_file.write("\n-------------------------------------------\n")
            debug_file.write("Kontrol tamamlandı.\n")
            debug_file.close()

        return results

    # ======================================================
    # İÇ KAPAK SAYFASI - TEZ TÜRÜ SONRASI BOŞLUK KONTROLÜ
    # ======================================================
    elif check["check"] == "inner_cover_post_thesis_type_spacing":
        import os

        expected_blank_min = 9
        expected_blank_max = 9
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 16))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        debug_mode = check.get("debug", False)

        rule_title = check.get(
            "description",
            f"Tez türü satırından sonra {expected_blank_min}-{expected_blank_max} satır boşluk bulunmalı "
            f"({expected_name}, {int(expected_size)} punto, {expected_spacing} satır aralığı, "
            f"önce={expected_before}, sonra={expected_after})"
        )

        # -------------------------------------------------------
        # 1️⃣ DEBUG dosyası
        # -------------------------------------------------------
        debug_file = None
        if debug_mode:
            debug_path = dbg_path("debug_inner_cover.txt")
            new_file = not os.path.exists(debug_path)
            debug_file = open(debug_path, "a", encoding="utf-8")
            if new_file:
                debug_file.write("[DEBUG] Tez Türü Sonrası Boşluk Kontrolü\n")
                debug_file.write("===========================================\n\n")
            else:
                debug_file.write("\n-------------------------------------------\n")
                debug_file.write("Yeni kontrol çalıştırması başlatıldı (Tez türü sonrası boşluk).\n\n")

        # -------------------------------------------------------
        # 2️⃣ Tez türü satırını MEMO’dan oku  ✅ (text opsiyonel, index zorunlu)
        # -------------------------------------------------------
        thesis_idx = memo.get("inner_cover_thesis_type_index", None)
        thesis_text = (memo.get("inner_cover_thesis_type_text", "") or "").strip()

        if debug_file:
            debug_file.write(f"ℹ️ Memo'dan alınan tez türü metni: '{thesis_text}'\n")
            debug_file.write(f"ℹ️ Memo'dan alınan tez türü index: {thesis_idx}\n")

        if thesis_idx is None or thesis_idx >= len(paragraphs):
            results.append((0, False, rule_title, "Tez türü satırı index bilgisi memo’da bulunamadı."))
            if debug_file:
                debug_file.write("❌ Tez türü satırı index bilgisi memo’da yok.\n")
                debug_file.close()
            return results

        # -------------------------------------------------------
        # 3️⃣ Tez türü satırından sonraki boş satırları say  ✅ DİNAMİK OFFSET MANTIĞI
        # -------------------------------------------------------
        blank_count = 0
        for p in paragraphs[thesis_idx + 1:]:
            if not p.text.strip():
                blank_count += 1
            else:
                break

        # -------------------------------------------------------
        # 3B) Boş satırların biçim kontrolü ✅ (GÜNCELLENDİ - BOLD KALDIRILDI)
        #     - font size: run → (run yoksa) XML pPr/rPr/w:sz → stil
        #     - font name: run → (run yoksa) XML pPr/rPr/w:rFonts → stil
        # -------------------------------------------------------
        errors = []

        for k in range(1, blank_count + 1):
            p_blank = paragraphs[thesis_idx + k]
            pf = p_blank.paragraph_format

            # line_spacing (pf/stil)
            ls = pf.line_spacing or getattr(p_blank.style.paragraph_format, "line_spacing", None)

            # space_before/after (sadece pf varsa, yoksa 0.0)
            sb = (pf.space_before.pt if pf.space_before else 0.0)
            sa = (pf.space_after.pt if pf.space_after else 0.0)

            # font adı (run/stil) + size (run → XML(pPr/rPr) → stil)
            fn = None
            fs = None

            # 1) Run'dan oku
            for r in p_blank.runs:
                if r.font:
                    if not fn and r.font.name:
                        fn = r.font.name
                    if fs is None and r.font.size:
                        fs = r.font.size.pt

            # 2) Run yoksa / run'da size yoksa: XML pPr/rPr oku
            if fs is None or fn is None:
                try:
                    # Size: w:pPr/w:rPr/w:sz/@w:val  (yarım punto → /2)
                    if fs is None:
                        sz_vals = p_blank._p.xpath("./w:pPr/w:rPr/w:sz/@w:val")
                        if sz_vals:
                            fs = float(int(sz_vals[0])) / 2.0

                    # Font name: w:pPr/w:rPr/w:rFonts/@w:ascii (yoksa hAnsi)
                    if fn is None:
                        font_vals = p_blank._p.xpath("./w:pPr/w:rPr/w:rFonts/@w:ascii")
                        if not font_vals:
                            font_vals = p_blank._p.xpath("./w:pPr/w:rPr/w:rFonts/@w:hAnsi")
                        if font_vals:
                            fn = font_vals[0]
                except Exception:
                    pass

            # 3) Stil fallback
            style_font = getattr(p_blank.style, "font", None)
            style_name = style_font.name if style_font and style_font.name else None
            style_size = style_font.size.pt if style_font and style_font.size else None

            actual_font = fn or style_name
            actual_size = fs if fs is not None else style_size

            if actual_font and actual_font != expected_name:
                errors.append(f"{k}. boş satırda yazı tipi {actual_font} yerine {expected_name} olmalı")
            if actual_size is not None and round(float(actual_size), 1) != round(expected_size, 1):
                errors.append(f"{k}. boş satırda punto {actual_size} yerine {expected_size} olmalı")
            if ls is not None and round(float(ls), 2) != expected_spacing:
                errors.append(f"{k}. boş satırda satır aralığı {ls} yerine {expected_spacing} olmalı")
            if round(sb, 1) != round(expected_before, 1):
                errors.append(f"{k}. boş satırda paragraf öncesi {sb} yerine {expected_before} olmalı")
            if round(sa, 1) != round(expected_after, 1):
                errors.append(f"{k}. boş satırda paragraf sonrası {sa} yerine {expected_after} olmalı")

            if debug_file:
                debug_file.write(
                    f"  {k}. boş satır: Font={actual_font}, Size={actual_size}, "
                    f"LS={ls}, Before={sb}, After={sa}\n"
                )

        # -------------------------------------------------------
        # 4️⃣ Bilgi amaçlı biçimsel özellikler
        # -------------------------------------------------------
        p = paragraphs[thesis_idx]
        fn = None
        for r in p.runs:
            if r.font and r.font.name:
                fn = r.font.name
                break

        style_font = getattr(p.style, "font", None)
        style_name = style_font.name if style_font and style_font.name else None
        actual_font = fn or style_name

        fs = None
        for r in p.runs:
            if r.font and r.font.size:
                fs = r.font.size.pt
                break
        style_size = style_font.size.pt if style_font and style_font.size else None
        actual_size = fs or style_size

        pf = p.paragraph_format
        ls = pf.line_spacing or getattr(p.style.paragraph_format, "line_spacing", None)
        sb = pf.space_before.pt if pf.space_before else 0.0
        sa = pf.space_after.pt if pf.space_after else 0.0

        if debug_file:
            debug_file.write(f"🎯 Tez türü: '{thesis_text}' (index={thesis_idx})\n")
            debug_file.write(f"Boş satır sayısı: {blank_count}\n\n")
            debug_file.write("Biçimsel Özellikler (bilgi amaçlı):\n")
            debug_file.write(f"  Yazı tipi: {actual_font}\n")
            debug_file.write(f"  Punto: {actual_size}\n")
            debug_file.write(f"  Satır aralığı: {ls}\n")
            debug_file.write(f"  Önce: {sb} pt, Sonra: {sa} pt\n")

        # -------------------------------------------------------
        # 5️⃣ Sonuç değerlendirmesi
        # -------------------------------------------------------
        if blank_count < expected_blank_min or blank_count > expected_blank_max:
            msg = f"{blank_count} satır boşluk var, {expected_blank_min}-{expected_blank_max} satır olmalı."
            results.append((thesis_idx, False, rule_title, msg + (("; " + "; ".join(errors)) if errors else "")))
            if debug_file:
                debug_file.write(f"\n❌ {msg}\n")
                if errors:
                    debug_file.write("❌ Boş satır biçim hataları:\n")
                    for e in errors:
                        debug_file.write(f" - {e}\n")
        else:
            if errors:
                results.append((thesis_idx, False, rule_title, "; ".join(errors)))
                if debug_file:
                    debug_file.write("\n❌ Boş satır biçim hataları:\n")
                    for e in errors:
                        debug_file.write(f" - {e}\n")
            else:
                results.append((thesis_idx, True, rule_title, ""))
                if debug_file:
                    debug_file.write("\n✅ Boşluk sayısı ve biçimi uygun.\n")

        if debug_file:
            debug_file.write("-------------------------------------------\n")
            debug_file.write("Kontrol tamamlandı.\n")
            debug_file.close()

        return results

    # ======================================================
    # İÇ KAPAK SAYFASI - MEZUNİYET TARİHİ (örnek: HAZİRAN 2025) KONTROLÜ
    # ======================================================
    elif check["check"] == "inner_cover_graduation_date":
        import os, re

        expected_font = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 16))
        expected_bold = check.get("bold", True)
        expected_align = check.get("alignment", "center").lower()
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        must_be_uppercase = check.get("must_be_uppercase", True)
        debug_mode = check.get("debug", False)

        rule_title = check.get(
            "description",
            "Mezuniyet tarihi satırı: ay adı büyük harfle, ardından yıl (örnek: HAZİRAN 2025)."
        )

        # -------------------------------------------------------
        # 1️⃣ DEBUG dosyası
        # -------------------------------------------------------
        debug_file = None
        if debug_mode:
            debug_path = dbg_path("debug_inner_cover.txt")
            new_file = not os.path.exists(debug_path)
            debug_file = open(debug_path, "a", encoding="utf-8")
            if new_file:
                debug_file.write("[DEBUG] Mezuniyet Tarihi Kontrolü\n")
                debug_file.write("===========================================\n\n")
            else:
                debug_file.write("\n-------------------------------------------\n")
                debug_file.write("Yeni kontrol çalıştırması başlatıldı (Mezuniyet tarihi).\n\n")

        # -------------------------------------------------------
        # 2️⃣ Önceki kural: tez türü satırını bul
        #    ✅ Memo yoksa fallback: dokümanda tez türü satırını ara (dinamik offset uyumlu)
        # -------------------------------------------------------
        thesis_type_idx = memo.get("inner_cover_thesis_type_index", None)

        if thesis_type_idx is None:
            # fallback arama: YÜKSEK LİSANS TEZİ / DOKTORA TEZİ satırını bul
            valid_texts = ["YÜKSEK LİSANS TEZİ", "DOKTORA TEZİ"]
            found = None
            for idx, pp in enumerate(paragraphs):
                t = (pp.text or "").strip().upper()
                if t in valid_texts:
                    found = idx
                    break

            if found is None:
                results.append((0, False, rule_title, "Tez türü satırı bulunamadı (önceki kural çalışmadı)."))
                if debug_file:
                    debug_file.write("❌ Tez türü satırı memo’da yok ve dokümanda da bulunamadı.\n")
                    debug_file.close()
                return results

            thesis_type_idx = found
            memo["inner_cover_thesis_type_index"] = thesis_type_idx
            memo["inner_cover_thesis_type_text"] = paragraphs[thesis_type_idx].text.strip()

            if debug_file:
                debug_file.write(f"⚠️ Memo'da tez türü yoktu; dokümanda bulundu (index={thesis_type_idx}).\n")

        # Tez türünden sonraki dolu satırı bul (mezuniyet tarihi satırı) ✅ DİNAMİK OFFSET
        i = thesis_type_idx + 1
        while i < len(paragraphs) and not paragraphs[i].text.strip():
            i += 1

        if i >= len(paragraphs):
            results.append((thesis_type_idx, False, rule_title, "Mezuniyet tarihi satırı bulunamadı."))
            if debug_file:
                debug_file.write("❌ Mezuniyet tarihi satırı bulunamadı.\n")
                debug_file.close()
            return results

        p = paragraphs[i]
        text = p.text.strip()
        errors = []

        # -------------------------------------------------------
        # 3️⃣ Metin biçimi kontrolü
        # -------------------------------------------------------
        pattern = re.compile(r"^(OCAK|ŞUBAT|MART|NİSAN|MAYIS|HAZİRAN|TEMMUZ|AĞUSTOS|EYLÜL|EKİM|KASIM|ARALIK)\s+20\d{2}$")
        if not pattern.match(text):
            errors.append("Metin biçimi hatalı (örnek: 'HAZİRAN 2025' şeklinde olmalı).")

        if must_be_uppercase and text.upper() != text:
            errors.append("Tarih satırı tamamen büyük harf olmalı.")

        # -------------------------------------------------------
        # 4️⃣ Biçim kontrolleri
        # -------------------------------------------------------
        fn = None
        for r in p.runs:
            if r.font and r.font.name:
                fn = r.font.name
                break
        style_font = getattr(p.style, "font", None)
        style_name = style_font.name if style_font and style_font.name else None
        actual_font = fn or style_name
        if actual_font and actual_font != expected_font:
            errors.append(f"Yazı tipi {actual_font} yerine {expected_font} olmalı")

        fs = None
        for r in p.runs:
            if r.font and r.font.size:
                fs = r.font.size.pt
                break
        style_size = style_font.size.pt if style_font and style_font.size else None
        actual_size = fs or style_size
        if actual_size and round(float(actual_size), 1) != round(expected_size, 1):
            errors.append(f"Punto {actual_size} yerine {expected_size} olmalı")

        has_bold_run = any(r.bold for r in p.runs if r.text.strip())
        style_bold = style_font.bold if style_font else None
        if expected_bold and not (has_bold_run or style_bold):
            errors.append("Metin kalın değil")

        align_map = {
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }
        expected_align_enum = align_map.get(expected_align, WD_PARAGRAPH_ALIGNMENT.CENTER)
        style_align = getattr(p.style.paragraph_format, "alignment", None)
        if p.alignment not in [expected_align_enum, None] and style_align != expected_align_enum:
            errors.append("Metin ortalanmamış")

        pf = p.paragraph_format
        ls = pf.line_spacing or getattr(p.style.paragraph_format, "line_spacing", None)
        sb = (pf.space_before.pt if pf.space_before else 0.0)
        sa = (pf.space_after.pt if pf.space_after else 0.0)
        if ls is not None and round(float(ls), 2) != expected_spacing:
            errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
        if round(sb, 1) != round(expected_before, 1):
            errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
        if round(sa, 1) != round(expected_after, 1):
            errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

        # -------------------------------------------------------
        # 5️⃣ Her durumda memo’ya kaydet (satır bulunduğu için)
        # -------------------------------------------------------
        memo["inner_cover_graduation_date_index"] = i
        memo["inner_cover_graduation_date_text"] = text

        if debug_file:
            debug_file.write(f"🎯 Mezuniyet tarihi satırı bulundu (index={i}): '{text}'\n")

        # -------------------------------------------------------
        # 6️⃣ Sonuç yazımı
        # -------------------------------------------------------
        if errors:
            results.append((i, False, rule_title, "; ".join(errors)))
            if debug_file:
                debug_file.write("\n❌ Hatalar:\n")
                for e in errors:
                    debug_file.write(f" - {e}\n")
        else:
            results.append((i, True, rule_title, ""))
            if debug_file:
                debug_file.write("\n✅ Mezuniyet tarihi kontrolleri başarıyla geçti.\n")

        if debug_file:
            debug_file.write("-------------------------------------------\n")
            debug_file.write("Kontrol tamamlandı.\n")
            debug_file.close()

        return results

    # ======================================================
    # İÇ KAPAK SAYFASI - ÜNİVERSİTENİN BULUNDUĞU İL (SON SATIR)
    # ======================================================
    elif check["check"] == "inner_cover_city_name":
        import os, re

        expected_font = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 16))
        expected_bold = check.get("bold", True)
        expected_align = check.get("alignment", "center").lower()
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        must_be_uppercase = check.get("must_be_uppercase", True)
        debug_mode = check.get("debug", False)

        rule_title = check.get(
            "description",
            f"İl adı satırı: Tamamı büyük harf olmalı (örnek: OSMANİYE). "
            f"({expected_font}, {expected_size} pt, kalın={expected_bold}, ortalı, "
            f"satır aralığı {expected_spacing}, önce={expected_before}, sonra={expected_after})"
        )

        # -------------------------------------------------------
        # 1️⃣ DEBUG dosyası
        # -------------------------------------------------------
        debug_file = None
        if debug_mode:
            debug_path = dbg_path("debug_inner_cover.txt")
            new_file = not os.path.exists(debug_path)
            debug_file = open(debug_path, "a", encoding="utf-8")
            if new_file:
                debug_file.write("[DEBUG] İl Adı Kontrolü\n")
                debug_file.write("===========================================\n\n")
            else:
                debug_file.write("\n-------------------------------------------\n")
                debug_file.write("Yeni kontrol çalıştırması başlatıldı (İl adı).\n\n")

        # -------------------------------------------------------
        # 2️⃣ Mezuniyet tarihi satırı index’ini memo’dan al
        #    - Yoksa fallback: son dolu satırı il adı kabul et
        # -------------------------------------------------------
        grad_idx = memo.get("inner_cover_graduation_date_index", None)

        if grad_idx is not None:
            # Mezuniyet tarihinden sonraki dolu satırı bul (asıl mantık)
            i = grad_idx + 1
            while i < len(paragraphs) and not paragraphs[i].text.strip():
                i += 1

            if i >= len(paragraphs):
                results.append((grad_idx, False, rule_title, "İl adı satırı bulunamadı."))
                if debug_file:
                    debug_file.write("❌ İl adı satırı bulunamadı (mezuniyet tarihi sonrası dolu satır yok).\n")
                    debug_file.close()
                return results

            if debug_file:
                debug_file.write(f"ℹ️ Mezuniyet tarihi index memo'dan alındı: {grad_idx}\n")
                debug_file.write("ℹ️ İl adı satırı mezuniyet tarihinden sonra aranacak.\n")

        else:
            # Fallback: iç kapakta son dolu satır = il adı (mezuniyet tarihi kuralı çalışmamış olsa da)
            i = None
            for idx in range(len(paragraphs) - 1, -1, -1):
                if paragraphs[idx].text.strip():
                    i = idx
                    break

            if i is None:
                results.append((0, False, rule_title, "İl adı satırı bulunamadı (belge tamamen boş görünüyor)."))
                if debug_file:
                    debug_file.write("❌ Belgedeki tüm paragraflar boş görünüyor.\n")
                    debug_file.close()
                return results

            if debug_file:
                debug_file.write("⚠️ Memo'da inner_cover_graduation_date_index yok.\n")
                debug_file.write(f"⚠️ Fallback uygulandı: son dolu satır il adı kabul edildi (index={i}).\n")

        # -------------------------------------------------------
        # 3️⃣ İl adı satırını kontrol et
        # -------------------------------------------------------
        p = paragraphs[i]
        text = p.text.strip()
        errors = []

        # Metin kontrolü
        if must_be_uppercase and text.upper() != text:
            errors.append("İl adı tamamı büyük harf olmalı.")
        if not re.match(r"^[A-ZÇĞİÖŞÜ\s]+$", text):
            errors.append("İl adı yalnızca harflerden oluşmalı.")
        if len(text.split()) > 1:
            errors.append("İl adı tek kelime olmalı (örnek: OSMANİYE).")

        # -------------------------------------------------------
        # 4️⃣ Biçimsel kontroller
        # -------------------------------------------------------
        fn = None
        for r in p.runs:
            if r.font and r.font.name:
                fn = r.font.name
                break
        style_font = getattr(p.style, "font", None)
        style_name = style_font.name if style_font and style_font.name else None
        actual_font = fn or style_name
        if actual_font and actual_font != expected_font:
            errors.append(f"Yazı tipi {actual_font} yerine {expected_font} olmalı")

        fs = None
        for r in p.runs:
            if r.font and r.font.size:
                fs = r.font.size.pt
                break
        style_size = style_font.size.pt if style_font and style_font.size else None
        actual_size = fs or style_size
        if actual_size and round(float(actual_size), 1) != round(expected_size, 1):
            errors.append(f"Punto {actual_size} yerine {expected_size} olmalı")

        has_bold_run = any(r.bold for r in p.runs if r.text.strip())
        style_bold = style_font.bold if style_font else None
        if expected_bold and not (has_bold_run or style_bold):
            errors.append("Metin kalın değil")

        align_map = {
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }
        expected_align_enum = align_map.get(expected_align, WD_PARAGRAPH_ALIGNMENT.CENTER)
        style_align = getattr(p.style.paragraph_format, "alignment", None)
        if p.alignment not in [expected_align_enum, None] and style_align != expected_align_enum:
            errors.append("Metin ortalanmamış")

        pf = p.paragraph_format
        ls = pf.line_spacing or getattr(p.style.paragraph_format, "line_spacing", None)
        sb = (pf.space_before.pt if pf.space_before else 0.0)
        sa = (pf.space_after.pt if pf.space_after else 0.0)
        if ls is not None and round(float(ls), 2) != expected_spacing:
            errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
        if round(sb, 1) != round(expected_before, 1):
            errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
        if round(sa, 1) != round(expected_after, 1):
            errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

        # -------------------------------------------------------
        # 5️⃣ Sonuç ve memo’ya kaydetme (her durumda)
        # -------------------------------------------------------
        memo["inner_cover_city_text"] = text
        memo["inner_cover_city_index"] = i

        if not errors:
            results.append((i, True, rule_title, ""))
            if debug_file:
                debug_file.write(f"✅ İl adı '{text}' olarak bulundu ve geçerli (index={i}).\n")
        else:
            results.append((i, False, rule_title, "; ".join(errors)))
            if debug_file:
                debug_file.write(f"⚠️ İl adı bulundu (index={i}): '{text}'\n")
                debug_file.write("❌ Ancak hatalar var:\n")
                for e in errors:
                    debug_file.write(f" - {e}\n")

        if debug_file:
            debug_file.write("-------------------------------------------\n")
            debug_file.write("Kontrol tamamlandı.\n")
            debug_file.close()

        return results



    # ===============================================================================================================#
    # ===============================================================================================================#    
    # ===============================================================================================================#
    # ONAY SAYFASI BAŞLIĞI ("TEZ ONAYI") KONTROLÜ
    # - markers ile toleranslı bul
    # - metin expected_text ile tam eşleşmiyorsa ihlal üret (ama yine de bulundu kabul et)
    # - memo["approval_heading_idx"] kaydet
    # - stil kontrolleri effective_* ile
    # ======================================================
    elif check["check"] == "approval_heading":
        import re

        expected_text = (check.get("text", "TEZ ONAYI") or "").strip()
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after  = float(check.get("space_after", 24))
        must_be_upper = bool(check.get("must_be_upper", True))

        # alignment yaml: "center/left/right/justify" veya enum verilmiş olabilir
        expected_align = (check.get("alignment", check.get("align", "center")) or "center").lower().strip()

        # markers: YAML list veya string
        markers = check.get("markers", [r"^TEZ\s*ONAYI$"])
        if isinstance(markers, str):
            markers = [markers]

        def clean_text(s: str) -> str:
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())


        def norm_tr_upper(s: str) -> str:
            # 1) whitespace temizliği
            s = clean_text(s)

            # 2) Word bazen "İ"yi iki karakter olarak saklayabiliyor: "I" + "̇" (combining dot)
            #    Bunu tek karakterli dönüşüme girmeden önce normalize edelim.
            s = s.replace("I\u0307", "I")   # "İ" -> "I"  (U+0307: combining dot above)

            # 3) Türkçe duyarsızlaştırma (tek karakterli map şart!)
            trans = str.maketrans({
                "ı": "i", "İ": "i",
                "ç": "c", "Ç": "c",
                "ğ": "g", "Ğ": "g",
                "ö": "o", "Ö": "o",
                "ş": "s", "Ş": "s",
                "ü": "u", "Ü": "u",
            })
            s = s.translate(trans)

            return s.upper()


        rule_title = check.get(
            "description",
            f"{expected_text}\n{expected_name}, {int(expected_size)} punto, "
            f"{'BÜYÜK HARF' if must_be_upper else 'Normal'}, "
            f"{expected_align}, {expected_spacing} satır aralığı, "
            f"önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        # 1) Toleranslı bulma (markers)
        found = None
        for i, para in enumerate(paragraphs):
            cand = clean_text(para.text)
            for m in markers:
                if re.match(m, cand, re.IGNORECASE):
                    found = (i, para, m)
                    break
            if found:
                break

        if not found:
            results.append((0, False, rule_title, f"'{expected_text}' başlığı bulunamadı (markers eşleşmedi)"))
        else:
            idx, para, used_marker = found
            memo["approval_heading_idx"] = idx

            errors = []

            # 2) Katı metin doğrulama (tam metin bekleniyor)
            raw_title = clean_text(para.text)
            if norm_tr_upper(raw_title) != norm_tr_upper(expected_text):
                errors.append(f"Başlık metni '{expected_text}' olmalı (bulunan: {raw_title})")

            # 3) Büyük harf kontrolü (opsiyon)
            if must_be_upper and raw_title and (raw_title != raw_title.upper()):
                errors.append("Başlık BÜYÜK HARF olmalı")

            # 4) Font / Size (effective)
            fn = effective_font_name(para)
            fs = effective_font_size_pt(para)
            if fn and fn != expected_name:
                errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")
            if fs is not None and round(float(fs), 1) != round(float(expected_size), 1):
                errors.append(f"Yazı boyutu {fs}pt yerine {expected_size}pt olmalı")

            # 5) Alignment (effective)
            align_map = {
                "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            }
            expected_enum = align_map.get(expected_align)
            actual_enum = effective_alignment(para)

            if expected_enum is not None:
                # actual None ise Word default LEFT gibi davranır
                if actual_enum is None:
                    if expected_enum != WD_PARAGRAPH_ALIGNMENT.LEFT:
                        errors.append("Başlık ortalanmış olmalı" if expected_enum == WD_PARAGRAPH_ALIGNMENT.CENTER else "Başlık hizalaması yanlış")
                elif actual_enum != expected_enum:
                    errors.append("Başlık ortalanmış olmalı" if expected_enum == WD_PARAGRAPH_ALIGNMENT.CENTER else "Başlık hizalaması yanlış")

            # 6) Satır aralığı / boşluklar (effective)
            ls = effective_line_spacing(para, default=expected_spacing)
            sb = effective_space_pt(para, "before")
            sa = effective_space_pt(para, "after")

            if ls is not None and round(float(ls), 2) != round(float(expected_spacing), 2):
                errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
            if round(float(sb), 1) != round(float(expected_before), 1):
                errors.append(f"Paragraf öncesi boşluk {sb} yerine {expected_before} olmalı")
            if round(float(sa), 1) != round(float(expected_after), 1):
                errors.append(f"Paragraf sonrası boşluk {sa} yerine {expected_after} olmalı")

            results.append((idx, len(errors) == 0, rule_title, "; ".join(errors)))

    # ======================================================
    # ONAY SAYFASI TEZ BAŞLIĞI KONTROLÜ
    # - memo["approval_heading_idx"] sonrası boşları geç → ilk dolu satır = tez başlığı
    # - memo["approval_title_idx"] kaydet
    # - no_italic: tez başlığında italik karakter bulunmamalı
    # ======================================================
    elif check["check"] == "approval_title_block":
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after  = float(check.get("space_after", 0))
        must_be_upper = bool(check.get("must_be_upper", False))

        expected_align = (check.get("alignment", check.get("align", "center")) or "center").lower().strip()

        # ✅ yeni opsiyon: italik olmasın
        no_italic = bool(check.get("no_italic", False))

        # ✅ yeni opsiyon: kalın olmasın
        no_bold = bool(check.get("no_bold", False))

        rule_title = (
            f"TEZ BAŞLIĞI\n"
            f"{expected_name}, {int(expected_size)} punto, "
            f"{'BÜYÜK HARF' if must_be_upper else 'Normal'}, "
            f"{expected_align}, {expected_spacing} satır aralığı, "
            f"önce {int(expected_before)}, sonra {int(expected_after)}"
            + ("; italik & kalın font YOK" if no_italic else "")
            + ("; " if no_bold else "")

        )

        # 1) Heading idx: memo öncelikli
        heading_idx = memo.get("approval_heading_idx")

        if heading_idx is None:
            results.append((0, False, rule_title, "TEZ ONAYI başlığı memo’da yok, tez başlığı kontrolü yapılamadı"))
        else:
            # 2) Başlıktan sonraki ilk dolu satırı bul (dinamik)
            j = heading_idx + 1
            while j < len(paragraphs):
                txt = (paragraphs[j].text or "").replace("\u00A0", " ").strip()
                if txt != "":
                    break
                j += 1

            if j >= len(paragraphs):
                results.append((heading_idx, False, rule_title, "Tez başlığı satırı bulunamadı"))
            else:
                p = paragraphs[j]
                memo["approval_title_idx"] = j

                errors = []

                # Metin boş olmamalı
                para_text = (p.text or "").strip()
                if not para_text:
                    errors.append("Tez başlığı satırı boş")

                # Font / size (effective)
                fn = effective_font_name(p)
                fs = effective_font_size_pt(p)
                if fn and fn != expected_name:
                    errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")
                if fs is not None and round(float(fs), 1) != round(float(expected_size), 1):
                    errors.append(f"Yazı boyutu {fs}pt yerine {expected_size}pt olmalı")

                # ✅ Italic yasaksa (EFFECTIVE)
                # Word'de italik run'da None görünebilir; effective_italic bunu yakalar.
                if no_italic and effective_italic(p):
                    errors.append("Tez başlığında italik karakter bulunmamalı")
                    
                # ✅ Bold yasaksa (EFFECTIVE)
                if no_bold and effective_bold(p):
                    errors.append("Tez başlığı kalın (bold) olmamalı")


                # Alignment (effective)
                align_map = {
                    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                    "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                }
                expected_enum = align_map.get(expected_align)
                actual_enum = effective_alignment(p)
                if expected_enum is not None:
                    if actual_enum is None:
                        if expected_enum != WD_PARAGRAPH_ALIGNMENT.LEFT:
                            errors.append("Başlık ortalanmış olmalı" if expected_enum == WD_PARAGRAPH_ALIGNMENT.CENTER else "Başlık hizalaması yanlış")
                    elif actual_enum != expected_enum:
                        errors.append("Başlık ortalanmış olmalı" if expected_enum == WD_PARAGRAPH_ALIGNMENT.CENTER else "Başlık hizalaması yanlış")

                # Satır aralığı / boşluklar (effective)
                ls = effective_line_spacing(p, default=expected_spacing)
                sb = effective_space_pt(p, "before")
                sa = effective_space_pt(p, "after")

                if ls is not None and round(float(ls), 2) != round(float(expected_spacing), 2):
                    errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
                if round(float(sb), 1) != round(float(expected_before), 1):
                    errors.append(f"Paragraf öncesi boşluk {sb} yerine {expected_before} olmalı")
                if round(float(sa), 1) != round(float(expected_after), 1):
                    errors.append(f"Paragraf sonrası boşluk {sa} yerine {expected_after} olmalı")

                # Büyük harf opsiyonu
                if must_be_upper and para_text and (para_text != para_text.upper()):
                    errors.append("Başlık BÜYÜK HARF olmalı")

                results.append((j, len(errors) == 0, rule_title, "; ".join(errors)))

    # ======================================================
    # TEZ BAŞLIĞI ile ONAY METNİ arasında HİÇ SATIR OLMAMALI
    # (approval_title_idx + 1 == approval_text_idx)
    # ======================================================
    elif check["check"] == "approval_no_line_between_title_and_text":
        rule_title = check.get(
            "description",
            "TEZ BAŞLIĞI satırından hemen sonra ONAY METNİ gelmeli (arada satır olmamalı)"
        )

        # 1) Tez başlığı idx (memo’dan)
        title_idx = memo.get("approval_title_idx")


        def _preview(idx: int, take_words: int = 7, max_chars: int = 70) -> str:
            if idx is None or idx < 0 or idx >= len(paragraphs):
                return ""
            txt = (paragraphs[idx].text or "").strip()
            if not txt:
                return ""
            pv = " ".join(txt.split()[:take_words])
            if len(pv) > max_chars:
                pv = pv[:max_chars].rstrip() + "…"
            return pv

        # 2) Onay metni idx (memo’dan, yoksa fallback arama)
        text_idx = memo.get("approval_text_idx")
        if text_idx is None:
            # approval_text_spacing ile AYNI mantık (fallback)
            def normalize_tr_full(s: str) -> str:
                s = (s or "").strip().lower()
                trans = str.maketrans({
                    "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                    "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u",
                    "â":"a","Â":"a","î":"i","Î":"i","û":"u","Û":"u",
                })
                s = s.translate(trans)
                return " ".join(s.split())

            required_keywords = [
                "tarafindan",
                "danismanliginda",
                "universitesi",
                "anabilim",
                "kabul edilmistir"
            ]

            for idx, para in enumerate(paragraphs):
                norm = normalize_tr_full(para.text or "")
                if all(kw in norm for kw in required_keywords):
                    text_idx = idx
                    memo["approval_text_idx"] = idx
                    break

        # 3) Ön şart kontrolleri
        if title_idx is None:
            results.append(
                (0, False, rule_title,
                "Tez başlığı bulunamadı (approval_title_idx yok)")
            )

        elif text_idx is None:
            results.append(
                (title_idx, False, rule_title,
                "Onay metni bulunamadı (approval_text_idx yok)")
            )

        else:
            # 4) KATI KURAL: hemen alt satır olmak zorunda
            expected_idx = title_idx + 1

            if text_idx != expected_idx:
                # Aradaki satır sayısını raporlayalım (diagnostic için)
                gap = text_idx - title_idx - 1
                
                title_pv = _preview(title_idx)
                expected_pv = _preview(expected_idx)   # başlığın hemen altındaki satır (olması gereken onay metni)
                actual_pv = _preview(text_idx)         # bulunan onay metni satırı      
                          
                results.append(
                    (
                        title_idx,
                        False,
                        rule_title,                       
                        f"Tez başlığı sonrası ONAY METNİ hemen gelmeli; ancak arada {gap} satır var. "
                        f"Başlık(idx={title_idx}): '{title_pv}' | "
                        f"Beklenen satır(idx={expected_idx}): '{expected_pv}' | "
                        f"Bulunan onay metni(idx={text_idx}): '{actual_pv}'"
                        
                    )
                )
            else:
                results.append(
                    (title_idx, True, rule_title, "")
                )

    # ======================================================
    # ONAY SAYFASI ONAY METNİ SONRASI BOŞLUK KONTROLÜ (memo destekli)
    # - onay metni idx bulunur → memo["approval_text_idx"]
    # - boşluklardan sonra ilk içerik idx → memo["approval_text_after_blank_first_content_idx"]
    # ======================================================
    elif check["check"] == "approval_text_spacing":
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 0))

        rule_title = (
            f"ONAY METNİ SONRASI 2 SATIR BOŞLUK\n"
            f"Satır aralığı {expected_spacing}, önce {expected_before}, sonra {expected_after}"
        )

        def normalize_tr_full(s: str) -> str:
            s = (s or "").strip().lower()
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g","ö":"o","Ö":"o",
                "ş":"s","Ş":"s","ü":"u","Ü":"u","â":"a","Â":"a","î":"i","Î":"i","û":"u","Û":"u",
            })
            s = s.translate(trans).replace("’","'").replace("“","\"").replace("”","\"")
            return " ".join(s.split())

        def is_blank_para_text(txt: str) -> bool:
            return (not txt) or (txt.strip() == "") or ("".join(txt.split()) == "")

        # Onay paragrafını tespit: aynı anahtarlar
        required_keywords = ["tarafindan", "danismanliginda", "universitesi", "anabilim", "kabul edilmistir"]

        found_idx = memo.get("approval_text_idx")
        if found_idx is None:
            for idx, para in enumerate(paragraphs):
                norm = normalize_tr_full(para.text or "")
                if all(kw in norm for kw in required_keywords):
                    found_idx = idx
                    memo["approval_text_idx"] = idx
                    break

        if found_idx is None:
            results.append((0, False, rule_title, "Onay metni paragrafı bulunamadı, boşluk kontrolü yapılamadı"))
        else:
            errors = []

            # Sonraki 2 paragraf boş mu ve formatı doğru mu?
            for k in range(1, 3):
                
                idx_real = found_idx + k
                loc = f"{k}. satır (belge:{idx_real})"

                if idx_real >= len(paragraphs):
                    errors.append(f"{k}. satır (belge:{idx_real}) — eksik")
                    continue
                p = paragraphs[idx_real]
                
                txt = (p.text or "").strip()
                pv = " ".join(txt.split()[:7])
                if len(pv) > 70:
                    pv = pv[:70].rstrip() + "…"
                loc = loc + (f" ('{pv}')" if pv else "")

                
                
                if not is_blank_para_text(p.text):
                    errors.append(f"{loc} - boş değil (boş olmalı)")

                ls = effective_line_spacing(p, default=expected_spacing)
                if isinstance(ls, (int, float)) and round(float(ls), 2) != round(float(expected_spacing), 2):
                    errors.append(f"{loc}. satır aralığı {ls} yerine {expected_spacing} olmalı")

                before_pt = effective_space_pt(p, "before")
                after_pt  = effective_space_pt(p, "after")
                if round(float(before_pt), 1) != round(float(expected_before), 1):
                    errors.append(f"{loc}. satır öncesi {before_pt} pt yerine {expected_before} pt olmalı")
                if round(float(after_pt), 1) != round(float(expected_after), 1):
                    errors.append(f"{loc}. satır sonrası {after_pt} pt yerine {expected_after} pt olmalı")

            # ✅ 2 boşluktan sonra ilk içerik satırını memo’ya yaz
            j = found_idx + 1
            blank_count = 0
            while j < len(paragraphs):
                if is_blank_para_text(paragraphs[j].text):
                    blank_count += 1
                    j += 1
                    continue
                break
            memo["approval_text_after_blank_first_content_idx"] = j

            results.append((found_idx, len(errors) == 0, rule_title, "; ".join(errors)))

    # ======================================================
    # ONAY MAKAMI BLOĞU (dinamik) - 5846 referanslı ama OFFSET yok
    # - 5846 satırını bul
    # - yukarıdan expected_lines adet DOLU paragrafı (boşları atlayarak) topla
    # - bulunan bloğun başlangıcı memo["authority_start_idx"]
    # ======================================================
    elif check["check"] == "approval_authority_block":
        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 0))
        expected_lines   = int(check.get("expected_lines", 3))

        rule_title = (
            f"ONAY MAKAMI BİLGİLERİ\n"
            f"{expected_lines} satır, {expected_name}, {int(expected_size)} punto, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        def is_effectively_blank(text: str) -> bool:
            if not text:
                return True
            cleaned = (
                text.replace(" ", "")
                    .replace("\t", "")
                    .replace("\n", "")
                    .replace("\r", "")
                    .replace("\xa0", "")
                    .replace("\u00A0", "")
                    .replace(".", "")
                    .replace("…", "")
                    .replace("-", "")
            )
            return cleaned.strip() == ""

        # --- 5846 satırını bul
        onay_son = memo.get("approval_5846_idx")
        if onay_son is None:
            for i, p in enumerate(paragraphs):
                if "5846" in (p.text or ""):
                    onay_son = i
                    memo["approval_5846_idx"] = i
                    break

        if onay_son is None:
            results.append((0, False, rule_title, "5846 bulunamadı"))
        else:
            # --- yukarı doğru expected_lines dolu satır topla (boşları atla)
            collected = []  # (idx, paragraph)
            j = onay_son - 1
            while j >= 0 and len(collected) < expected_lines:
                if not is_effectively_blank(paragraphs[j].text or ""):
                    collected.append((j, paragraphs[j]))
                j -= 1

            if len(collected) < expected_lines:
                results.append((onay_son, False, rule_title, f"Onay makamı bloğu için {expected_lines} dolu satır bulunamadı"))
            else:
                # collected ters yönde toplandı; en üst satır = min idx
                collected_sorted = sorted(collected, key=lambda x: x[0])
                authority_start = collected_sorted[0][0]
                memo["authority_start_idx"] = authority_start

                errors = []

                # --- her satırı format kontrolü
                for k, (idx, p) in enumerate(collected_sorted, start=1):
                    # --- satır önizleme (ilk birkaç kelime)
                    txt = (p.text or "").strip()
                    pv = " ".join(txt.split()[:7])
                    if len(pv) > 70:
                        pv = pv[:70].rstrip() + "…"
                    loc = f"{k}. satır (belge:{idx})" + (f" ('{pv}')" if pv else "")
                    
                    # 1) Satır boş olmamalı
                    if is_effectively_blank(p.text or ""):
                        errors.append(f"{loc}. satır boş olmamalı")
                        continue

                    # 2) Font/punto (effective)
                    fn = effective_font_name(p)
                    fs = effective_font_size_pt(p)
                    if fn and fn != expected_name:
                        errors.append(f"{loc}. satır yazı tipi {fn} yerine {expected_name} olmalı")
                    if fs is not None and round(float(fs), 1) != round(float(expected_size), 1):
                        errors.append(f"{loc}. satır punto {fs} yerine {expected_size} olmalı")

                    # 3) Spacing (effective)
                    ls = effective_line_spacing(p, default=expected_spacing)
                    sb = effective_space_pt(p, "before")
                    sa = effective_space_pt(p, "after")

                    if ls is not None and round(float(ls), 2) != round(float(expected_spacing), 2):
                        errors.append(f"{loc}. satırın satır aralığı {ls} yerine {expected_spacing} olmalı")
                    if round(float(sb), 1) != round(float(expected_before), 1):
                        errors.append(f"{loc}. satır öncesi {sb} yerine {expected_before} olmalı")
                    if round(float(sa), 1) != round(float(expected_after), 1):
                        errors.append(f"{loc}. satır sonrası {sa} yerine {expected_after} olmalı")

                if errors:
                    results.append((authority_start, False, rule_title, "| ".join(errors)))
                else:
                    results.append((authority_start, True, rule_title, ""))

    # ======================================================
    # MAKAM BLOĞU ÖNCESİ BOŞLUK KONTROLÜ (TERS YÖNLÜ, MAKSİMUM KISITLI)
    # ======================================================
    elif check["check"] == "approval_authority_spacing_reverse":
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        expected_blank_min = int(check.get("expected_blank_min", 2))
        expected_blank_max = int(check.get("expected_blank_max", 3))
        debug_mode = check.get("debug", False)

        rule_title = (
            f"MAKAM BLOĞU ÖNCESİ {expected_blank_min}-{expected_blank_max} SATIR BOŞLUK KONTROLÜ\n"
            f"{expected_name}, {int(expected_size)} punto, {expected_spacing} satır aralığı, "
            f"önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        start_idx = memo.get("authority_start_idx", None)
        if start_idx is None:
            results.append((0, False, rule_title, "Makam bloğu satırı bulunamadı, kontrol yapılamadı"))
        else:
            # --- yardımcılar ---
            def is_blank_para(p) -> bool:
                t = (p.text or "")
                t = t.replace("\u200b", "").replace("\xa0", "").replace("\u00A0", "").strip()
                return (t == "") or ("".join(t.split()) == "")

            def eff_font_name(p):
                for r in p.runs:
                    if r.font and r.font.name:
                        return r.font.name
                if p.style and p.style.font and p.style.font.name:
                    return p.style.font.name
                return None

            def eff_font_size_pt(p):
                for r in p.runs:
                    if r.font and r.font.size:
                        return r.font.size.pt
                if p.style and p.style.font and p.style.font.size:
                    return p.style.font.size.pt
                return None

            def eff_spacing(p):
                pf = p.paragraph_format
                ls = pf.line_spacing
                sb = pf.space_before.pt if pf.space_before else 0.0
                sa = pf.space_after.pt if pf.space_after else 0.0
                return ls, sb, sa

            # --- debug ---
            dbg = None
            if debug_mode:
                dbg = open(dbg_path("debug_approval_page.txt"), "a", encoding="utf-8")
                dbg.write("\n\n===========================================\n")
                dbg.write("[DEBUG] Makam Bloğu Öncesi Boşluk Kontrolü (max-kısıtlı)\n")
                dbg.write(f"Makam bloğu başlangıç satırı: {start_idx}\n")
                dbg.write(f"Alt-üst limit: {expected_blank_min}-{expected_blank_max}\n")
                dbg.write("===========================================\n")

            errors = []
            blank_count = 0

            # 🔴 kritik değişiklik: yalnızca en fazla expected_blank_max kadar geriye bak
            for k in range(1, expected_blank_max + 1):
                idx = start_idx - k
                if idx < 0:
                    break

                p = paragraphs[idx]
                is_blank = is_blank_para(p)

                if dbg:
                    preview = (p.text or "").strip().replace("\n", " ")
                    if len(preview) > 60:
                        preview = preview[:60] + "..."
                    dbg.write(f"[{idx}. satır] {'(boş)' if is_blank else '(dolu)'} → {preview or '(boş paragraf)'}\n")

                if not is_blank:
                    # makam bloğuna en yakın dolu satıra çarptık; geriye bakmayı bırak
                    if dbg:
                        dbg.write("→ DOLU satıra gelindi, sayım durduruldu.\n")
                    break

                # boşsa say
                blank_count += 1

                # maksimuma ulaşıldıysa daha yukarıyı (ör. onay sonrası 2 boşluk) bilerek sayma
                if blank_count >= expected_blank_max:
                    if dbg:
                        dbg.write(f"→ Maksimum {expected_blank_max} boş satıra ulaşıldı, sayım durduruldu.\n")
                    break

            if dbg:
                dbg.write(f"Toplam sayılan (sadece en yakındaki) boş satır: {blank_count}\n")
                dbg.write("-------------------------------------------\n\n")
                dbg.close()

            # --- sonuç kararı ---
            if expected_blank_min <= blank_count <= expected_blank_max:
                results.append((start_idx, True, rule_title, "Boşluk sayısı uygun"))
            else:
                results.append(
                    (
                        start_idx,
                        False,
                        rule_title,
                        f"Boşluk sayısı {blank_count}, {expected_blank_min}-{expected_blank_max} aralığında olmalı"
                    )
                )

    # ======================================================
    # 5846 SATIRININ ÜSTÜNDE 2 SATIR BOŞLUK (dinamik + memo)
    # ======================================================
    elif check["check"] == "approval_authority_spacing":
        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 0))

        rule_title = (
            f"ONAY MAKAMI SONRASI BOŞLUK\n"
            f"2 satır, {expected_name}, {int(expected_size)} punto, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        def is_blank_para(p) -> bool:
            t = (p.text or "")
            t = t.replace("\u200b", "").replace("\xa0", "").replace("\u00A0", "").strip()
            return (t == "") or ("".join(t.split()) == "")

        # --- 5846 satırını bul (memo öncelikli)
        onay_son = memo.get("approval_5846_idx")
        if onay_son is None:
            for i, p in enumerate(paragraphs):
                if "5846" in (p.text or ""):
                    onay_son = i
                    memo["approval_5846_idx"] = i
                    break

        if onay_son is None:
            results.append((0, False, rule_title, "5846 bulunamadı"))
        else:
            # 2 satır üst kontrol: onay_son-1 ve onay_son-2
            errors = []

            for k in range(1, 3):
                idx = onay_son - k
                loc = f"{k}. boşluk satırı (belge:{idx})"

                if idx < 0:
                    errors.append(f"{loc} — eksik")
                    continue

                p = paragraphs[idx]
                
                txt = (p.text or "").strip()
                pv = " ".join(txt.split()[:7])
                if len(pv) > 70:
                    pv = pv[:70].rstrip() + "…"
                loc = loc + (f" ('{pv}')" if pv else "")

                # 1) Satır boş olmalı
                if not is_blank_para(p):
                    errors.append(f"{loc} - boş değil")
                    continue

                # 2) Biçim kontrolü (effective)
                fn = effective_font_name(p)
                fs = effective_font_size_pt(p)
                if fn and fn != expected_name:
                    errors.append(f"{loc}- yazı tipi {fn} yerine {expected_name} olmalı")
                if fs is not None and round(float(fs), 1) != round(float(expected_size), 1):
                    errors.append(f"{loc}- punto {fs} yerine {expected_size} olmalı")

                ls = effective_line_spacing(p, default=expected_spacing)
                sb = effective_space_pt(p, "before")
                sa = effective_space_pt(p, "after")

                if ls is not None and round(float(ls), 2) != round(float(expected_spacing), 2):
                    errors.append(f"{loc}- satır aralığı {ls} yerine {expected_spacing} olmalı")
                if round(float(sb), 1) != round(float(expected_before), 1):
                    errors.append(f"{loc}- öncesi {sb} yerine {expected_before} olmalı")
                if round(float(sa), 1) != round(float(expected_after), 1):
                    errors.append(f"{loc}- sonrası {sa} yerine {expected_after} olmalı")

            results.append(((onay_son - 2) if onay_son >= 2 else onay_son, len(errors) == 0, rule_title, "; ".join(errors)))

    # ======================================================
    # COPYRIGHT NOTICE (TEZ ONAYI SAYFASI SONU)
    # ======================================================
    elif check["check"] == "copyright_notice":
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 10))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        expected_italic = check.get("italic", True)
        target_text = check.get("text_contains", "").strip()

        rule_title = (
            f"COPYRIGHT NOTICE\n"
            f"{expected_name}, {int(expected_size)} punto, italik, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        # ------------------------------------------------------
        # Normalizasyon (KORUNDU)
        # Amaç: Türkçe karakterleri normalize edip aramayı sağlamlaştırmak
        # ------------------------------------------------------
        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı": "i", "İ": "i", "ç": "c", "Ç": "c", "ğ": "g", "Ğ": "g",
                "ö": "o", "Ö": "o", "ş": "s", "Ş": "s", "ü": "u", "Ü": "u"
            })
            return (s or "").lower().translate(trans).strip()

        # ------------------------------------------------------
        # NOT: effective_italic artık burada TANIMLANMIYOR.
        # run_check yardımcı fonksiyonlarında global olarak var:
        #   def effective_italic(para): ...
        # Böylece tüm kurallar aynı italik okuma standardını kullanır.
        # ------------------------------------------------------

        target_norm = norm_tr(target_text).replace(" ", "")

        notice_idx = None
        for i, p in enumerate(paragraphs):
            txt_norm = norm_tr(p.text).replace(" ", "")
            if target_norm in txt_norm:  # text_contains mantığı (KORUNDU)
                notice_idx = i
                break

        if notice_idx is None:
            results.append((0, False, rule_title, "Copyright metni bulunamadı"))
        else:
            # 🔹 Telif hakkı satırı bulunduğunda memo’ya kaydet (AYNI İSİM)
            memo["copyright_notice_index"] = notice_idx

            p = paragraphs[notice_idx]
            errors = []

            # ------------------------------------------------------
            # Yazı tipi / punto / italik (EFFECTIVE)
            # effective_* fonksiyonları:
            #   - run + style zincirinden okur
            #   - run'da None olsa bile Word'deki görünen biçimi yakalar
            # ------------------------------------------------------
            fn = effective_font_name(p)         # run + style zinciri
            fs = effective_font_size_pt(p)      # run + style zinciri

            if fn and fn != expected_name:
                errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")
            if fs is not None and round(float(fs), 1) != round(float(expected_size), 1):
                errors.append(f"Punto {fs} yerine {expected_size} olmalı")

            # İtalik kontrolü (EFFECTIVE) ✅ KRİTİK
            # Artık lokal fonksiyon değil, helper bölümündeki global effective_italic kullanılır.
            if expected_italic and not effective_italic(p):
                errors.append("Metin italik değil")

            # ------------------------------------------------------
            # Satır aralığı / boşluklar (EFFECTIVE)
            # ------------------------------------------------------
            ls = effective_line_spacing(p, default=expected_spacing)
            sb = effective_space_pt(p, "before")
            sa = effective_space_pt(p, "after")

            if ls is not None and round(float(ls), 2) != round(float(expected_spacing), 2):
                errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
            if round(float(sb), 1) != round(float(expected_before), 1):
                errors.append(f"Öncesi {sb} yerine {expected_before} olmalı")
            if round(float(sa), 1) != round(float(expected_after), 1):
                errors.append(f"Sonrası {sa} yerine {expected_after} olmalı")

            if errors:
                results.append((notice_idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((notice_idx, True, rule_title, ""))

    # ===============================================================================================================#
    # ===============================================================================================================#    
    # ===============================================================================================================#
    # ======================================================
    # ======================================================
    # TEZ ÇALIŞMASI ETİK BEYANI - BAŞLIK
    # ======================================================
    elif check["check"] == "ethics_statement_heading":
        import re
        import unicodedata

        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.5))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))

        # YAML’dan expected_text okumak istemiyorsan default’u burada zaten gömülü:
        expected_text = check.get("expected_text", "TEZ ÇALIŞMASI ETİK BEYANI")
        markers = check.get("markers", [])

        rule_title = (
            f"TEZ ÇALIŞMASI ETİK BEYANI - BAŞLIK\n"
            f"{expected_name}, {int(expected_size)} punto, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        # -------------------------------------------------------
        # Türkçe normalizasyon (metin karşılaştırması için)
        # - i + combining dot (i̇) gibi durumları güvenle temizler
        # -------------------------------------------------------
        def norm_tr(s: str) -> str:
            if s is None:
                return ""

            s = str(s)

            # i̇ gibi birleşik/combining karakterleri parçala
            s = unicodedata.normalize("NFKD", s)
            # combining dot above (U+0307) temizle
            s = s.replace("\u0307", "")

            s = " ".join(s.strip().split()).lower()
            s = (s.replace("ı", "i").replace("İ", "i")
                .replace("ç", "c").replace("ğ", "g").replace("ö", "o")
                .replace("ş", "s").replace("ü", "u"))
            return s

        # ✅ expected_norm’u güvenli şekilde koda göm (YAML’dan ayrıca okumaya gerek yok)
        expected_norm = norm_tr(expected_text)

        # -------------------------------------------------------
        # Başlığı regex marker ile bul
        # -------------------------------------------------------
        found_idx = None
        used_marker = None

        for i, p in enumerate(paragraphs):
            raw = (p.text or "").strip()
            for m in markers:
                if re.match(m, raw, re.IGNORECASE):
                    found_idx = i
                    used_marker = m
                    break
            if found_idx is not None:
                break

        if found_idx is None:
            results.append((
                0,
                False,
                rule_title,
                f"Başlık '{expected_text}' bulunamadı"
            ))
            return results

        # -------------------------------------------------------
        # Memo kaydı (ZİNCİR KOPMASIN DİYE HER ZAMAN)
        # -------------------------------------------------------
        memo["ethics_statement_heading_idx"] = found_idx

        p = paragraphs[found_idx]
        raw_heading = (p.text or "").strip()
        heading_norm = norm_tr(raw_heading)

        errors = []

        # -------------------------------------------------------
        # Metin doğruluğu (SADECE BURADA KATI)
        # -------------------------------------------------------
        if heading_norm != expected_norm:
            errors.append(
                f"Başlık '{raw_heading}' yerine '{expected_text}' olmalı"
            )

        # -------------------------------------------------------
        # Yazı tipi (run → stil)
        # -------------------------------------------------------
        fn = None
        for r in p.runs:
            if r.font and r.font.name:
                fn = r.font.name
                break

        style_font = getattr(p.style, "font", None)
        style_name = style_font.name if style_font and style_font.name else None
        actual_font = fn or style_name

        if actual_font and actual_font != expected_name:
            errors.append(f"Yazı tipi {actual_font} yerine {expected_name} olmalı")

        # -------------------------------------------------------
        # Punto (run → stil)
        # -------------------------------------------------------
        fs = None
        for r in p.runs:
            if r.font and r.font.size:
                fs = r.font.size.pt
                break

        style_size = style_font.size.pt if style_font and style_font.size else None
        actual_size = fs if fs is not None else style_size

        if actual_size is not None and round(float(actual_size), 1) != round(expected_size, 1):
            errors.append(f"Punto {actual_size} yerine {expected_size} olmalı")

        # -------------------------------------------------------
        # Satır aralığı ve boşluklar (pf)
        # -------------------------------------------------------
        pf = p.paragraph_format
        ls = pf.line_spacing
        sb = pf.space_before.pt if pf.space_before else 0.0
        sa = pf.space_after.pt if pf.space_after else 0.0

        if ls is not None and round(float(ls), 2) != round(expected_spacing, 2):
            errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
        if round(sb, 1) != round(expected_before, 1):
            errors.append(f"Öncesi {sb} yerine {expected_before} olmalı")
        if round(sa, 1) != round(expected_after, 1):
            errors.append(f"Sonrası {sa} yerine {expected_after} olmalı")

        # -------------------------------------------------------
        # Sonuç
        # -------------------------------------------------------
        if errors:
            results.append((found_idx, False, rule_title, "; ".join(errors)))
        else:
            results.append((found_idx, True, rule_title, ""))

        return results



    #==============================================
    # TEZ ÇALIŞMASI ETİK BEYANI - BAŞLIK SONRASI BOŞLUK KONTROLÜ
    # - Başlıktan sonra 2 boş satır olmalı
    # - Boş satırlar: 1.5 satır aralığı, önce/sonra 0
    # - Boş satır font/punto kontrolü: run → XML(pPr/rPr) → stil
    # - Dinamik offset: memo["ethics_statement_heading_idx"] sonrası say
    # ======================================================
    elif check["check"] == "ethics_statement_post_heading_spacing":
        import os

        # ✅ rules.yaml uyumlu: expected_lines (tek değer) öncelikli
        expected_lines = check.get("expected_lines", None)
        if expected_lines is None:
            # geriye dönük uyumluluk: blank_lines kullanan varsa
            expected_lines = check.get("blank_lines", 2)

        expected_blank_min = int(expected_lines)
        expected_blank_max = int(expected_lines)

        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))

        expected_spacing = float(check.get("line_spacing", 1.5))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 0))

        debug_mode = check.get("debug", False)

        rule_title = check.get(
            "description",
            f"TEZ ÇALIŞMASI ETİK BEYANI - BAŞLIK SONRASI BOŞLUK\n"
            f"Başlıktan sonra {expected_blank_min}-{expected_blank_max} satır boşluk olmalı. "
            f"Boş satırlar: {expected_name}, {int(expected_size)} punto; "
            f"satır aralığı={expected_spacing}, önce={expected_before}, sonra={expected_after}."
        )

        # -------------------------------------------------------
        # 1️⃣ DEBUG dosyası (append)
        # -------------------------------------------------------
        debug_file = None
        if debug_mode:
            debug_path = dbg_path("debug_ethics_statement.txt")
            new_file = not os.path.exists(debug_path)
            debug_file = open(debug_path, "a", encoding="utf-8")
            if new_file:
                debug_file.write("[DEBUG] Etik Beyan - Başlık Sonrası Boşluk Kontrolü\n")
                debug_file.write("===========================================\n\n")
            else:
                debug_file.write("\n-------------------------------------------\n")
                debug_file.write("Yeni kontrol çalıştırması başlatıldı (Etik beyan boşluk).\n\n")

        # -------------------------------------------------------
        # 2️⃣ Başlık index'ini MEMO'dan oku
        # -------------------------------------------------------
        heading_idx = memo.get("ethics_statement_heading_idx", None)

        if debug_file:
            debug_file.write(f"ℹ️ Memo'dan alınan heading_idx: {heading_idx}\n")

        if heading_idx is None or heading_idx >= len(paragraphs):
            results.append((0, False, rule_title, "Etik beyan başlık index'i memo’da bulunamadı."))
            if debug_file:
                debug_file.write("❌ heading_idx memo’da yok veya geçersiz.\n")
                debug_file.close()
            return results

        # -------------------------------------------------------
        # 3️⃣ Başlıktan sonraki boş satırları say (Dinamik offset)
        # -------------------------------------------------------
        blank_count = 0
        for p in paragraphs[heading_idx + 1:]:
            if not p.text.strip():
                blank_count += 1
            else:
                break

        if debug_file:
            debug_file.write(f"ℹ️ Başlıktan sonra bulunan boş satır sayısı: {blank_count}\n")

        # -------------------------------------------------------
        # 3B) Boş satırların biçim kontrolü ✅
        #     - line_spacing: pf → stil
        #     - before/after: sadece pf varsa, yoksa 0.0
        #     - font/punto: run → XML(pPr/rPr) → stil
        # -------------------------------------------------------
        errors = []

        for k in range(1, blank_count + 1):
            p_blank = paragraphs[heading_idx + k]
            pf = p_blank.paragraph_format

            # line_spacing (pf/stil)
            ls = pf.line_spacing or getattr(p_blank.style.paragraph_format, "line_spacing", None)

            # space_before/after (sadece pf varsa, yoksa 0.0)
            sb = (pf.space_before.pt if pf.space_before else 0.0)
            sa = (pf.space_after.pt if pf.space_after else 0.0)

            # font adı / punto: run → XML(pPr/rPr) → stil
            fn = None
            fs = None

            # 1) Run'dan oku
            for r in p_blank.runs:
                if r.font:
                    if not fn and r.font.name:
                        fn = r.font.name
                    if fs is None and r.font.size:
                        fs = r.font.size.pt

            # 2) XML pPr/rPr fallback
            if fs is None or fn is None:
                try:
                    if fs is None:
                        sz_vals = p_blank._p.xpath("./w:pPr/w:rPr/w:sz/@w:val")
                        if sz_vals:
                            fs = float(int(sz_vals[0])) / 2.0

                    if fn is None:
                        font_vals = p_blank._p.xpath("./w:pPr/w:rPr/w:rFonts/@w:ascii")
                        if not font_vals:
                            font_vals = p_blank._p.xpath("./w:pPr/w:rPr/w:rFonts/@w:hAnsi")
                        if font_vals:
                            fn = font_vals[0]
                except Exception:
                    pass

            # 3) Stil fallback
            style_font = getattr(p_blank.style, "font", None)
            style_name = style_font.name if style_font and style_font.name else None
            style_size = style_font.size.pt if style_font and style_font.size else None

            actual_font = fn or style_name
            actual_size = fs if fs is not None else style_size

            # --- Hata kontrolü
            if actual_font and actual_font != expected_name:
                errors.append(f"{k}. boş satırda yazı tipi {actual_font} yerine {expected_name} olmalı")
            if actual_size is not None and round(float(actual_size), 1) != round(expected_size, 1):
                errors.append(f"{k}. boş satırda punto {actual_size} yerine {expected_size} olmalı")
            if ls is not None and round(float(ls), 2) != round(expected_spacing, 2):
                errors.append(f"{k}. boş satırda satır aralığı {ls} yerine {expected_spacing} olmalı")
            if round(sb, 1) != round(expected_before, 1):
                errors.append(f"{k}. boş satırda paragraf öncesi {sb} yerine {expected_before} olmalı")
            if round(sa, 1) != round(expected_after, 1):
                errors.append(f"{k}. boş satırda paragraf sonrası {sa} yerine {expected_after} olmalı")

            if debug_file:
                debug_file.write(
                    f"  {k}. boş satır: Font={actual_font}, Size={actual_size}, "
                    f"LS={ls}, Before={sb}, After={sa}\n"
                )

        # -------------------------------------------------------
        # 4️⃣ Boş satır sayısı değerlendirmesi
        # -------------------------------------------------------
        if blank_count < expected_blank_min or blank_count > expected_blank_max:
            msg = f"{blank_count} satır boşluk var, {expected_blank_min}-{expected_blank_max} satır olmalı."
            # Sayı hatası varsa, biçim hatalarını da ekle
            results.append((heading_idx, False, rule_title, msg + (("; " + "; ".join(errors)) if errors else "")))
            if debug_file:
                debug_file.write(f"\n❌ {msg}\n")
                if errors:
                    debug_file.write("❌ Boş satır biçim hataları:\n")
                    for e in errors:
                        debug_file.write(f" - {e}\n")
        else:
            # Sayı doğruysa sadece biçim hataları kaldı mı bak
            if errors:
                results.append((heading_idx, False, rule_title, "; ".join(errors)))
                if debug_file:
                    debug_file.write("\n❌ Boş satır biçim hataları:\n")
                    for e in errors:
                        debug_file.write(f" - {e}\n")
            else:
                results.append((heading_idx, True, rule_title, ""))
                if debug_file:
                    debug_file.write("\n✅ Boşluk sayısı ve biçimi uygun.\n")

        if debug_file:
            debug_file.write("-------------------------------------------\n")
            debug_file.write("Kontrol tamamlandı.\n")
            debug_file.close()

        return results

    # ======================================================
    # TEZ ÇALIŞMASI ETİK BEYANI - GÖVDE
    # ======================================================
    elif check["check"] == "ethics_statement_body":
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.5))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        markers = check.get("markers", ["^TEZ ÇALIŞMASI ETİK BEYANI$"])
        until_marker = check.get("until_marker")  # genelde "^ÖZET$"

        import re

        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u"
            })
            return " ".join(((s or "").lower().translate(trans)).split())

        # YAML markers/regex başlarını/sonlarını temizleyip normalize et
        norm_markers = [norm_tr(re.sub(r'^\^|\$$', '', m)) for m in markers]
        norm_until = norm_tr(re.sub(r'^\^|\$$', '', until_marker)) if until_marker else None

        # until_marker yoksa ÖZET/ABSTRACT yedeği
        variant_names = ["ÖZET", "ABSTRACT"]
        norm_variants = [norm_tr(v) for v in variant_names]

        # -------------------------------------------------------
        # ✅ 1) start_idx: memo öncelikli, yoksa fallback arama
        # -------------------------------------------------------
        start_idx = memo.get("ethics_statement_heading_idx", None)
        if start_idx is not None:
            if not (0 <= int(start_idx) < len(paragraphs)):
                start_idx = None  # geçersizse fallback'a düş

        end_idx = None

        if start_idx is None:
            for i, p in enumerate(paragraphs):
                raw = (p.text or "").strip()
                nrm = norm_tr(raw)

                if start_idx is None:
                    for nm in norm_markers:
                        if nrm == nm:
                            start_idx = i
                            break

                if end_idx is None:
                    if norm_until:
                        if nrm == norm_until:
                            end_idx = i
                    else:
                        if nrm in norm_variants:
                            end_idx = i

                if start_idx is not None and end_idx is not None:
                    break
        else:
            for i in range(start_idx + 1, len(paragraphs)):
                raw = (paragraphs[i].text or "").strip()
                nrm = norm_tr(raw)

                if end_idx is None:
                    if norm_until:
                        if nrm == norm_until:
                            end_idx = i
                            break
                    else:
                        if nrm in norm_variants:
                            end_idx = i
                            break

        rule_title = (
            f"ETİK BEYAN GÖVDE: {expected_name}, {int(expected_size)} punto, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        if start_idx is None or end_idx is None:
            msgs = []
            if start_idx is None:
                msgs.append("Etik Beyan başlığı bulunamadı (aranan: " + ", ".join(markers) + ")")
            if end_idx is None:
                if until_marker:
                    msgs.append(f"ÖZET başlığı bulunamadı (aranan until_marker: {until_marker})")
                else:
                    msgs.append("ÖZET/ABSTRACT başlığı bulunamadı (aranan: " + ", ".join(variant_names) + ")")
            results.append((0, False, rule_title, " ; ".join(msgs)))
        else:
            errors = []

            # -------------------------------------------------------
            # ✅ 2) Dinamik offset: başlıktan sonra boşları geç
            # -------------------------------------------------------
            j0 = start_idx + 1
            while j0 < end_idx and norm_tr((paragraphs[j0].text or "").strip()) == "":
                j0 += 1

            for j in range(j0, end_idx):
                p = paragraphs[j]
                txt = (p.text or "").strip()

                if norm_tr(txt) == "":
                    continue

                preview = txt[:30].replace("\n", " ") + ("..." if len(txt) > 30 else "")

                # Yazı tipi
                fn = None
                for r in p.runs:
                    if r.font and r.font.name:
                        fn = r.font.name
                        break
                style_font = getattr(p.style, "font", None)
                style_name = style_font.name if style_font and style_font.name else None
                actual_font = fn or style_name

                if actual_font and actual_font != expected_name:
                    errors.append(
                        f"Satır {j+1} ('{preview}'): yazı tipi {actual_font} yerine {expected_name} olmalı"
                    )

                # Punto
                fs = None
                for r in p.runs:
                    if r.font and r.font.size:
                        fs = r.font.size.pt
                        break
                style_size = style_font.size.pt if style_font and style_font.size else None
                actual_size = fs if fs is not None else style_size

                if actual_size is not None and round(float(actual_size), 1) != round(expected_size, 1):
                    errors.append(
                        f"Satır {j+1} ('{preview}'): punto {actual_size} yerine {expected_size} olmalı"
                    )

                # Satır aralığı
                pf = p.paragraph_format
                ls = pf.line_spacing or getattr(p.style.paragraph_format, "line_spacing", None)

                sb = pf.space_before.pt if pf.space_before else 0.0
                sa = pf.space_after.pt if pf.space_after else 0.0

                if ls is not None and round(float(ls), 2) != round(expected_spacing, 2):
                    errors.append(
                        f"Satır {j+1} ('{preview}'): satır aralığı {ls} yerine {expected_spacing} olmalı"
                    )
                if round(sb, 1) != round(expected_before, 1):
                    errors.append(
                        f"Satır {j+1} ('{preview}'): öncesi {sb} yerine {expected_before} olmalı"
                    )
                if round(sa, 1) != round(expected_after, 1):
                    errors.append(
                        f"Satır {j+1} ('{preview}'): sonrası {sa} yerine {expected_after} olmalı"
                    )

            if errors:
                results.append((start_idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((start_idx, True, rule_title, ""))

    # ======================================================
    # TEZ ÇALIŞMASI ETİK BEYANI - ÜYZ BEYANLARI ÇAKIŞMA KONTROLÜ
    # - "Kullanmadım" ve "Etik rehber çerçevesinde..." iki beyan
    # - İkisi aynı anda varsa ihlal
    # ======================================================
    elif check["check"] == "ethics_statement_ai_statement_exclusive":
        import re

        # --- başlık bulma için (memo yoksa)
        markers = check.get("markers", ["^TEZ ÇALIŞMASI ETİK BEYANI$"])
        until_marker = check.get("until_marker")  # örn "^ÖZET$"

        # --- iki beyanın yakalama regexleri (YAML'dan gelsin)
        # ai_statement_patterns:
        #   no_ai: [...]
        #   ai_under_guidelines: [...]
        patterns_cfg = check.get("ai_statement_patterns", {}) or {}
        pat_no_ai_list = patterns_cfg.get("no_ai", []) or []
        pat_ai_guidelines_list = patterns_cfg.get("ai_under_guidelines", []) or []

        rule_title = check.get(
            "description",
            "ETİK BEYAN GÖVDE - ÜYZ Beyanları: 'kullanmadım' ve 'etik rehber çerçevesinde' beyanları aynı anda yer alamaz."
        )

        # -------------------------------------------------------
        # Türkçe normalizasyon (HATA GİDERİLDİ)
        # - maketrans() sadece 1-char key kabul eder.
        # - Word bazen 'i' + combining dot (U+0307) şeklinde getirir.
        # -------------------------------------------------------
        def norm_tr(s: str) -> str:
            if s is None:
                return ""

            s = str(s)

            # 1) combining dot above'ı temizle: "i\u0307" -> "i"
            #    Böylece "i̇" gibi iki codepointli durumlar güvenli hale gelir.
            s = s.replace("\u0307", "")

            # 2) tek karakterlik dönüşümler
            trans = str.maketrans({
                "ı": "i", "İ": "i",
                "ç": "c", "Ç": "c",
                "ğ": "g", "Ğ": "g",
                "ö": "o", "Ö": "o",
                "ş": "s", "Ş": "s",
                "ü": "u", "Ü": "u"
            })

            # 3) whitespace normalize
            return " ".join(s.lower().translate(trans).split())

        # YAML markers/regex başlarını/sonlarını temizleyip normalize et (eşitlik karşılaştırması için)
        norm_markers = [norm_tr(re.sub(r'^\^|\$$', '', m)) for m in markers]
        norm_until = norm_tr(re.sub(r'^\^|\$$', '', until_marker)) if until_marker else None
        variant_names = ["ÖZET", "ABSTRACT"]
        norm_variants = [norm_tr(v) for v in variant_names]

        # -------------------------------------------------------
        # 1) Başlık index (memo → fallback arama)
        # -------------------------------------------------------
        start_idx = memo.get("ethics_statement_heading_idx", None)

        if start_idx is None:
            # fallback: markers ile başlığı bul
            found = None
            for i, p in enumerate(paragraphs):
                raw = (p.text or "").strip()
                nrm = norm_tr(raw)
                for nm in norm_markers:
                    if nrm == nm:
                        found = i
                        break
                if found is not None:
                    break

            if found is None:
                results.append((0, False, rule_title, "Etik Beyan başlığı bulunamadı (ÜYZ çakışma kontrolü yapılamadı)."))
                return results

            start_idx = found
            memo["ethics_statement_heading_idx"] = start_idx

        # -------------------------------------------------------
        # 2) Bitiş index (until_marker → yoksa ÖZET/ABSTRACT)
        # -------------------------------------------------------
        end_idx = None
        for i in range(start_idx + 1, len(paragraphs)):
            raw = (paragraphs[i].text or "").strip()
            nrm = norm_tr(raw)
            if norm_until:
                if nrm == norm_until:
                    end_idx = i
                    break
            else:
                if nrm in norm_variants:
                    end_idx = i
                    break

        if end_idx is None:
            # bitiş yoksa: belge sonuna kadar tara (toleranslı)
            end_idx = len(paragraphs)

        # -------------------------------------------------------
        # 3) Gövde içinde iki beyanın varlığını ara
        # -------------------------------------------------------
        found_no_ai = False
        found_ai_guidelines = False
        where_no_ai = []
        where_ai_guidelines = []

        # Regex derle (CASE-INSENSITIVE)
        re_no_ai = [re.compile(p, re.IGNORECASE) for p in pat_no_ai_list if p]
        re_ai_gl = [re.compile(p, re.IGNORECASE) for p in pat_ai_guidelines_list if p]

        for j in range(start_idx + 1, end_idx):
            txt = (paragraphs[j].text or "").strip()
            if not txt:
                continue

            for rx in re_no_ai:
                if rx.search(txt):
                    found_no_ai = True
                    where_no_ai.append(j)
                    break

            for rx in re_ai_gl:
                if rx.search(txt):
                    found_ai_guidelines = True
                    where_ai_guidelines.append(j)
                    break

            if found_no_ai and found_ai_guidelines:
                break

        # -------------------------------------------------------
        # 4) Kural değerlendirmesi  (✅ XOR: tam olarak 1 beyan olmalı)
        #    ✅ Bu blok DÖNGÜNÜN DIŞINDA olmalı!
        # -------------------------------------------------------
        if found_no_ai and found_ai_guidelines:
            msg = (
                "Etik beyan gövdesinde ÜYZ ile ilgili iki alternatif beyan aynı anda bulunuyor. "
                "Öğrenci yalnızca birini seçmeli."
            )
            if where_no_ai:
                msg += f" | 'kullanmadım' benzeri beyan: paragraf index {where_no_ai[0]}"
            if where_ai_guidelines:
                msg += f" | 'etik rehber çerçevesinde' benzeri beyan: paragraf index {where_ai_guidelines[0]}"
            results.append((start_idx, False, rule_title, msg))

        elif (not found_no_ai) and (not found_ai_guidelines):
            msg = (
                "Etik beyan gövdesinde ÜYZ ile ilgili zorunlu beyan bulunamadı. "
                "Öğrenci iki alternatif metinden birini seçip bırakmalı (diğerini silmeli)."
            )
            results.append((start_idx, False, rule_title, msg))

        else:
            results.append((start_idx, True, rule_title, ""))




        return results



    # ===============================================================================================================#
    # ===============================================================================================================#    
    # ===============================================================================================================#
    # TÜRKÇE ÖZET BAŞLIĞI  (toleranslı bul, katı doğrula + memo)
    # - markers: YAML’den çoklu regex listesi okuyabilir
    # - Yanlış yazımı yakalarsa: "başlık metni ÖZET olmalı" ihlali yazar
    # - Bulduğu index’i memo["abstract_tr_heading_idx"] olarak kaydeder
    # - Font/size/bold/alignment/spacing: effective_* zinciri ile kontrol eder
    # ======================================================
    elif check["check"] == "abstract_tr_heading":
        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 0))
        expected_bold    = check.get("bold", False)

        # alignment: "center/left/right/justify" (yaml)
        expected_align = (check.get("alignment", check.get("align", "")) or "").lower().strip()

        # markers: YAML list veya tek string olabilir
        markers = check.get("markers", [r"^ÖZET$"])
        if isinstance(markers, str):
            markers = [markers]

        import re

        # --- Yardımcılar (bu fonksiyon içinde lokal) ---
        def clean_text(s: str) -> str:
            # NBSP → space, tab → space, fazla boşlukları sadeleştir
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        def norm_tr_lower(s: str) -> str:
            # Türkçe karakter duyarsızlaştırma + lower
            trans = str.maketrans({
                "ı": "i", "İ": "i", "ç": "c", "Ç": "c",
                "ğ": "g", "Ğ": "g", "ö": "o", "Ö": "o",
                "ş": "s", "Ş": "s", "ü": "u", "Ü": "u"
            })
            return (s or "").translate(trans).lower()

        # ------------------------------------------------------
        # 1) Toleranslı bulma: markers ile eşleşen ilk paragrafı bul
        # ------------------------------------------------------
        found = None
        found_marker = None
        for i, p in enumerate(paragraphs):
            cand = clean_text(p.text)
            # marker’lar regex; boşluk/tab/NBSP toleransı cand’da zaten temizlendi
            for m in markers:
                if re.match(m, cand, re.IGNORECASE):
                    found = (i, p)
                    found_marker = m
                    break
            if found:
                break

        rule_title = (
            f"TÜRKÇE ÖZET Başlığı: {expected_name}, {int(expected_size)} punto, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}, "
            f"{'kalın' if expected_bold else 'normal'}, "
            f"{'center' if expected_align=='center' else (expected_align or 'any')}"
        )

        if not found:
            results.append((0, False, rule_title, "ÖZET başlığı bulunamadı (markers eşleşmedi)"))
        else:
            idx, p = found

            # ✅ memo’ya kaydet
            memo["abstract_tr_heading_idx"] = idx

            errors = []

            # ------------------------------------------------------
            # 2) Katı metin doğrulama: olması gereken tam metin "ÖZET"
            # (Bulduk ama yanlış yazılmışsa ihlal üret)
            # ------------------------------------------------------
            raw_title = clean_text(p.text)
            # Boşlukları sadeleştirip kontrol edelim
            # Türkçe case hassasiyeti için normalize ederek karşılaştıralım
            if norm_tr_lower(raw_title) != norm_tr_lower("ÖZET"):
                errors.append(f"Başlık metni ÖZET olmalı (bulunan: {raw_title})")

            # ------------------------------------------------------
            # 3) Font / Size (effective_* zinciri)
            # ------------------------------------------------------
            fn = effective_font_name(p)
            fs = effective_font_size_pt(p)

            if fn and fn != expected_name:
                errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")
            if fs is not None and round(float(fs), 1) != round(float(expected_size), 1):
                errors.append(f"Punto {fs} yerine {expected_size} olmalı")

            # ------------------------------------------------------
            # 4) Bold (effective)
            # ------------------------------------------------------
            if expected_bold and not effective_bold(p):
                errors.append("Başlık kalın değil")

            # ------------------------------------------------------
            # 5) Alignment (effective)
            # ------------------------------------------------------
            if expected_align:
                align_map = {
                    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                    "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                }
                expected_enum = align_map.get(expected_align)
                actual_enum = effective_alignment(p)

                if expected_enum is not None:
                    if actual_enum is None:
                        # Word default genelde LEFT sayılır; beklenen LEFT değilse hata
                        if expected_enum != WD_PARAGRAPH_ALIGNMENT.LEFT:
                            if expected_enum == WD_PARAGRAPH_ALIGNMENT.CENTER:
                                errors.append("Başlık ortalı değil")
                            elif expected_enum == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                                errors.append("Başlık sağa yaslı değil")
                            elif expected_enum == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                                errors.append("Başlık iki yana yaslı değil")
                    elif actual_enum != expected_enum:
                        if expected_enum == WD_PARAGRAPH_ALIGNMENT.CENTER:
                            errors.append("Başlık ortalı değil")
                        elif expected_enum == WD_PARAGRAPH_ALIGNMENT.LEFT:
                            errors.append("Başlık sola yaslı değil")
                        elif expected_enum == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                            errors.append("Başlık sağa yaslı değil")
                        elif expected_enum == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                            errors.append("Başlık iki yana yaslı değil")

            # ------------------------------------------------------
            # 6) Line spacing / before / after (effective)
            # ------------------------------------------------------
            ls = effective_line_spacing(p, default=expected_spacing)
            sb = effective_space_pt(p, "before")
            sa = effective_space_pt(p, "after")

            if ls is not None and round(float(ls), 2) != round(float(expected_spacing), 2):
                errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
            if round(float(sb), 1) != round(float(expected_before), 1):
                errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
            if round(float(sa), 1) != round(float(expected_after), 1):
                errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

            # ------------------------------------------------------
            # Sonuç
            # ------------------------------------------------------
            if errors:
                results.append((idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((idx, True, rule_title, ""))

    # ======================================================
    # TÜRKÇE ÖZET BAŞLIĞINDAN SONRA BOŞLUK (en az N + fazla boşluk ihlali + memo)
    # ======================================================
    elif check["check"] == "abstract_tr_spacing_after_heading":
        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 0))
        required_blank   = int(check.get("required_blank_lines", 5))

        # YAML markers desteği (yanlış yazımlar / varyasyonlar)
        markers = check.get("markers", [r"^ÖZET$"])

        import re

        # 1) Başlık idx: önce memo, yoksa markers ile ara
        heading_idx = memo.get("abstract_tr_heading_idx")
        if heading_idx is None:
            found = None
            for i, p in enumerate(paragraphs):
                cand = " ".join((p.text or "").replace("\u00A0", " ").strip().split())
                for m in markers:
                    if re.match(m, cand, re.IGNORECASE):
                        found = i
                        break
                if found is not None:
                    heading_idx = found
                    memo["abstract_tr_heading_idx"] = heading_idx
                    break

        rule_title = (
            f"ÖZET başlığından sonra en az {required_blank} satır boşluk: "
            f"{expected_name}, {int(expected_size)} punto, {expected_spacing} satır aralığı, "
            f"önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        if heading_idx is None:
            results.append((0, False, rule_title, "ÖZET başlığı bulunamadı"))
        else:
            errors = []

            # 2) Başlıktan sonra ardışık boş satırları say
            blank_count = 0
            j = heading_idx + 1
            while j < len(paragraphs):
                txt = (paragraphs[j].text or "").replace("\u00A0", " ").strip()
                if txt != "":
                    break
                blank_count += 1
                j += 1

            # 3) Bir sonraki dolu satırı memo’ya yaz (sonraki fonksiyonlar bunu kullansın)
            memo["abstract_tr_after_heading_first_content_idx"] = (heading_idx + 1 + blank_count)

            # 4) En az N kuralı
            if blank_count < required_blank:
                errors.append(f"Başlıktan sonra {required_blank} boş satır olmalı (bulunan: {blank_count})")

            # 5) İlk required_blank satırın biçimini kontrol et (varsa)
            for k in range(1, min(required_blank, blank_count) + 1):
                
                idx = heading_idx + k
                loc = f"{k}. satır (belge:{idx})"

                p = paragraphs[idx]
                
                txt_full = (p.text or "").replace("\u00A0", " ").strip()
                pv = " ".join(txt_full.split()[:7])
                if len(pv) > 70:
                    pv = pv[:70].rstrip() + "…"
                loc = loc + (f" ('{pv}')" if pv else "")

                

                # Boş olmalı
                if (p.text or "").replace("\u00A0", " ").strip() != "":
                    errors.append(f"{loc}- boş değil (boş olmalı)")
                    continue

                # Format kontrolleri (paragraf formatı + run/stil)
                fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
                fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
                if not fn and p.style and p.style.font and p.style.font.name:
                    fn = p.style.font.name
                if not fs and p.style and p.style.font and p.style.font.size:
                    fs = p.style.font.size.pt

                pf = p.paragraph_format
                ls = pf.line_spacing
                sb = pf.space_before.pt if pf.space_before else 0.0
                sa = pf.space_after.pt if pf.space_after else 0.0

                if fn and fn != expected_name:
                    errors.append(f"{loc}- yazı tipi {fn} yerine {expected_name} olmalı")
                if fs and round(float(fs), 1) != round(expected_size, 1):
                    errors.append(f"{loc}- punto {fs} yerine {expected_size} olmalı")
                if ls is not None and round(float(ls), 2) != round(expected_spacing, 2):
                    errors.append(f"{loc}-satır aralığı {ls} yerine {expected_spacing} olmalı")
                if round(sb, 1) != round(expected_before, 1):
                    errors.append(f"{loc}- öncesi {sb} yerine {expected_before} olmalı")
                if round(sa, 1) != round(expected_after, 1):
                    errors.append(f"{loc}- sonrası {sa} yerine {expected_after} olmalı")

            # 6) Fazla boşluk ihlali (zinciri bozmaz)
            if blank_count > required_blank:

                errors.append(
                    f"Fazladan boş satır var (+{blank_count - required_blank}); "
                    f"başlık sonrası boşluk sayısı={blank_count}, beklenen={required_blank}"
                )
            if errors:
                results.append((heading_idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((heading_idx, True, rule_title, ""))

    # ======================================================
    # TÜRKÇE ÖZET - TEZ BAŞLIĞI (sabit offset yok, memo kullan)
    # ======================================================
    elif check["check"] == "abstract_tr_thesis_title":
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_bold = check.get("bold", True)
        expected_align = (check.get("alignment", "center") or "").lower()
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        no_italic = bool(check.get("no_italic", False))  # yeni opsiyon: italik olmasın

        rule_title = (
            f"TÜRKÇE ÖZET - Tez Başlığı: {expected_name}, {int(expected_size)} punto, "
            f"kalın={expected_bold}, alignment={expected_align}, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}"
            + ("; italik YOK" if no_italic else "")
        )

        import re

        # 1) Başlık idx’yi bul (memo öncelikli)
        heading_idx = memo.get("abstract_tr_heading_idx")
        if heading_idx is None:
            # fallback: sadece ÖZET’i bul
            for i, p in enumerate(paragraphs):
                if re.match(r"^ÖZET$", (p.text or "").strip(), re.IGNORECASE):
                    heading_idx = i
                    memo["abstract_tr_heading_idx"] = i
                    break

        if heading_idx is None:
            results.append((0, False, rule_title, "ÖZET başlığı bulunamadı, tez başlığı kontrolü yapılamadı"))
        else:
            # 2) Tez başlığı satırı: heading sonrası boşlukları sayıp ilk dolu satır
            title_idx = memo.get("abstract_tr_after_heading_first_content_idx")
            if title_idx is None:
                j = heading_idx + 1
                while j < len(paragraphs):
                    if (paragraphs[j].text or "").replace("\u00A0", " ").strip() != "":
                        break
                    j += 1
                title_idx = j
                memo["abstract_tr_after_heading_first_content_idx"] = title_idx

            if title_idx >= len(paragraphs):
                results.append((heading_idx, False, rule_title, "Tez başlığı satırı yok"))
            else:
                p = paragraphs[title_idx]
                errors = []

                if not (p.text or "").strip():
                    errors.append("Tez başlığı satırı boş")

                # Font/punto
                fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
                fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
                if not fn and p.style and p.style.font and p.style.font.name:
                    fn = p.style.font.name
                if not fs and p.style and p.style.font and p.style.font.size:
                    fs = p.style.font.size.pt

                if fn and fn != expected_name:
                    errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")
                if fs and round(float(fs), 1) != round(expected_size, 1):
                    errors.append(f"Punto {fs} yerine {expected_size} olmalı")

                # Bold
                if expected_bold:
                    any_bold = any((r.bold or (r.font and r.font.bold)) for r in p.runs if (r.text or "").strip())
                    if not any_bold and p.style and p.style.font and getattr(p.style.font, "bold", None):
                        any_bold = True
                    if not any_bold:
                        errors.append("Tez başlığı kalın değil")

                # Italic yasaksa
                if no_italic:
                    any_italic = any((r.italic or (r.font and r.font.italic)) for r in p.runs if (r.text or "").strip())
                    if not any_italic and p.style and p.style.font and getattr(p.style.font, "italic", None):
                        any_italic = True
                    if any_italic:
                        errors.append("Tez başlığında italik karakter bulunmamalı")

                # Alignment
                align_val = p.alignment
                if not align_val and p.style and p.style.paragraph_format and p.style.paragraph_format.alignment:
                    align_val = p.style.paragraph_format.alignment

                if expected_align == "center" and align_val != WD_PARAGRAPH_ALIGNMENT.CENTER:
                    errors.append("Tez başlığı ortalı değil")
                elif expected_align == "left" and align_val != WD_PARAGRAPH_ALIGNMENT.LEFT:
                    errors.append("Tez başlığı sola yaslı değil")

                # Satır aralığı / boşluklar
                pf = p.paragraph_format
                ls = pf.line_spacing
                sb = pf.space_before.pt if pf.space_before else 0.0
                sa = pf.space_after.pt if pf.space_after else 0.0
                if ls is not None and round(float(ls), 2) != round(expected_spacing, 2):
                    errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
                if round(sb, 1) != round(expected_before, 1):
                    errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
                if round(sa, 1) != round(expected_after, 1):
                    errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

                # memo
                memo["abstract_tr_thesis_title_idx"] = title_idx

                if errors:
                    results.append((title_idx, False, rule_title, "; ".join(errors)))
                else:
                    results.append((title_idx, True, rule_title, ""))

    # ======================================================
    # TÜRKÇE ÖZET - TEZ BAŞLIĞI SONRASI BOŞLUK (en az N + fazla boşluk ihlali + memo)
    # ======================================================
    elif check["check"] == "abstract_tr_spacing_after_title":
        expected_blank_lines = int(check.get("expected_blank_lines", 2))
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))

        rule_title = (
            f"TÜRKÇE ÖZET - Tez Başlığı Sonrası: en az {expected_blank_lines} satır boşluk, "
            f"{expected_name}, {int(expected_size)} punto, {expected_spacing} satır aralığı, "
            f"önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        # 1) title_idx: memo’dan al, yoksa fail (zincir mantığı)
        title_idx = memo.get("abstract_tr_thesis_title_idx")
        if title_idx is None:
            results.append((0, False, rule_title, "Tez başlığı bulunamadı (memo yok). Önce abstract_tr_thesis_title çalışmalı."))
        else:
            errors = []

            # 2) Tez başlığından sonra ardışık boş satır say
            blank_count = 0
            j = title_idx + 1
            while j < len(paragraphs):
                txt = (paragraphs[j].text or "").replace("\u00A0", " ").strip()
                if txt != "":
                    break
                blank_count += 1
                j += 1

            # 3) Bir sonraki dolu satırı memo’ya yaz (yazar adı için)
            memo["abstract_tr_after_title_first_content_idx"] = (title_idx + 1 + blank_count)

            # 4) En az N kuralı
            if blank_count < expected_blank_lines:
                errors.append(f"Tez başlığından sonra {expected_blank_lines} boş satır olmalı (bulunan: {blank_count})")

            # 5) İlk expected_blank_lines boş satır format kontrolü (varsa)
            for k in range(1, min(expected_blank_lines, blank_count) + 1):
                idx = title_idx + k
                loc = f"{k}. satır (belge:{idx})"
                
                p = paragraphs[idx]
                
                txt_full = (p.text or "").replace("\u00A0", " ").strip()
                pv = " ".join(txt_full.split()[:7])
                if len(pv) > 70:
                    pv = pv[:70].rstrip() + "…"
                loc = loc + (f" ('{pv}')" if pv else "")


                if (p.text or "").replace("\u00A0", " ").strip() != "":
                    errors.append(f"{loc}- boş değil  (boş olmalı)")
                    continue

                fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
                fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
                if not fn and p.style and p.style.font and p.style.font.name:
                    fn = p.style.font.name
                if not fs and p.style and p.style.font and p.style.font.size:
                    fs = p.style.font.size.pt

                pf = p.paragraph_format
                ls = pf.line_spacing
                sb = pf.space_before.pt if pf.space_before else 0.0
                sa = pf.space_after.pt if pf.space_after else 0.0

                if fn and fn != expected_name:
                    errors.append(f"{loc}- yazı tipi {fn} yerine {expected_name} olmalı")
                if fs and round(float(fs), 1) != round(expected_size, 1):
                    errors.append(f"{loc}- punto {fs} yerine {expected_size} olmalı")
                if ls is not None and round(float(ls), 2) != round(expected_spacing, 2):
                    errors.append(f"{loc}- satır aralığı {ls} yerine {expected_spacing} olmalı")
                if round(sb, 1) != round(expected_before, 1):
                    errors.append(f"{loc}- öncesi {sb} yerine {expected_before} olmalı")
                if round(sa, 1) != round(expected_after, 1):
                    errors.append(f"{loc}-sonrası {sa} yerine {expected_after} olmalı")

            # 6) Fazladan boşluk ihlali (zinciri bozmaz)
            if blank_count > expected_blank_lines:

                errors.append(
                    f"Fazladan boş satır var (+{blank_count - expected_blank_lines}); "
                    f"başlangıç idx={title_idx}, boşluk sayısı={blank_count}"
                )


            if errors:
                results.append((title_idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((title_idx, True, rule_title, ""))

    # ======================================================
    # TÜRKÇE ÖZET - TEZ YAZARI ADI SOYADI (sabit offset yok, memo kullan)
    # ======================================================
    elif check["check"] == "abstract_tr_author_name":
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_bold = check.get("bold", True)
        expected_align = (check.get("alignment", "center") or "").lower()
        expected_uppercase = bool(check.get("uppercase", True))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))

        rule_title = (
            f"TÜRKÇE ÖZET - Tez Yazarı Adı Soyadı: {expected_name}, {int(expected_size)} punto, "
            f"kalın={expected_bold}, alignment={expected_align}, büyükharf={expected_uppercase}, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        import re

        # Türkçe büyük harf normalizasyonu (mevcut yaklaşımı koruyalım)
        def norm_tr_upper(s: str) -> str:
            trans = str.maketrans({"ı":"I","i":"İ","ç":"Ç","ğ":"Ğ","ö":"Ö","ş":"Ş","ü":"Ü"})
            return (s or "").translate(trans).upper().strip()

        # 1) Tez başlığı idx (memo)
        title_idx = memo.get("abstract_tr_thesis_title_idx")
        if title_idx is None:
            results.append((0, False, rule_title, "Tez başlığı bulunamadı (memo yok). Önce abstract_tr_thesis_title çalışmalı."))
        else:
            # 2) Yazar adı satırı: tez başlığından sonraki boşluklardan sonra ilk dolu satır
            author_idx = memo.get("abstract_tr_after_title_first_content_idx")
            if author_idx is None:
                j = title_idx + 1
                # önce boşlukları geç
                while j < len(paragraphs) and (paragraphs[j].text or "").replace("\u00A0", " ").strip() == "":
                    j += 1
                author_idx = j
                memo["abstract_tr_after_title_first_content_idx"] = author_idx

            if author_idx >= len(paragraphs):
                results.append((title_idx, False, rule_title, "Tez yazarı adı satırı yok"))
            else:
                p = paragraphs[author_idx]
                errors = []

                raw_text = (p.text or "").replace("\u00A0", " ").strip()
                if raw_text == "":
                    errors.append("Yazar adı satırı boş")

                # Metin büyük harf kontrolü
                if expected_uppercase and raw_text:
                    if norm_tr_upper(raw_text) != raw_text:
                        errors.append("Yazar adı tamamen büyük harf değil")

                # Font/punto (run + style zinciri)
                fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
                fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
                if not fn and p.style and p.style.font and p.style.font.name:
                    fn = p.style.font.name
                if not fs and p.style and p.style.font and p.style.font.size:
                    fs = p.style.font.size.pt

                if fn and fn != expected_name:
                    errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")
                if fs and round(float(fs), 1) != round(expected_size, 1):
                    errors.append(f"Punto {fs} yerine {expected_size} olmalı")

                # Bold
                if expected_bold:
                    any_bold = any((r.bold or (r.font and r.font.bold)) for r in p.runs if (r.text or "").strip())
                    if not any_bold and p.style and p.style.font and getattr(p.style.font, "bold", None):
                        any_bold = True
                    if not any_bold:
                        errors.append("Yazar adı kalın değil")

                # Alignment (para + style fallback)
                align_val = p.alignment
                if align_val is None and p.style and p.style.paragraph_format:
                    align_val = p.style.paragraph_format.alignment

                if expected_align == "center" and align_val != WD_PARAGRAPH_ALIGNMENT.CENTER:
                    errors.append("Yazar adı ortalı değil")
                elif expected_align == "left" and align_val != WD_PARAGRAPH_ALIGNMENT.LEFT:
                    errors.append("Yazar adı sola yaslı değil")

                # Satır aralığı / boşluklar
                pf = p.paragraph_format
                ls = pf.line_spacing
                sb = pf.space_before.pt if pf.space_before else 0.0
                sa = pf.space_after.pt if pf.space_after else 0.0
                if ls is not None and round(float(ls), 2) != round(expected_spacing, 2):
                    errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
                if round(sb, 1) != round(expected_before, 1):
                    errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
                if round(sa, 1) != round(expected_after, 1):
                    errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

                # memo
                memo["abstract_tr_author_idx"] = author_idx

                if errors:
                    results.append((author_idx, False, rule_title, "; ".join(errors)))
                else:
                    results.append((author_idx, True, rule_title, ""))

    # ======================================================
    # TÜRKÇE ÖZET - TEZ YAZARI SONRASI BOŞLUK (en az N + fazla boşluk ihlali + memo)
    # ======================================================
    elif check["check"] == "abstract_tr_spacing_after_author":
        expected_blank_lines = int(check.get("expected_blank_lines", 2))
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))

        rule_title = (
            f"TÜRKÇE ÖZET - Yazar Adı Sonrası: en az {expected_blank_lines} satır boşluk, "
            f"{expected_name}, {int(expected_size)} punto, {expected_spacing} satır aralığı, "
            f"önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        author_idx = memo.get("abstract_tr_author_idx")
        if author_idx is None:
            results.append((0, False, rule_title, "Yazar adı bulunamadı (memo yok). Önce abstract_tr_author_name çalışmalı."))
        else:
            errors = []

            # 1) Ardışık boş satır say
            blank_count = 0
            j = author_idx + 1
            while j < len(paragraphs):
                txt = (paragraphs[j].text or "").replace("\u00A0", " ").strip()
                if txt != "":
                    break
                blank_count += 1
                j += 1

            # 2) Sonraki ilk dolu satırı memo’ya yaz (program bloğu için)
            memo["abstract_tr_after_author_first_content_idx"] = (author_idx + 1 + blank_count)

            # 3) En az N
            if blank_count < expected_blank_lines:
                errors.append(f"Yazar adından sonra {expected_blank_lines} boş satır olmalı (bulunan: {blank_count})")

            # 4) İlk N boş satır format kontrolü (varsa)
            for k in range(1, min(expected_blank_lines, blank_count) + 1):
                idx = author_idx + k
                loc = f"{k}. satır (belge:{idx})"
                
                p = paragraphs[idx]
                
                txt_full = (p.text or "").replace("\u00A0", " ").strip()
                pv = " ".join(txt_full.split()[:7])
                if len(pv) > 70:
                    pv = pv[:70].rstrip() + "…"
                loc = loc + (f" ('{pv}')" if pv else "")


                if (p.text or "").replace("\u00A0", " ").strip() != "":
                    errors.append(f"{loc} boş değil")
                    continue

                fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
                fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
                if not fn and p.style and p.style.font and p.style.font.name:
                    fn = p.style.font.name
                if not fs and p.style and p.style.font and p.style.font.size:
                    fs = p.style.font.size.pt

                pf = p.paragraph_format
                ls = pf.line_spacing
                sb = pf.space_before.pt if pf.space_before else 0.0
                sa = pf.space_after.pt if pf.space_after else 0.0

                if fn and fn != expected_name:
                    errors.append(f"{loc}- yazı tipi {fn} yerine {expected_name} olmalı")
                if fs and round(float(fs), 1) != round(expected_size, 1):
                    errors.append(f"{loc}- punto {fs} yerine {expected_size} olmalı")
                if ls is not None and round(float(ls), 2) != round(expected_spacing, 2):
                    errors.append(f"{loc}- satır aralığı {ls} yerine {expected_spacing} olmalı")
                if round(sb, 1) != round(expected_before, 1):
                    errors.append(f"{loc}- öncesi {sb} yerine {expected_before} olmalı")
                if round(sa, 1) != round(expected_after, 1):
                    errors.append(f"{loc}- sonrası {sa} yerine {expected_after} olmalı")

            # 5) Fazla boşluk ihlali (zinciri bozmaz)
            if blank_count > expected_blank_lines:
                errors.append(f"Fazladan boş satır var (+{blank_count - expected_blank_lines})")

            if errors:
                results.append((author_idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((author_idx, True, rule_title, ""))

    # ======================================================
    # TÜRKÇE ÖZET - PROGRAM VE DANIŞMAN BLOĞU (sabit offset yok, memo + optional lines)
    # ======================================================
    elif check["check"] == "abstract_tr_program_block":
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_bold = check.get("bold", True)
        expected_align = (check.get("alignment", "center") or "").lower()
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        expected_lines = int(check.get("expected_lines", 3))     # toplam blok satırı
        optional_lines = check.get("optional_lines", [])         # örn: [3] (3. satır opsiyonel)

        rule_title = (
            f"TÜRKÇE ÖZET - Program/Danışman Bloğu: {expected_lines} satır, "
            f"{expected_name}, {int(expected_size)} punto, kalın={expected_bold}, "
            f"alignment={expected_align}, {expected_spacing} satır aralığı, "
            f"önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        start_idx = memo.get("abstract_tr_after_author_first_content_idx")
        if start_idx is None:
            results.append((0, False, rule_title, "Program/Danışman bloğu başlangıcı bulunamadı (memo yok). Önce abstract_tr_spacing_after_author çalışmalı."))
        else:
            errors = []

            # Program bloğu: start_idx’ten itibaren expected_lines satır
            for k in range(expected_lines):
                idx = start_idx + k
                line_no = k + 1  # 1-based
                loc = f"{line_no}. satır (belge:{idx})"

                if idx >= len(paragraphs):
                    # satır yoksa: opsiyonelse geç, değilse hata
                    if line_no in optional_lines:
                        continue
                    errors.append(f"{loc}- eksik (program/danışman bloğu)")
                    continue

                p = paragraphs[idx]
                raw = (p.text or "").replace("\u00A0", " ").strip()
                
                pv = " ".join(raw.split()[:7])
                if len(pv) > 70:
                    pv = pv[:70].rstrip() + "…"
                loc = loc + (f" ('{pv}')" if pv else "")


                # opsiyonel satır tamamen boşsa sorun değil
                if line_no in optional_lines and raw == "":
                    continue

                # boş olmaması bekleniyor (opsiyonel değilse)
                if raw == "" and (line_no not in optional_lines):
                    errors.append(f"{loc}- boş olmamalı")
                    continue

                # Font/punto (run + style zinciri)
                fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
                fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
                if not fn and p.style and p.style.font and p.style.font.name:
                    fn = p.style.font.name
                if not fs and p.style and p.style.font and p.style.font.size:
                    fs = p.style.font.size.pt

                if fn and fn != expected_name:
                    errors.append(f"{loc}- yazı tipi {fn} yerine {expected_name} olmalı")
                if fs and round(float(fs), 1) != round(expected_size, 1):
                    errors.append(f"{loc}- punto {fs} yerine {expected_size} olmalı")

                # Bold
                if expected_bold:
                    any_bold = any((r.bold or (r.font and r.font.bold)) for r in p.runs if (r.text or "").strip())
                    if not any_bold and p.style and p.style.font and getattr(p.style.font, "bold", None):
                        any_bold = True
                    if not any_bold:
                        errors.append(f"{loc}- kalın değil")

                # Alignment (para + style fallback)
                align_val = p.alignment
                if align_val is None and p.style and p.style.paragraph_format:
                    align_val = p.style.paragraph_format.alignment

                if expected_align == "center" and align_val != WD_PARAGRAPH_ALIGNMENT.CENTER:
                    errors.append(f"{loc}- ortalı değil")
                elif expected_align == "left" and align_val != WD_PARAGRAPH_ALIGNMENT.LEFT:
                    errors.append(f"{loc}- sola yaslı değil")

                # Satır aralığı / boşluklar
                pf = p.paragraph_format
                ls = pf.line_spacing
                sb = pf.space_before.pt if pf.space_before else 0.0
                sa = pf.space_after.pt if pf.space_after else 0.0
                if ls is not None and round(float(ls), 2) != round(expected_spacing, 2):
                    errors.append(f"{loc}- satır aralığı {ls} yerine {expected_spacing} olmalı")
                if round(sb, 1) != round(expected_before, 1):
                    errors.append(f"{loc}- öncesi {sb} yerine {expected_before} olmalı")
                if round(sa, 1) != round(expected_after, 1):
                    errors.append(f"{loc}- sonrası {sa} yerine {expected_after} olmalı")

            # memo: program bloğu aralığı
            memo["abstract_tr_program_block_start_idx"] = start_idx

            # ✅ Etkin (effective) blok bitişini hesapla:
            effective_end_idx = start_idx + expected_lines

            # Sondan geriye doğru: opsiyonel satır boşsa bloktan düş
            for line_no in sorted(optional_lines, reverse=True):
                idx = start_idx + (line_no - 1)
                if idx < len(paragraphs):
                    txt = (paragraphs[idx].text or "").replace("\u00A0", " ").strip()
                    # sadece BLOĞUN SONUNDAKİ opsiyonelleri düşmek için:
                    if idx == effective_end_idx - 1 and txt == "":
                        effective_end_idx -= 1

            memo["abstract_tr_program_block_end_idx"] = effective_end_idx

            
            
            

            if errors:
                results.append((start_idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((start_idx, True, rule_title, ""))

    # ======================================================
    # TÜRKÇE ÖZET - PROGRAM/DANIŞMAN BLOĞU SONRASI BOŞLUK
    # (en az N + fazla boşluk ihlali + memo)
    # ======================================================
    elif check["check"] == "abstract_tr_spacing_after_program_block":
        expected_blank_lines = int(check.get("expected_blank_lines", 2))
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))

        rule_title = (
            f"TÜRKÇE ÖZET - Program/Danışman Bloğu Sonrası: en az {expected_blank_lines} satır boşluk, "
            f"{expected_name}, {int(expected_size)} punto, {expected_spacing} satır aralığı, "
            f"önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        end_idx = memo.get("abstract_tr_program_block_end_idx")
        if end_idx is None:
            results.append((0, False, rule_title,
                            "Program/Danışman bloğu bulunamadı (memo yok). Önce abstract_tr_program_block çalışmalı."))
        else:
            errors = []

            # 1) Ardışık boş satır say
            blank_count = 0
            j = end_idx
            while j < len(paragraphs):
                txt = (paragraphs[j].text or "").replace("\u00A0", " ").strip()
                if txt != "":
                    break
                blank_count += 1
                j += 1

            # 2) Sonraki ilk dolu satırı memo’ya yaz (tarih için)
            memo["abstract_tr_after_program_first_content_idx"] = end_idx + blank_count

            # 3) En az N
            if blank_count < expected_blank_lines:
                errors.append(f"Program bloğundan sonra {expected_blank_lines} boş satır olmalı (bulunan: {blank_count})")

            # 4) İlk N boş satır format kontrolü (varsa)
            for k in range(1, min(expected_blank_lines, blank_count) + 1):
                idx = end_idx + (k - 1)
                p = paragraphs[idx]

                if (p.text or "").replace("\u00A0", " ").strip() != "":
                    errors.append(f"{k}. satır boş değil")
                    continue

                fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
                fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
                if not fn and p.style and p.style.font and p.style.font.name:
                    fn = p.style.font.name
                if not fs and p.style and p.style.font and p.style.font.size:
                    fs = p.style.font.size.pt

                pf = p.paragraph_format
                ls = pf.line_spacing
                sb = pf.space_before.pt if pf.space_before else 0.0
                sa = pf.space_after.pt if pf.space_after else 0.0

                if fn and fn != expected_name:
                    errors.append(f"{k}. boş satır yazı tipi {fn} yerine {expected_name} olmalı")
                if fs and round(float(fs), 1) != round(expected_size, 1):
                    errors.append(f"{k}. boş satır punto {fs} yerine {expected_size} olmalı")
                if ls is not None and round(float(ls), 2) != round(expected_spacing, 2):
                    errors.append(f"{k}. boş satır satır aralığı {ls} yerine {expected_spacing} olmalı")
                if round(sb, 1) != round(expected_before, 1):
                    errors.append(f"{k}. boş satır öncesi {sb} yerine {expected_before} olmalı")
                if round(sa, 1) != round(expected_after, 1):
                    errors.append(f"{k}. boş satır sonrası {sa} yerine {expected_after} olmalı")

            # 5) Fazla boşluk ihlali (zinciri bozmaz)
            if blank_count > expected_blank_lines:
                errors.append(f"Fazladan boş satır var (+{blank_count - expected_blank_lines})")

            if errors:
                results.append((end_idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((end_idx, True, rule_title, ""))

    # ======================================================
    # TÜRKÇE ÖZET - TARİH (Ay + Yıl) (memo tabanlı arama)
    # ======================================================
    elif check["check"] == "abstract_tr_date":
        import re
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_bold = check.get("bold", True)
        expected_align = (check.get("alignment", "center") or "").lower()
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        pattern = check.get("pattern")
        normalize_tr = bool(check.get("normalize_tr", True))

        rule_title = check.get("description", "ÖZET sayfasındaki tarih (Ay + Yıl)")

        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u",
                "\u00A0":" "
            })
            return (s or "").translate(trans).strip()

        # 1) arama başlangıcı: program bloğu sonrası ilk dolu satır
        start_hint = memo.get("abstract_tr_after_program_first_content_idx")
        search_ranges = []

        if isinstance(start_hint, int) and 0 <= start_hint < len(paragraphs):
            search_ranges.append(range(start_hint, len(paragraphs)))

        # fallback: tüm belge
        search_ranges.append(range(0, len(paragraphs)))

        found = None
        for rr in search_ranges:
            for i in rr:
                p = paragraphs[i]
                txt = (p.text or "").replace("\u00A0", " ").strip()
                test_txt = norm_tr(txt) if normalize_tr else txt
                if pattern and re.match(pattern, test_txt, re.IGNORECASE):
                    found = (i, p)
                    break
            if found:
                break

        if not found:
            results.append((0, False, rule_title, "Tarih satırı bulunamadı veya format hatalı"))
        else:
            idx, p = found
            memo["abstract_tr_date_idx"] = idx
            errors = []

            # Font/punto (run + style zinciri)
            fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
            fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
            if not fn and p.style and p.style.font and p.style.font.name:
                fn = p.style.font.name
            if not fs and p.style and p.style.font and p.style.font.size:
                fs = p.style.font.size.pt

            if fn and fn != expected_name:
                errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")
            if fs and round(float(fs), 1) != round(expected_size, 1):
                errors.append(f"Punto {fs} yerine {expected_size} olmalı")

            # Bold
            if expected_bold:
                any_bold = any((r.bold or (r.font and r.font.bold)) for r in p.runs if (r.text or "").strip())
                if not any_bold and p.style and p.style.font and getattr(p.style.font, "bold", None):
                    any_bold = True
                if not any_bold:
                    errors.append("Tarih satırı kalın değil")

            # Alignment (para + style fallback)
            align_val = p.alignment
            if align_val is None and p.style and p.style.paragraph_format:
                align_val = p.style.paragraph_format.alignment

            if expected_align == "center" and align_val != WD_PARAGRAPH_ALIGNMENT.CENTER:
                errors.append("Tarih satırı ortalı değil")

            # Satır aralığı / boşluklar
            pf = p.paragraph_format
            ls = pf.line_spacing
            sb = pf.space_before.pt if pf.space_before else 0.0
            sa = pf.space_after.pt if pf.space_after else 0.0
            if ls is not None and round(float(ls), 2) != round(expected_spacing, 2):
                errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
            if round(sb, 1) != round(expected_before, 1):
                errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
            if round(sa, 1) != round(expected_after, 1):
                errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

            if errors:
                results.append((idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((idx, True, rule_title, ""))

    # ======================================================
    # TÜRKÇE ÖZET - SAYFA BİLGİSİ (memo tabanlı arama)
    # ======================================================
    elif check["check"] == "abstract_tr_pageinfo":
        import re
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_bold = check.get("bold", True)
        expected_align = (check.get("alignment", "center") or "").lower()
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        pattern = check.get("pattern")
        normalize_tr = bool(check.get("normalize_tr", True))

        rule_title = check.get("description", "ÖZET sayfasındaki sayfa bilgisi")

        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u",
                "\u00A0":" "
            })
            return (s or "").translate(trans).strip()

        # 1) arama başlangıcı: tarih satırından sonra
        date_idx = memo.get("abstract_tr_date_idx")
        search_ranges = []

        if isinstance(date_idx, int) and 0 <= date_idx + 1 < len(paragraphs):
            search_ranges.append(range(date_idx + 1, len(paragraphs)))

        # fallback: tüm belge
        search_ranges.append(range(0, len(paragraphs)))

        found = None
        for rr in search_ranges:
            for i in rr:
                p = paragraphs[i]
                txt = (p.text or "").replace("\u00A0", " ").strip()
                test_txt = norm_tr(txt) if normalize_tr else txt
                if pattern and re.search(pattern, test_txt, re.IGNORECASE):
                    found = (i, p)
                    break
            if found:
                break

        if not found:
            results.append((0, False, rule_title, "Sayfa bilgisi satırı bulunamadı veya format hatalı"))
        else:
            idx, p = found
            memo["abstract_tr_pageinfo_idx"] = idx
            errors = []

            # Font/punto (run + style zinciri)
            fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
            fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
            if not fn and p.style and p.style.font and p.style.font.name:
                fn = p.style.font.name
            if not fs and p.style and p.style.font and p.style.font.size:
                fs = p.style.font.size.pt

            if fn and fn != expected_name:
                errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")
            if fs and round(float(fs), 1) != round(expected_size, 1):
                errors.append(f"Punto {fs} yerine {expected_size} olmalı")

            # Bold
            if expected_bold:
                any_bold = any((r.bold or (r.font and r.font.bold)) for r in p.runs if (r.text or "").strip())
                if not any_bold and p.style and p.style.font and getattr(p.style.font, "bold", None):
                    any_bold = True
                if not any_bold:
                    errors.append("Sayfa bilgisi satırı kalın değil")

            # Alignment (para + style fallback)
            align_val = p.alignment
            if align_val is None and p.style and p.style.paragraph_format:
                align_val = p.style.paragraph_format.alignment

            if expected_align == "center" and align_val != WD_PARAGRAPH_ALIGNMENT.CENTER:
                errors.append("Sayfa bilgisi satırı ortalı değil")

            # Satır aralığı / boşluklar
            pf = p.paragraph_format
            ls = pf.line_spacing
            sb = pf.space_before.pt if pf.space_before else 0.0
            sa = pf.space_after.pt if pf.space_after else 0.0
            if ls is not None and round(float(ls), 2) != round(expected_spacing, 2):
                errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
            if round(sb, 1) != round(expected_before, 1):
                errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
            if round(sa, 1) != round(expected_after, 1):
                errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

            if errors:
                results.append((idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((idx, True, rule_title, ""))

    # ======================================================
    # TÜRKÇE ÖZET - TARİH+SAYFA BİLGİSİ SONRASI BOŞLUK
    # (en az N + fazla boşluk ihlali + memo)
    # ======================================================
    elif check["check"] == "abstract_tr_spacing_after_pageinfo":
        expected_blank_lines = int(check.get("expected_blank_lines", 3))
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))

        rule_title = (
            f"TÜRKÇE ÖZET - Tarih ve Sayfa Bilgisi Sonrası: en az {expected_blank_lines} satır boşluk, "
            f"{expected_name}, {int(expected_size)} punto, {expected_spacing} satır aralığı, "
            f"önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        pageinfo_idx = memo.get("abstract_tr_pageinfo_idx")
        if pageinfo_idx is None:
            results.append((0, False, rule_title,
                            "Sayfa bilgisi satırı bulunamadı (memo yok). Önce abstract_tr_pageinfo çalışmalı."))
        else:
            errors = []

            # 1) Sayfa bilgisinden sonra ardışık boş satır say
            blank_count = 0
            j = pageinfo_idx + 1
            while j < len(paragraphs):
                txt = (paragraphs[j].text or "").replace("\u00A0", " ").strip()
                if txt != "":
                    break
                blank_count += 1
                j += 1

            # 2) Sonraki ilk dolu satırı memo’ya yaz (gövde için)
            memo["abstract_tr_after_pageinfo_first_content_idx"] = pageinfo_idx + 1 + blank_count

            # 3) En az N
            if blank_count < expected_blank_lines:
                errors.append(f"Sayfa bilgisinden sonra {expected_blank_lines} boş satır olmalı (bulunan: {blank_count})")

            # 4) İlk N boş satır biçim kontrolü
            for k in range(1, min(expected_blank_lines, blank_count) + 1):
                idx = (pageinfo_idx + k)
                p = paragraphs[idx]

                if (p.text or "").replace("\u00A0", " ").strip() != "":
                    errors.append(f"{k}. satır boş değil")
                    continue

                fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
                fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
                if not fn and p.style and p.style.font and p.style.font.name:
                    fn = p.style.font.name
                if not fs and p.style and p.style.font and p.style.font.size:
                    fs = p.style.font.size.pt

                pf = p.paragraph_format
                ls = pf.line_spacing
                sb = pf.space_before.pt if pf.space_before else 0.0
                sa = pf.space_after.pt if pf.space_after else 0.0

                if fn and fn != expected_name:
                    errors.append(f"{k}. boş satır yazı tipi {fn} yerine {expected_name} olmalı")
                if fs and round(float(fs), 1) != round(expected_size, 1):
                    errors.append(f"{k}. boş satır punto {fs} yerine {expected_size} olmalı")
                if ls is not None and round(float(ls), 2) != round(expected_spacing, 2):
                    errors.append(f"{k}. boş satır satır aralığı {ls} yerine {expected_spacing} olmalı")
                if round(sb, 1) != round(expected_before, 1):
                    errors.append(f"{k}. boş satır öncesi {sb} yerine {expected_before} olmalı")
                if round(sa, 1) != round(expected_after, 1):
                    errors.append(f"{k}. boş satır sonrası {sa} yerine {expected_after} olmalı")

            # 5) Fazla boşluk ihlali (zinciri bozmaz)
            if blank_count > expected_blank_lines:
                errors.append(f"Fazladan boş satır var (+{blank_count - expected_blank_lines})")

            if errors:
                results.append((pageinfo_idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((pageinfo_idx, True, rule_title, ""))

    # ======================================================
    # TÜRKÇE ÖZET GÖVDE
    # (başlangıç: pageinfo sonrası boşluklardan sonra ilk dolu satır (memo),
    #  bitiş: Anahtar Kelimeler marker satırı,
    #  sabit offset yok)
    # ======================================================
    elif check["check"] == "abstract_tr_body":
        expected_name   = check.get("font_name", "Times New Roman")
        expected_size   = float(check.get("font_size_pt", 12))
        expected_spacing= float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after  = float(check.get("space_after", 0))

        rule_title = (
            f"TÜRKÇE ÖZET Gövdesi: {expected_name}, {int(expected_size)} punto, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}, iki yana yaslı"
        )

        import re

        # YAML marker desteği (yanlış yazımlar vs.)
        # örn: ["^ANAHTAR\\s*KELIMELER\\s*:", "^ANAHTAR\\s*KELİMELER\\s*:", "^KEYWORDS\\s*:"] gibi
        kw_markers = check.get("keywords_markers", [r"^ANAHTAR\s*KELIMELER\s*:\s*"])

        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u",
                "\u00A0":" ", "\t":" "
            })
            return (s or "").translate(trans)

        def clean_compact(s: str) -> str:
            # boşlukları sadeleştir (ama tamamen silme, yazım kontrolünde gerekebilir)
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        def is_effectively_blank(text: str) -> bool:
            if not text:
                return True
            cleaned = (text or "").strip()
            cleaned = cleaned.replace(".", "").replace("…", "").replace("-", "")
            return cleaned.strip() == ""

        
        from docx.oxml.ns import qn
        def _get_effective_alignment(p, doc):
            """
            Return WD_PARAGRAPH_ALIGNMENT value or None if not determinable.
            Priority: direct pPr/jc -> style chain -> docDefaults -> None
            """

            # 1) Direct formatting on paragraph: <w:pPr><w:jc w:val="both"/>
            pPr = p._p.pPr
            if pPr is not None and pPr.jc is not None and pPr.jc.val is not None:
                return pPr.jc.val  # this is already an enum-like value in python-docx

            # 2) Style chain (basedOn)
            style = p.style
            visited = set()
            while style is not None and style.style_id not in visited:
                visited.add(style.style_id)

                pf = style.paragraph_format
                if pf is not None and pf.alignment is not None:
                    return pf.alignment

                style = style.base_style

            # 3) docDefaults in styles.xml (paragraph defaults)
            # document.styles.element is <w:styles> root
            styles_elm = doc.styles.element
            docDefaults = styles_elm.find(qn('w:docDefaults'))
            if docDefaults is not None:
                pPrDefault = docDefaults.find(qn('w:pPrDefault'))
                if pPrDefault is not None:
                    pPr = pPrDefault.find(qn('w:pPr'))
                    if pPr is not None:
                        jc = pPr.find(qn('w:jc'))
                        if jc is not None:
                            val = jc.get(qn('w:val'))
                            # map common Word values to python-docx enum if you want;
                            # but often comparing string works too (e.g., "both")
                            return val  # "both", "left", "right", "center", "distribute", etc.

            return None


        # 1) Başlangıç: memo’dan al
        start_idx = memo.get("abstract_tr_after_pageinfo_first_content_idx")
        if not isinstance(start_idx, int) or start_idx < 0 or start_idx >= len(paragraphs):
            start_idx = None

        # 2) Bitiş: Anahtar Kelimeler satırı (markers ile)
        end_idx = None
        end_para = None
        for i, p in enumerate(paragraphs):
            cand = clean_compact(norm_tr(p.text)).upper()
            for m in kw_markers:
                if re.match(m, cand, re.IGNORECASE):
                    end_idx = i
                    end_para = p
                    break
            if end_idx is not None:
                break

        if start_idx is None or end_idx is None or start_idx >= end_idx:
            results.append((0, False, rule_title,
                            "Özet gövdesi aralığı bulunamadı (pageinfo sonrası başlangıç veya Anahtar Kelimeler bulunamadı)."))
        else:
            errors = []

            # 2.a) Anahtar Kelimeler yazımı doğru mu? (bulduk ama yanlış olabilir)
            # doğru kabul: "Anahtar Kelimeler:" (Türkçe karakter + boşluk toleranslı)
            raw_kw = clean_compact(end_para.text)
            raw_kw_norm = clean_compact(norm_tr(raw_kw)).lower().replace(" ", "")
            # "anahtarkelimeler:" ile başlamıyorsa yazım hatası ama end_idx olarak kullanmaya devam
            if not raw_kw_norm.startswith("anahtarkelimeler:"):
                errors.append(f"Anahtar Kelimeler yazımı hatalı (olması gereken: 'Anahtar Kelimeler:'; bulunan: '{raw_kw}')")

            # keyword satırı idx memo (bir sonraki fonksiyon kullanacak)
            memo["abstract_tr_keywords_idx"] = end_idx

            # 3) Gövde biçim kontrolleri
            for j in range(start_idx, end_idx):
                p = paragraphs[j]
                if is_effectively_blank(p.text):
                    continue

                # --- Hata mesajı için: "satır" + ilk birkaç kelime önizleme (yeni fonksiyon yazmadan, burada çözüm) ---
                _t = clean_compact(p.text)
                _w = _t.split()
                _pv = " ".join(_w[:7])
                if len(_pv) > 60:
                    _pv = _pv[:60].rstrip()
                if _t and len(_t) > len(_pv):
                    _pv += "…"
                loc = f"{j}. satır" + (f" ('{_pv}')" if _pv else "")
                # --------------------------------------------------------------------------------------------------

                # Font/punto (run + style)
                fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
                fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
                if not fn and p.style and p.style.font and p.style.font.name:
                    fn = p.style.font.name
                if not fs and p.style and p.style.font and p.style.font.size:
                    fs = p.style.font.size.pt

                if fn and fn != expected_name:
                    errors.append(f"{loc} — yazı tipi {fn} (beklenen: {expected_name})")
                if fs is not None and round(float(fs), 1) != round(expected_size, 1):
                    errors.append(f"{loc} — punto {fs} (beklenen: {expected_size})")

                # Satır aralığı / boşluklar
                pf = p.paragraph_format
                ls = pf.line_spacing
                sb = pf.space_before.pt if pf.space_before else 0.0
                sa = pf.space_after.pt  if pf.space_after  else 0.0

                if ls is not None and round(float(ls), 2) != round(expected_spacing, 2):
                    errors.append(f"{loc} — satır aralığı {ls} (beklenen: {expected_spacing})")
                if round(sb, 1) != round(expected_before, 1):
                    errors.append(f"{loc} — öncesi {sb} (beklenen: {expected_before})")
                if round(sa, 1) != round(expected_after, 1):
                    errors.append(f"{loc} — sonrası {sa} (beklenen: {expected_after})")

                # Justify (iki yana yaslı) – effective (direct + style chain + docDefaults)
                align_val = _get_effective_alignment(p, doc)

                # python-docx enum gelirse: WD_PARAGRAPH_ALIGNMENT.JUSTIFY ile kıyaslanır
                # docDefaults'tan string gelirse: "both" Word'de justify demektir.
                is_justify = (align_val == WD_PARAGRAPH_ALIGNMENT.JUSTIFY) or (str(align_val).lower() in {"both", "distribute"})


                if align_val is not None and not is_justify:
                    errors.append(f"{loc} — iki yana yaslı değil")


            if errors:
                results.append((start_idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((start_idx, True, rule_title, ""))


    # ======================================================
    # TÜRKÇE ÖZET - ANAHTAR KELİMELER SATIRI
    # (memo + YAML markers + yazım hatası ihlali)
    # ======================================================
    elif check["check"] == "abstract_tr_keywords_line":
        expected_name   = check.get("font_name", "Times New Roman")
        expected_size   = float(check.get("font_size_pt", 12))
        expected_spacing= float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after  = float(check.get("space_after", 0))

        # ------------------------------------------------------
        # YAML marker desteği
        # Örn: ["^ANAHTAR\\s*KELIMELER\\s*:\\s*", "^ANAHTAR\\s*KELİMELER\\s*:\\s*"]
        #
        # DÜZELTME:
        # - Default marker artık ':' zorunlu değil ve farklı ayraçları da yakalar.
        #   Böylece satır "Anahtar Kelimeler" / "Anahtar Kelimeler：" / "Anahtar Kelimeler -"
        #   gibi yazımlarda "bulunamadı" düşer.
        # - Kılavuz gereği ':' şartını yine yazım kontrolünde ihlal olarak raporluyoruz.
        # ------------------------------------------------------
        markers = check.get("markers", [
            r"^ANAHTAR\s*KELIMELER\s*([:：;；\-–—])?\s*"
        ])

        rule_title = (
            f"Anahtar Kelimeler satırı: {expected_name}, {int(expected_size)} punto, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}, sola yaslı"
        )

        import re

        # ------------------------------------------------------
        # Türkçe normalize + görünmez karakter toleransı
        # DÜZELTME:
        # - Zero-width space / word joiner / BOM gibi karakterler
        #   regex match’i bozabildiği için temizleniyor.
        # - NBSP ve tab -> space
        # ------------------------------------------------------
        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u",
                "\u00A0":" ",   # NBSP
                "\t":" ",
                "\u200b":"",    # zero-width space
                "\u2060":"",    # word joiner
                "\ufeff":"",    # BOM
            })
            return (s or "").translate(trans)

        def clean(s: str) -> str:
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        # 1) Önce memo idx
        keyword_idx = memo.get("abstract_tr_keywords_idx")
        found = None

        if isinstance(keyword_idx, int) and 0 <= keyword_idx < len(paragraphs):
            p = paragraphs[keyword_idx]
            # cand: normalize + temizle + UPPER
            cand = clean(norm_tr(p.text)).upper()
            for m in markers:
                # DÜZELTME:
                # - marker match toleransı artırıldı (marker zaten opsiyonel ayraç içeriyor olabilir)
                if re.match(m, cand, re.IGNORECASE):
                    found = (keyword_idx, p)
                    break

        # 2) Fallback: tüm dokümanda ara
        if not found:
            for i, p in enumerate(paragraphs):
                cand = clean(norm_tr(p.text)).upper()
                for m in markers:
                    if re.match(m, cand, re.IGNORECASE):
                        found = (i, p)
                        break
                if found:
                    break

        if not found:
            results.append((0, False, rule_title, "Anahtar Kelimeler satırı bulunamadı"))
        else:
            idx, p = found
            memo["abstract_tr_keywords_idx"] = idx  # tekrar yaz (güncel kalsın)

            errors = []

            # ------------------------------------------------------
            # 0) Yazım doğrulama (bulduk ama kılavuza uygun mu?)
            # DÜZELTME:
            # - Önce "anahtarkelimeler" ile başlıyor mu diye bak.
            # - Ardından kılavuz gereği ayraç olarak MUTLAKA ":" var mı diye kontrol et.
            #   (Farklı ayraç varsa veya hiç ayraç yoksa yazım ihlali olarak raporlanır.)
            # ------------------------------------------------------
            raw = clean(p.text)
            raw_norm_compact = clean(norm_tr(raw)).lower().replace(" ", "")

            # En azından başlık doğru mu?
            if not raw_norm_compact.startswith("anahtarkelimeler"):
                errors.append(
                    f"Anahtar Kelimeler yazımı hatalı (olması gereken: 'Anahtar Kelimeler:'; bulunan: '{raw}')"
                )
            else:
                # Başlıktan hemen sonra gelen ayraç karakteri (varsa)
                head_len = len("anahtarkelimeler")
                sep = raw_norm_compact[head_len:head_len+1] if len(raw_norm_compact) > head_len else ""

                # Kılavuz gereği sadece ":" kabul edilecek
                if sep != ":":
                    if sep == "":
                        errors.append("Anahtar Kelimeler ifadesinden sonra ':' bulunmalı.")
                    else:
                        errors.append(
                            f"Anahtar Kelimeler ayıracı ':' olmalı (bulunan ayraç: '{sep}')."
                        )

            # ------------------------------------------------------
            # 1) Hizalama sola yaslı mı? (style fallback toleranslı)
            # ------------------------------------------------------
            align_val = p.alignment
            if align_val is None and p.style and p.style.paragraph_format:
                align_val = p.style.paragraph_format.alignment
            if align_val not in (None, WD_PARAGRAPH_ALIGNMENT.LEFT):
                errors.append("Satır sola yaslı değil")

            # ------------------------------------------------------
            # 2) Font/punto (run + style)
            # ------------------------------------------------------
            fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
            fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
            if not fn and p.style and p.style.font and p.style.font.name:
                fn = p.style.font.name
            if not fs and p.style and p.style.font and p.style.font.size:
                fs = p.style.font.size.pt

            if fn and fn != expected_name:
                errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")
            if fs is not None and round(float(fs), 1) != round(expected_size, 1):
                errors.append(f"Punto {fs} yerine {expected_size} olmalı")

            # ------------------------------------------------------
            # 3) Satır aralığı ve boşluklar
            # ------------------------------------------------------
            pf = p.paragraph_format
            ls = pf.line_spacing
            sb = pf.space_before.pt if pf.space_before else 0.0
            sa = pf.space_after.pt if pf.space_after else 0.0
            if ls is not None and round(float(ls), 2) != round(expected_spacing, 2):
                errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
            if round(sb, 1) != round(expected_before, 1):
                errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
            if round(sa, 1) != round(expected_after, 1):
                errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

            # ------------------------------------------------------
            # 4) Bold kontrolü: etiket kalın, sonrası kalın değil
            # DÜZELTME:
            # - ":" yerine farklı ayraçlar varsa da "sonrası" algılansın diye ayraç seti eklendi.
            # - Run'lar bölünmüş olabileceği için toleranslı mantık korunuyor.
            # ------------------------------------------------------
            seps = {":", "：", ";", "；", "-", "–", "—"}

            label_bold_ok = False
            after_sep_seen = False

            for r in p.runs:
                r_text_raw = r.text or ""
                rt = clean(norm_tr(r_text_raw)).lower()
                rt_compact = rt.replace(" ", "")

                # Ayraç görüldüyse bundan sonrası "etiket sonrası" kabul edilir
                if any(ch in r_text_raw for ch in seps):
                    after_sep_seen = True

                # Etiket parçaları (anahtar/kelimeler) en az birinde bold olmalı
                if ("anahtar" in rt_compact) or ("kelimeler" in rt_compact):
                    if r.bold or (r.font and r.font.bold):
                        label_bold_ok = True

                # Ayraç sonrası içerik bold olmamalı
                if after_sep_seen:
                    if r_text_raw.strip() and (r.bold or (r.font and r.font.bold)):
                        # Etiket parçası değilse ihlal
                        if "anahtar" not in rt_compact and "kelimeler" not in rt_compact:
                            errors.append("Anahtar kelimeler kısmı kalın olmamalı")
                            break

            if not label_bold_ok:
                errors.append("'Anahtar Kelimeler:' etiketi kalın değil")

            if errors:
                results.append((idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((idx, True, rule_title, ""))


    # ===============================================================================================================#    
    # ===============================================================================================================#
    # ======================================================
    # ABSTRACT (İngilizce Özet) Başlığı
    # ======================================================
    elif check["check"] == "abstract_en_heading":
        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 0))
        must_bold        = check.get("bold", True)

        # 🔧 ÖNEMLİ: rules.yaml bu kural için "alignment" kullanıyor (align değil).
        # Geriye dönük uyum için ikisini de destekleyelim.
        expected_align = (check.get("alignment", check.get("align", "center")) or "").lower()
        must_center = (expected_align == "center")

        markers = check.get("markers", ["^ABSTRACT$"])
        import re

        found = None
        for i, p in enumerate(paragraphs):
            for marker in markers:
                candidate = " ".join((p.text or "").split())
                if re.match(marker, candidate, re.IGNORECASE):
                    found = (i, p)
                    break
            if found:
                break

        rule_title = (
            f"ENGLISH ABSTRACT Heading: {expected_name}, {int(expected_size)} pt, "
            f"{expected_spacing} line spacing, before {int(expected_before)}, after {int(expected_after)}, "
            f"{'bold' if must_bold else 'normal'}, "
            f"{'centered' if must_center else ('left aligned' if expected_align == 'left' else expected_align or 'any')}"
        )

        if not found:
            results.append((0, False, rule_title, "ABSTRACT başlığı yok yada yanlış yazılımış (büyük/küçük harf duyarsız veya fazladan boşluklar var)"))
        else:
            idx, p = found

            # ✅ Başlık bulunduğunda memo’ya kaydet (yeni ek; diğer yerleri bozmaz)
            memo["abstract_en_heading_idx"] = idx

            errors = []

            # ------------------------------------------------------
            # EK ÖZELLİK: Başlık metni doğrulama (toleranslı bul, katı doğrula)
            # Amaç: Yanlış yazımı da yakalayıp "ABSTRACT olmalı" ihlali üretmek
            # ------------------------------------------------------
            raw_title = (p.text or "").strip()
            norm_title = " ".join(raw_title.split()).upper()  # boşluk normalize + büyük harf
            if norm_title != "ABSTRACT":
                # Bulduk ama yanlış yazılmış: format kontrolü + metin ihlali birlikte raporlanır
                errors.append(f"Başlık metni ABSTRACT olmalı (bulunan: {raw_title})")

            # ------------------------------------------------------
            # EFFECTIVE font / size (run + style zinciri)
            # ------------------------------------------------------
            fn = effective_font_name(p)
            fs = effective_font_size_pt(p)

            if fn and fn != expected_name:
                errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")
            if fs is not None and round(float(fs), 1) != round(float(expected_size), 1):
                errors.append(f"Punto {fs} yerine {expected_size} olmalı")

            # ------------------------------------------------------
            # EFFECTIVE bold (run + run-style + style zinciri)
            # ------------------------------------------------------
            if must_bold and not effective_bold(p):
                errors.append("Başlık Kalın değil")

            # ------------------------------------------------------
            # EFFECTIVE alignment (paragraph + style zinciri)
            # ------------------------------------------------------
            if expected_align == "center":
                if effective_alignment(p) != WD_PARAGRAPH_ALIGNMENT.CENTER:
                    errors.append("Başlık Ortalı değil")
            elif expected_align == "left":
                if effective_alignment(p) != WD_PARAGRAPH_ALIGNMENT.LEFT:
                    errors.append("Başlık Sola yaslı değil")
            # expected_align boş/any ise hizalama kontrolü yapmıyoruz.

            # ------------------------------------------------------
            # EFFECTIVE line spacing / before-after (paragraph + style zinciri)
            # ------------------------------------------------------
            ls = effective_line_spacing(p, default=expected_spacing)
            sb = effective_space_pt(p, "before")
            sa = effective_space_pt(p, "after")

            if ls is not None and round(float(ls), 2) != round(float(expected_spacing), 2):
                errors.append(f"Satır boşluğu {ls} yerine {expected_spacing} olmalı")
            if round(float(sb), 1) != round(float(expected_before), 1):
                errors.append(f"Öncesi {sb} yerine {expected_before} nk olmalı")
            if round(float(sa), 1) != round(float(expected_after), 1):
                errors.append(f"Sonrası {sa} yerine {expected_after} nk olmalı")

            if errors:
                results.append((idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((idx, True, rule_title, ""))

    # ======================================================
    # ABSTRACT Başlığından Sonra 5 Satır Boşluk
    # ======================================================
    elif check["check"] == "abstract_en_spacing_after_heading":
        expected_blank_lines = int(check.get("expected_blank_lines", 5))
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))

        import re

        def clean(s: str) -> str:
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        # ✅ 1) Heading index: önce memo, yoksa markers ile bul
        heading_idx = memo.get("abstract_en_heading_idx", None)
        if heading_idx is None:
            markers = check.get("markers", [r"^ABSTRACT$"])
            for i, p in enumerate(paragraphs):
                candidate = clean(p.text)
                for m in markers:
                    if re.match(m, candidate, re.IGNORECASE):
                        heading_idx = i
                        memo["abstract_en_heading_idx"] = i
                        break
                if heading_idx is not None:
                    break

        rule_title = (
            f"ABSTRACT sonrası {expected_blank_lines} satır boşluk: "
            f"{expected_name}, {int(expected_size)} punto, {expected_spacing} satır aralığı, "
            f"önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        if heading_idx is None:
            results.append((0, False, rule_title, "ABSTRACT başlığı bulunamadı"))
        else:
            errors = []

            # ✅ 2) Heading'den sonra ardışık boş satır say
            j = heading_idx + 1
            blank_count = 0
            while j < len(paragraphs) and clean(paragraphs[j].text) == "":
                blank_count += 1
                j += 1

            first_nonblank_idx = j if j < len(paragraphs) else None

            # ✅ 3) Memo: gerçek boşluk sayısı + tez başlığı idx (ilk dolu satır)
            memo["abstract_en_blank_after_heading_count"] = blank_count
            if first_nonblank_idx is not None:
                memo["abstract_en_thesis_title_idx"] = first_nonblank_idx

            # ✅ 4) Kural: en az 5 olsun, fazlası da ihlal
            if blank_count < expected_blank_lines:
                errors.append(f"Boş satır sayısı {blank_count}; en az {expected_blank_lines} olmalı")
            elif blank_count > expected_blank_lines:
                errors.append(f"Fazladan boş satır var: {blank_count} satır (beklenen: {expected_blank_lines})")

            # ✅ 5) İlk expected_blank_lines satırının biçim kontrolleri (mevcut mantık korunarak)
            # Not: Fazladan boşluklar varsa dahi, ilk 5 satırın biçimi doğru mu diye yine bakıyoruz.
            for offset in range(1, min(blank_count, expected_blank_lines) + 1):
                idx = heading_idx + offset
                if idx >= len(paragraphs):
                    break
                p = paragraphs[idx]

                # Biçim denetimi (senin mevcut kodunla aynı mantık)
                fn = None
                for r in p.runs:
                    if r.font and r.font.name:
                        fn = r.font.name
                        break
                if not fn and p.style and p.style.font and p.style.font.name:
                    fn = p.style.font.name
                if fn and fn != expected_name:
                    errors.append(f"{offset}. satır yazı tipi {fn} yerine {expected_name} olmalı")

                fs = None
                for r in p.runs:
                    if r.font and r.font.size:
                        fs = r.font.size.pt
                        break
                if not fs and p.style and p.style.font and p.style.font.size:
                    fs = p.style.font.size.pt
                if fs and round(float(fs), 1) != round(expected_size, 1):
                    errors.append(f"{offset}. satır punto {fs} yerine {expected_size} olmalı")

                pf = p.paragraph_format
                ls = pf.line_spacing
                sb = pf.space_before.pt if pf.space_before else 0.0
                sa = pf.space_after.pt if pf.space_after else 0.0
                if ls is not None and round(float(ls), 2) != expected_spacing:
                    errors.append(f"{offset}. satırın satır aralığı {ls} yerine {expected_spacing} olmalı")
                if round(sb, 1) != round(expected_before, 1):
                    errors.append(f"{offset}. satır öncesi {sb} yerine {expected_before} olmalı")
                if round(sa, 1) != round(expected_after, 1):
                    errors.append(f"{offset}. satır sonrası {sa} yerine {expected_after} olmalı")

            if errors:
                results.append((heading_idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((heading_idx, True, rule_title, ""))

    # ======================================================
    # ABSTRACT Sonrası Tez Başlığı
    # ======================================================
    elif check["check"] == "abstract_en_thesis_title":
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        must_bold = check.get("bold", True)
        must_center = check.get("align", "center") == "center"

        import re

        def clean(s: str) -> str:
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        # ✅ Heading idx: memo öncelikli
        heading_idx = memo.get("abstract_en_heading_idx", None)
        if heading_idx is None:
            markers = check.get("markers", [r"^ABSTRACT$"])
            for i, p in enumerate(paragraphs):
                candidate = clean(p.text)
                for m in markers:
                    if re.match(m, candidate, re.IGNORECASE):
                        heading_idx = i
                        memo["abstract_en_heading_idx"] = i
                        break
                if heading_idx is not None:
                    break

        rule_title = (
            f"ABSTRACT sonrası tez başlığı: {expected_name}, {int(expected_size)} punto, "
            f"{'kalın' if must_bold else 'normal'}, {'ortalı' if must_center else 'sol'}, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        if heading_idx is None:
            results.append((0, False, rule_title, "ABSTRACT başlığı bulunamadı"))
        else:
            # ✅ TEZ BAŞLIĞI idx: memo varsa onu kullan, yoksa heading sonrası ilk dolu satırı bul
            title_idx = memo.get("abstract_en_thesis_title_idx", None)
            if title_idx is None:
                j = heading_idx + 1
                while j < len(paragraphs) and clean(paragraphs[j].text) == "":
                    j += 1
                title_idx = j if j < len(paragraphs) else None
                if title_idx is not None:
                    memo["abstract_en_thesis_title_idx"] = title_idx

            if title_idx is None:
                results.append((heading_idx, False, rule_title, "Tez başlığı satırı yok"))
            else:
                p = paragraphs[title_idx]
                errors = []

                if not clean(p.text):
                    errors.append("Tez başlığı satırı boş")

                # Yazı tipi
                fn = None
                for r in p.runs:
                    if r.font and r.font.name:
                        fn = r.font.name
                        break
                if not fn and p.style and p.style.font and p.style.font.name:
                    fn = p.style.font.name
                if fn and fn != expected_name:
                    errors.append(f"Yazı tipi {fn} yerine {expected_name}")

                # Punto
                fs = None
                for r in p.runs:
                    if r.font and r.font.size:
                        fs = r.font.size.pt
                        break
                if not fs and p.style and p.style.font and p.style.font.size:
                    fs = p.style.font.size.pt
                if fs and round(float(fs), 1) != round(expected_size, 1):
                    errors.append(f"Punto {fs} yerine {expected_size} olmalı")

                # Kalınlık
                if must_bold:
                    any_bold = any(r.bold for r in p.runs if r.text.strip())
                    if not any_bold and p.style and p.style.font and p.style.font.bold:
                        any_bold = True
                    if not any_bold:
                        errors.append("Tez başlığı kalın değil")

                # ------------------------------------------------------
                # ✅ EK KURAL: Tez başlığında italik OLMAMALI
                # - run italic / run style / paragraph style zinciri dahil "effective" kontrol
                # ------------------------------------------------------
                if check.get("no_italic", True):  # yaml’dan kapatılabilir; default True
                    # 1) Paragraf genelinde effective italic varsa → hata
                    if effective_italic(p):
                        errors.append("Tez başlığında italik kullanılmamalı")

                    # 2) Daha hassas: herhangi bir text run’ı italik mi?
                    # (Bazı durumlarda effective_italic False dönebilir; yine de run’da italic True yakalanır)
                    any_run_italic = False
                    for r in p.runs:
                        if not (r.text and r.text.strip()):
                            continue
                        if (r.font and r.font.italic) or (getattr(r, "italic", None) is True) or (r.italic is True):
                            any_run_italic = True
                            break
                        rs = getattr(r, "style", None)
                        if rs and rs.font and rs.font.italic:
                            any_run_italic = True
                            break
                    if any_run_italic:
                        # aynı mesajı iki kez yazmamak için:
                        if "Tez başlığında italik kullanılmamalı" not in errors:
                            errors.append("Tez başlığında italik kullanılmamalı")


                # Ortalanma
                align_val = p.alignment
                if not align_val and p.style and p.style.paragraph_format.alignment:
                    align_val = p.style.paragraph_format.alignment
                if must_center and align_val != WD_PARAGRAPH_ALIGNMENT.CENTER:
                    errors.append("Tez başlığı ortalı değil")

                # Satır aralığı, önce/sonra boşluk
                pf = p.paragraph_format
                ls = pf.line_spacing
                sb = pf.space_before.pt if pf.space_before else 0.0
                sa = pf.space_after.pt if pf.space_after else 0.0
                if ls is not None and round(float(ls), 2) != expected_spacing:
                    errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
                if round(sb, 1) != round(expected_before, 1):
                    errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
                if round(sa, 1) != round(expected_after, 1):
                    errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

                if errors:
                    results.append((title_idx, False, rule_title, "; ".join(errors)))
                else:
                    results.append((title_idx, True, rule_title, ""))
    
    # ======================================================
    # ABSTRACT Tez Başlığından Sonra 2 Satır Boşluk
    # ======================================================
    elif check["check"] == "abstract_en_spacing_after_title":
        expected_blank_lines = int(check.get("expected_blank_lines", 2))
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))

        import re

        def clean(s: str) -> str:
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        # ✅ Heading idx: memo öncelikli (yoksa markers fallback)
        heading_idx = memo.get("abstract_en_heading_idx", None)
        if heading_idx is None:
            markers = check.get("markers", [r"^ABSTRACT$"])
            for i, p in enumerate(paragraphs):
                candidate = clean(p.text)
                for m in markers:
                    if re.match(m, candidate, re.IGNORECASE):
                        heading_idx = i
                        memo["abstract_en_heading_idx"] = i
                        break
                if heading_idx is not None:
                    break

        rule_title = (
            f"Tez başlığından sonra {expected_blank_lines} satır boşluk: "
            f"{expected_name}, {int(expected_size)} punto, {expected_spacing} satır aralığı, "
            f"önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        if heading_idx is None:
            results.append((0, False, rule_title, "ABSTRACT başlığı bulunamadı"))
        else:
            # ✅ Title idx: memo öncelikli
            title_idx = memo.get("abstract_en_thesis_title_idx", None)
            if title_idx is None:
                # fallback: heading sonrası ilk dolu satır
                j = heading_idx + 1
                while j < len(paragraphs) and clean(paragraphs[j].text) == "":
                    j += 1
                title_idx = j if j < len(paragraphs) else None
                if title_idx is not None:
                    memo["abstract_en_thesis_title_idx"] = title_idx

            if title_idx is None:
                results.append((heading_idx, False, rule_title, "Tez başlığı bulunamadı (title_idx yok)"))
            else:
                errors = []

                # ✅ Title'dan sonra ardışık boş satır say
                j = title_idx + 1
                blank_count = 0
                while j < len(paragraphs) and clean(paragraphs[j].text) == "":
                    blank_count += 1
                    j += 1

                first_nonblank_idx = j if j < len(paragraphs) else None

                # ✅ memo: title sonrası boşluk sayısı + yazar idx (ilk dolu satır)
                memo["abstract_en_blank_after_title_count"] = blank_count
                if first_nonblank_idx is not None:
                    memo["abstract_en_author_idx"] = first_nonblank_idx

                # ✅ kural: en az N, fazlası ihlal
                if blank_count < expected_blank_lines:
                    errors.append(f"Boş satır sayısı {blank_count}; en az {expected_blank_lines} olmalı")
                elif blank_count > expected_blank_lines:
                    errors.append(f"Fazladan boş satır var: {blank_count} satır (beklenen: {expected_blank_lines})")

                # ✅ ilk expected_blank_lines satırının biçim kontrolleri
                for offset in range(1, min(blank_count, expected_blank_lines) + 1):
                    idx = title_idx + offset
                    if idx >= len(paragraphs):
                        break
                    p = paragraphs[idx]

                    # satır zaten boş olmalı
                    if clean(p.text) != "":
                        errors.append(f"{offset}. satır boş değil (text='{clean(p.text)}')")

                    # biçim kontrolü (mevcut mantık)
                    fn = None
                    for r in p.runs:
                        if r.font and r.font.name:
                            fn = r.font.name
                            break
                    if not fn and p.style and p.style.font and p.style.font.name:
                        fn = p.style.font.name
                    if fn and fn != expected_name:
                        errors.append(f"{offset}. satır yazı tipi {fn} yerine {expected_name} olmalı")

                    fs = None
                    for r in p.runs:
                        if r.font and r.font.size:
                            fs = r.font.size.pt
                            break
                    if not fs and p.style and p.style.font and p.style.font.size:
                        fs = p.style.font.size.pt
                    if fs and round(float(fs), 1) != round(expected_size, 1):
                        errors.append(f"{offset}. satır punto {fs} yerine {expected_size} olmalı")

                    pf = p.paragraph_format
                    ls = pf.line_spacing
                    sb = pf.space_before.pt if pf.space_before else 0.0
                    sa = pf.space_after.pt if pf.space_after else 0.0
                    if ls is not None and round(float(ls), 2) != expected_spacing:
                        errors.append(f"{offset}. satırın satır aralığı {ls} yerine {expected_spacing} olmalı")
                    if round(sb, 1) != round(expected_before, 1):
                        errors.append(f"{offset}. satır öncesi {sb} yerine {expected_before} olmalı")
                    if round(sa, 1) != round(expected_after, 1):
                        errors.append(f"{offset}. satır sonrası {sa} yerine {expected_after} olmalı")

                if errors:
                    results.append((title_idx, False, rule_title, "; ".join(errors)))
                else:
                    results.append((title_idx, True, rule_title, ""))


    # ======================================================
    # ABSTRACT Sonrası Tez Yazarının Adı Soyadı
    # ======================================================
    elif check["check"] == "abstract_en_author_name":
        expected_name = check.get("font_name", "Times New Roman")          # Beklenen yazı tipi
        expected_size = float(check.get("font_size_pt", 12))              # Beklenen punto
        expected_spacing = float(check.get("line_spacing", 1.0))          # Beklenen satır aralığı
        expected_before = float(check.get("space_before", 0))             # Beklenen paragraf öncesi boşluk
        expected_after = float(check.get("space_after", 0))               # Beklenen paragraf sonrası boşluk
        must_bold = check.get("bold", True)                               # Kalın olmalı mı?
        must_center = check.get("align", "center") == "center"            # Ortalı olmalı mı?
        must_upper = check.get("uppercase", True)                         # Büyük harf olmalı mı?

        import re

        def clean(s: str) -> str:
            # Boşluk/tab/nbsp temizliği + çoklu boşlukları tek boşluğa indirgeme
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        # ✅ Heading idx: memo öncelikli
        heading_idx = memo.get("abstract_en_heading_idx", None)
        if heading_idx is None:
            markers = check.get("markers", [r"^ABSTRACT$"])
            for i, p in enumerate(paragraphs):
                candidate = clean(p.text)
                for m in markers:
                    if re.match(m, candidate, re.IGNORECASE):
                        heading_idx = i
                        memo["abstract_en_heading_idx"] = i
                        break
                if heading_idx is not None:
                    break

        # Rapor satırında gözükecek kural açıklaması
        rule_title = (
            f"Tez yazarı adı-soyadı: {expected_name}, {int(expected_size)} punto, "
            f"{'kalın' if must_bold else 'normal'}, {'ortalı' if must_center else 'sol'}, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}, "
            f"{'BÜYÜK HARF' if must_upper else 'normal'}"
        )

        if heading_idx is None:
            results.append((0, False, rule_title, "ABSTRACT başlığı bulunamadı"))
        else:
            # ✅ author_idx: memo öncelikli (title sonrası ilk dolu satır)
            author_idx = memo.get("abstract_en_author_idx", None)

            # fallback: title idx varsa title’dan sonraki ilk dolu satır; yoksa heading sonrası ilk dolu satır
            if author_idx is None:
                title_idx = memo.get("abstract_en_thesis_title_idx", None)
                start = (title_idx + 1) if (title_idx is not None) else (heading_idx + 1)
                j = start
                while j < len(paragraphs) and clean(paragraphs[j].text) == "":
                    j += 1
                author_idx = j if j < len(paragraphs) else None
                if author_idx is not None:
                    memo["abstract_en_author_idx"] = author_idx

            if author_idx is None or author_idx >= len(paragraphs):
                results.append((heading_idx, False, rule_title, "Tez yazarı satırı yok"))
            else:
                p = paragraphs[author_idx]
                errors = []

                # Satır boş mu?
                if not clean(p.text):
                    errors.append("Tez yazarı satırı boş")

                # ======================================================
                # Büyük harf kontrolü (Türkçe karakter normalize + ALL CAPS desteği)  ✅ GÜNCELLENDİ
                # ======================================================
                def norm_tr(s: str) -> str:
                    # Türkçe karakterleri ASCII karşılıklarına indirger (İ/ı dahil) -> karşılaştırmayı stabil yapar
                    trans = str.maketrans({"ı":"I","i":"İ","ç":"Ç","ğ":"Ğ","ö":"Ö","ş":"Ş","ü":"Ü"})
                    return (s or "").translate(trans)

                def is_effectively_all_caps(para, txt_clean: str) -> bool:
                    """
                    Metin "efektif" olarak büyük harf mi?
                    1) Metin gerçekten büyük harf (normalize edilmiş karşılaştırma ile)
                    2) Word'de ALL CAPS biçimi (run veya style) açık
                    """
                    if not txt_clean:
                        return False

                    # 1) Metin düzeyinde büyük harf kontrolü (normalize edilmiş metin üzerinde)
                    if norm_tr(txt_clean) == norm_tr(txt_clean).upper():
                        return True

                    # 2) Run bazında ALL CAPS kontrolü
                    for r in para.runs:
                        if (r.text or "").strip() and r.font and r.font.all_caps:
                            return True

                    # 3) Stil zinciri üzerinden ALL CAPS kontrolü
                    s = para.style
                    while s is not None:
                        if s.font and s.font.all_caps:
                            return True
                        s = getattr(s, "base_style", None)

                    return False

                txt = clean(p.text)

                if must_upper and not is_effectively_all_caps(p, txt):
                    errors.append("Tez yazarı adı-soyadı büyük harflerle yazılmamış")

                # Yazı tipi
                fn = None
                for r in p.runs:
                    if r.font and r.font.name:
                        fn = r.font.name
                        break
                if not fn and p.style and p.style.font and p.style.font.name:
                    fn = p.style.font.name
                if fn and fn != expected_name:
                    errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")

                # Punto
                fs = None
                for r in p.runs:
                    if r.font and r.font.size:
                        fs = r.font.size.pt
                        break
                if not fs and p.style and p.style.font and p.style.font.size:
                    fs = p.style.font.size.pt
                if fs and round(float(fs), 1) != round(expected_size, 1):
                    errors.append(f"Punto {fs} yerine {expected_size} olmalı")

                # Kalınlık
                if must_bold:
                    any_bold = any(r.bold for r in p.runs if r.text.strip())
                    if not any_bold and p.style and p.style.font and p.style.font.bold:
                        any_bold = True
                    if not any_bold:
                        errors.append("Yazar adı kalın değil")

                # Ortalanma
                align_val = p.alignment
                if not align_val and p.style and p.style.paragraph_format.alignment:
                    align_val = p.style.paragraph_format.alignment
                if must_center and align_val != WD_PARAGRAPH_ALIGNMENT.CENTER:
                    errors.append("Yazar adı ortalı değil")

                # Satır aralığı, önce/sonra boşluk
                pf = p.paragraph_format
                ls = pf.line_spacing
                sb = pf.space_before.pt if pf.space_before else 0.0
                sa = pf.space_after.pt if pf.space_after else 0.0
                if ls is not None and round(float(ls), 2) != expected_spacing:
                    errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
                if round(sb, 1) != round(expected_before, 1):
                    errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
                if round(sa, 1) != round(expected_after, 1):
                    errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

                # Sonuç
                if errors:
                    results.append((author_idx, False, rule_title, "; ".join(errors)))
                else:
                    results.append((author_idx, True, rule_title, ""))

    # ======================================================
    # ABSTRACT Sonrası Tez Yazarının Adı- Soyadı Sonra 2 Satır Boşluk
    # ======================================================
    elif check["check"] == "abstract_en_spacing_after_author":
        expected_blank_lines = int(check.get("expected_blank_lines", 2))
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))

        import re

        def clean(s: str) -> str:
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        # ✅ Heading idx: memo öncelikli
        heading_idx = memo.get("abstract_en_heading_idx", None)
        if heading_idx is None:
            markers = check.get("markers", [r"^ABSTRACT$"])
            for i, p in enumerate(paragraphs):
                candidate = clean(p.text)
                for m in markers:
                    if re.match(m, candidate, re.IGNORECASE):
                        heading_idx = i
                        memo["abstract_en_heading_idx"] = i
                        break
                if heading_idx is not None:
                    break

        rule_title = (
            f"Tez yazarı sonrası {expected_blank_lines} satır boşluk: "
            f"{expected_name}, {int(expected_size)} punto, {expected_spacing} satır aralığı, "
            f"önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        if heading_idx is None:
            results.append((0, False, rule_title, "ABSTRACT başlığı bulunamadı"))
        else:
            # ✅ author_idx: memo öncelikli
            author_idx = memo.get("abstract_en_author_idx", None)
            if author_idx is None:
                # fallback: heading sonrası ilk dolu satır
                j = heading_idx + 1
                while j < len(paragraphs) and clean(paragraphs[j].text) == "":
                    j += 1
                author_idx = j if j < len(paragraphs) else None
                if author_idx is not None:
                    memo["abstract_en_author_idx"] = author_idx

            if author_idx is None or author_idx >= len(paragraphs):
                results.append((heading_idx, False, rule_title, "Tez yazarı satırı bulunamadı (author_idx yok)"))
            else:
                errors = []

                # ✅ author’dan sonra ardışık boş satır say
                j = author_idx + 1
                blank_count = 0
                while j < len(paragraphs) and clean(paragraphs[j].text) == "":
                    blank_count += 1
                    j += 1

                first_nonblank_idx = j if j < len(paragraphs) else None

                # ✅ memo: yazar sonrası boşluk sayısı + program bloğu başlangıcı (ilk dolu satır)
                memo["abstract_en_blank_after_author_count"] = blank_count
                if first_nonblank_idx is not None:
                    memo["abstract_en_program_start_idx"] = first_nonblank_idx

                # ✅ kural: en az N, fazlası ihlal
                if blank_count < expected_blank_lines:
                    errors.append(f"Boş satır sayısı {blank_count}; en az {expected_blank_lines} olmalı")
                elif blank_count > expected_blank_lines:
                    errors.append(f"Fazladan boş satır var: {blank_count} satır (beklenen: {expected_blank_lines})")

                # ✅ ilk expected_blank_lines satır biçim kontrolleri
                for offset in range(1, min(blank_count, expected_blank_lines) + 1):
                    idx = author_idx + offset
                    if idx >= len(paragraphs):
                        break
                    p = paragraphs[idx]

                    if clean(p.text) != "":
                        errors.append(f"{offset}. satır boş değil (text='{clean(p.text)}')")

                    fn = None
                    for r in p.runs:
                        if r.font and r.font.name:
                            fn = r.font.name
                            break
                    if not fn and p.style and p.style.font and p.style.font.name:
                        fn = p.style.font.name
                    if fn and fn != expected_name:
                        errors.append(f"{offset}. satır yazı tipi {fn} yerine {expected_name} olmalı")

                    fs = None
                    for r in p.runs:
                        if r.font and r.font.size:
                            fs = r.font.size.pt
                            break
                    if not fs and p.style and p.style.font and p.style.font.size:
                        fs = p.style.font.size.pt
                    if fs and round(float(fs), 1) != round(expected_size, 1):
                        errors.append(f"{offset}. satır punto {fs} yerine {expected_size} olmalı")

                    pf = p.paragraph_format
                    ls = pf.line_spacing
                    sb = pf.space_before.pt if pf.space_before else 0.0
                    sa = pf.space_after.pt if pf.space_after else 0.0
                    if ls is not None and round(float(ls), 2) != expected_spacing:
                        errors.append(f"{offset}. satırın satır aralığı {ls} yerine {expected_spacing} olmalı")
                    if round(sb, 1) != round(expected_before, 1):
                        errors.append(f"{offset}. satır öncesi {sb} yerine {expected_before} olmalı")
                    if round(sa, 1) != round(expected_after, 1):
                        errors.append(f"{offset}. satır sonrası {sa} yerine {expected_after} olmalı")

                if errors:
                    results.append((author_idx, False, rule_title, "; ".join(errors)))
                else:
                    results.append((author_idx, True, rule_title, ""))

    # ======================================================
    # ABSTRACT Sonrası Program / Danışman Bloğu
    # ======================================================
    elif check["check"] == "abstract_en_program_block":
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        expected_lines = int(check.get("expected_lines", 3))
        optional_lines = check.get("optional_lines", [3])  # 3. satır opsiyonel

        import re

        def clean(s: str) -> str:
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        # ✅ Heading idx: memo öncelikli (fallback markers)
        heading_idx = memo.get("abstract_en_heading_idx", None)
        if heading_idx is None:
            markers = check.get("markers", [r"^ABSTRACT$"])
            for i, p in enumerate(paragraphs):
                candidate = clean(p.text)
                for m in markers:
                    if re.match(m, candidate, re.IGNORECASE):
                        heading_idx = i
                        memo["abstract_en_heading_idx"] = i
                        break
                if heading_idx is not None:
                    break

        rule_title = (
            f"Program/Danışman Bloğu: {expected_lines} satır "
            f"({expected_name}, {int(expected_size)} punto, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)})"
        )

        if heading_idx is None:
            results.append((0, False, rule_title, "ABSTRACT başlığı bulunamadı"))
        else:
            # ✅ Program bloğu başlangıcı: memo öncelikli
            start_idx = memo.get("abstract_en_program_start_idx", None)

            # Fallback: author_idx sonrası ilk dolu satır
            if start_idx is None:
                author_idx = memo.get("abstract_en_author_idx", None)
                if author_idx is not None:
                    j = author_idx + 1
                    while j < len(paragraphs) and clean(paragraphs[j].text) == "":
                        j += 1
                    start_idx = j if j < len(paragraphs) else None
                    if start_idx is not None:
                        memo["abstract_en_program_start_idx"] = start_idx

            if start_idx is None or start_idx >= len(paragraphs):
                results.append((heading_idx, False, rule_title, "Program/Danışman bloğu başlangıcı bulunamadı"))
            else:
                errors = []

                # ✅ expected_lines kadar satır kontrol et (opsiyonel satır mantığı korunur)
                for line_num in range(1, expected_lines + 1):
                    idx = start_idx + line_num - 1
                    loc = f"{line_num}. satır (belge:{idx})"
                    
                    if idx >= len(paragraphs):
                        if line_num in optional_lines:
                            continue
                        errors.append(f"{loc}- yok")
                        continue

                    p = paragraphs[idx]
                    
                    raw = clean(p.text)
                    pv = " ".join(raw.split()[:7])
                    if len(pv) > 70:
                        pv = pv[:70].rstrip() + "…"
                    loc = loc + (f" ('{pv}')" if pv else "")


                    # Opsiyonel satır boş olabilir
                    if line_num in optional_lines and not clean(p.text):
                        continue

                    # Boş satır olamaz (opsiyonel hariç)
                    if not raw:
                        errors.append(f"{loc}- boş")
                        continue

                    # Yazı tipi
                    fn = None
                    for r in p.runs:
                        if r.font and r.font.name:
                            fn = r.font.name
                            break
                    if not fn and p.style and p.style.font and p.style.font.name:
                        fn = p.style.font.name
                    if fn and fn != expected_name:
                        errors.append(f"{loc}- yazı tipi {fn} yerine {expected_name} olmalı")

                    # Punto
                    fs = None
                    for r in p.runs:
                        if r.font and r.font.size:
                            fs = r.font.size.pt
                            break
                    if not fs and p.style and p.style.font and p.style.font.size:
                        fs = p.style.font.size.pt
                    if fs and round(float(fs), 1) != round(expected_size, 1):
                        errors.append(f"{loc}- punto {fs} yerine {expected_size} olmalı")

                    # Satır aralığı ve boşluklar
                    pf = p.paragraph_format
                    ls = pf.line_spacing
                    sb = pf.space_before.pt if pf.space_before else 0.0
                    sa = pf.space_after.pt if pf.space_after else 0.0
                    if ls is not None and round(float(ls), 2) != expected_spacing:
                        errors.append(f"{loc}- satır aralığı {ls} yerine {expected_spacing} olmalı")
                    if round(sb, 1) != round(expected_before, 1):
                        errors.append(f"{loc}- öncesi {sb} yerine {expected_before} olmalı")
                    if round(sa, 1) != round(expected_after, 1):
                        errors.append(f"{loc}- sonrası {sa} yerine {expected_after} olmalı")

                # ✅ Bloğun bittiği ilk boş satırı tespit edip memo’ya yaz
                # (Bundan sonrası spacing_after_program_block ve date/pages için anchor)
                j = start_idx
                while j < len(paragraphs) and clean(paragraphs[j].text) != "":
                    j += 1
                program_block_end_idx = j if j < len(paragraphs) else None
                if program_block_end_idx is not None:
                    memo["abstract_en_program_block_end_idx"] = program_block_end_idx

                    # boşlardan sonra gelen ilk dolu satır (date/pages genelde burada)
                    k = program_block_end_idx
                    while k < len(paragraphs) and clean(paragraphs[k].text) == "":
                        k += 1
                    memo["abstract_en_after_program_block_first_nonblank_idx"] = (k if k < len(paragraphs) else None)

                if errors:
                    results.append((start_idx, False, rule_title, "; ".join(errors)))
                else:
                    results.append((start_idx, True, rule_title, ""))

    # ======================================================
    # ABSTRACT Sonrası Program/Danışman Bloğundan Sonra 2 Satır Boşluk
    # ======================================================
    elif check["check"] == "abstract_en_spacing_after_program_block":
        expected_blank_lines = int(check.get("expected_blank_lines", 2))
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))

        import re

        def clean(s: str) -> str:
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        # ✅ Heading idx: memo öncelikli
        heading_idx = memo.get("abstract_en_heading_idx", None)
        if heading_idx is None:
            markers = check.get("markers", [r"^ABSTRACT$"])
            for i, p in enumerate(paragraphs):
                candidate = clean(p.text)
                for m in markers:
                    if re.match(m, candidate, re.IGNORECASE):
                        heading_idx = i
                        memo["abstract_en_heading_idx"] = i
                        break
                if heading_idx is not None:
                    break

        rule_title = (
            f"Program/Danışman bloğundan sonra {expected_blank_lines} satır boşluk: "
            f"{expected_name}, {int(expected_size)} punto, {expected_spacing} satır aralığı, "
            f"önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        if heading_idx is None:
            results.append((0, False, rule_title, "ABSTRACT başlığı bulunamadı"))
        else:
            # ✅ program_block_end_idx: memo öncelikli
            block_end_idx = memo.get("abstract_en_program_block_end_idx", None)

            # fallback: program_start_idx’den boş satıra kadar git
            if block_end_idx is None:
                start_idx = memo.get("abstract_en_program_start_idx", None)
                if start_idx is not None:
                    j = start_idx
                    while j < len(paragraphs) and clean(paragraphs[j].text) != "":
                        j += 1
                    block_end_idx = j if j < len(paragraphs) else None
                    if block_end_idx is not None:
                        memo["abstract_en_program_block_end_idx"] = block_end_idx

            if block_end_idx is None:
                results.append((heading_idx, False, rule_title, "Program bloğu bitişi bulunamadı"))
            else:
                errors = []

                # ✅ block_end_idx zaten ilk boş satır; buradan itibaren ardışık boş say
                j = block_end_idx
                blank_count = 0
                while j < len(paragraphs) and clean(paragraphs[j].text) == "":
                    blank_count += 1
                    j += 1

                first_nonblank_idx = j if j < len(paragraphs) else None

                # ✅ memo: program bloğu sonrası boşluk sayısı + date/pages anchor
                memo["abstract_en_blank_after_program_block_count"] = blank_count
                if first_nonblank_idx is not None:
                    memo["abstract_en_date_pages_anchor_idx"] = first_nonblank_idx

                # ✅ kural: en az N, fazlası ihlal
                if blank_count < expected_blank_lines:
                    errors.append(f"Boş satır sayısı {blank_count}; en az {expected_blank_lines} olmalı")
                elif blank_count > expected_blank_lines:
                    errors.append(f"Fazladan boş satır var: {blank_count} satır (beklenen: {expected_blank_lines})")

                # ✅ ilk expected_blank_lines boş satır biçim kontrolleri
                for offset in range(min(blank_count, expected_blank_lines)):
                    idx = block_end_idx + offset
                    if idx >= len(paragraphs):
                        break
                    p = paragraphs[idx]

                    if clean(p.text) != "":
                        errors.append(f"{offset+1}. satır boş değil (text='{clean(p.text)}')")

                    fn = None
                    for r in p.runs:
                        if r.font and r.font.name:
                            fn = r.font.name
                            break
                    if not fn and p.style and p.style.font and p.style.font.name:
                        fn = p.style.font.name
                    if fn and fn != expected_name:
                        errors.append(f"{offset+1}. satır yazı tipi {fn} yerine {expected_name} olmalı")

                    fs = None
                    for r in p.runs:
                        if r.font and r.font.size:
                            fs = r.font.size.pt
                            break
                    if not fs and p.style and p.style.font and p.style.font.size:
                        fs = p.style.font.size.pt
                    if fs and round(float(fs), 1) != round(expected_size, 1):
                        errors.append(f"{offset+1}. satır punto {fs} yerine {expected_size} olmalı")

                    pf = p.paragraph_format
                    ls = pf.line_spacing
                    sb = pf.space_before.pt if pf.space_before else 0.0
                    sa = pf.space_after.pt if pf.space_after else 0.0
                    if ls is not None and round(float(ls), 2) != expected_spacing:
                        errors.append(f"{offset+1}. satırın satır aralığı {ls} yerine {expected_spacing} olmalı")
                    if round(sb, 1) != round(expected_before, 1):
                        errors.append(f"{offset+1}. satır öncesi {sb} yerine {expected_before} olmalı")
                    if round(sa, 1) != round(expected_after, 1):
                        errors.append(f"{offset+1}. satır sonrası {sa} yerine {expected_after} olmalı")

                if errors:
                    results.append((block_end_idx, False, rule_title, "; ".join(errors)))
                else:
                    results.append((block_end_idx, True, rule_title, ""))

    # ======================================================
    # ABSTRACT Tarih Satırı (örn: "June 2025")
    # ======================================================
    elif check["check"] == "abstract_en_date":
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_bold = check.get("bold", True)
        expected_align = check.get("alignment", "center")
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))

        import re

        def clean(s: str) -> str:
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        # İngilizce aylar (kısa ve uzun)
        months_full = ["january","february","march","april","may","june","july","august","september","october","november","december"]
        months_abbr = ["jan","feb","mar","apr","may","jun","jul","aug","sep","oct","nov","dec"]
        month_pattern = r"\b(" + "|".join(months_full + months_abbr) + r")\s+\d{4}\b"

        rule_title = "ABSTRACT Tarih Satırı (örn: 'June 2025')"

        # ✅ anchor: program bloğu sonrası boşluklardan sonra gelen ilk dolu satır
        anchor_idx = memo.get("abstract_en_date_pages_anchor_idx", None)

        # fallback: heading_idx
        if anchor_idx is None:
            heading_idx = memo.get("abstract_en_heading_idx", None)
            anchor_idx = heading_idx if heading_idx is not None else 0

        found = None
        search_to = min(len(paragraphs), anchor_idx + 40)  # makul pencere
        for i in range(anchor_idx, search_to):
            txt = clean(paragraphs[i].text).lower()
            if re.search(month_pattern, txt, re.IGNORECASE):
                found = (i, paragraphs[i])
                break

        if not found:
            results.append((anchor_idx, False, rule_title,
                            "Tarih satırı bulunamadı (örn: 'June 2025'). Not: Tarih sayfa bilgisi ile aynı satırda olabilir."))
        else:
            idx, p = found

            # ✅ memo
            memo["abstract_en_date_idx"] = idx
            # date/pages info aynı satırdaysa pageinfo olarak da işaretleyebiliriz
            memo.setdefault("abstract_en_pageinfo_idx", idx)

            errors = []

            # Font
            fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
            if not fn and p.style and p.style.font and p.style.font.name:
                fn = p.style.font.name
            if fn and fn != expected_name:
                errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")

            # Size
            fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
            if not fs and p.style and p.style.font and p.style.font.size:
                fs = p.style.font.size.pt
            if fs and round(float(fs), 1) != round(expected_size, 1):
                errors.append(f"Punto {fs} yerine {expected_size} olmalı")

            # Bold
            if expected_bold:
                any_bold = any(r.bold for r in p.runs if r.text.strip())
                if not any_bold and p.style and p.style.font and getattr(p.style.font, "bold", None):
                    any_bold = True
                if not any_bold:
                    errors.append("Metin kalın değil")

            # Alignment
            align_val = p.alignment
            if not align_val and p.style and p.style.paragraph_format and p.style.paragraph_format.alignment:
                align_val = p.style.paragraph_format.alignment
            if expected_align == "center" and align_val != WD_PARAGRAPH_ALIGNMENT.CENTER:
                errors.append("Ortalanmamış")

            # Spacing
            pf = p.paragraph_format
            ls = pf.line_spacing
            sb = pf.space_before.pt if pf.space_before else 0.0
            sa = pf.space_after.pt if pf.space_after else 0.0
            if ls is not None and round(float(ls), 2) != expected_spacing:
                errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
            if round(sb, 1) != round(expected_before, 1):
                errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
            if round(sa, 1) != round(expected_after, 1):
                errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

            if errors:
                results.append((idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((idx, True, rule_title, ""))

    # ======================================================
    # ABSTRACT Sayfa Bilgisi (örn: "viii + 150 pages")
    # ======================================================
    elif check["check"] == "abstract_en_pages":
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_bold = check.get("bold", True)
        expected_align = check.get("alignment", "center")
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))

        import re

        def clean(s: str) -> str:
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        page_pattern = r"\b[ivxlcdm]+\b\s*\+\s*\d+\s*pages?\b"

        rule_title = "ABSTRACT Sayfa Bilgisi (örn: 'viii + 150 pages')"

        anchor_idx = memo.get("abstract_en_date_pages_anchor_idx", None)
        if anchor_idx is None:
            heading_idx = memo.get("abstract_en_heading_idx", None)
            anchor_idx = heading_idx if heading_idx is not None else 0

        found = None
        search_to = min(len(paragraphs), anchor_idx + 50)
        for i in range(anchor_idx, search_to):
            txt = clean(paragraphs[i].text).lower()
            if re.search(page_pattern, txt, re.IGNORECASE):
                found = (i, paragraphs[i])
                break

        if not found:
            results.append((anchor_idx, False, rule_title,
                            "Sayfa bilgisi satırı bulunamadı (örn: 'viii + 150 pages'). Not: Tarihle aynı satırda olabilir."))
        else:
            idx, p = found

            # ✅ memo
            memo["abstract_en_pages_idx"] = idx
            # pageinfo = date/pages satırı; aynı satırda olabilir, tek satırsa bu idx olur
            memo["abstract_en_pageinfo_idx"] = idx

            errors = []

            fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
            if not fn and p.style and p.style.font and p.style.font.name:
                fn = p.style.font.name
            if fn and fn != expected_name:
                errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")

            fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
            if not fs and p.style and p.style.font and p.style.font.size:
                fs = p.style.font.size.pt
            if fs and round(float(fs), 1) != round(expected_size, 1):
                errors.append(f"Punto {fs} yerine {expected_size} olmalı")

            if expected_bold:
                any_bold = any(r.bold for r in p.runs if r.text.strip())
                if not any_bold and p.style and p.style.font and getattr(p.style.font, "bold", None):
                    any_bold = True
                if not any_bold:
                    errors.append("Metin kalın değil")

            align_val = p.alignment
            if not align_val and p.style and p.style.paragraph_format and p.style.paragraph_format.alignment:
                align_val = p.style.paragraph_format.alignment
            if expected_align == "center" and align_val != WD_PARAGRAPH_ALIGNMENT.CENTER:
                errors.append("Ortalanmamış")

            pf = p.paragraph_format
            ls = pf.line_spacing
            sb = pf.space_before.pt if pf.space_before else 0.0
            sa = pf.space_after.pt if pf.space_after else 0.0
            if ls is not None and round(float(ls), 2) != expected_spacing:
                errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
            if round(sb, 1) != round(expected_before, 1):
                errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
            if round(sa, 1) != round(expected_after, 1):
                errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

            if errors:
                results.append((idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((idx, True, rule_title, ""))

    # ======================================================
    # ABSTRACT Tarih + Sayfa Bilgisinden Sonra 3 Satır Boşluk
    # ======================================================
    elif check["check"] == "abstract_en_spacing_after_pageinfo":
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))
        expected_lines = int(check.get("expected_lines", 3))

        def clean(s: str) -> str:
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        rule_title = (
            f"ABSTRACT sayfa bilgisinden sonra {expected_lines} satır boşluk: "
            f"{expected_name}, {int(expected_size)} punto, {expected_spacing} satır aralığı, "
            f"önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        # ✅ pageinfo idx: memo öncelikli
        pageinfo_idx = memo.get("abstract_en_pageinfo_idx", None)

        # fallback: pages_idx, yoksa date_idx
        if pageinfo_idx is None:
            pageinfo_idx = memo.get("abstract_en_pages_idx", None) or memo.get("abstract_en_date_idx", None)

        if pageinfo_idx is None:
            anchor_idx = memo.get("abstract_en_date_pages_anchor_idx", 0)
            results.append((anchor_idx, False, rule_title, "Sayfa bilgisi (pageinfo) bulunamadı"))
        else:
            errors = []

            # ✅ pageinfo sonrası ardışık boş satırları say (en az N, fazlası ihlal)
            j = pageinfo_idx + 1
            blank_count = 0
            while j < len(paragraphs) and clean(paragraphs[j].text) == "":
                blank_count += 1
                j += 1

            first_nonblank_idx = j if j < len(paragraphs) else None

            # ✅ memo
            memo["abstract_en_blank_after_pageinfo_count"] = blank_count
            if first_nonblank_idx is not None:
                memo["abstract_en_body_start_candidate_idx"] = first_nonblank_idx

            if blank_count < expected_lines:
                errors.append(f"Boş satır sayısı {blank_count}; en az {expected_lines} olmalı")
            elif blank_count > expected_lines:
                errors.append(f"Fazladan boş satır var: {blank_count} satır (beklenen: {expected_lines})")

            # ✅ ilk expected_lines satırın biçim kontrolü
            for k in range(1, min(blank_count, expected_lines) + 1):
                idx = pageinfo_idx + k
                if idx >= len(paragraphs):
                    break
                p = paragraphs[idx]

                if clean(p.text) != "":
                    errors.append(f"{k}. satır boş değil: '{clean(p.text)}'")

                pf = p.paragraph_format
                ls = pf.line_spacing
                sb = pf.space_before.pt if pf.space_before else 0.0
                sa = pf.space_after.pt if pf.space_after else 0.0

                fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
                fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)

                if not fn and p.style and p.style.font and p.style.font.name:
                    fn = p.style.font.name
                if not fs and p.style and p.style.font and p.style.font.size:
                    fs = p.style.font.size.pt

                if fn and fn != expected_name:
                    errors.append(f"{k}. satır yazı tipi {fn} yerine {expected_name} olmalı")
                if fs and round(float(fs), 1) != round(expected_size, 1):
                    errors.append(f"{k}. satır punto {fs} yerine {expected_size} olmalı")
                if ls is not None and round(float(ls), 2) != expected_spacing:
                    errors.append(f"{k}. satır aralığı {ls} yerine {expected_spacing} olmalı")
                if round(sb, 1) != round(expected_before, 1):
                    errors.append(f"{k}. satır öncesi {sb} yerine {expected_before} olmalı")
                if round(sa, 1) != round(expected_after, 1):
                    errors.append(f"{k}. satır sonrası {sa} yerine {expected_after} olmalı")

            if errors:
                results.append((pageinfo_idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((pageinfo_idx, True, rule_title, ""))

    # ======================================================
    # ABSTRACT Gövde Metni
    # ======================================================
    elif check["check"] == "abstract_en_body":
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))

        rule_title = (
            f"ABSTRACT Gövde: {expected_name}, {int(expected_size)} punto, "
            f"Justify, {expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        import re

        def clean(s: str) -> str:
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        # ✅ body_start: memo öncelikli
        body_start = memo.get("abstract_en_body_start_candidate_idx", None)

        # fallback: pageinfo idx + expected blanks (ama artık spacing_after_pageinfo memo yazıyor olmalı)
        if body_start is None:
            pageinfo_idx = memo.get("abstract_en_pageinfo_idx", None)
            blank_expected = int(check.get("blanks_after_pageinfo", 3))
            if pageinfo_idx is not None:
                j = pageinfo_idx + 1
                # boşları geç
                while j < len(paragraphs) and clean(paragraphs[j].text) == "":
                    j += 1
                body_start = j if j < len(paragraphs) else None

        # ✅ keywords idx: memo öncelikli
        keywords_idx = memo.get("abstract_en_keywords_idx", None)

        # fallback: keywords ara
        if keywords_idx is None:
            def is_keywords_para(p):
                t = clean(p.text).lower().replace(" ", "")
                return t.startswith("keywords:")
            start_scan = body_start if body_start is not None else 0
            for k in range(start_scan, min(len(paragraphs), start_scan + 400)):
                if is_keywords_para(paragraphs[k]):
                    keywords_idx = k
                    memo["abstract_en_keywords_idx"] = k
                    break

        if body_start is None:
            anchor = memo.get("abstract_en_pageinfo_idx", memo.get("abstract_en_date_pages_anchor_idx", 0))
            results.append((anchor, False, rule_title, "Gövde başlangıcı bulunamadı"))
        elif keywords_idx is None:
            results.append((body_start, False, rule_title, "Keywords satırı bulunamadı"))
        else:
            start_idx = body_start
            end_idx = keywords_idx - 1

            errors = []
            for k in range(start_idx, end_idx + 1):
                p = paragraphs[k]
                txt = clean(p.text)
                if txt == "":
                    continue

                # --- Hata mesajı için: "satır" + ilk birkaç kelime önizleme (fonksiyon yazmadan, burada) ---
                _t = txt
                _w = _t.split()
                _pv = " ".join(_w[:7])
                if len(_pv) > 60:
                    _pv = _pv[:60].rstrip()
                if _t and len(_t) > len(_pv):
                    _pv += "…"
                loc = f"{k}. satır" + (f" ('{_pv}')" if _pv else "")
                # ------------------------------------------------------------------------------------------------

                # Font
                fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
                if not fn and p.style and p.style.font and p.style.font.name:
                    fn = p.style.font.name
                if fn and fn != expected_name:
                    errors.append(f"{loc} — yazı tipi {fn} yerine {expected_name} olmalı")

                # Size
                fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
                if not fs and p.style and p.style.font and p.style.font.size:
                    fs = p.style.font.size.pt
                if fs and round(float(fs), 1) != round(expected_size, 1):
                    errors.append(f"{loc} — punto {fs} yerine {expected_size} olmalı")

                # Justify (stil dahil) + (None ise false positive üretme)
                align_val = p.alignment
                if align_val is None and p.style and p.style.paragraph_format:
                    align_val = p.style.paragraph_format.alignment
                if align_val is not None and align_val != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                    errors.append(f"{loc} — iki yana yaslı değil")

                # Spacing
                pf = p.paragraph_format
                ls = pf.line_spacing
                sb = pf.space_before.pt if pf.space_before else 0.0
                sa = pf.space_after.pt if pf.space_after else 0.0
                if ls is not None and round(float(ls), 2) != expected_spacing:
                    errors.append(f"{loc} — satır aralığı {ls} yerine {expected_spacing} olmalı")
                if round(sb, 1) != round(expected_before, 1):
                    errors.append(f"{loc} — öncesi {sb} yerine {expected_before} olmalı")
                if round(sa, 1) != round(expected_after, 1):
                    errors.append(f"{loc} — sonrası {sa} yerine {expected_after} olmalı")

            if errors:
                results.append((start_idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((start_idx, True, rule_title, ""))

    # ======================================================
    # ABSTRACT - KEYWORDS satırı
    # ======================================================
    elif check["check"] == "abstract_en_keywords":
        expected_name = check.get("font_name", "Times New Roman")
        expected_size = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before = float(check.get("space_before", 0))
        expected_after = float(check.get("space_after", 0))

        # ✅ YAML’den marker desteği (yanlış yazımlar / farklı formatlar için)
        # Örn: "^KEY\\s*WORDS\\s*:$", "^KEYWORDS\\s*:$", "^KEY\\-WORDS\\s*:$" gibi
        markers = check.get("markers", [r"^KEYWORDS\s*:\s*"])  # default: Keywords: ile başlayan

        rule_title = (
            f"ABSTRACT Keywords satırı: {expected_name}, {int(expected_size)} punto, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        import re

        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı": "i", "İ": "i", "ç": "c", "Ç": "c", "ğ": "g", "Ğ": "g",
                "ö": "o", "Ö": "o", "ş": "s", "Ş": "s", "ü": "u", "Ü": "u"
            })
            return (s or "").translate(trans)

        # Yardımcı: satır başını normalize ederek kontrol edelim
        def clean(s: str) -> str:
            # NBSP → space, tab → space, fazla boşlukları sadeleştir
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        found = None
        for i, p in enumerate(paragraphs):
            cand = clean(norm_tr(p.text)).upper()  # normalize + uppercase
            # marker match (toleranslı bulma)
            for m in markers:
                if re.match(m, cand, re.IGNORECASE):
                    found = (i, p)
                    break
            if found:
                break

        if not found:
            results.append((0, False, rule_title, "Keywords satırı bulunamadı"))
        else:
            idx, p = found

            # ✅ memo’ya kaydet (ileride başka kontroller kullanmak isterse)
            memo["abstract_en_keywords_idx"] = idx

            errors = []

            # ------------------------------------------------------
            # ✅ EK ÖZELLİK: Yazım doğrulama (bulduk ama doğru mu?)
            # Kural: satır "Keywords:" ile başlamalı (case-insensitive + boşluk toleranslı)
            # - "Key Words:" gibi yakalanır ama ihlal üretir.
            # ------------------------------------------------------
            raw = clean(p.text)
            raw_norm = clean(norm_tr(raw)).lower().replace(" ", "")  # "key words:" → "keywords:"
            # Doğru kabul: tam olarak "keywords:" ile başlamalı
            # (İstersen burada "keywords :" da kabul edilir; zaten boşlukları kaldırıyoruz)
            if not raw_norm.startswith("keywords:"):
                errors.append(f"Keywords yazımı hatalı (olması gereken: 'Keywords:'; bulunan: '{raw}')")

            # ---------- Hizalama ----------
            align_val = p.alignment
            style_align = p.style.paragraph_format.alignment if (p.style and p.style.paragraph_format) else None
            if not (align_val in (None, WD_PARAGRAPH_ALIGNMENT.LEFT) or style_align in (None, WD_PARAGRAPH_ALIGNMENT.LEFT)):
                errors.append("Satır sola yaslı değil")

            # ---------- Font & Punto ----------
            fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
            fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
            if not fn and p.style and p.style.font and p.style.font.name:
                fn = p.style.font.name
            if not fs and p.style and p.style.font and p.style.font.size:
                fs = p.style.font.size.pt

            if fn and fn != expected_name:
                errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")
            if fs and round(float(fs), 1) != round(expected_size, 1):
                errors.append(f"Punto {fs} yerine {expected_size} olmalı")

            # ---------- Satır aralığı & boşluklar ----------
            pf = p.paragraph_format
            ls = pf.line_spacing
            sb = pf.space_before.pt if pf.space_before else 0.0
            sa = pf.space_after.pt if pf.space_after else 0.0
            if ls is not None and round(float(ls), 2) != expected_spacing:
                errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
            if round(sb, 1) != round(expected_before, 1):
                errors.append(f"Paragraf öncesi {sb} yerine {expected_before} olmalı")
            if round(sa, 1) != round(expected_after, 1):
                errors.append(f"Paragraf sonrası {sa} yerine {expected_after} olmalı")

            # ---------- "Keywords:" bold kontrolü ----------
            bold_keywords = False
            for r in p.runs:
                # Burada da "key words" gibi varyasyonları yakalayalım
                rt = clean(norm_tr(r.text)).lower().replace(" ", "")
                if "keywords" in rt:
                    if r.bold or (r.font and r.font.bold):
                        bold_keywords = True
            if p.style and p.style.font and p.style.font.bold:
                bold_keywords = True

            if not bold_keywords:
                errors.append("Keywords: metni kalın değil")

            # ---------- Sonuç ----------
            if errors:
                results.append((idx, False, rule_title, "; ".join(errors)))
            else:
                results.append((idx, True, rule_title, ""))
    # ======================================================


    # ===============================================================================================================#
    # ===============================================================================================================#    
    # ===============================================================================================================#
    # ÖN SÖZ / TEŞEKKÜR SAYFASI - BAŞLIK KONTROLÜ (EFFECTIVE: run + style zinciri)
    # ===============================================================================================================#
    elif check["check"] == "acknowledgements_heading":
        import re

        expected_name     = check.get("font_name", "Times New Roman")
        expected_size     = float(check.get("font_size_pt", 12))
        expected_bold     = bool(check.get("bold", True))
        expected_all_caps = bool(check.get("all_caps", True))
        expected_align    = (check.get("alignment", "center") or "center").lower().strip()
        expected_spacing  = float(check.get("line_spacing", 1.5))
        expected_before   = float(check.get("space_before", 0))
        expected_after    = float(check.get("space_after", 24))
        markers           = check.get("markers", ["^ÖN SÖZ$", "^TEŞEKKÜR$", "^ÖN SÖZ/TEŞEKKÜR$", "^ÖN SÖZ / TEŞEKKÜR$"])

        rule_title = (
            f"ÖN SÖZ / TEŞEKKÜR başlığı: {expected_name}, {int(expected_size)} punto, "
            f"{'kalın' if expected_bold else 'normal'}, {'büyük harf' if expected_all_caps else 'normal harf'}, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        def clean(s: str) -> str:
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u"
            })
            return (clean(s)).lower().translate(trans).strip()

        # --- Stil zincirinden "etkili" değer çözümleme ---
        def resolve_from_styles(para, attr_name):
            val = getattr(para.paragraph_format, attr_name)
            if val is not None:
                return val
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None:
                    v = getattr(pf, attr_name)
                    if v is not None:
                        return v
                s = getattr(s, "base_style", None)
            return None

        def effective_alignment(para):
            if para.alignment is not None:
                return para.alignment
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None and pf.alignment is not None:
                    return pf.alignment
                s = getattr(s, "base_style", None)
            return None

        def effective_space_pt(para, which: str) -> float:
            attr = "space_before" if which == "before" else "space_after"
            length = resolve_from_styles(para, attr)
            return 0.0 if length is None else round(length.pt, 1)

        def effective_line_spacing(para, default=1.0) -> float:
            ls = resolve_from_styles(para, "line_spacing")
            if ls is None:
                return float(default)
            if hasattr(ls, "pt"):
                return round(ls.pt, 1)
            try:
                return round(float(ls), 2)
            except Exception:
                return ls

        def effective_bold(para) -> bool:
            # 1) run/r.font.bold açıkça True ise
            for r in para.runs:
                if r.bold is True:
                    return True
                if r.font and r.font.bold is True:
                    return True
            # 2) stil zinciri
            s = para.style
            while s is not None:
                if getattr(s, "font", None) and s.font.bold is True:
                    return True
                s = getattr(s, "base_style", None)
            return False

        def effective_font_name(para):
            fn = next((r.font.name for r in para.runs if r.font and r.font.name), None)
            if not fn and para.style and para.style.font and para.style.font.name:
                fn = para.style.font.name
            return fn

        def effective_font_size_pt(para):
            fs = next((r.font.size.pt for r in para.runs if r.font and r.font.size), None)
            if (fs is None) and para.style and para.style.font and para.style.font.size:
                fs = para.style.font.size.pt
            return fs

        # --- Marker’ları normalize edip compile et ---
        compiled_markers = []
        for m in markers:
            m_norm = norm_tr(m)
            try:
                compiled_markers.append(re.compile(m_norm, re.IGNORECASE))
            except re.error:
                compiled_markers.append(re.compile(re.escape(m_norm), re.IGNORECASE))

        # --- Başlığı ara ---
        found = None
        for i, p in enumerate(paragraphs):
            raw = clean(p.text)
            if raw == "":
                continue
            n = norm_tr(raw)

            # marker eşleşmesi
            if any(pat.match(n) for pat in compiled_markers):
                found = (i, p)
                break

            # esnek fallback (bazı şablonlar için)
            if re.match(r"^on\s*soz\s*/?\s*tesekkur$", n, flags=re.IGNORECASE):
                found = (i, p)
                break

        if not found:
            results.append((0, False, rule_title, "ÖN SÖZ/TEŞEKKÜR başlığı bulunamadı"))
            return results

        idx, p = found

        # ✅ memo'ya yaz (diğer kontroller buradan okuyacak)
        memo["acknowledgements_heading_idx"] = idx
        memo["acknowledgements_heading_text"] = clean(p.text)

        # --- Effective değerler ---
        eff_align = effective_alignment(p)
        eff_ls    = effective_line_spacing(p)
        eff_sb    = effective_space_pt(p, "before")
        eff_sa    = effective_space_pt(p, "after")

        errors = []

        fn = effective_font_name(p)
        fs = effective_font_size_pt(p)

        if fn and fn != expected_name:
            errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")
        if fs is not None and round(float(fs), 1) != round(expected_size, 1):
            errors.append(f"Punto {fs} yerine {expected_size} olmalı")

        if expected_bold and not effective_bold(p):
            errors.append("Başlık kalın değil")

        if expected_all_caps and clean(p.text) and not clean(p.text).isupper():
            errors.append("Başlık büyük harflerle yazılmamış")

        if expected_align == "center" and eff_align != WD_PARAGRAPH_ALIGNMENT.CENTER:
            errors.append("Başlık ortalı değil")
        elif expected_align == "left" and eff_align != WD_PARAGRAPH_ALIGNMENT.LEFT:
            errors.append("Başlık sola dayalı değil")
        elif expected_align == "right" and eff_align != WD_PARAGRAPH_ALIGNMENT.RIGHT:
            errors.append("Başlık sağa dayalı değil")

        # toleranslı aralık karşılaştırmaları
        if isinstance(eff_ls, (int, float)) and abs(eff_ls - expected_spacing) > 0.1:
            errors.append(f"Satır aralığı {eff_ls} yerine {expected_spacing} olmalı")
        if abs(eff_sb - expected_before) > 1:
            errors.append(f"Paragraf öncesi {eff_sb} yerine {expected_before} olmalı")
        if abs(eff_sa - expected_after) > 1:
            errors.append(f"Paragraf sonrası {eff_sa} yerine {expected_after} olmalı")

        if errors:
            results.append((idx, False, rule_title, "; ".join(errors)))
        else:
            results.append((idx, True, rule_title, ""))

    # ======================================================
    # ÖN SÖZ / TEŞEKKÜR - BAŞLIKTAN SONRA BOŞLUK KONTROLÜ (EFFECTIVE + memo)
    # ======================================================
    elif check["check"] == "acknowledgements_no_blank":
        rule_title = "ÖN SÖZ / TEŞEKKÜR başlığı ile gövde arasında boşluk kontrolü"

        def clean(s: str) -> str:
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u"
            })
            return (clean(s)).lower().translate(trans).strip()

        # ✅ 1) önce memo'dan al
        heading_idx = memo.get("acknowledgements_heading_idx", None)

        # ✅ 2) memo yoksa fallback ara (aynı mantık)
        if heading_idx is None:
            import re
            markers = check.get("markers", [
                "^ÖN SÖZ$", "^TEŞEKKÜR$", "^ÖN SÖZ/TEŞEKKÜR$", "^ÖN SÖZ / TEŞEKKÜR$", "^ÖNSÖZ/TEŞEKKÜR$"
            ])
            compiled = []
            for m in markers:
                m_norm = norm_tr(m)
                try:
                    compiled.append(re.compile(m_norm, re.IGNORECASE))
                except re.error:
                    compiled.append(re.compile(re.escape(m_norm), re.IGNORECASE))

            for i, p in enumerate(paragraphs):
                t = norm_tr(p.text)
                if t and any(pat.match(t) for pat in compiled):
                    heading_idx = i
                    break

        if heading_idx is None:
            results.append((0, False, rule_title, "ÖN SÖZ/TEŞEKKÜR başlığı bulunamadı"))
            return results

        # --- başlıktan sonra ilk dolu paragrafı bul ---
        first_body_idx = None
        had_blank_between = False

        j = heading_idx + 1
        while j < len(paragraphs):
            t = clean(paragraphs[j].text)
            if t == "":
                had_blank_between = True
                j += 1
                continue
            first_body_idx = j
            break

        if first_body_idx is None:
            results.append((heading_idx, False, rule_title, "Başlıktan sonra gövde paragrafı bulunamadı"))
            return results

        # ✅ memo: gövde başlangıcını kaydet (body_format burada tekrar aramasın)
        memo["acknowledgements_body_start_idx"] = first_body_idx

        if had_blank_between:
            # boş satırın ilk görüldüğü yere işaret etmek daha anlamlı olur
            results.append((heading_idx + 1, False, rule_title, "Başlık ile gövde arasında boş satır(lar) var"))
        else:
            results.append((first_body_idx, True, rule_title, ""))

    # ======================================================
    # ÖN SÖZ / TEŞEKKÜR - GÖVDE METNİ BİÇİM KONTROLÜ (EFFECTIVE: run + style zinciri)
    # ======================================================
    elif check["check"] == "acknowledgements_body_format":
        import re

        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_align   = (check.get("alignment", "justify") or "justify").lower().strip()
        expected_spacing = float(check.get("line_spacing", 1.5))
        expected_before  = float(check.get("space_before", 6))
        expected_after   = float(check.get("space_after", 6))

        # Opsiyonel: gövdeyi nerede durduracağız? (bir sonraki bölüm başlığı vb.)
        stop_markers = check.get("stop_markers", [
            r"^İÇİNDEKİLER\b", r"^OZET\b", r"^ÖZET\b", r"^ABSTRACT\b",
            r"^SİMGELER\b", r"^KISALTMALAR\b", r"^TABLOLAR\b", r"^ŞEKİLLER\b"
        ])

        rule_title = (
            f"ÖN SÖZ / TEŞEKKÜR gövdesi: {expected_name}, {int(expected_size)} punto, "
            f"{expected_align} hizalı, {expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        def clean(s: str) -> str:
            return " ".join((s or "").replace("\u00A0", " ").replace("\t", " ").strip().split())

        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u"
            })
            return (clean(s)).lower().translate(trans).strip()

        # --- Stil zinciri (effective) ---
        def resolve_from_styles(para, attr_name):
            val = getattr(para.paragraph_format, attr_name)
            if val is not None:
                return val
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None:
                    v = getattr(pf, attr_name)
                    if v is not None:
                        return v
                s = getattr(s, "base_style", None)
            return None

        def effective_alignment(para):
            if para.alignment is not None:
                return para.alignment
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None and pf.alignment is not None:
                    return pf.alignment
                s = getattr(s, "base_style", None)
            return None

        def effective_space_pt(para, which: str) -> float:
            attr = "space_before" if which == "before" else "space_after"
            length = resolve_from_styles(para, attr)
            return 0.0 if length is None else round(length.pt, 1)

        def effective_line_spacing(para, default=1.0) -> float:
            ls = resolve_from_styles(para, "line_spacing")
            if ls is None:
                return float(default)
            if hasattr(ls, "pt"):
                return round(ls.pt, 1)
            try:
                return round(float(ls), 2)
            except Exception:
                return ls

        def effective_font_name(para):
            fn = next((r.font.name for r in para.runs if r.font and r.font.name), None)
            if not fn and para.style and para.style.font and para.style.font.name:
                fn = para.style.font.name
            return fn

        def effective_font_size_pt(para):
            fs = next((r.font.size.pt for r in para.runs if r.font and r.font.size), None)
            if (fs is None) and para.style and para.style.font and para.style.font.size:
                fs = para.style.font.size.pt
            return fs

        # 1) body_start: memo öncelikli
        body_start = memo.get("acknowledgements_body_start_idx", None)

        # 2) body_start yoksa: heading’den sonra ilk dolu paragrafı bul
        if body_start is None:
            heading_idx = memo.get("acknowledgements_heading_idx", None)

            # heading_idx yoksa esnekçe ara
            if heading_idx is None:
                for i, p in enumerate(paragraphs):
                    if re.match(r"^on\s*soz\s*/?\s*tesekkur$", norm_tr(p.text), flags=re.IGNORECASE):
                        heading_idx = i
                        memo["acknowledgements_heading_idx"] = i
                        break

            if heading_idx is None:
                results.append((0, False, rule_title, "ÖN SÖZ/TEŞEKKÜR başlığı bulunamadı, gövde kontrolü atlandı"))
                return results

            j = heading_idx + 1
            while j < len(paragraphs) and clean(paragraphs[j].text) == "":
                j += 1
            body_start = j if j < len(paragraphs) else None
            if body_start is not None:
                memo["acknowledgements_body_start_idx"] = body_start

        if body_start is None:
            anchor = memo.get("acknowledgements_heading_idx", 0)
            results.append((anchor, False, rule_title, "Gövde başlangıcı bulunamadı"))
            return results

        # 3) body_end: stop_markers ile bul (yoksa 200 paragrafla sınırla)
        compiled_stop = []
        for pat in stop_markers:
            try:
                compiled_stop.append(re.compile(pat, re.IGNORECASE))
            except re.error:
                compiled_stop.append(re.compile(re.escape(pat), re.IGNORECASE))

        body_end = min(len(paragraphs) - 1, body_start + 200)
        for k in range(body_start, min(len(paragraphs), body_start + 200)):
            t = clean(paragraphs[k].text)
            if t == "":
                continue
            if any(cp.match(t) for cp in compiled_stop):
                body_end = k - 1
                break

        # --- Gövde paragraflarını tek tek kontrol et (ABSTRACT mantığı) ---
        errors = []
        for k in range(body_start, body_end + 1):
            
            loc = f"{k}. satır (belge:{k})"

            p = paragraphs[k]
            txt = clean(p.text)
            if txt == "":
                continue

            pv = " ".join(txt.split()[:10])  # gövde için 10 kelime daha iyi
            if len(pv) > 90:
                pv = pv[:90].rstrip() + "…"
            loc = loc + (f" ('{pv}')" if pv else "")


            fn = effective_font_name(p)
            fs = effective_font_size_pt(p)

            if fn and fn != expected_name:
                errors.append(f"{loc}- yazı tipi {fn} yerine {expected_name} olmalı")
            if fs is not None and round(float(fs), 1) != round(expected_size, 1):
                errors.append(f"{loc}- punto {fs} yerine {expected_size} olmalı")

            eff_align = effective_alignment(p)
            eff_ls    = effective_line_spacing(p)
            eff_sb    = effective_space_pt(p, "before")
            eff_sa    = effective_space_pt(p, "after")

            if expected_align == "justify" and eff_align != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                errors.append(f"{loc}- iki yana yaslı değil")
            elif expected_align == "center" and eff_align != WD_PARAGRAPH_ALIGNMENT.CENTER:
                errors.append(f"{loc}- ortalı değil")
            elif expected_align == "left" and eff_align != WD_PARAGRAPH_ALIGNMENT.LEFT:
                errors.append(f"{loc}- sola dayalı değil")
            elif expected_align == "right" and eff_align != WD_PARAGRAPH_ALIGNMENT.RIGHT:
                errors.append(f"{loc}- sağa dayalı değil")

            if isinstance(eff_ls, (int, float)) and abs(eff_ls - expected_spacing) > 0.1:
                errors.append(f"{loc}- satır aralığı {eff_ls} yerine {expected_spacing} olmalı")
            if abs(eff_sb - expected_before) > 1:
                errors.append(f"{loc}- öncesi {eff_sb} yerine {expected_before} olmalı")
            if abs(eff_sa - expected_after) > 1:
                errors.append(f"{loc}- sonrası {eff_sa} yerine {expected_after} olmalı")

        if errors:
            results.append((body_start, False, rule_title, "; ".join(errors)))
        else:
            results.append((body_start, True, rule_title, ""))



    # ===============================================================================================================#    
    # ===============================================================================================================#
    # ===============================================================================================================#
    # ======================================================
    # ======================================================
    # İÇİNDEKİLER - BAŞLIK (ÖNSÖZ/TEŞEKKÜR'DEN SONRA, BİÇİM AYNI)
    # ======================================================
    elif check["check"] == "toc_heading":
        import re
        from collections import Counter
        import xml.etree.ElementTree as ET

        expected_name     = check.get("font_name", "Times New Roman")
        expected_size     = float(check.get("font_size_pt", 12))
        expected_bold     = bool(check.get("bold", True))
        expected_all_caps = bool(check.get("all_caps", True))
        expected_align    = (check.get("alignment", "center") or "center").lower()
        expected_spacing  = float(check.get("line_spacing", 1.5))
        expected_before   = float(check.get("space_before", 0))
        expected_after    = float(check.get("space_after", 24))
        markers           = check.get("markers", ["^İÇİNDEKİLER$"])

        rule_title = (
            f"İÇİNDEKİLER başlığı: {expected_name}, {int(expected_size)} punto, "
            f"{'kalın' if expected_bold else 'normal'}, "
            f"{'BÜYÜK HARF' if expected_all_caps else 'normal'}, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        # ---------- yardımcılar ----------
        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u"
            })
            return (s or "").lower().translate(trans).strip()

        def resolve_from_styles_pf(para, attr):
            val = getattr(para.paragraph_format, attr)
            if val is not None:
                return val
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None:
                    v = getattr(pf, attr)
                    if v is not None:
                        return v
                s = getattr(s, "base_style", None)
            return None

        def effective_space_pt(para, which: str) -> float:
            attr = "space_before" if which == "before" else "space_after"
            length = resolve_from_styles_pf(para, attr)
            return 0.0 if length is None else round(length.pt, 1)

        def effective_line_spacing(para, default=1.0) -> float:
            ls = resolve_from_styles_pf(para, "line_spacing")
            if ls is None:
                return float(default)
            if hasattr(ls, "pt"):
                return round(ls.pt, 1)
            try:
                return round(float(ls), 2)
            except Exception:
                return ls

        def effective_alignment(para):
            if para.alignment is not None:
                return para.alignment
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None and pf.alignment is not None:
                    return pf.alignment
                s = getattr(s, "base_style", None)
            return None

        def effective_bold(para):
            for r in para.runs:
                if r.bold is True or (r.font and r.font.bold is True):
                    return True
            s = para.style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.bold is True:
                    return True
                s = getattr(s, "base_style", None)
            return False

        # ========== THEME + XML tabanlı "effective font" ==========
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"

        def _get_theme_fonts_from_package():
            # memo cache
            if memo.get("_theme_fonts"):
                return memo["_theme_fonts"]
            minor_latin = None
            major_latin = None
            try:
                pkg = paragraphs[0]._parent.part.package
                theme_part = None
                for part in pkg.parts:
                    if str(part.partname).endswith("/word/theme/theme1.xml"):
                        theme_part = part
                        break
                if theme_part is not None:
                    root = ET.fromstring(theme_part.blob)
                    ns = {"a": A_NS}
                    minor = root.find(".//a:minorFont/a:latin", ns)
                    major = root.find(".//a:majorFont/a:latin", ns)
                    minor_latin = minor.get("typeface") if minor is not None else None
                    major_latin = major.get("typeface") if major is not None else None
            except Exception:
                pass
            memo["_theme_fonts"] = {"minor_latin": minor_latin, "major_latin": major_latin}
            return memo["_theme_fonts"]

        def _theme_to_font(theme_key: str, theme_fonts: dict):
            if not theme_key:
                return None
            k = theme_key.lower()
            # minorHAnsi / minorAscii / minorEastAsia / minorBidi
            if k.startswith("minor"):
                return theme_fonts.get("minor_latin")
            if k.startswith("major"):
                return theme_fonts.get("major_latin")
            return None

        def effective_font_name(para) -> str:
            theme_fonts = _get_theme_fonts_from_package()

            # 1) Para XML içinde w:rFonts (explicit veya theme) ara
            try:
                p_xml = ET.fromstring(para._element.xml)
                fonts = []

                for rfonts in p_xml.iter():
                    if rfonts.tag == f"{{{W_NS}}}rFonts":
                        # explicit
                        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
                            v = rfonts.get(f"{{{W_NS}}}{attr}")
                            if v:
                                fonts.append(v)
                        # theme
                        for attr in ("asciiTheme", "hAnsiTheme", "csTheme", "eastAsiaTheme", "cstheme"):
                            v = rfonts.get(f"{{{W_NS}}}{attr}")
                            if v:
                                tv = _theme_to_font(v, theme_fonts)
                                if tv:
                                    fonts.append(tv)

                if fonts:
                    return Counter(fonts).most_common(1)[0][0]
            except Exception:
                pass

            # 2) Stil zinciri (font.name)
            s = para.style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.name:
                    return f.name
                s = getattr(s, "base_style", None)

            # 3) Son çare: theme minor
            return theme_fonts.get("minor_latin") or None

        # ---------- 1) ÖNSÖZ/TEŞEKKÜR başlığını bul ----------
        ack_idx = memo.get("acknowledgements_heading_idx")
        if ack_idx is None:
            for i, p in enumerate(paragraphs):
                if re.match(r"^on\s*soz\s*/?\s*tesekkur$", norm_tr(p.text)):
                    ack_idx = i
                    memo["acknowledgements_heading_idx"] = i
                    break

        if ack_idx is None:
            results.append((0, False, rule_title, "ÖN SÖZ/TEŞEKKÜR başlığı bulunamadı; İÇİNDEKİLER konum kontrolü yapılamadı"))
            return results

        # ---------- 2) İÇİNDEKİLER başlığını sadece ack'ten sonra ara ----------
        compiled = []
        for m in markers:
            try:
                compiled.append(re.compile(norm_tr(m), re.IGNORECASE))
            except re.error:
                compiled.append(re.compile("^" + re.escape(norm_tr(m).strip("^$")) + "$", re.IGNORECASE))

        found = None
        for i in range(ack_idx + 1, len(paragraphs)):
            raw  = (paragraphs[i].text or "").strip()
            norm = norm_tr(raw)
            for pat in compiled:
                if pat.match(norm):
                    found = (i, paragraphs[i])
                    break
            if found:
                break

        if not found:
            results.append((ack_idx, False, rule_title, "İÇİNDEKİLER başlığı bulunamadı (ÖN SÖZ/TEŞEKKÜR’den sonra)"))
            return results

        # ---------- 3) Biçim kontrolleri ----------
        idx, p = found
        errors = []

        fn_eff = effective_font_name(p)
        if fn_eff and re.sub(r"\s+", "", fn_eff.lower()) != re.sub(r"\s+", "", expected_name.lower()):
            errors.append(f"Yazı tipi {fn_eff} yerine {expected_name} olmalı")

        fs = None
        for r in p.runs:
            if r.font and r.font.size:
                fs = r.font.size.pt
                break
        if not fs and p.style and p.style.font and p.style.font.size:
            fs = p.style.font.size.pt
        if fs and round(float(fs), 1) != round(expected_size, 1):
            errors.append(f"Punto {fs} yerine {expected_size} olmalı")

        if expected_bold and not effective_bold(p):
            errors.append("Başlık kalın değil")
        if expected_all_caps and not (p.text or "").isupper():
            errors.append("Başlık BÜYÜK HARF değil")

        eff_align = effective_alignment(p)
        eff_ls    = effective_line_spacing(p, default=expected_spacing)
        eff_sb    = effective_space_pt(p, "before")
        eff_sa    = effective_space_pt(p, "after")

        if expected_align == "center" and eff_align != WD_PARAGRAPH_ALIGNMENT.CENTER:
            errors.append("Başlık ortalı değil")
        if isinstance(eff_ls, (int, float)) and abs(eff_ls - expected_spacing) > 0.1:
            errors.append(f"Satır aralığı {eff_ls} yerine {expected_spacing} olmalı")
        if abs(eff_sb - expected_before) > 1:
            errors.append(f"Paragraf öncesi {eff_sb} yerine {expected_before} olmalı")
        if abs(eff_sa - expected_after) > 1:
            errors.append(f"Paragraf sonrası {eff_sa} yerine {expected_after} olmalı")

        if errors:
            results.append((idx, False, rule_title, "; ".join(errors)))
        else:
            results.append((idx, True, rule_title, ""))

        memo["toc_heading_idx"] = idx

    # ======================================================
    # İÇİNDEKİLER – BAŞLIKTAN SONRA BOŞ SATIR YOK (hemen 'Sayfa' satırı gelmeli)
    # ======================================================
    elif check["check"] == "toc_one_blank_with_format":
        import re

        # Artık bu check'in amacı: 1 boş satır değil, 0 boş satır
        markers = check.get("markers", ["^İÇİNDEKİLER$"])

        rule_title = "İÇİNDEKİLER – başlıktan sonra boş satır olmamalı (hemen 'Sayfa' satırı gelmeli)"

        def is_blank_para_text(txt: str) -> bool:
            if txt is None:
                return True
            t = txt.replace("\xa0", " ")
            return (t.strip() == "") or ("".join(t.split()) == "")

        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u"
            })
            return (s or "").replace("\u00A0", " ").strip().lower().translate(trans)

        # --- TOC başlığını memo’dan al (en güvenlisi)
        toc_idx = memo.get("toc_heading_idx", None)
        if toc_idx is None:
            return [(0, False, rule_title, "memo'da toc_heading_idx yok; önce 'toc_heading' çalışmalı")]

        first_after = toc_idx + 1
        if first_after >= len(paragraphs):
            return [(toc_idx, False, rule_title, "Başlıktan sonra paragraf yok")]

        # ✅ Yeni kural: başlıktan hemen sonra gelen satır boş olamaz
        if is_blank_para_text(paragraphs[first_after].text):
            # kaç tane boş var (mesajı zenginleştirmek için)
            j = first_after
            blanks = 0
            while j < len(paragraphs) and is_blank_para_text(paragraphs[j].text):
                blanks += 1
                j += 1
            return [(first_after, False, rule_title, f"Başlıktan sonra boş satır olmamalı (bulunan boş satır sayısı: {blanks})")]

        # (Opsiyonel ama pratik) İlk satırın 'Sayfa' olduğunu burada kontrol etmek istiyorsan:
        txt = (paragraphs[first_after].text or "").replace("\xa0", " ").strip()
        core = re.sub(r"[^\w]+", "", norm_tr(txt))
        if core not in ("sayfa", "page"):
            # Biçimi zaten 'toc_page_label_line' kontrol ediyor; burada sadece konumu garanti ediyoruz.
            return [(first_after, False, rule_title, f"Boş satır yok ama hemen sonra 'Sayfa' satırı gelmeli (bulunan: '{txt}')")]

        return [(first_after, True, rule_title, "")]


    # ======================================================
    # İÇİNDEKİLER – 'Sayfa' satırı (başlıktan sonra tarayarak bulur)
    # ======================================================
    elif check["check"] == "toc_page_label_line":
        import re
        from collections import Counter
        import xml.etree.ElementTree as ET

        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_align   = (check.get("alignment", "right") or "right").lower().strip()
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 0))
        expected_bold    = bool(check.get("bold", True))

        rule_title = (
            f"İÇİNDEKİLER – 'Sayfa' satırı: {expected_name}, {int(expected_size)} pt, "
            f"{'sağa dayalı' if expected_align=='right' else expected_align}, "
            f"{expected_spacing} satır, önce {int(expected_before)}, sonra {int(expected_after)}, "
            f"{'kalın' if expected_bold else 'normal'}"
        )

        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u"
            })
            return (s or "").replace("\u00A0", " ").strip().lower().translate(trans)

        def resolve_from_styles_pf(para, attr):
            val = getattr(para.paragraph_format, attr)
            if val is not None:
                return val
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None:
                    v = getattr(pf, attr)
                    if v is not None:
                        return v
                s = getattr(s, "base_style", None)
            return None

        def effective_line_spacing(para, default=1.0):
            ls = resolve_from_styles_pf(para, "line_spacing")
            if ls is None:
                return float(default)
            if hasattr(ls, "pt"):
                return round(ls.pt, 1)
            try:
                return round(float(ls), 2)
            except Exception:
                return ls

        def effective_space_pt(para, which: str) -> float:
            attr = "space_before" if which == "before" else "space_after"
            length = resolve_from_styles_pf(para, attr)
            return 0.0 if length is None else round(length.pt, 1)

        def effective_alignment(para):
            if para.alignment is not None:
                return para.alignment
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None and pf.alignment is not None:
                    return pf.alignment
                s = getattr(s, "base_style", None)
            return None

        def effective_bold(para) -> bool:
            for r in para.runs:
                if r.bold is True or (r.font and r.font.bold is True):
                    return True
            s = para.style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.bold is True:
                    return True
                s = getattr(s, "base_style", None)
            return False

        def norm_font(n: str) -> str:
            return re.sub(r"\s+", "", (n or "").lower())

        def run_font_name(run):
            if run.font and run.font.name:
                return run.font.name
            rpr = getattr(run._element, "rPr", None)
            if rpr is not None and getattr(rpr, "rFonts", None) is not None:
                rf = rpr.rFonts
                for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                    val = getattr(rf, attr, None)
                    if val:
                        return val
            return None

        def style_font_name(style):
            s = style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.name:
                    return f.name
                el = getattr(s, "element", None)
                if el is not None:
                    rpr = getattr(el, "rPr", None)
                    if rpr is not None and getattr(rpr, "rFonts", None) is not None:
                        rf = rpr.rFonts
                        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                            val = getattr(rf, attr, None)
                            if val:
                                return val
                s = getattr(s, "base_style", None)
            return None

        def para_font_size_pt(para):
            for run in para.runs:
                if run.font and run.font.size:
                    return round(run.font.size.pt, 1)
            s = para.style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.size:
                    return round(f.size.pt, 1)
                s = getattr(s, "base_style", None)
            return None

        def is_blank(txt: str) -> bool:
            t = (txt or "").replace("\xa0", " ").replace("\t", " ")
            return t.strip() == ""

        # ======================================================
        # ✅ YENİ (SADECE FONT İÇİN): Theme + paragraf XML üzerinden "etkili font adı"
        # - Arama mantığını / diğer kontrolleri değiştirmez
        # - 'Yazı tipi tespit edilemedi' hatasını azaltır
        # ======================================================
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"

        def _get_theme_fonts_from_package():
            # memo cache
            if memo.get("_theme_fonts"):
                return memo["_theme_fonts"]

            minor_latin = None
            major_latin = None
            try:
                pkg = paragraphs[0]._parent.part.package
                theme_part = None
                for part in pkg.parts:
                    if str(part.partname).endswith("/word/theme/theme1.xml"):
                        theme_part = part
                        break

                if theme_part is not None:
                    root = ET.fromstring(theme_part.blob)
                    ns = {"a": A_NS}
                    minor = root.find(".//a:minorFont/a:latin", ns)
                    major = root.find(".//a:majorFont/a:latin", ns)
                    minor_latin = minor.get("typeface") if minor is not None else None
                    major_latin = major.get("typeface") if major is not None else None
            except Exception:
                pass

            memo["_theme_fonts"] = {"minor_latin": minor_latin, "major_latin": major_latin}
            return memo["_theme_fonts"]

        def _theme_to_font(theme_key: str, theme_fonts: dict):
            if not theme_key:
                return None
            k = theme_key.lower()
            if k.startswith("minor"):
                return theme_fonts.get("minor_latin")
            if k.startswith("major"):
                return theme_fonts.get("major_latin")
            return None

        def effective_font_name(para) -> str:
            theme_fonts = _get_theme_fonts_from_package()

            # 1) Paragraf XML içindeki rFonts (explicit veya theme) topla
            try:
                p_xml = ET.fromstring(para._element.xml)
                fonts = []

                for node in p_xml.iter():
                    if node.tag == f"{{{W_NS}}}rFonts":
                        # explicit
                        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
                            v = node.get(f"{{{W_NS}}}{attr}")
                            if v:
                                fonts.append(v)
                        # theme (Word çoğu zaman TOC'da bunu kullanır)
                        for attr in ("asciiTheme", "hAnsiTheme", "csTheme", "eastAsiaTheme", "cstheme"):
                            v = node.get(f"{{{W_NS}}}{attr}")
                            if v:
                                tv = _theme_to_font(v, theme_fonts)
                                if tv:
                                    fonts.append(tv)

                if fonts:
                    return Counter(fonts).most_common(1)[0][0]
            except Exception:
                pass

            # 2) Run -> rFonts (senin mevcut mantığın)
            run_names = []
            for r in para.runs:
                nm = run_font_name(r)
                if nm:
                    run_names.append(nm)
            if run_names:
                return Counter(run_names).most_common(1)[0][0]

            # 3) Stil zinciri
            sn = style_font_name(para.style)
            if sn:
                return sn

            # 4) Son çare: tema minor latin
            return theme_fonts.get("minor_latin") or None

        # 1) TOC başlığı memo’dan
        toc_idx = memo.get("toc_heading_idx", None)
        if toc_idx is None:
            return [(0, False, rule_title, "memo'da toc_heading_idx yok; önce 'toc_heading' çalışmalı")]

        # 2) 'Sayfa' satırını toc başlığından sonra, ilk 15 paragraf içinde ARA
        sayfa_idx = None
        scan_limit = min(len(paragraphs), toc_idx + 1 + 15)

        for i in range(toc_idx + 1, scan_limit):
            txt = (paragraphs[i].text or "").replace("\xa0", " ").strip()
            if is_blank(txt):
                continue

            core = re.sub(r"[^\w]+", "", norm_tr(txt))
            if core in ("sayfa", "page"):
                sayfa_idx = i
                break

            if ("\t" in (paragraphs[i].text or "")) and re.search(r"(\d+|[ivxlcdm]{1,8})\s*$", txt, flags=re.I):
                break

        if sayfa_idx is None:
            return [(toc_idx, False, rule_title, "İÇİNDEKİLER bölümünde 'Sayfa' satırı bulunamadı (başlıktan sonra)")]

        p3 = paragraphs[sayfa_idx]
        errors = []

        # Metin
        txt3 = (p3.text or "").replace("\xa0", " ").strip()

        core3 = re.sub(r"[^\w]+", "", norm_tr(txt3))
        if core3 not in ("sayfa", "page"):
            errors.append("'Sayfa' metni bekleniyordu")

        # ✅ Bulunduysa, biçimden bağımsız tam yazım (S büyük, ayfa küçük)
        only_letters = re.sub(r"[^A-Za-zÇĞİÖŞÜçğıöşü]+", "", txt3)
        if norm_tr(only_letters) == "sayfa" and only_letters != "Sayfa":
            errors.append("Metin tam olarak 'Sayfa' yazılmalı (S büyük, ayfa küçük)")

        # Hizalama
        eff_align = effective_alignment(p3)
        if expected_align == "right" and eff_align != WD_PARAGRAPH_ALIGNMENT.RIGHT:
            errors.append("Sağa dayalı değil")

        # Aralıklar
        eff_ls = effective_line_spacing(p3, default=expected_spacing)
        eff_sb = effective_space_pt(p3, "before")
        eff_sa = effective_space_pt(p3, "after")

        if isinstance(eff_ls, (int, float)) and abs(eff_ls - expected_spacing) > 0.1:
            errors.append(f"Satır aralığı {eff_ls} yerine {expected_spacing} olmalı")
        if abs(eff_sb - expected_before) > 1:
            errors.append(f"Paragraf öncesi {eff_sb} yerine {expected_before} olmalı")
        if abs(eff_sa - expected_after) > 1:
            errors.append(f"Paragraf sonrası {eff_sa} yerine {expected_after} olmalı")

        # ✅ FONT adı: artık theme + XML + stil zinciri destekli (sadece bu kısmı güçlendirdik)
        eff_name = effective_font_name(p3)
        fs = para_font_size_pt(p3)

        if eff_name is None:
            errors.append(f"Yazı tipi tespit edilemedi; {expected_name} olmalı")
        elif norm_font(eff_name) != norm_font(expected_name):
            errors.append(f"Yazı tipi {eff_name} yerine {expected_name} olmalı")

        if fs is not None and round(float(fs), 1) != round(expected_size, 1):
            errors.append(f"Punto {fs} yerine {expected_size} olmalı")

        if expected_bold and not effective_bold(p3):
            errors.append("Kalın değil")

        # memo
        memo["toc_sayfa_idx"] = sayfa_idx

        return [(sayfa_idx, len(errors) == 0, rule_title, "; ".join(errors))]


    # ======================================================
    # İÇİNDEKİLER – GÖVDE BİÇİMİ (Sayfa satırından sonra)
    # TNR 12pt, JUSTIFY, 1.5, 0/0 — Bitiş: 'ÖZGEÇMİŞ'
    # YALNIZCA memo["toc_heading_idx"] ve memo["toc_sayfa_idx"] kullanır.
    # Sibling taraması: <w:p>, <w:tbl> VE <w:sdt>
    # ======================================================
    elif check["check"] == "toc_table_body_format":
        import re, datetime
        from collections import Counter
        import xml.etree.ElementTree as ET

        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_align   = (check.get("alignment", "justify") or "justify").lower()
        expected_spacing = float(check.get("line_spacing", 1.5))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 0))
        debug_enabled    = bool(check.get("debug", False))
        DEBUG_F          = "debug_toc_table_end.txt"

        rule_title = (
            f"İÇİNDEKİLER – gövde biçimi: {expected_name}, {int(expected_size)} pt, "
            f"{'iki yana yaslı' if expected_align=='justify' else expected_align}, "
            f"{expected_spacing} satır, önce {int(expected_before)}, sonra {int(expected_after)} (Bitiş: ÖZGEÇMİŞ)"
        )

        # ---------- yardımcılar ----------
        def norm_tr(s: str) -> str:
            trans = str.maketrans({"ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g","ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u"})
            return (s or "").lower().translate(trans).strip()

        def resolve_from_styles_pf(para, attr):
            val = getattr(para.paragraph_format, attr)
            if val is not None: return val
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None:
                    v = getattr(pf, attr)
                    if v is not None: return v
                s = getattr(s, "base_style", None)
            return None

        def effective_alignment(para):
            if para.alignment is not None: return para.alignment
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None and pf.alignment is not None:
                    return pf.alignment
                s = getattr(s, "base_style", None)
            return None

        def effective_line_spacing(para, default=1.0):
            ls = resolve_from_styles_pf(para, "line_spacing")
            if ls is None: return float(default)
            if hasattr(ls, "pt"): return round(ls.pt, 1)
            try: return round(float(ls), 2)
            except Exception: return ls

        def effective_space_pt(para, which: str) -> float:
            length = resolve_from_styles_pf(para, "space_before" if which=="before" else "space_after")
            return 0.0 if length is None else round(length.pt, 1)

        # ========== THEME + XML tabanlı "effective font" ==========
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"

        def _get_theme_fonts_from_package():
            if memo.get("_theme_fonts"):
                return memo["_theme_fonts"]
            minor_latin = None
            major_latin = None
            try:
                pkg = paragraphs[0]._parent.part.package
                theme_part = None
                for part in pkg.parts:
                    if str(part.partname).endswith("/word/theme/theme1.xml"):
                        theme_part = part
                        break
                if theme_part is not None:
                    root = ET.fromstring(theme_part.blob)
                    ns = {"a": A_NS}
                    minor = root.find(".//a:minorFont/a:latin", ns)
                    major = root.find(".//a:majorFont/a:latin", ns)
                    minor_latin = minor.get("typeface") if minor is not None else None
                    major_latin = major.get("typeface") if major is not None else None
            except Exception:
                pass
            memo["_theme_fonts"] = {"minor_latin": minor_latin, "major_latin": major_latin}
            return memo["_theme_fonts"]

        def _theme_to_font(theme_key: str, theme_fonts: dict):
            if not theme_key:
                return None
            k = theme_key.lower()
            if k.startswith("minor"):
                return theme_fonts.get("minor_latin")
            if k.startswith("major"):
                return theme_fonts.get("major_latin")
            return None

        def effective_font_name(para) -> str:
            theme_fonts = _get_theme_fonts_from_package()

            # 1) Para XML içinde w:rFonts (explicit veya theme) ara
            try:
                p_xml = ET.fromstring(para._element.xml)
                fonts = []
                for node in p_xml.iter():
                    if node.tag == f"{{{W_NS}}}rFonts":
                        # explicit
                        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
                            v = node.get(f"{{{W_NS}}}{attr}")
                            if v:
                                fonts.append(v)
                        # theme
                        for attr in ("asciiTheme", "hAnsiTheme", "csTheme", "eastAsiaTheme", "cstheme"):
                            v = node.get(f"{{{W_NS}}}{attr}")
                            if v:
                                tv = _theme_to_font(v, theme_fonts)
                                if tv:
                                    fonts.append(tv)
                if fonts:
                    return Counter(fonts).most_common(1)[0][0]
            except Exception:
                pass

            # 2) Stil zinciri
            s = para.style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.name:
                    return f.name
                s = getattr(s, "base_style", None)

            # 3) Son çare: theme minor
            return theme_fonts.get("minor_latin") or None

        def para_font_size_pt(para):
            for run in para.runs:
                if run.font and run.font.size:
                    return round(run.font.size.pt, 1)
            s = para.style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.size:
                    return round(f.size.pt, 1)
                s = getattr(s, "base_style", None)
            return None

        def strip_toc_label(text: str) -> str:
            t = (text or "").replace("\xa0", " ").strip()
            t = re.split(r"\.{2,}|\t", t)[0]
            t = re.sub(r"\s*\d+$", "", t)
            t = re.sub(r"\s*[ivxlcdm]{2,8}$", "", t, flags=re.I)
            n = norm_tr(t)
            n = re.sub(r"[^\w ]+", "", n)
            n = re.sub(r"\s+", " ", n).strip()
            return n

        def p_text_from_p_el(p_el):
            texts = []
            for node in p_el.iter():
                if node.tag.endswith('}t') and node.text:
                    texts.append(node.text)
            return "".join(texts) if texts else ""

        # ---------- 1) Yalnızca MEMO kullan ----------
        toc_idx   = memo.get("toc_heading_idx")
        sayfa_idx = memo.get("toc_sayfa_idx")

        if debug_enabled:
            with open(DEBUG_F, "w", encoding="utf-8") as dbg:
                dbg.write(f"[TOC-END-DEBUG] {datetime.datetime.now():%Y-%m-%d %H:%M:%S}\n")
                dbg.write(f"memo.toc_heading_idx={toc_idx}, memo.toc_sayfa_idx={sayfa_idx}\n")

        if toc_idx is None:
            if debug_enabled:
                with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                    dbg.write("[ERR] memo'da toc_heading_idx yok (önce 'toc_heading' çalışmalı).\n")
            return [(0, False, rule_title, "İÇİNDEKİLER başlığı (memo) yok; önce 'toc_heading' çalışmalı")]

        # ---------- 2) Başlangıç: 'Sayfa' varsa ondan sonra; yoksa TOC başlığından sonra devam et ----------
        def is_sayfa_line(p):
            t = (p.text or "").replace("\xa0", " ").strip()
            t = re.sub(r"[^\w]+", "", t)      # noktalama vs temizle
            return norm_tr(t) == "sayfa"

        if sayfa_idx is not None:
            start = sayfa_idx + 1
            anchor_for_siblings = paragraphs[sayfa_idx]._element
        else:
            # 'Sayfa' satırı tespit edilemedi -> kontrol DURMASIN
            # TOC başlığından sonra ilk dolu paragrafı bul
            j = toc_idx + 1
            while j < len(paragraphs) and (paragraphs[j].text or "").replace("\xa0", " ").strip() == "":
                j += 1

            if j >= len(paragraphs):
                if debug_enabled:
                    with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                        dbg.write("[ERR] TOC başlığından sonra içerik yok.\n")
                return [(toc_idx, False, rule_title, "İÇİNDEKİLER gövdesi yok")]

            # Eğer ilk dolu satır aslında 'Sayfa' ise (memo yazılmamış olabilir) onu atla
            if is_sayfa_line(paragraphs[j]):
                start = j + 1
                anchor_for_siblings = paragraphs[j]._element
                # (İstersen burada memo'ya da yazabilirsin ama şart değil)
                # memo["toc_sayfa_idx"] = j
            else:
                start = j
                anchor_for_siblings = paragraphs[toc_idx]._element  # sibling taramasına TOC başlığından başla

        if start >= len(paragraphs):
            if debug_enabled:
                with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                    dbg.write(f"[ERR] start={start} paragrafları aşıyor; gövde yok.\n")
            return [(toc_idx, False, rule_title, "İÇİNDEKİLER gövdesi yok")]

        if debug_enabled:
            with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                dbg.write(f"[LOC] toc_idx={toc_idx}, sayfa_idx={sayfa_idx}\n")
                dbg.write(f"[LOC] start(paragraph idx)={start}\n")

        # ---------- 3) Sibling gezer: <w:p>, <w:tbl>, <w:sdt> ----------
        end_found_in_table = False
        end_found_in_sdt   = False
        end_idx = None

        cur = anchor_for_siblings.getnext()  # <-- ÖNEMLİ: artık sayfa yoksa toc'tan başlıyor

        while cur is not None:
            tag = cur.tag.split('}')[-1]

            if tag == "p":
                j = None
                for i_pp, pp in enumerate(paragraphs):
                    if pp._element is cur:
                        j = i_pp
                        break
                raw = (paragraphs[j].text or "") if j is not None else p_text_from_p_el(cur)
                lab = strip_toc_label(raw)

                if lab in {"ozgecmis", "oz gecmis"}:
                    end_idx = j
                    break

            elif tag == "tbl":
                p_els = [el for el in cur.iter() if el.tag.endswith('}p')]
                found_here = False
                for p_el in p_els:
                    raw = p_text_from_p_el(p_el)
                    lab = strip_toc_label(raw)
                    if lab in {"ozgecmis", "oz gecmis"}:
                        end_found_in_table = True
                        found_here = True
                        break
                if found_here:
                    break

            elif tag == "sdt":
                p_els = [el for el in cur.iter() if el.tag.endswith('}p')]
                found_here = False
                for p_el in p_els:
                    raw = p_text_from_p_el(p_el)
                    lab = strip_toc_label(raw)
                    if lab in {"ozgecmis", "oz gecmis"}:
                        end_found_in_sdt = True
                        found_here = True
                        break
                if found_here:
                    break

            cur = cur.getnext()

        if not (end_found_in_table or end_found_in_sdt or end_idx is not None):
            return [(start, False, rule_title, "Bitiş tespiti için 'ÖZGEÇMİŞ' bulunamadı")]

        # ---------- 4) Biçim doğrulama ----------
        errors = []
        checked = 0

        if end_idx is not None:
            for j in range(start, end_idx):
                
                loc = f"{j}. satır (belge:{j})"
                
                p = paragraphs[j]
                txt = (p.text or "").replace("\xa0", " ").strip()
                if txt == "":
                    pv = " ".join(txt.split()[:10])
                    if len(pv) > 90:
                        pv = pv[:90].rstrip() + "…"
                    loc = loc + (f" ('{pv}')" if pv else "")

                    continue

                eff_align = effective_alignment(p)
                eff_ls    = effective_line_spacing(p, default=expected_spacing)
                eff_sb    = effective_space_pt(p, "before")
                eff_sa    = effective_space_pt(p, "after")

                # hizalama
                if expected_align == "justify" and eff_align != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                    errors.append(f"{j}: Paragraf iki yana yaslı değil")

                # spacing
                if isinstance(eff_ls,(int,float)) and abs(eff_ls - expected_spacing) > 0.1:
                    errors.append(f"{loc}: Satır aralığı {eff_ls} yerine {expected_spacing} olmalı")
                if abs(eff_sb - expected_before) > 1:
                    errors.append(f"{loc}: Öncesi {eff_sb} yerine {expected_before} olmalı")
                if abs(eff_sa - expected_after) > 1:
                    errors.append(f"{loc}: Sonrası {eff_sa} yerine {expected_after} olmalı")

                # FONT: artık theme + hyperlink/field uyumlu
                fn_eff = effective_font_name(p)
                if fn_eff is None:
                    errors.append(f"{loc}: Yazı tipi tespit edilemedi; {expected_name} olmalı")
                else:
                    if re.sub(r"\s+","",fn_eff.lower()) != re.sub(r"\s+","",expected_name.lower()):
                        errors.append(f"{j}: Yazı tipi {fn_eff} yerine {expected_name} olmalı")

                fs = para_font_size_pt(p)
                if fs is not None and round(float(fs),1) != round(expected_size,1):
                    errors.append(f"{loc}: Punto {fs} yerine {expected_size} olmalı")

                checked += 1

        ok = (len(errors) == 0)
        return [(start, ok, rule_title, "; ".join(errors))]

    # ===============================================================================================================#
    # ===============================================================================================================#    
    # ===============================================================================================================#
    # ======================================================
    # ÇİZELGELER DİZİNİ - BAŞLIK (İÇİNDEKİLER'DEN SONRA, BİÇİM AYNI)
    # ======================================================
    elif check["check"] == "list_of_tables_heading":
        import re

        expected_name     = check.get("font_name", "Times New Roman")
        expected_size     = float(check.get("font_size_pt", 12))
        expected_bold     = bool(check.get("bold", True))
        expected_all_caps = bool(check.get("all_caps", True))
        expected_align    = check.get("alignment", "center").lower()
        expected_spacing  = float(check.get("line_spacing", 1.5))
        expected_before   = float(check.get("space_before", 0))
        expected_after    = float(check.get("space_after", 24))
        markers           = check.get("markers", ["^ÇİZELGELER DİZİNİ$"])
        must_exist        = check.get("must_exist", False)

        rule_title = (
            f"ÇİZELGELER DİZİNİ başlığı: {expected_name}, {int(expected_size)} punto, "
            f"{'kalın' if expected_bold else 'normal'}, "
            f"{'BÜYÜK HARF' if expected_all_caps else 'normal'}, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u"
            })
            return (s or "").lower().translate(trans).strip()

        def resolve_from_styles(para, attr_name):
            val = getattr(para.paragraph_format, attr_name)
            if val is not None:
                return val
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None:
                    v = getattr(pf, attr_name)
                    if v is not None:
                        return v
                s = getattr(s, "base_style", None)
            return None

        def effective_space_pt(para, which: str) -> float:
            attr = "space_before" if which == "before" else "space_after"
            length = resolve_from_styles(para, attr)
            return 0.0 if length is None else round(length.pt, 1)

        def effective_line_spacing(para, default=1.0) -> float:
            ls = resolve_from_styles(para, "line_spacing")
            if ls is None:
                return float(default)
            if hasattr(ls, "pt"):
                return round(ls.pt, 1)
            try:
                return round(float(ls), 2)
            except Exception:
                return ls

        def effective_alignment(para):
            if para.alignment is not None:
                return para.alignment
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None and pf.alignment is not None:
                    return pf.alignment
                s = getattr(s, "base_style", None)
            return None

        def effective_bold(para):
            for r in para.runs:
                if r.bold is True or (r.font and r.font.bold is True):
                    return True
            s = para.style
            while s is not None:
                if getattr(s, "font", None) and s.font.bold is True:
                    return True
                s = getattr(s, "base_style", None)
            return False

        # 1) İÇİNDEKİLER başlığı referansı
        toc_idx = memo.get("toc_heading_idx")
        if toc_idx is None:
            return [(0, False, rule_title, "İÇİNDEKİLER başlığı bulunamadı; ÇİZELGELER DİZİNİ konum kontrolü yapılamadı")]

        # 2) ÇİZELGELER DİZİNİ başlığını İÇİNDEKİLER’den sonra ara
        compiled = []
        for m in markers:
            try:
                compiled.append(re.compile(norm_tr(m), re.IGNORECASE))
            except re.error:
                compiled.append(re.compile("^" + re.escape(norm_tr(m).strip("^$")) + "$", re.IGNORECASE))

        found = None
        for i in range(toc_idx + 1, len(paragraphs)):
            raw = (paragraphs[i].text or "").strip()
            norm = norm_tr(raw)
            for pat in compiled:
                if pat.match(norm):
                    found = (i, paragraphs[i])
                    break
            if found:
                break

        # 2.a) Bulunamadı → zorunlu/isteğe bağlı ayrımı
        if not found:
            if must_exist:
                return [(toc_idx, False, rule_title, "Zorunlu sayfa bulunamadı (İÇİNDEKİLER’den sonra bekleniyordu).")]
            else:
                return [(toc_idx, False, rule_title, "Bu sayfa isteğe bağlı ve tezde bulunmadı.")]

        # 3) Biçim kontrolleri
        idx, p = found
        errors = []

        fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
        fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
        if fn and fn != expected_name:
            errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")
        if fs and round(float(fs), 1) != round(expected_size, 1):
            errors.append(f"Punto {fs} yerine {expected_size} olmalı")

        if expected_bold and not effective_bold(p):
            errors.append("Başlık kalın değil")
        if expected_all_caps and not p.text.isupper():
            errors.append("Başlık BÜYÜK HARF değil")

        eff_align = effective_alignment(p)
        eff_ls   = effective_line_spacing(p)
        eff_sb   = effective_space_pt(p, "before")
        eff_sa   = effective_space_pt(p, "after")

        if expected_align == "center" and eff_align != WD_PARAGRAPH_ALIGNMENT.CENTER:
            errors.append("Başlık ortalı değil")
        if isinstance(eff_ls, (int, float)) and abs(eff_ls - expected_spacing) > 0.1:
            errors.append(f"Satır aralığı {eff_ls} yerine {expected_spacing} olmalı")
        if abs(eff_sb - expected_before) > 1:
            errors.append(f"Paragraf öncesi {eff_sb} yerine {expected_before} olmalı")
        if abs(eff_sa - expected_after) > 1:
            errors.append(f"Paragraf sonrası {eff_sa} yerine {expected_after} olmalı")

        if errors:
            result = [(idx, False, rule_title, "; ".join(errors))]
        else:
            result = [(idx, True, rule_title, "")]

        memo["list_of_tables_idx"] = idx  # başlık paragraf indeksi
        return result

    # ======================================================
    # ÇİZELGELER DİZİNİ – 2. SATIR 'Sayfa'
    # (TNR 12pt, RIGHT, 1.0, 0/0, BOLD; başlıktan hemen sonra, boş satır yok)
    # ======================================================
    elif check["check"] == "list_of_tables_page_label_line":
        import re
        from collections import Counter

        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_align   = check.get("alignment", "right").lower()
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 0))
        expected_bold    = bool(check.get("bold", True))
        must_exist       = check.get("must_exist", False)  # YAML'a ekleyebilirsin

        rule_title = (
            f"ÇİZELGELER DİZİNİ – 2. satır 'Sayfa' satırı: {expected_name}, {int(expected_size)} pt, "
            f"{'sağa dayalı' if expected_align=='right' else expected_align}, "
            f"{expected_spacing} satır, önce {int(expected_before)}, sonra {int(expected_after)}, "
            f"{'kalın' if expected_bold else 'normal'}"
        )

        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u"
            })
            return (s or "").lower().translate(trans).strip()

        def resolve_from_styles_pf(para, attr):
            val = getattr(para.paragraph_format, attr)
            if val is not None:
                return val
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None:
                    v = getattr(pf, attr)
                    if v is not None:
                        return v
                s = getattr(s, "base_style", None)
            return None

        def effective_line_spacing(para, default=1.0):
            ls = resolve_from_styles_pf(para, "line_spacing")
            if ls is None:
                return float(default)
            if hasattr(ls, "pt"):
                return round(ls.pt, 1)
            try:
                return round(float(ls), 2)
            except Exception:
                return ls

        def effective_space_pt(para, which: str) -> float:
            attr = "space_before" if which == "before" else "space_after"
            length = resolve_from_styles_pf(para, attr)
            return 0.0 if length is None else round(length.pt, 1)

        def effective_alignment(para):
            if para.alignment is not None:
                return para.alignment
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None and pf.alignment is not None:
                    return pf.alignment
                s = getattr(s, "base_style", None)
            return None

        def effective_bold(para):
            for r in para.runs:
                if r.bold is True or (r.font and r.font.bold is True):
                    return True
            s = para.style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.bold is True:
                    return True
                s = getattr(s, "base_style", None)
            return False

        def run_font_name(run):
            if run.font and run.font.name:
                return run.font.name
            rpr = getattr(run._element, "rPr", None)
            if rpr is not None and getattr(rpr, "rFonts", None) is not None:
                rf = rpr.rFonts
                for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                    val = getattr(rf, attr, None)
                    if val:
                        return val
            return None

        def style_font_name(style):
            s = style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.name:
                    return f.name
                el = getattr(s, "element", None)
                if el is not None:
                    rpr = getattr(el, "rPr", None)
                    if rpr is not None and getattr(rpr, "rFonts", None) is not None:
                        rf = rpr.rFonts
                        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                            val = getattr(rf, attr, None)
                            if val:
                                return val
                s = getattr(s, "base_style", None)
            return None

        def para_font_size_pt(para):
            for run in para.runs:
                if run.font and run.font.size:
                    return round(run.font.size.pt, 1)
            s = para.style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.size:
                    return round(f.size.pt, 1)
                s = getattr(s, "base_style", None)
            return None

        # 1) Başlık indeksi yoksa, bu kontrolü atla (isteğe bağlı sayfa)
        lot_idx = memo.get("list_of_tables_idx")
        if lot_idx is None:
            if must_exist:
                return [(0, False, rule_title, "ÇİZELGELER DİZİNİ başlığı bulunamadı; 'Sayfa' konumu doğrulanamadı.")]
            else:
                # Atlandı: aynı eksik için ikinci kez kırmızı göstermeyelim.
                return [(0, True, rule_title, "Atlandı: Çizelgeler Dizini isteğe bağlı ve tezde bulunmadı.")]

        # 2) Hemen sonraki paragraf 'Sayfa' olmalı (boş satır olmadan)
        sayfa_idx = lot_idx + 1
        if sayfa_idx >= len(paragraphs):
            return [(lot_idx, False, rule_title, "2. satır (Sayfa) bulunamadı")]

        p2 = paragraphs[sayfa_idx]
        text2 = (p2.text or "").replace("\xa0", " ").strip()
        errors = []

        # Metin kontrolü
        if norm_tr(re.sub(r"[^\w]+", "", text2)) != "sayfa":
            errors.append("2. satır metni 'Sayfa' olmalı")

        # Hizalama: RIGHT
        eff_align = effective_alignment(p2)
        if expected_align == "right" and eff_align != WD_PARAGRAPH_ALIGNMENT.RIGHT:
            errors.append("2. satır sağa dayalı değil")

        # Satır aralığı ve paragraf boşlukları
        eff_ls = effective_line_spacing(p2, default=1.0)
        eff_sb = effective_space_pt(p2, "before")
        eff_sa = effective_space_pt(p2, "after")
        if isinstance(eff_ls, (int, float)) and abs(eff_ls - expected_spacing) > 0.1:
            errors.append(f"Satır aralığı {eff_ls} yerine {expected_spacing} olmalı")
        if abs(eff_sb - expected_before) > 1:
            errors.append(f"Paragraf öncesi {eff_sb} yerine {expected_before} olmalı")
        if abs(eff_sa - expected_after) > 1:
            errors.append(f"Paragraf sonrası {eff_sa} yerine {expected_after} olmalı")

        # Font adı ve punto
        run_names = [run_font_name(r) for r in p2.runs if run_font_name(r)]
        eff_name = Counter(run_names).most_common(1)[0][0] if run_names else style_font_name(p2.style)
        fs = para_font_size_pt(p2)

        def norm_font(n: str) -> str:
            return re.sub(r"\s+", "", (n or "").lower())

        expected_norm = norm_font(expected_name)
        actual_norm   = norm_font(eff_name) if eff_name else None

        if actual_norm is None:
            errors.append(f"Yazı tipi tespit edilemedi; {expected_name} olmalı")
        elif actual_norm != expected_norm:
            errors.append(f"Yazı tipi {eff_name} yerine {expected_name} olmalı")

        if fs is not None and round(float(fs), 1) != round(expected_size, 1):
            errors.append(f"Punto {fs} yerine {expected_size} olmalı")

        if expected_bold and not effective_bold(p2):
            errors.append("2. satır kalın değil")

        memo["list_of_tables_sayfa_idx"] = sayfa_idx
        return [(sayfa_idx, len(errors) == 0, rule_title, "; ".join(errors))]

    # ======================================================
    # ÇİZELGELER DİZİNİ – GÖVDE (Sayfa satırından sonra tablo)
    # TNR 12pt, LEFT, 1.5, 0/0 — Bitiş: sonraki başlık veya ilk ORTALI paragraf
    # Sibling taraması: <w:p>, <w:tbl>, <w:sdt>
    # Debug log: debug_lot_table_end.txt
    # ======================================================
    elif check["check"] == "list_of_tables_body_format":
        import re, datetime
        from collections import Counter
        

        # --------------- Parametreler ---------------
        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_align   = check.get("alignment", "left").lower()
        expected_spacing = float(check.get("line_spacing", 1.5))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 0))
        must_exist       = check.get("must_exist", False)
        debug_enabled    = bool(check.get("debug", False))
        DEBUG_F          = dbg_path("debug_lot_table_end.txt")

        rule_title = (
            f"ÇİZELGELER DİZİNİ – gövde biçimi: {expected_name}, {int(expected_size)} pt, "
            f"{'sola yaslı' if expected_align=='left' else expected_align}, "
            f"{expected_spacing} satır, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        # --------------- Yardımcılar ---------------
        def norm_tr(s: str) -> str:
            trans = str.maketrans({"ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g","ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u"})
            return (s or "").lower().translate(trans).strip()

        def strip_label_like_toc(text: str) -> str:
            """TOC satırı gibi: dotted leader/tab öncesini al, sonda sayfa no/roma rakamlarını sök, normalize et."""
            t = (text or "").replace("\xa0", " ").strip()
            t = re.split(r"\.{2,}|\t", t)[0]                 # dotted leader / tab öncesi
            t = re.sub(r"\s*\d+$", "", t)                    # sonda sayı
            t = re.sub(r"\s*[ivxlcdm]{2,8}$", "", t, flags=re.I)  # sonda roma
            n = norm_tr(t)
            n = re.sub(r"[^\w ]+", "", n)
            n = re.sub(r"\s+", " ", n).strip()
            return n

        def resolve_from_styles_pf(para, attr):
            val = getattr(para.paragraph_format, attr)
            if val is not None: return val
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None:
                    v = getattr(pf, attr)
                    if v is not None: return v
                s = getattr(s, "base_style", None)
            return None

        def effective_alignment(para):
            if para.alignment is not None: return para.alignment
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf and pf.alignment is not None:
                    return pf.alignment
                s = getattr(s, "base_style", None)
            return None

        def effective_line_spacing(para, default=1.0):
            ls = resolve_from_styles_pf(para, "line_spacing")
            if ls is None: return float(default)
            if hasattr(ls, "pt"): return round(ls.pt, 1)
            try: return round(float(ls), 2)
            except Exception: return ls

        def effective_space_pt(para, which: str) -> float:
            length = resolve_from_styles_pf(para, "space_before" if which=="before" else "space_after")
            return 0.0 if length is None else round(length.pt, 1)

        def norm_font(n: str) -> str:
            return re.sub(r"\s+", "", (n or "").lower())

        def run_font_name(run):
            if run.font and run.font.name:
                return run.font.name
            rpr = getattr(run._element, "rPr", None)
            if rpr is not None and getattr(rpr, "rFonts", None) is not None:
                rf = rpr.rFonts
                for attr in ("ascii","hAnsi","eastAsia","cs"):
                    val = getattr(rf, attr, None)
                    if val: return val
            return None

        def style_font_name(style):
            s = style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.name: return f.name
                el = getattr(s, "element", None)
                if el is not None:
                    rpr = getattr(el, "rPr", None)
                    if rpr is not None and getattr(rpr, "rFonts", None) is not None:
                        rf = rpr.rFonts
                        for attr in ("ascii","hAnsi","eastAsia","cs"):
                            val = getattr(rf, attr, None)
                            if val: return val
                s = getattr(s, "base_style", None)
            return None

        def para_font_size_pt(para):
            for run in para.runs:
                if run.font and run.font.size:
                    return round(run.font.size.pt, 1)
            s = para.style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.size:
                    return round(f.size.pt, 1)
                s = getattr(s, "base_style", None)
            return None

        def p_text_from_p_el(p_el):
            """Namespace fark etmeksizin <w:p> içindeki tüm <w:t> metinlerini birleştir."""
            texts = []
            for node in p_el.iter():
                if node.tag.endswith('}t') and node.text:
                    texts.append(node.text)
            return "".join(texts) if texts else ""

        def find_para_index_by_element(p_el):
            """Verilen XML <w:p> elementini paragraphs listesinde bulup indeksini döndürür (yoksa None)."""
            for i_pp, pp in enumerate(paragraphs):
                if pp._element is p_el:
                    return i_pp
            return None

        # XML'den hizalama okumak için (fallback)
        def align_from_p_el(p_el):
            """<w:pPr><w:jc w:val='left|right|center|both'> üzerinden hizalama yakalamaya çalış."""
            try:
                pPr = getattr(p_el, "pPr", None)
                if pPr is not None and getattr(pPr, "jc", None) is not None:
                    val = pPr.jc.val
                    if not val: return None
                    val = str(val).lower()
                    if val == "left":   return WD_PARAGRAPH_ALIGNMENT.LEFT
                    if val == "right":  return WD_PARAGRAPH_ALIGNMENT.RIGHT
                    if val == "center": return WD_PARAGRAPH_ALIGNMENT.CENTER
                    if val == "both":   return WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except Exception:
                pass
            return None

        # --------------- 1) Başlangıç ---------------
        sayfa_idx = memo.get("list_of_tables_sayfa_idx")
        if sayfa_idx is None:
            if must_exist:
                return [(0, False, rule_title, "ÇİZELGELER DİZİNİ bulunamadı (zorunlu sayfa).")]
            else:
                return [(0, True, rule_title, "Atlandı: Çizelgeler Dizini isteğe bağlı ve tezde bulunmadı.")]

        start_para = paragraphs[sayfa_idx]
        cur = start_para._element.getnext()  # 'Sayfa'nın hemen sonraki sibling

        if debug_enabled:
            with open(DEBUG_F, "w", encoding="utf-8") as dbg:
                dbg.write(f"[LOT-TABLE-DEBUG] {datetime.datetime.now():%Y-%m-%d %H:%M:%S}\n")
                dbg.write(f"memo.list_of_tables_sayfa_idx={sayfa_idx}\n")
                dbg.write("[SCAN] Kardeş (sibling) bloklar taranıyor...\n")

        # --------------- 2) Bitişi bul (başlık/ortalanmış paragraf) ---------------
        body_p_indices = []  # biçim denetimi yapılacak gerçek paragraph indeksleri
        end_idx = None
        end_found_reason = None

        while cur is not None:
            tag = cur.tag.split('}')[-1]  # 'p', 'tbl', 'sdt' vb.

            if tag == "p":
                j = find_para_index_by_element(cur)
                raw = p_text_from_p_el(cur) if j is None else (paragraphs[j].text or "")
                nrm = strip_label_like_toc(raw)

                # Debug satırı (P)
                if debug_enabled:
                    preview = raw.replace("\n", " ").strip()
                    if len(preview) > 140: preview = preview[:140] + "…"
                    if j is not None:
                        eff_align_dbg = effective_alignment(paragraphs[j])
                        if eff_align_dbg is None:
                            eff_align_dbg = align_from_p_el(cur)
                    else:
                        eff_align_dbg = align_from_p_el(cur)
                    align_str = str(eff_align_dbg).replace("WD_PARAGRAPH_ALIGNMENT.", "") if eff_align_dbg is not None else "None"
                    with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                        dbg.write(f"  [P] idx={j if j is not None else '-'}  label={nrm!r}  align={align_str}  raw={preview!r}\n")

                # Bitiş koşulu: bilinen başlıklar veya ilk ORTALI paragraf
                if nrm in {"sekiller dizini", "simge dizini", "kisaltmalar", "bolum 1", "giris", "girıs"}:
                    end_idx = j if j is not None else 0
                    end_found_reason = f"Başlık: {nrm}"
                    break
                if j is not None:
                    eff_align = effective_alignment(paragraphs[j])
                else:
                    eff_align = align_from_p_el(cur)
                if eff_align == WD_PARAGRAPH_ALIGNMENT.CENTER and raw.strip():
                    end_idx = j if j is not None else 0
                    end_found_reason = "Ortalanmış paragraf"
                    break

                if j is not None:
                    body_p_indices.append(j)

            elif tag == "tbl":
                # Tablo içindeki tüm <w:p> düğümlerini sırayla işle
                p_els = [el for el in cur.iter() if el.tag.endswith('}p')]
                if debug_enabled and not p_els:
                    with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                        dbg.write("  [TBL] (paragraf içermiyor)\n")
                for k, p_el in enumerate(p_els):
                    j = find_para_index_by_element(p_el)
                    raw = p_text_from_p_el(p_el)
                    nrm = strip_label_like_toc(raw)

                    if j is not None:
                        eff_align_dbg = effective_alignment(paragraphs[j])
                        if eff_align_dbg is None:
                            eff_align_dbg = align_from_p_el(p_el)
                    else:
                        eff_align_dbg = align_from_p_el(p_el)
                    align_str = str(eff_align_dbg).replace("WD_PARAGRAPH_ALIGNMENT.", "") if eff_align_dbg is not None else "None"

                    if debug_enabled:
                        preview = raw.replace("\n", " ").strip()
                        if len(preview) > 140: preview = preview[:140] + "…"
                        with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                            dbg.write(f"  [TBL] p={k} idx={j if j is not None else '-'} label={nrm!r} align={align_str} raw={preview!r}\n")

                    if nrm in {"sekiller dizini", "simge dizini", "kisaltmalar", "bolum 1", "giris", "girıs"}:
                        end_idx = j if j is not None else 0
                        end_found_reason = f"TBL Başlık: {nrm}"
                        break
                    if eff_align_dbg == WD_PARAGRAPH_ALIGNMENT.CENTER and raw.strip():
                        end_idx = j if j is not None else 0
                        end_found_reason = "TBL Ortalanmış paragraf"
                        break

                    if j is not None:
                        body_p_indices.append(j)

                if end_idx is not None:
                    break

            elif tag == "sdt":
                # SDT (content control) içindeki <w:p> düğümlerini sırayla işle
                p_els = [el for el in cur.iter() if el.tag.endswith('}p')]
                if debug_enabled and not p_els:
                    with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                        dbg.write("  [SDT] (paragraf içermiyor)\n")
                for k, p_el in enumerate(p_els):
                    j = find_para_index_by_element(p_el)
                    raw = p_text_from_p_el(p_el)
                    nrm = strip_label_like_toc(raw)

                    if j is not None:
                        eff_align_dbg = effective_alignment(paragraphs[j])
                        if eff_align_dbg is None:
                            eff_align_dbg = align_from_p_el(p_el)
                    else:
                        eff_align_dbg = align_from_p_el(p_el)
                    align_str = str(eff_align_dbg).replace("WD_PARAGRAPH_ALIGNMENT.", "") if eff_align_dbg is not None else "None"

                    if debug_enabled:
                        preview = raw.replace("\n", " ").strip()
                        if len(preview) > 140: preview = preview[:140] + "…"
                        with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                            dbg.write(f"  [SDT] p={k} idx={j if j is not None else '-'} label={nrm!r} align={align_str} raw={preview!r}\n")

                    if nrm in {"sekiller dizini", "simge dizini", "kisaltmalar", "bolum 1", "giris", "girıs"}:
                        end_idx = j if j is not None else 0
                        end_found_reason = f"SDT Başlık: {nrm}"
                        break
                    if eff_align_dbg == WD_PARAGRAPH_ALIGNMENT.CENTER and raw.strip():
                        end_idx = j if j is not None else 0
                        end_found_reason = "SDT Ortalanmış paragraf"
                        break

                    if j is not None:
                        body_p_indices.append(j)

                if end_idx is not None:
                    break

            # sonraki kardeşe geç
            cur = cur.getnext()

        if end_idx is None:
            end_idx = len(paragraphs)
            end_found_reason = "Belge sonu"

        if debug_enabled:
            with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                dbg.write(f"[END] reason={end_found_reason}, end_idx={end_idx}\n")
                dbg.write(f"[BODY] denetlenecek paragraph indexleri: {body_p_indices}\n")

        # --------------- 3) Biçim Denetimi ---------------
        errors = []
        checked = 0

        for j in body_p_indices:
            loc = f"{j}. satır (belge:{j})"

            if j is None or j < 0 or j >= len(paragraphs):
                continue
            p = paragraphs[j]
            txt = (p.text or "").replace("\xa0", " ").strip()
            if txt == "":
                pv = " ".join(txt.split()[:10])
                if len(pv) > 90:
                    pv = pv[:90].rstrip() + "…"
                loc = loc + (f" ('{pv}')" if pv else "")

                continue

            # --- HİZALAMA: çoklu kaynak + LOT için varsayılan LEFT ---
            eff_align = effective_alignment(p)
            if eff_align is None:
                eff_align = align_from_p_el(p._element)
            if eff_align is None:
                eff_align = WD_PARAGRAPH_ALIGNMENT.LEFT  # <<< kritik varsayılan

            eff_ls    = effective_line_spacing(p, default=1.5)
            eff_sb    = effective_space_pt(p, "before")
            eff_sa    = effective_space_pt(p, "after")

            run_names = [run_font_name(r) for r in p.runs if run_font_name(r)]
            eff_name  = Counter(run_names).most_common(1)[0][0] if run_names else style_font_name(p.style)
            fs        = para_font_size_pt(p)

            # --- Biçimsel kurallar ---
            if expected_align == "left" and eff_align != WD_PARAGRAPH_ALIGNMENT.LEFT:
                errors.append(f"{loc}: Paragraf sola yaslı değil")
            if isinstance(eff_ls, (int, float)) and abs(eff_ls - expected_spacing) > 0.1:
                errors.append(f"{loc}: Satır aralığı {eff_ls} yerine {expected_spacing} olmalı")
            if abs(eff_sb - expected_before) > 1:
                errors.append(f"{loc}: Öncesi {eff_sb} yerine {expected_before} olmalı")
            if abs(eff_sa - expected_after) > 1:
                errors.append(f"{loc}: Sonrası {eff_sa} yerine {expected_after} olmalı")

            exp_norm = norm_font(expected_name)
            act_norm = norm_font(eff_name) if eff_name else None
            if act_norm is None:
                errors.append(f"{loc}: Yazı tipi tespit edilemedi; {expected_name} olmalı")
            elif act_norm != exp_norm:
                errors.append(f"{loc}: Yazı tipi {eff_name} yerine {expected_name} olmalı")
            if fs is not None and round(float(fs),1) != round(expected_size,1):
                errors.append(f"{loc}: Punto {fs} yerine {expected_size} olmalı")

            # --- debug satırı (nihai align ile) ---
            if debug_enabled:
                preview = txt.replace("\n", " ")
                if len(preview) > 120:
                    preview = preview[:120] + "…"
                align_str = str(eff_align).replace("WD_PARAGRAPH_ALIGNMENT.", "")
                with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                    dbg.write(f"[P{j}] Align={align_str:8s}  LS={eff_ls:<4}  SB={eff_sb:<4}  SA={eff_sa:<4}  Font={eff_name or '?'}  Txt={preview!r}\n")

            checked += 1

        ok = (len(errors) == 0)
        if debug_enabled:
            with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                dbg.write(f"[DONE] checked={checked}, errors={len(errors)}\n")

        # Denetim başlangıcı: 'Sayfa' sonrası ilk içerik paragrafı (veya body_p_indices ilk eleman)
        begin_idx = body_p_indices[0] if body_p_indices else (sayfa_idx + 1)
        return [(begin_idx, ok, rule_title, "; ".join(errors))]


    # ===============================================================================================================#
    # ===============================================================================================================#    
    # ===============================================================================================================#
    # ======================================================
    # ŞEKİLLER DİZİNİ - BAŞLIK (İÇİNDEKİLER'DEN SONRA, BİÇİM AYNI)
    # ======================================================
    elif check["check"] == "list_of_figures_heading":
        import re

        expected_name     = check.get("font_name", "Times New Roman")
        expected_size     = float(check.get("font_size_pt", 12))
        expected_bold     = bool(check.get("bold", True))
        expected_all_caps = bool(check.get("all_caps", True))
        expected_align    = check.get("alignment", "center").lower()
        expected_spacing  = float(check.get("line_spacing", 1.5))
        expected_before   = float(check.get("space_before", 0))
        expected_after    = float(check.get("space_after", 24))
        markers           = check.get("markers", ["^ŞEKİLLER DİZİNİ$"])
        must_exist        = check.get("must_exist", False)

        rule_title = (
            f"ŞEKİLLER DİZİNİ başlığı: {expected_name}, {int(expected_size)} punto, "
            f"{'kalın' if expected_bold else 'normal'}, "
            f"{'BÜYÜK HARF' if expected_all_caps else 'normal'}, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u"
            })
            return (s or "").lower().translate(trans).strip()

        def resolve_from_styles(para, attr_name):
            val = getattr(para.paragraph_format, attr_name)
            if val is not None:
                return val
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None:
                    v = getattr(pf, attr_name)
                    if v is not None:
                        return v
                s = getattr(s, "base_style", None)
            return None

        def effective_space_pt(para, which: str) -> float:
            attr = "space_before" if which == "before" else "space_after"
            length = resolve_from_styles(para, attr)
            return 0.0 if length is None else round(length.pt, 1)

        def effective_line_spacing(para, default=1.0) -> float:
            ls = resolve_from_styles(para, "line_spacing")
            if ls is None:
                return float(default)
            if hasattr(ls, "pt"):
                return round(ls.pt, 1)
            try:
                return round(float(ls), 2)
            except Exception:
                return ls

        def effective_alignment(para):
            if para.alignment is not None:
                return para.alignment
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None and pf.alignment is not None:
                    return pf.alignment
                s = getattr(s, "base_style", None)
            return None

        def effective_bold(para):
            for r in para.runs:
                if r.bold is True or (r.font and r.font.bold is True):
                    return True
            s = para.style
            while s is not None:
                if getattr(s, "font", None) and s.font.bold is True:
                    return True
                s = getattr(s, "base_style", None)
            return False

        # 1) İÇİNDEKİLER başlığı referansı
        toc_idx = memo.get("toc_heading_idx")
        if toc_idx is None:
            return [(0, False, rule_title, "İÇİNDEKİLER başlığı bulunamadı; ŞEKİLLER DİZİNİ konum kontrolü yapılamadı")]

        # 2) ŞEKİLLER DİZİNİ başlığını İÇİNDEKİLER’den sonra ara
        compiled = []
        for m in markers:
            try:
                compiled.append(re.compile(norm_tr(m), re.IGNORECASE))
            except re.error:
                compiled.append(re.compile("^" + re.escape(norm_tr(m).strip("^$")) + "$", re.IGNORECASE))

        found = None
        for i in range(toc_idx + 1, len(paragraphs)):
            raw = (paragraphs[i].text or "").strip()
            norm = norm_tr(raw)
            for pat in compiled:
                if pat.match(norm):
                    found = (i, paragraphs[i])
                    break
            if found:
                break

        # 2.a) Bulunamadı → zorunlu/isteğe bağlı ayrımı
        if not found:
            if must_exist:
                return [(toc_idx, False, rule_title, "Zorunlu sayfa bulunamadı (İÇİNDEKİLER’den sonra bekleniyordu).")]
            else:
                return [(toc_idx, False, rule_title, "Bu sayfa isteğe bağlı ve tezde bulunmadı.")]

        # 3) Biçim kontrolleri
        idx, p = found
        errors = []

        fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
        fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
        if fn and fn != expected_name:
            errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")
        if fs and round(float(fs), 1) != round(expected_size, 1):
            errors.append(f"Punto {fs} yerine {expected_size} olmalı")

        if expected_bold and not effective_bold(p):
            errors.append("Başlık kalın değil")
        if expected_all_caps and not p.text.isupper():
            errors.append("Başlık BÜYÜK HARF değil")

        eff_align = effective_alignment(p)
        eff_ls   = effective_line_spacing(p)
        eff_sb   = effective_space_pt(p, "before")
        eff_sa   = effective_space_pt(p, "after")

        if expected_align == "center" and eff_align != WD_PARAGRAPH_ALIGNMENT.CENTER:
            errors.append("Başlık ortalı değil")
        if isinstance(eff_ls, (int, float)) and abs(eff_ls - expected_spacing) > 0.1:
            errors.append(f"Satır aralığı {eff_ls} yerine {expected_spacing} olmalı")
        if abs(eff_sb - expected_before) > 1:
            errors.append(f"Paragraf öncesi {eff_sb} yerine {expected_before} olmalı")
        if abs(eff_sa - expected_after) > 1:
            errors.append(f"Paragraf sonrası {eff_sa} yerine {expected_after} olmalı")

        memo["list_of_figures_idx"] = idx  # başlık paragraf indeksi
        return [(idx, len(errors) == 0, rule_title, "; ".join(errors))]

    # ======================================================
    # ŞEKİLLER DİZİNİ – 2. SATIR 'Sayfa'
    # (TNR 12pt, RIGHT, 1.0, 0/0, BOLD; başlıktan hemen sonra, boş satır yok)
    # ======================================================
    elif check["check"] == "list_of_figures_page_label_line":
        import re
        from collections import Counter

        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_align   = check.get("alignment", "right").lower()
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 0))
        expected_bold    = bool(check.get("bold", True))
        must_exist       = check.get("must_exist", False)

        rule_title = (
            f"ŞEKİLLER DİZİNİ – 2. satır 'Sayfa' satırı: {expected_name}, {int(expected_size)} pt, "
            f"{'sağa dayalı' if expected_align=='right' else expected_align}, "
            f"{expected_spacing} satır, önce {int(expected_before)}, sonra {int(expected_after)}, "
            f"{'kalın' if expected_bold else 'normal'}"
        )

        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u"
            })
            return (s or "").lower().translate(trans).strip()

        def resolve_from_styles_pf(para, attr):
            val = getattr(para.paragraph_format, attr)
            if val is not None:
                return val
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None:
                    v = getattr(pf, attr)
                    if v is not None:
                        return v
                s = getattr(s, "base_style", None)
            return None

        def effective_line_spacing(para, default=1.0):
            ls = resolve_from_styles_pf(para, "line_spacing")
            if ls is None:
                return float(default)
            if hasattr(ls, "pt"):
                return round(ls.pt, 1)
            try:
                return round(float(ls), 2)
            except Exception:
                return ls

        def effective_space_pt(para, which: str) -> float:
            attr = "space_before" if which == "before" else "space_after"
            length = resolve_from_styles_pf(para, attr)
            return 0.0 if length is None else round(length.pt, 1)

        def effective_alignment(para):
            if para.alignment is not None:
                return para.alignment
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None and pf.alignment is not None:
                    return pf.alignment
                s = getattr(s, "base_style", None)
            return None

        def effective_bold(para):
            for r in para.runs:
                if r.bold is True or (r.font and r.font.bold is True):
                    return True
            s = para.style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.bold is True:
                    return True
                s = getattr(s, "base_style", None)
            return False

        def run_font_name(run):
            if run.font and run.font.name:
                return run.font.name
            rpr = getattr(run._element, "rPr", None)
            if rpr is not None and getattr(rpr, "rFonts", None) is not None:
                rf = rpr.rFonts
                for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                    val = getattr(rf, attr, None)
                    if val:
                        return val
            return None

        def style_font_name(style):
            s = style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.name:
                    return f.name
                el = getattr(s, "element", None)
                if el is not None:
                    rpr = getattr(el, "rPr", None)
                    if rpr is not None and getattr(rpr, "rFonts", None) is not None:
                        rf = rpr.rFonts
                        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                            val = getattr(rf, attr, None)
                            if val:
                                return val
                s = getattr(s, "base_style", None)
            return None

        def para_font_size_pt(para):
            for run in para.runs:
                if run.font and run.font.size:
                    return round(run.font.size.pt, 1)
            s = para.style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.size:
                    return round(f.size.pt, 1)
                s = getattr(s, "base_style", None)
            return None

        lof_idx = memo.get("list_of_figures_idx")
        if lof_idx is None:
            if must_exist:
                return [(0, False, rule_title, "ŞEKİLLER DİZİNİ başlığı bulunamadı; 'Sayfa' konumu doğrulanamadı.")]
            else:
                return [(0, True, rule_title, "Atlandı: Şekiller Dizini isteğe bağlı ve tezde bulunmadı.")]

        sayfa_idx = lof_idx + 1
        if sayfa_idx >= len(paragraphs):
            return [(lof_idx, False, rule_title, "2. satır (Sayfa) bulunamadı")]

        p2 = paragraphs[sayfa_idx]
        text2 = (p2.text or "").replace("\xa0", " ").strip()
        errors = []

        if norm_tr(re.sub(r"[^\w]+", "", text2)) != "sayfa":
            errors.append("2. satır metni 'Sayfa' olmalı")

        eff_align = effective_alignment(p2)
        if expected_align == "right" and eff_align != WD_PARAGRAPH_ALIGNMENT.RIGHT:
            errors.append("2. satır sağa dayalı değil")

        eff_ls = effective_line_spacing(p2, default=1.0)
        eff_sb = effective_space_pt(p2, "before")
        eff_sa = effective_space_pt(p2, "after")
        if isinstance(eff_ls, (int, float)) and abs(eff_ls - expected_spacing) > 0.1:
            errors.append(f"Satır aralığı {eff_ls} yerine {expected_spacing} olmalı")
        if abs(eff_sb - expected_before) > 1:
            errors.append(f"Paragraf öncesi {eff_sb} yerine {expected_before} olmalı")
        if abs(eff_sa - expected_after) > 1:
            errors.append(f"Paragraf sonrası {eff_sa} yerine {expected_after} olmalı")

        run_names = [run_font_name(r) for r in p2.runs if run_font_name(r)]
        eff_name = Counter(run_names).most_common(1)[0][0] if run_names else style_font_name(p2.style)
        fs = para_font_size_pt(p2)

        def norm_font(n: str) -> str:
            return re.sub(r"\s+", "", (n or "").lower())

        expected_norm = norm_font(expected_name)
        actual_norm   = norm_font(eff_name) if eff_name else None

        if actual_norm is None:
            errors.append(f"Yazı tipi tespit edilemedi; {expected_name} olmalı")
        elif actual_norm != expected_norm:
            errors.append(f"Yazı tipi {eff_name} yerine {expected_name} olmalı")

        if fs is not None and round(float(fs), 1) != round(expected_size, 1):
            errors.append(f"Punto {fs} yerine {expected_size} olmalı")

        if expected_bold and not effective_bold(p2):
            errors.append("2. satır kalın değil")

        memo["list_of_figures_sayfa_idx"] = sayfa_idx
        return [(sayfa_idx, len(errors) == 0, rule_title, "; ".join(errors))]

    # ======================================================
    # ŞEKİLLER DİZİNİ – GÖVDE (Sayfa satırından sonra tablo)
    # TNR 12pt, LEFT, 1.5, 0/0 — Bitiş: sonraki başlık veya ilk ORTALI paragraf
    # Sibling taraması: <w:p>, <w:tbl>, <w:sdt>
    # Debug log: debug_lof_table_end.txt
    # ======================================================
    elif check["check"] == "list_of_figures_body_format":
        import re, datetime
        from collections import Counter
        

        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_align   = check.get("alignment", "left").lower()
        expected_spacing = float(check.get("line_spacing", 1.5))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 0))
        must_exist       = check.get("must_exist", False)
        debug_enabled    = bool(check.get("debug", False))
        DEBUG_F          = dbg_path("debug_lof_table_end.txt")

        rule_title = (
            f"ŞEKİLLER DİZİNİ – gövde biçimi: {expected_name}, {int(expected_size)} pt, "
            f"{'sola yaslı' if expected_align=='left' else expected_align}, "
            f"{expected_spacing} satır, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        # ---------- yardımcılar ----------
        def norm_tr(s: str) -> str:
            trans = str.maketrans({"ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g","ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u"})
            return (s or "").lower().translate(trans).strip()

        def strip_label_like_toc(text: str) -> str:
            t = (text or "").replace("\xa0", " ").strip()
            t = re.split(r"\.{2,}|\t", t)[0]
            t = re.sub(r"\s*\d+$", "", t)
            t = re.sub(r"\s*[ivxlcdm]{2,8}$", "", t, flags=re.I)
            n = norm_tr(t)
            n = re.sub(r"[^\w ]+", "", n)
            n = re.sub(r"\s+", " ", n).strip()
            return n

        def resolve_from_styles_pf(para, attr):
            val = getattr(para.paragraph_format, attr)
            if val is not None: return val
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None:
                    v = getattr(pf, attr)
                    if v is not None: return v
                s = getattr(s, "base_style", None)
            return None

        def effective_alignment(para):
            if para.alignment is not None: return para.alignment
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf and pf.alignment is not None:
                    return pf.alignment
                s = getattr(s, "base_style", None)
            return None

        def effective_line_spacing(para, default=1.0):
            ls = resolve_from_styles_pf(para, "line_spacing")
            if ls is None: return float(default)
            if hasattr(ls, "pt"): return round(ls.pt, 1)
            try: return round(float(ls), 2)
            except Exception: return ls

        def effective_space_pt(para, which: str) -> float:
            length = resolve_from_styles_pf(para, "space_before" if which=="before" else "space_after")
            return 0.0 if length is None else round(length.pt, 1)

        def norm_font(n: str) -> str:
            return re.sub(r"\s+", "", (n or "").lower())

        def run_font_name(run):
            if run.font and run.font.name:
                return run.font.name
            rpr = getattr(run._element, "rPr", None)
            if rpr is not None and getattr(rpr, "rFonts", None) is not None:
                rf = rpr.rFonts
                for attr in ("ascii","hAnsi","eastAsia","cs"):
                    val = getattr(rf, attr, None)
                    if val: return val
            return None

        def style_font_name(style):
            s = style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.name: return f.name
                el = getattr(s, "element", None)
                if el is not None:
                    rpr = getattr(el, "rPr", None)
                    if rpr is not None and getattr(rpr, "rFonts", None) is not None:
                        rf = rpr.rFonts
                        for attr in ("ascii","hAnsi","eastAsia","cs"):
                            val = getattr(rf, attr, None)
                            if val: return val
                s = getattr(s, "base_style", None)
            return None

        def para_font_size_pt(para):
            for run in para.runs:
                if run.font and run.font.size:
                    return round(run.font.size.pt, 1)
            s = para.style
            while s is not None:
                f = getattr(s, "font", None)
                if f and f.size:
                    return round(f.size.pt, 1)
                s = getattr(s, "base_style", None)
            return None

        def p_text_from_p_el(p_el):
            texts = []
            for node in p_el.iter():
                if node.tag.endswith('}t') and node.text:
                    texts.append(node.text)
            return "".join(texts) if texts else ""

        def find_para_index_by_element(p_el):
            for i_pp, pp in enumerate(paragraphs):
                if pp._element is p_el:
                    return i_pp
            return None

        def align_from_p_el(p_el):
            try:
                pPr = getattr(p_el, "pPr", None)
                if pPr is not None and getattr(pPr, "jc", None) is not None:
                    val = pPr.jc.val
                    if not val: return None
                    val = str(val).lower()
                    if val == "left":   return WD_PARAGRAPH_ALIGNMENT.LEFT
                    if val == "right":  return WD_PARAGRAPH_ALIGNMENT.RIGHT
                    if val == "center": return WD_PARAGRAPH_ALIGNMENT.CENTER
                    if val == "both":   return WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except Exception:
                pass
            return None

        # ---------- başlangıç ----------
        sayfa_idx = memo.get("list_of_figures_sayfa_idx")
        if sayfa_idx is None:
            if must_exist:
                return [(0, False, rule_title, "ŞEKİLLER DİZİNİ bulunamadı (zorunlu sayfa).")]
            else:
                return [(0, True, rule_title, "Atlandı: Şekiller Dizini isteğe bağlı ve tezde bulunmadı.")]

        start_para = paragraphs[sayfa_idx]
        cur = start_para._element.getnext()

        if debug_enabled:
            with open(DEBUG_F, "w", encoding="utf-8") as dbg:
                dbg.write(f"[LOF-TABLE-DEBUG] {datetime.datetime.now():%Y-%m-%d %H:%M:%S}\n")
                dbg.write(f"memo.list_of_figures_sayfa_idx={sayfa_idx}\n")
                dbg.write("[SCAN] Kardeş (sibling) bloklar taranıyor...\n")

        # ---------- bitiş tespiti ----------
        body_p_indices = []
        end_idx = None
        end_found_reason = None

        while cur is not None:
            tag = cur.tag.split('}')[-1]

            if tag == "p":
                j = find_para_index_by_element(cur)
                raw = p_text_from_p_el(cur) if j is None else (paragraphs[j].text or "")
                nrm = strip_label_like_toc(raw)

                if debug_enabled:
                    preview = raw.replace("\n", " ").strip()
                    if len(preview) > 140: preview = preview[:140] + "…"
                    if j is not None:
                        eff_align_dbg = effective_alignment(paragraphs[j]) or align_from_p_el(cur)
                    else:
                        eff_align_dbg = align_from_p_el(cur)
                    align_str = str(eff_align_dbg).replace("WD_PARAGRAPH_ALIGNMENT.", "") if eff_align_dbg is not None else "None"
                    with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                        dbg.write(f"  [P] idx={j if j is not None else '-'}  label={nrm!r}  align={align_str}  raw={preview!r}\n")

                if nrm in {"cizelgeler dizini", "simge dizini", "kisaltmalar", "bolum 1", "giris", "girıs"}:
                    end_idx = j if j is not None else 0
                    end_found_reason = f"Başlık: {nrm}"
                    break
                eff_align = (effective_alignment(paragraphs[j]) if j is not None else align_from_p_el(cur))
                if eff_align == WD_PARAGRAPH_ALIGNMENT.CENTER and raw.strip():
                    end_idx = j if j is not None else 0
                    end_found_reason = "Ortalanmış paragraf"
                    break

                if j is not None:
                    body_p_indices.append(j)

            elif tag == "tbl":
                p_els = [el for el in cur.iter() if el.tag.endswith('}p')]
                if debug_enabled and not p_els:
                    with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                        dbg.write("  [TBL] (paragraf içermiyor)\n")
                for k, p_el in enumerate(p_els):
                    j = find_para_index_by_element(p_el)
                    raw = p_text_from_p_el(p_el)
                    nrm = strip_label_like_toc(raw)
                    eff_align_dbg = (effective_alignment(paragraphs[j]) if j is not None else align_from_p_el(p_el)) or align_from_p_el(p_el)
                    align_str = str(eff_align_dbg).replace("WD_PARAGRAPH_ALIGNMENT.", "") if eff_align_dbg is not None else "None"

                    if debug_enabled:
                        preview = raw.replace("\n", " ").strip()
                        if len(preview) > 140: preview = preview[:140] + "…"
                        with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                            dbg.write(f"  [TBL] p={k} idx={j if j is not None else '-'} label={nrm!r} align={align_str} raw={preview!r}\n")

                    if nrm in {"cizelgeler dizini", "simge dizini", "kisaltmalar", "bolum 1", "giris", "girıs"}:
                        end_idx = j if j is not None else 0
                        end_found_reason = f"TBL Başlık: {nrm}"
                        break
                    if eff_align_dbg == WD_PARAGRAPH_ALIGNMENT.CENTER and raw.strip():
                        end_idx = j if j is not None else 0
                        end_found_reason = "TBL Ortalanmış paragraf"
                        break

                    if j is not None:
                        body_p_indices.append(j)
                if end_idx is not None:
                    break

            elif tag == "sdt":
                p_els = [el for el in cur.iter() if el.tag.endswith('}p')]
                if debug_enabled and not p_els:
                    with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                        dbg.write("  [SDT] (paragraf içermiyor)\n")
                for k, p_el in enumerate(p_els):
                    j = find_para_index_by_element(p_el)
                    raw = p_text_from_p_el(p_el)
                    nrm = strip_label_like_toc(raw)
                    eff_align_dbg = (effective_alignment(paragraphs[j]) if j is not None else align_from_p_el(p_el)) or align_from_p_el(p_el)
                    align_str = str(eff_align_dbg).replace("WD_PARAGRAPH_ALIGNMENT.", "") if eff_align_dbg is not None else "None"

                    if debug_enabled:
                        preview = raw.replace("\n", " ").strip()
                        if len(preview) > 140: preview = preview[:140] + "…"
                        with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                            dbg.write(f"  [SDT] p={k} idx={j if j is not None else '-'} label={nrm!r} align={align_str} raw={preview!r}\n")

                    if nrm in {"cizelgeler dizini", "simge dizini", "kisaltmalar", "bolum 1", "giris", "girıs"}:
                        end_idx = j if j is not None else 0
                        end_found_reason = f"SDT Başlık: {nrm}"
                        break
                    if eff_align_dbg == WD_PARAGRAPH_ALIGNMENT.CENTER and raw.strip():
                        end_idx = j if j is not None else 0
                        end_found_reason = "SDT Ortalanmış paragraf"
                        break

                    if j is not None:
                        body_p_indices.append(j)
                if end_idx is not None:
                    break

            cur = cur.getnext()

        if end_idx is None:
            end_idx = len(paragraphs)
            end_found_reason = "Belge sonu"

        if debug_enabled:
            with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                dbg.write(f"[END] reason={end_found_reason}, end_idx={end_idx}\n")
                dbg.write(f"[BODY] denetlenecek paragraph indexleri: {body_p_indices}\n")

        # ---------- biçim denetimi ----------
        errors = []
        checked = 0

        for j in body_p_indices:
            loc = f"{j}. satır (belge:{j})"
            
            if j is None or j < 0 or j >= len(paragraphs):
                continue
            p = paragraphs[j]
            txt = (p.text or "").replace("\xa0", " ").strip()
            if txt == "":
                pv = " ".join(txt.split()[:10])
                if len(pv) > 90:
                    pv = pv[:90].rstrip() + "…"
                loc = loc + (f" ('{pv}')" if pv else "")

                continue

            eff_align = effective_alignment(p) or align_from_p_el(p._element)
            if eff_align is None:  # LOF için varsayılan: LEFT
                eff_align = WD_PARAGRAPH_ALIGNMENT.LEFT

            eff_ls    = effective_line_spacing(p, default=1.5)
            eff_sb    = effective_space_pt(p, "before")
            eff_sa    = effective_space_pt(p, "after")

            run_names = [run_font_name(r) for r in p.runs if run_font_name(r)]
            eff_name  = Counter(run_names).most_common(1)[0][0] if run_names else style_font_name(p.style)
            fs        = para_font_size_pt(p)

            if expected_align == "left" and eff_align != WD_PARAGRAPH_ALIGNMENT.LEFT:
                errors.append(f"{loc}: Paragraf sola yaslı değil")
            if isinstance(eff_ls, (int, float)) and abs(eff_ls - expected_spacing) > 0.1:
                errors.append(f"{loc}: Satır aralığı {eff_ls} yerine {expected_spacing} olmalı")
            if abs(eff_sb - expected_before) > 1:
                errors.append(f"{loc}: Öncesi {eff_sb} yerine {expected_before} olmalı")
            if abs(eff_sa - expected_after) > 1:
                errors.append(f"{loc}: Sonrası {eff_sa} yerine {expected_after} olmalı")

            exp_norm = norm_font(expected_name)
            act_norm = norm_font(eff_name) if eff_name else None
            if act_norm is None:
                errors.append(f"{loc}: Yazı tipi tespit edilemedi; {expected_name} olmalı")
            elif act_norm != exp_norm:
                errors.append(f"{loc}: Yazı tipi {eff_name} yerine {expected_name} olmalı")
            if fs is not None and round(float(fs),1) != round(expected_size,1):
                errors.append(f"{loc}: Punto {fs} yerine {expected_size} olmalı")

            if debug_enabled:
                preview = txt.replace("\n", " ")
                if len(preview) > 120:
                    preview = preview[:120] + "…"
                align_str = str(eff_align).replace("WD_PARAGRAPH_ALIGNMENT.", "")
                with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                    dbg.write(f"[P{j}] Align={align_str:8s}  LS={eff_ls:<4}  SB={eff_sb:<4}  SA={eff_sa:<4}  Font={eff_name or '?'}  Txt={preview!r}\n")

            checked += 1

        ok = (len(errors) == 0)
        if debug_enabled:
            with open(DEBUG_F, "a", encoding="utf-8") as dbg:
                dbg.write(f"[DONE] checked={checked}, errors={len(errors)}\n")

        begin_idx = body_p_indices[0] if body_p_indices else (memo.get("list_of_figures_sayfa_idx", 0) + 1)
        return [(begin_idx, ok, rule_title, "; ".join(errors))]


    # ===============================================================================================================#
    # ===============================================================================================================#    
    # ===============================================================================================================#
    # SİMGELER VE KISALTMALAR – BAŞLIK (ÖN SAYFALAR SONU, BİÇİM KONTROLÜ)
    # Zorunlu değil: must_exist=false ise bulunamazsa bilgilendirme yapar.
    # Arama başlangıcı: (varsa) ŞEKİLLER → ÇİZELGELER → yoksa İÇİNDEKİLER’den sonra.
    # Bulunursa memoya yazılır: memo["symbols_abbreviations_idx"] = idx
    # ======================================================
    elif check["check"] == "symbols_abbreviations_heading":
        import re
        
        # ----------- Beklenen biçim (YAML'dan) -----------
        expected_name     = check.get("font_name", "Times New Roman")
        expected_size     = float(check.get("font_size_pt", 12))
        expected_bold     = bool(check.get("bold", True))
        expected_all_caps = bool(check.get("all_caps", True))
        expected_align    = check.get("alignment", "center").lower()
        expected_spacing  = float(check.get("line_spacing", 1.5))
        expected_before   = float(check.get("space_before", 0))
        expected_after    = float(check.get("space_after", 24))
        # >>> Toleranslar YAML'dan yönetilebilir (yoksa 2.0 pt)
        space_tol_pt      = float(check.get("space_tolerance_pt", 2.0))
        must_exist        = bool(check.get("must_exist", False))
        markers           = check.get("markers", ["^SİMGELER VE KISALTMALAR$"])

        rule_title = (
            f"SİMGELER VE KISALTMALAR başlığı: {expected_name}, {int(expected_size)} punto, "
            f"{'kalın' if expected_bold else 'normal'}, "
            f"{'BÜYÜK HARF' if expected_all_caps else 'normal'}, "
            f"{expected_spacing} satır aralığı, önce {int(expected_before)}, sonra {int(expected_after)}"
        )

        # ----------- Yardımcılar -----------
        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı":"i","İ":"i","ç":"c","Ç":"c","ğ":"g","Ğ":"g",
                "ö":"o","Ö":"o","ş":"s","Ş":"s","ü":"u","Ü":"u"
            })
            return (s or "").lower().translate(trans).strip()

        def resolve_from_styles(para, attr_name):
            """paragraph_format + stil zinciri üzerinden attr (line_spacing, space_before/after) çöz."""
            val = getattr(para.paragraph_format, attr_name)
            if val is not None:
                return val
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None:
                    v = getattr(pf, attr_name)
                    if v is not None:
                        return v
                s = getattr(s, "base_style", None)
            return None

        def effective_line_spacing(para, default=1.0) -> float:
            """Satır aralığı: paragraph_format/stilden al; yoksa default."""
            ls = resolve_from_styles(para, "line_spacing")
            if ls is None:
                return float(default)
            if hasattr(ls, "pt"):  # Length
                return round(ls.pt, 1)
            try:
                return round(float(ls), 2)  # numeric (1.0, 1.5)
            except Exception:
                return ls

        # ----------- YENİ: XML fallback'li boşluk çözücü -----------
        def effective_space_pt_strict(para, which: str):
            """
            Önce XML'den <w:pPr><w:spacing w:after|w:before> okur (twips→pt).
            Auto spacing açıksa (afterAutospacing/beforeAutospacing) (None, True) döner.
            Bulunamazsa stil zincirinden Length→pt okur.
            return: (pt: float|None, is_auto: bool)
            """
            try:
                p_el = para._element
                pPr  = getattr(p_el, "pPr", None)
                if pPr is not None and getattr(pPr, "spacing", None) is not None:
                    sp = pPr.spacing
                    if which == "after":
                        if getattr(sp, "afterAutospacing", None) in (True, 1, "1"):
                            return (None, True)
                        val = getattr(sp, "after", None)
                    else:
                        if getattr(sp, "beforeAutospacing", None) in (True, 1, "1"):
                            return (None, True)
                        val = getattr(sp, "before", None)
                    if val is not None:
                        try:
                            twips = int(val)     # twips
                            return (round(twips / 20.0, 1), False)  # pt
                        except Exception:
                            pass
            except Exception:
                pass

            # XML yoksa/stil üzerinden dene
            attr = "space_after" if which == "after" else "space_before"
            length = resolve_from_styles(para, attr)
            if length is None:
                return (None, False)  # None döndür, 0 varsayma!
            return (round(length.pt, 1), False)

        def effective_alignment(para):
            """Hizalamayı paragraph_format/stil zinciri üzerinden oku."""
            if para.alignment is not None:
                return para.alignment
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None and pf.alignment is not None:
                    return pf.alignment
                s = getattr(s, "base_style", None)
            return None

        def effective_bold(para):
            """Run veya stil zincirinde bold True mu?"""
            for r in para.runs:
                if r.bold is True or (r.font and r.font.bold is True):
                    return True
            s = para.style
            while s is not None:
                if getattr(s, "font", None) and s.font.bold is True:
                    return True
                s = getattr(s, "base_style", None)
            return False

        # ----------- 1) Arama başlangıç noktasını belirle (güvenli) -----------
        # Öncelik: ŞEKİLLER → ÇİZELGELER → İÇİNDEKİLER
        candidates = [
            memo.get("list_of_figures_idx"),
            memo.get("list_of_tables_idx"),
            memo.get("toc_heading_idx"),
        ]
        # mevcut olanların maksimumunu al; hiç yoksa -1
        start_from = max([i for i in candidates if isinstance(i, int)], default=-1)

        if start_from < 0 and must_exist:
            return [(0, False, rule_title, "Ön sayfa referansı yok; konum kontrolü yapılamadı (zorunlu sayfa).")]
        elif start_from < 0:
            return [(0, True, rule_title, "Atlandı: Ön sayfa referansı bulunamadı; bu sayfa isteğe bağlı.")]

        # ----------- 2) Başlığı yalnızca start_from'dan SONRA ara -----------
        compiled = []
        for m in markers:
            try:
                compiled.append(re.compile(norm_tr(m), re.IGNORECASE))
            except re.error:
                compiled.append(re.compile("^" + re.escape(norm_tr(m).strip("^$")) + "$", re.IGNORECASE))

        found = None
        for i in range(start_from + 1, len(paragraphs)):
            raw = (paragraphs[i].text or "").strip()
            norm = norm_tr(raw)
            for pat in compiled:
                if pat.match(norm):
                    found = (i, paragraphs[i])
                    break
            if found:
                break

        if not found:
            if must_exist:
                return [(start_from, False, rule_title,
                         "Zorunlu sayfa bulunamadı (ön sayfalardan sonra bekleniyordu).")]
            else:
                return [(start_from, True, rule_title,
                         "Atlandı: 'SİMGELER VE KISALTMALAR' isteğe bağlı ve tezde bulunmadı.")]

        # ----------- 3) Biçim kontrolleri -----------
        idx, p = found
        errors = []

        # Font adı/punto (run veya stil)
        fn = next((r.font.name for r in p.runs if r.font and r.font.name), None)
        fs = next((r.font.size.pt for r in p.runs if r.font and r.font.size), None)
        if fn and fn != expected_name:
            errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")
        if fs and round(float(fs), 1) != round(expected_size, 1):
            errors.append(f"Punto {fs} yerine {expected_size} olmalı")

        # Bold & ALL CAPS
        if expected_bold and not effective_bold(p):
            errors.append("Başlık kalın değil")
        if expected_all_caps and not p.text.isupper():
            errors.append("Başlık BÜYÜK HARF değil")

        # Hizalama
        eff_align = effective_alignment(p)
        if expected_align == "center" and eff_align != WD_PARAGRAPH_ALIGNMENT.CENTER:
            errors.append("Başlık ortalı değil")

        # Satır aralığı
        eff_ls = effective_line_spacing(p, default=expected_spacing)
        if isinstance(eff_ls, (int, float)) and abs(eff_ls - expected_spacing) > 0.1:
            errors.append(f"Satır aralığı {eff_ls} yerine {expected_spacing} olmalı")

        # Paragraf boşlukları (XML fallback + tolerans)
        eff_sb, sb_auto = effective_space_pt_strict(p, "before")
        eff_sa, sa_auto = effective_space_pt_strict(p, "after")

        # Auto spacing açık ise hata yazma (istersen bilgi notu ekleyebilirsin)
        if not sb_auto and eff_sb is not None and abs(eff_sb - expected_before) > space_tol_pt:
            errors.append(f"Satır öncesi {eff_sb} yerine {expected_before} olmalı")
        if not sa_auto and eff_sa is not None and abs(eff_sa - expected_after) > space_tol_pt:
            errors.append(f"Satır sonrası {eff_sa} yerine {expected_after} olmalı")

        # Bulundu: memoya konum yaz
        memo["symbols_abbreviations_idx"] = idx

        return [(idx, len(errors) == 0, rule_title, "; ".join(errors))]


    # ===============================================================================================================#
    # ===============================================================================================================#    
    # ===============================================================================================================#
    # ======================================================
    # BÖLÜM BAŞLIĞI BLOĞU (ör: BÖLÜM I / BÖLÜM 1 + GİRİŞ)
    # ======================================================
    elif check["check"] == "chapter_heading_block":
        import re

        # --- Beklenen biçimsel parametreler ---
        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_bold    = check.get("bold", True)
        expected_align   = check.get("alignment", "center").lower()
        expected_spacing = float(check.get("line_spacing", 1.5))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 24))
        expected_style   = check.get("expected_style", "Heading 1")
        debug_mode       = check.get("debug", False)

        rule_title = (
            f"BÖLÜM Başlığı Bloğu\n"
            f"{expected_name}, {expected_size} punto, kalın={expected_bold}, "
            f"{expected_align}, {expected_spacing} satır aralığı, "
            f"önce {expected_before}, sonra {expected_after}, stil={expected_style}"
        )

        chapter_blocks = []
        forbidden_idxs = set()
        errors = []
        debug_file = None

        if debug_mode:
            debug_file = open(dbg_path("chapter_debug.txt"), "w", encoding="utf-8")
            debug_file.write("=== BÖLÜM BAŞLIKLARI DEBUG KAYDI ===\n\n")

        # ✅ Sıkı bölüm satırı: sadece "BÖLÜM I" veya "BÖLÜM 1" gibi TEK satır
        # - Roman: ivxlcdm+
        # - Arabik: \d+
        chapter_line_pat = re.compile(r"^b[oö]l[uü]m\s+([ivxlcdm]+|\d+)\s*$", re.IGNORECASE)

        # ✅ "GİRİŞ" (tam satır)
        giris_exact_pat = re.compile(r"^g[iı]r[iı]ş\s*$", re.IGNORECASE)

        # ❌ Hatalı tek-satır giriş varyasyonları (bölüm satırı yerine yazılmış):
        # 1.GİRİŞ, 1. GİRİŞ, 1 GİRİŞ, I.GİRİŞ, I) GİRİŞ, I - GİRİŞ vb.
        giris_numbered_pat = re.compile(
            r"^(?:\d+|[ivxlcdm]+)\s*[\.\)\-–—]?\s*g[iı]r[iı]ş\s*$",
            re.IGNORECASE
        )

        # --- Etkin özellikleri hesaplayan yardımcı fonksiyonlar ---
        def dbg_para_props(p):
            fn_eff   = effective_font_name(p)
            fs_eff   = effective_font_size_pt(p)
            bold_eff = effective_bold(p)
            ls_eff   = effective_line_spacing(p, default=expected_spacing)
            sb_eff   = effective_space_pt(p, "before")
            sa_eff   = effective_space_pt(p, "after")
            al_eff   = effective_alignment(p)
            style_name = p.style.name if p.style else ""
            return fn_eff, fs_eff, bold_eff, ls_eff, sb_eff, sa_eff, al_eff, style_name

        # --- Biçimsel kontrol + stil kontrolü ---
        def check_format(p, label, idx_for_msg):
            fn_eff, fs_eff, bold_eff, ls_eff, sb_eff, sa_eff, al_eff, style_name = dbg_para_props(p)

            # Stil kontrolü
            if expected_style.lower() not in (style_name or "").lower():
                errors.append(f"{idx_for_msg}. {label}: Stil '{style_name}' yerine '{expected_style}' olmalı")

            # Biçim kontrolleri
            if fn_eff and fn_eff != expected_name:
                errors.append(f"{idx_for_msg}. {label} yazı tipi {fn_eff} yerine {expected_name} olmalı")
            if fs_eff and abs(fs_eff - expected_size) > 0.1:
                errors.append(f"{idx_for_msg}. {label} punto {fs_eff} yerine {expected_size} olmalı")
            if bool(bold_eff) != bool(expected_bold):
                errors.append(f"{idx_for_msg}. {label} kalın olmalı")

            # Hizalama kontrolü
            want = expected_align
            if want == "center":
                if al_eff != WD_PARAGRAPH_ALIGNMENT.CENTER:
                    errors.append(f"{idx_for_msg}. {label} ortalı olmalı")
            elif want == "left":
                if al_eff != WD_PARAGRAPH_ALIGNMENT.LEFT:
                    errors.append(f"{idx_for_msg}. {label} sola hizalı olmalı")
            elif want == "justify":
                if al_eff != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                    errors.append(f"{idx_for_msg}. {label} iki yana yaslı olmalı")

            if abs(ls_eff - expected_spacing) > 0.1:
                errors.append(f"{idx_for_msg}. {label} satır aralığı {ls_eff} yerine {expected_spacing} olmalı")
            if abs(sb_eff - expected_before) > 1.0:
                errors.append(f"{idx_for_msg}. {label} öncesi {sb_eff} yerine {expected_before} olmalı")
            if abs(sa_eff - expected_after) > 1.0:
                errors.append(f"{idx_for_msg}. {label} sonrası {sa_eff} yerine {expected_after} olmalı")

            return fn_eff, fs_eff, bold_eff, ls_eff, sb_eff, sa_eff, al_eff, style_name

        # ✅ Boş paragrafları atlayıp bir sonraki dolu paragrafı bul
        def next_nonempty_idx(start_i: int):
            j = start_i
            while j < len(paragraphs):
                if (paragraphs[j].text or "").strip():
                    return j
                j += 1
            return None

        # --- Ana tarama döngüsü: Normal "BÖLÜM X" + alt satır başlık bloğu ---
        for i, p in enumerate(paragraphs):
            text = (p.text or "").strip()
            if not chapter_line_pat.match(text):
                continue

            # Alt satır: boşları atlayarak bul
            title_idx = next_nonempty_idx(i + 1)
            next_p = paragraphs[title_idx] if title_idx is not None else None

            chapter_blocks.append((i, title_idx))
            forbidden_idxs.add(i)
            if title_idx is not None:
                forbidden_idxs.add(title_idx)

            # --- Debug kaydı ---
            if debug_file:
                debug_file.write(f"[{i+1}. paragraf] BÖLÜM satırı: {text}\n")
                fn, fs, bb, ls, sb, sa, al, st = dbg_para_props(p)
                debug_file.write(
                    f"  Font (eff): {fn}\n"
                    f"  Size (eff): {fs}\n"
                    f"  Bold (eff): {bb}\n"
                    f"  Line spacing (eff): {ls}\n"
                    f"  Before (eff): {sb}\n"
                    f"  After (eff): {sa}\n"
                    f"  Alignment (eff): {al}\n"
                    f"  Style: {st}\n\n"
                )
                if next_p:
                    debug_file.write(f"    Bölüm başlığı → {next_p.text.strip()}\n")
                    fn2, fs2, bb2, ls2, sb2, sa2, al2, st2 = dbg_para_props(next_p)
                    debug_file.write(
                        f"  Font (eff): {fn2}\n"
                        f"  Size (eff): {fs2}\n"
                        f"  Bold (eff): {bb2}\n"
                        f"  Line spacing (eff): {ls2}\n"
                        f"  Before (eff): {sb2}\n"
                        f"  After (eff): {sa2}\n"
                        f"  Alignment (eff): {al2}\n"
                        f"  Style: {st2}\n\n"
                    )
                else:
                    debug_file.write("    Bölüm başlığı bulunamadı.\n\n")

            # --- Biçim + Stil denetimi (iki satır da Heading stili ve aynı format) ---
            check_format(p, "BÖLÜM satırı", i + 1)

            if not next_p:
                loc = format_location_by_page(i + 1)
                errors.append(f"{loc} BÖLÜM satırının altında bölüm başlığı bulunamadı")
                continue

            check_format(next_p, "Bölüm başlığı", title_idx + 1)

            # ✅ İlk bölüm için içerik kuralı: başlık "GİRİŞ" olmalı (numarasız)
            if len(chapter_blocks) == 1:
                title_text = (next_p.text or "").strip()

                # 1) numaralı girişleri özellikle yakala
                if giris_numbered_pat.match(title_text):
                    loc = format_location_by_page(title_idx + 1)
                    errors.append(
                        f"{loc} İlk bölüm başlığı 'GİRİŞ' numarasız olmalı (örn: 'GİRİŞ'). Siz: '{title_text}'"
                    )
                # 2) giriş değilse hata
                elif not giris_exact_pat.match(title_text):
                    loc = format_location_by_page(title_idx + 1)
                    errors.append(
                        f"{loc} İlk bölüm başlığı 'GİRİŞ' olmalı. Siz: '{title_text}'"
                    )

                # ✅ Memo: tez gövdesi başlangıcı = "GİRİŞ" başlığından sonraki paragraf
                memo["first_chapter_heading_idx"] = i
                memo["first_chapter_title_idx"] = title_idx
                memo["chapter_first_idx"] = title_idx + 1

        # --- FALLBACK: Öğrenci "BÖLÜM X" satırı yazmadan tek satırda GİRİŞ / 1.GİRİŞ / 1 GİRİŞ yazıp altına metne geçmişse ---
        if not chapter_blocks:
            fallback_idx = None
            fallback_text = ""

            for i, p in enumerate(paragraphs):
                t = ((p.text or "").strip())
                if not t:
                    continue
                if giris_exact_pat.match(t) or giris_numbered_pat.match(t):
                    fallback_idx = i
                    fallback_text = t
                    break

            if fallback_idx is not None:
                body_idx = next_nonempty_idx(fallback_idx + 1)

                # ✅ Bu hatayı ihlal say, ama memoyu set edip kontrollerin devam etmesini sağla
                loc = format_location_by_page(fallback_idx + 1)
                errors.append(
                    f"{loc} Bölüm bloğu hatalı: 'BÖLÜM I/1' satırı + alt satır bölüm başlığı beklenirken "
                    f"'{fallback_text}' tek satır yazılmış ve bölüm başlığı satırı atlanmış görünüyor. "
                    f"Kontrollere devam etmek için gövde başlangıcı bu satırdan sonra kabul edildi."
                )

                # İsteğe bağlı: bu satırın formatını da kontrol et (Heading 1 bekleniyorsa faydalı olur)
                check_format(paragraphs[fallback_idx], "GİRİŞ satırı (yanlış konum)", fallback_idx + 1)

                # Gövde satırı bulunamadıysa da ihlal yaz
                if body_idx is None:
                    loc2 = format_location_by_page(fallback_idx + 1)
                    errors.append(f"{loc2} 'GİRİŞ' satırından sonra tez gövdesi (metin) bulunamadı")
                    # Yine de indeksleri güvenli bir değere çekelim
                    body_idx = fallback_idx + 1

                # Memo set: normal blok yok, ama devam edebilmek için başlangıçları kaydet
                memo["chapter_detect_mode"] = "fallback_giris_line"
                memo["first_chapter_heading_idx"] = None
                memo["first_chapter_title_idx"] = fallback_idx
                memo["chapter_first_idx"] = body_idx

                # Bu satırı diğer aramalardan hariç tut
                forbidden_idxs.add(fallback_idx)
            else:
                # Ne normal "BÖLÜM X" bloğu ne de fallback "GİRİŞ" satırı bulundu
                # (İstersen burada hata üretilebilir; şimdilik sessiz bırakmıyoruz)
                errors.append("BÖLÜM başlığı bloğu bulunamadı (ne 'BÖLÜM X' ne de 'GİRİŞ' satırı tespit edildi).")

        # --- Debug dosyasını kapat ---
        if debug_file:
            debug_file.write("\n=== KONTROL TAMAMLANDI ===\n")
            debug_file.close()

        # --- memo’ya kaydet ---
        memo["chapter_blocks"] = chapter_blocks
        memo["chapter_forbidden_idxs"] = forbidden_idxs

        # --- Sonuç ---
        if errors:
            results.append((0, False, rule_title, "; ".join(errors)))
        else:
            results.append((0, True, rule_title, ""))

    # ======================================================
    # BÖLÜM BAŞLIĞI – BOŞ SATIR OLMAMALI
    # (Sayfa/satır bilgisi yok, doğrudan başlıklar yazılır)
    # ======================================================
    elif check["check"] == "chapter_heading_no_blank":

        rule_title = "BÖLÜM satırı ile bölüm başlığı arasında boş satır olmamalı"

        chapter_blocks = memo.get("chapter_blocks", [])
        errors = []

        def short_text(s: str, n: int = 60) -> str:
            s = " ".join((s or "").strip().split())
            return (s[:n] + "...") if len(s) > n else s

        for bolum_idx, title_idx in chapter_blocks:
            # Bölüm satırı metni
            if not (0 <= bolum_idx < len(paragraphs)):
                continue
            bolum_text = short_text(paragraphs[bolum_idx].text)

            # Alt başlık yoksa bu kuralın konusu değil
            if title_idx is None or not (0 <= title_idx < len(paragraphs)):
                continue

            title_text = short_text(paragraphs[title_idx].text)

            # Arada boş satır varsa
            if title_idx != bolum_idx + 1:
                errors.append(
                    f"'{bolum_text}' ile '{title_text}' arasında boş satır bulunmamalı"
                )

        if errors:
            results.append((0, False, rule_title, "; ".join(errors)))
        else:
            results.append((0, True, rule_title, ""))

    # ======================================================
    # KAYNAKLAR BAŞLIĞI TESPİTİ ve BİÇİMSEL KONTROLÜ (VARYANT KABULLÜ)
    # - "KAYNAKLAR" yerine "KAYNAÇA" vb. yakın yazımları da tespit eder
    # - memo'ya kaydeder, akışı bozmaz
    # - Ancak kural ihlali olarak doğru metnin "KAYNAKLAR" olması gerektiğini yazar
    # ======================================================
    elif check["check"] == "references_heading_block" and check.get("enabled", True):
        import re

        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_bold    = check.get("bold", True)
        expected_caps    = check.get("all_caps", True)
        expected_align   = check.get("alignment", "center").lower()
        expected_spacing = float(check.get("line_spacing", 1.5))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 24))
        debug_mode       = check.get("debug", False)

        # Opsiyonel: stil adı kontrolü (verilmezse kontrol etmez)
        expected_style   = check.get("expected_style", None)

        # Opsiyonel: kabul edilecek varyantlar (tespit için)
        # -> Bunlar bulunursa memo set edilir ama "KAYNAKLAR olmalı" ihlali yazılır.
        accepted_variants = check.get("accepted_variants", None)
        if not accepted_variants:
            accepted_variants = [
                "KAYNAÇA",       # yaygın yanlış
                "KAYNAKÇA",      # şapka/ç varyasyonu
                "KAYNAKCA",      # ç yerine c
                "KAYNAKLARÇA",   # bazen ek hatası
                "KAYNAKLAR.",    # noktalı yazım (kılavuz izin vermiyorsa ihlal say)
                "KAYNAKLAR:",    # iki nokta
            ]

        rule_title = (
            f"KAYNAKLAR Başlığı Biçimsel Kontrolü\n"
            f"{expected_name}, {expected_size} punto, kalın={expected_bold}, "
            f"BÜYÜK HARF={expected_caps}, {expected_align}, "
            f"{expected_spacing} satır aralığı, önce {expected_before}, sonra {expected_after}"
            + (f", stil={expected_style}" if expected_style else "")
        )

        # --- Normalizasyon fonksiyonu (TR duyarsız) ---
        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı": "i", "İ": "i", "ç": "c", "Ç": "c", "ğ": "g", "Ğ": "g",
                "ö": "o", "Ö": "o", "ş": "s", "Ş": "s", "ü": "u", "Ü": "u"
            })
            return (s or "").lower().translate(trans).strip()

        # --- Son bölüm bloğundan sonra başlığı ara ---
        last_chapter_end = -1
        ch_blocks = memo.get("chapter_blocks", [])
        if ch_blocks:
            last_chapter_end = max([(b[1] if b[1] is not None else b[0]) for b in ch_blocks])

        start_search = last_chapter_end + 1
        if start_search < 0:
            start_search = 0

        # Aranan metinler (normalize)
        target_norm = norm_tr("KAYNAKLAR")
        variants_norm = [norm_tr(v) for v in (accepted_variants or [])]

        found_idx = None
        found_raw_text = None
        found_is_variant = False

        # 1) Önce tam "KAYNAKLAR" ara
        for i in range(start_search, len(paragraphs)):
            raw = (paragraphs[i].text or "").strip()
            if not raw:
                continue
            if norm_tr(raw) == target_norm:
                found_idx = i
                found_raw_text = raw
                found_is_variant = False
                break

        # 2) Bulunamadıysa: varyantları ara (memo için kabul et)
        if found_idx is None:
            for i in range(start_search, len(paragraphs)):
                raw = (paragraphs[i].text or "").strip()
                if not raw:
                    continue
                n = norm_tr(raw)
                if n in variants_norm:
                    found_idx = i
                    found_raw_text = raw
                    found_is_variant = True
                    break

        # --- Debug dosyası ---
        debug_file = None
        if debug_mode:
            debug_file = open(dbg_path("chapter_debug.txt"), "a", encoding="utf-8")
            debug_file.write("\n=== KAYNAKLAR BAŞLIĞI DEBUG KAYDI ===\n\n")
            debug_file.write(f"Arama başlangıcı idx={start_search+1} (paragraf no)\n")
            debug_file.write(f"Bulunan idx={(found_idx+1) if found_idx is not None else 'None'}\n")
            if found_idx is not None:
                debug_file.write(f"Bulunan metin: {found_raw_text}\n")
                debug_file.write(f"Varyant mı?: {found_is_variant}\n")
            debug_file.write("\n")

        if found_idx is None:
            msg = "⚠️ 'KAYNAKLAR' başlığı bulunamadı. (KAYNAÇA vb. varyantlar da dahil aranmıştır.)"
            if debug_file:
                debug_file.write(msg + "\n")
                debug_file.write("=== KAYNAKLAR BAŞLIĞI TARAMASI TAMAMLANDI ===\n")
                debug_file.close()
            results.append((0, False, rule_title, msg))
            return results

        p = paragraphs[found_idx]

        # Effective okuma (run + stil zinciri)
        fn = effective_font_name(p)
        fs = effective_font_size_pt(p)
        bb = effective_bold(p)
        ls = effective_line_spacing(p, default=expected_spacing)
        sb = effective_space_pt(p, "before")
        sa = effective_space_pt(p, "after")
        al = effective_alignment(p)
        txt = (p.text or "").strip()

        style_name = p.style.name if getattr(p, "style", None) else ""

        if debug_file:
            debug_file.write(f"[{found_idx+1}. paragraf] {txt}\n")
            debug_file.write(
                f"  Font (eff): {fn}\n"
                f"  Size (eff): {fs}\n"
                f"  Bold (eff): {bb}\n"
                f"  Line Spacing (eff): {ls}\n"
                f"  Before (eff): {sb}\n"
                f"  After (eff): {sa}\n"
                f"  Alignment (eff): {al}\n"
                f"  Style: {style_name}\n\n"
            )
            debug_file.write("=== KAYNAKLAR BAŞLIĞI TARAMASI TAMAMLANDI ===\n")
            debug_file.close()

        errors = []

        # ✅ Varyant bulunduysa: devam et ama "KAYNAKLAR olmalı" ihlali yaz
        if found_is_variant:
            errors.append(f"Başlık metni 'KAYNAKLAR' olmalı (siz: '{txt}')")

        # Opsiyonel stil kontrolü
        if expected_style:
            if expected_style.lower() not in (style_name or "").lower():
                errors.append(f"Stil '{style_name}' yerine '{expected_style}' olmalı")

        if fn and fn != expected_name:
            errors.append(f"Yazı tipi {fn} yerine {expected_name} olmalı")
        if fs and abs(fs - expected_size) > 0.1:
            errors.append(f"Punto {fs} yerine {expected_size} olmalı")
        if bool(bb) != bool(expected_bold):
            errors.append("Kalın olmalı")
        if expected_caps and txt != txt.upper():
            errors.append("Tüm harfler büyük olmalı")
        if abs(ls - expected_spacing) > 0.1:
            errors.append(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
        if abs(sb - expected_before) > 1.0:
            errors.append(f"Öncesi {sb} yerine {expected_before} olmalı")
        if abs(sa - expected_after) > 1.0:
            errors.append(f"Sonrası {sa} yerine {expected_after} olmalı")

        if expected_align == "center" and al != WD_PARAGRAPH_ALIGNMENT.CENTER:
            errors.append("Ortalanmış olmalı")
        elif expected_align == "left" and al != WD_PARAGRAPH_ALIGNMENT.LEFT:
            errors.append("Sola hizalı olmalı")
        elif expected_align == "justify" and al != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            errors.append("İki yana yaslı olmalı")

        # --- memo’ya kaydet (bulunduğu sürece set edilir) ---
        memo["references_heading_idx"] = found_idx
        memo["references_heading_block_idx"] = found_idx

        # --- Sonuç ---
        if errors:
            results.append((found_idx, False, rule_title, "; ".join(errors)))
        else:
            results.append((found_idx, True, rule_title, "KAYNAKLAR başlığı biçimsel olarak uygun."))

    # ======================================================
    # ALT BAŞLIK TESPİTİ ve BİÇİMSEL KONTROLÜ (Tez Metni İçinde)
    # ======================================================
    elif check["check"] == "subheading_detector":
        import re

        # --- YAML parametreleri ---
        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_bold    = check.get("bold", True)
        expected_spacing = float(check.get("line_spacing", 1.5))
        expected_before  = float(check.get("space_before", 18))
        expected_after   = float(check.get("space_after", 18))
        expected_align   = check.get("alignment", None)

        # ✅ Kılavuz: başlık 6 derinliğe kadar inebilir
        # Heading 2 (x.y) → Heading 7 (x.y.z.a.b.c) olacak şekilde varsayılanı genişletiyoruz.
        heading_styles = check.get(
            "heading_styles",
            ["Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6", "Heading 7"]
        )

        debug_mode       = check.get("debug", False)

        # --- Hizalama parametresini normalize et (tek veya çoklu) ---
        if isinstance(expected_align, str):
            expected_align = [expected_align.lower()]
        elif isinstance(expected_align, list):
            expected_align = [a.lower() for a in expected_align]
        else:
            expected_align = []

        align_label = ", ".join(expected_align) if expected_align else "hizalama (belirtilmedi)"
        rule_title = (
            f"Alt Başlıklar\n"
            f"{expected_name}, {expected_size} punto, kalın={expected_bold}, "
            f"{align_label}, {expected_spacing} satır aralığı, "
            f"önce {expected_before}, sonra {expected_after}"
        )

        # --- Türkçe normalizasyon ---
        def norm_tr(s):
            trans = str.maketrans("çğıöşüÇĞİÖŞÜ", "cgiosuCGIOSU")
            return (s or "").translate(trans).lower().strip()

        # --- Numara biçimli başlık deseni (kılavuz: 2.1 ile başlıyor) ---
        numbered_heading_pattern = re.compile(r"^\d+(\.\d+)+\.?\s+[A-Za-zÇÖŞÜĞİçöşüğı]")

        # Raporu şişirmemek için
        def short_text(s: str, n: int = 80) -> str:
            s = " ".join((s or "").strip().split())
            return (s[:n] + "...") if len(s) > n else s

        # =====================================================
        # 🔹 Tarama sınırları (memo öncelikli)
        # =====================================================
        start_idx = None
        end_idx = None
        start_marker_text = None
        end_marker_text = None

        # ✅ Başlangıç önceliği
        if memo.get("chapter_first_idx") is not None:
            start_idx = int(memo["chapter_first_idx"])
            start_marker_text = paragraphs[start_idx].text.strip() if 0 <= start_idx < len(paragraphs) else "(chapter_first_idx out of range)"
        elif memo.get("first_chapter_title_idx") is not None:
            start_idx = int(memo["first_chapter_title_idx"])
            start_marker_text = paragraphs[start_idx].text.strip() if 0 <= start_idx < len(paragraphs) else "(first_chapter_title_idx out of range)"
        elif "chapter_blocks" in memo and memo["chapter_blocks"]:
            start_idx = min([b[0] for b in memo["chapter_blocks"] if b[0] is not None])
            start_marker_text = paragraphs[start_idx].text.strip() if 0 <= start_idx < len(paragraphs) else "(chapter_blocks start out of range)"
        else:
            start_marker_text = "(Bölüm başlığı bulunamadı)"

        # ✅ Bitiş: references_heading_idx (KAYNAKLAR)
        if "references_heading_idx" in memo:
            end_idx = memo["references_heading_idx"]
            end_marker_text = paragraphs[end_idx].text.strip() if 0 <= end_idx < len(paragraphs) else "(references_heading_idx out of range)"
        else:
            for i, p in enumerate(paragraphs):
                if norm_tr(p.text) == "kaynaklar":
                    end_idx = i
                    end_marker_text = p.text.strip()
                    break
            if end_idx is None:
                end_marker_text = "(KAYNAKLAR bulunamadı)"

        # --- Debug dosyası ---
        debug_file = open(dbg_path("subheading_debug.txt"), "w", encoding="utf-8") if debug_mode else None
        if debug_file:
            debug_file.write("=== ALT BAŞLIK DEBUG KAYDI ===\n\n")
            debug_file.write("🔹 Tarama Aralığı:\n")
            debug_file.write(f"   Başlangıç idx: {start_idx} ({start_marker_text if start_marker_text else 'YOK'})\n")
            debug_file.write(f"   Bitiş idx: {end_idx} ({end_marker_text if end_marker_text else 'YOK'})\n\n")
            debug_file.write(f"   Beklenen satır aralığı: {expected_spacing}\n")
            debug_file.write(f"   Heading styles (birebir): {heading_styles}\n\n")

        if start_idx is None or end_idx is None or start_idx >= end_idx:
            if debug_file:
                debug_file.write("⚠️ Tez metni sınırları belirlenemedi — tarama durduruldu.\n")
                debug_file.close()
            results.append((0, False, rule_title, "Tez metni sınırları belirlenemedi (chapter/references eksik)."))
            return results

        # =====================================================
        # 🔹 Hizalamayı etkin belirleyen yardımcılar
        # =====================================================
        def resolve_from_styles_parfmt(para, attr_name):
            val = getattr(para.paragraph_format, attr_name, None)
            if val is not None:
                return val
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None:
                    v = getattr(pf, attr_name, None)
                    if v is not None:
                        return v
                s = getattr(s, "base_style", None)
            return None

        def effective_alignment_strict(para):
            al = para.paragraph_format.alignment
            if al is not None:
                return al
            al = resolve_from_styles_parfmt(para, "alignment")
            if al is not None:
                return al
            jc = para._element.xpath(".//w:jc")
            if jc:
                val = jc[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                if val == "both":
                    return WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                elif val == "left":
                    return WD_PARAGRAPH_ALIGNMENT.LEFT
                elif val == "right":
                    return WD_PARAGRAPH_ALIGNMENT.RIGHT
                elif val == "center":
                    return WD_PARAGRAPH_ALIGNMENT.CENTER
            return WD_PARAGRAPH_ALIGNMENT.LEFT

        def dbg_para_props(p):
            fn_eff   = effective_font_name(p)
            fs_eff   = effective_font_size_pt(p)
            bold_eff = effective_bold(p)
            ls_eff   = effective_line_spacing(p, default=expected_spacing)  # ✅ YAML’dan
            sb_eff   = effective_space_pt(p, "before")
            sa_eff   = effective_space_pt(p, "after")
            al_eff   = effective_alignment_strict(p)
            style_name = p.style.name if p.style else ""
            return fn_eff, fs_eff, bold_eff, ls_eff, sb_eff, sa_eff, al_eff, style_name

        # =====================================================
        # 🔹 Alt başlık tespiti ve kontrol
        # =====================================================
        errors = []
        subheading_idxs = set()

        # ✅ Artık her numaralı başlık için depth + style da saklıyoruz
        # [(idx, nums_tuple, raw_text, depth, style_name)]
        subheading_numbered_items = []

        def parse_heading_number(s: str):
            m = re.match(r"^(\d+(?:\.\d+)+)\.?\s+", s.strip())
            if not m:
                return None
            return tuple(int(x) for x in m.group(1).split("."))

        def style_to_expected_depth(style_name: str):
            """
            'Heading 2' -> 2, 'Heading 3' -> 3 ... gibi beklenen numara derinliği.
            Eşleşmezse None döner.
            """
            m = re.match(r"^\s*Heading\s+(\d+)\s*$", style_name or "", re.IGNORECASE)
            if not m:
                return None
            try:
                return int(m.group(1))
            except Exception:
                return None

        for i in range(start_idx, end_idx):
            p = paragraphs[i]
            text = (p.text or "").strip()
            if not text:
                continue

            style_name = p.style.name if p.style else ""

            # ✅ Alt başlık kabul kriteri:
            # - Style birebir heading_styles içinde
            # - veya numara desenine uyuyor
            if style_name in heading_styles or numbered_heading_pattern.match(text):
                subheading_idxs.add(i)

                nums = parse_heading_number(text)
                if nums:
                    depth = len(nums)
                    subheading_numbered_items.append((i, nums, text.strip(), depth, style_name))

                    # ✅ Kılavuz: 6 derinlikten fazla olmamalı
                    if depth > 6:
                        head = short_text(text)
                        errors.append(f"'{head}' başlığı 6 seviyeden fazla derinliğe inemez (gelen derinlik={depth}).")

                fn, fs, bb, ls, sb, sa, al, st = dbg_para_props(p)

                if debug_file:
                    debug_file.write(
                        f"[{i+1}. paragraf] ALT BAŞLIK: {text}\n"
                        f"  Style: {st}\n"
                        f"  Font (eff): {fn}\n"
                        f"  Size (eff): {fs}\n"
                        f"  Bold (eff): {bb}\n"
                        f"  Line spacing (eff): {ls}\n"
                        f"  Before (eff): {sb}\n"
                        f"  After (eff): {sa}\n"
                        f"  Alignment (eff): {al}\n\n"
                    )

                head = short_text(text)

                # ✅ Biçim kontrolleri (mevcut yaklaşımı bozmuyoruz)
                if fn and fn != expected_name:
                    errors.append(f"'{head}' alt başlık fontu {fn} yerine {expected_name} olmalı")
                if fs and abs(fs - expected_size) > 0.1:
                    errors.append(f"'{head}' alt başlık punto {fs} yerine {expected_size} olmalı")
                if bool(bb) != bool(expected_bold):
                    errors.append(f"'{head}' alt başlık kalınlık uygun değil")
                if abs(ls - expected_spacing) > 0.1:
                    errors.append(f"'{head}' alt başlık satır aralığı {ls} yerine {expected_spacing} olmalı")
                if abs(sb - expected_before) > 1.0:
                    errors.append(f"'{head}' alt başlık öncesi {sb} yerine {expected_before} olmalı")
                if abs(sa - expected_after) > 1.0:
                    errors.append(f"'{head}' alt başlık sonrası {sa} yerine {expected_after} olmalı")

                if expected_align:
                    allowed = set(expected_align)
                    current = (
                        "left" if al == WD_PARAGRAPH_ALIGNMENT.LEFT else
                        "right" if al == WD_PARAGRAPH_ALIGNMENT.RIGHT else
                        "center" if al == WD_PARAGRAPH_ALIGNMENT.CENTER else
                        "justify" if al == WD_PARAGRAPH_ALIGNMENT.JUSTIFY else
                        "left"
                    )
                    if current not in allowed:
                        errors.append(f"'{head}' alt başlık hizalaması ({current}) yalnızca {', '.join(sorted(allowed))} olmalı")

        if debug_file:
            debug_file.write(f"=== KONTROL TAMAMLANDI ({len(subheading_idxs)} alt başlık bulundu) ===\n")
            debug_file.close()

        memo["subheading_forbidden_idxs"] = subheading_idxs
        memo["subheading_numbered_items"] = subheading_numbered_items

        if errors:
            results.append((0, False, rule_title, "; ".join(errors)))
        else:
            results.append((0, True, rule_title, f"{len(subheading_idxs)} alt başlık bulundu."))


    # ======================================================
    # ALT BAŞLIK NUMARALARI SIRA KONTROLÜ
    # (örn. 2.1→2.2, bölüm geçişi 2.x→3.1, 4.1.2→4.1.3, ... 6 derinliğe kadar)
    # - Hata mesajında: hangi başlıkta, ne bekleniyordu, ne görüldü
    # ======================================================
    elif check["check"] == "subheading_number_sequence":
        rule_title = "Alt başlık numaraları sıralı olmalı (örn. 2.1→2.2, bölüm geçişi 2.x→3.1)"

        items = memo.get("subheading_numbered_items", [])
        errors = []

        def fmt(nums):
            return ".".join(str(x) for x in nums)

        def short_text(s: str, n: int = 90) -> str:
            s = " ".join((s or "").strip().split())
            return (s[:n] + "...") if len(s) > n else s

        if not items:
            results.append((0, True, rule_title, "Numaralı alt başlık bulunmadı."))
            return results

        # ------------------------------------------------------
        # ✅ Geriye uyumluluk:
        # items elemanları ya (idx, nums, raw) ya da (idx, nums, raw, depth, style_name)
        # ------------------------------------------------------
        norm_items = []
        for t in items:
            if len(t) >= 5:
                idx, nums, raw = t[0], t[1], t[2]
                depth = t[3]
                style_name = t[4]
            else:
                idx, nums, raw = t[0], t[1], t[2]
                depth = len(nums) if nums else 0
                style_name = ""
            norm_items.append((idx, nums, raw, depth, style_name))

        # Paragraf sırasına göre sırala
        norm_items = sorted(norm_items, key=lambda x: x[0])

        # ------------------------------------------------------
        # ✅ Hiyerarşik sayaç mantığı (6 derinliğe kadar)
        #
        # Her başlık için:
        #   parent = nums[:-1]
        #   cur    = nums[-1]
        #
        # Kurallar:
        # 1) Aynı parent altında child sayacı 1,2,3,... artmalı
        # 2) Yeni bir parent altında ilk child = 1 olmalı
        # 3) Derine inmek (örn 3.1.1 → 3.1.1.1) üst seviyeyi artırmaz
        # 4) Derinlik 6'yı aşamaz
        # ------------------------------------------------------
        last_child = {}   # { parent_tuple: last_seen_child_int }
        prev_item = None  # (idx, nums, raw, depth, style_name)

        for idx, nums, raw, depth, style_name in norm_items:
            head = short_text(raw)

            # En az 2 seviye olmalı (2.1 gibi)
            if not nums or len(nums) < 2:
                prev_item = (idx, nums, raw, depth, style_name)
                continue

            # 6 derinlik kuralı
            if len(nums) > 6:
                errors.append(
                    f"Başlık: '{head}' → numara derinliği 6'yı aşamaz (gelen: {len(nums)})."
                )
                prev_item = (idx, nums, raw, depth, style_name)
                continue

            parent = tuple(nums[:-1])
            cur = int(nums[-1])

            # İlk kez görülen parent → child 1 olmalı
            if parent not in last_child:
                if cur != 1:
                    expected = fmt(parent + (1,))
                    got = fmt(nums)
                    errors.append(
                        f"Başlık: '{head}' → '{fmt(parent)}' altında ilk alt başlık {expected} olmalı, "
                        f"ama {got} geldi."
                    )
                last_child[parent] = cur
                prev_item = (idx, nums, raw, depth, style_name)
                continue

            # Aynı parent altında artış kontrolü
            expected_cur = last_child[parent] + 1
            if cur != expected_cur:
                expected = fmt(parent + (expected_cur,))
                got = fmt(nums)
                prev_head = short_text(prev_item[2]) if prev_item else ""
                errors.append(
                    f"Başlık: '{head}' → alt başlık sırası bozuk. "
                    f"Beklenen: {expected}, Gelen: {got}. "
                    f"Önceki: '{prev_head}'."
                )

            last_child[parent] = cur
            prev_item = (idx, nums, raw, depth, style_name)

        if errors:
            results.append((0, False, rule_title, "; ".join(errors)))
        else:
            results.append((0, True, rule_title, "Numaralı alt başlık sıralaması uygun."))



    # ======================================================
    # ŞEKİL BAŞLIKLARI TESPİTİ ve BİÇİMSEL KONTROLÜ - RESİM ODAKLI (CAPTION SADECE ALTA)
    # ======================================================
    elif check["check"] == "figure_caption_detector" and check.get("enabled", True):
        import re

        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 10))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 0))
        expected_align   = (check.get("alignment", "left") or "left").lower().strip()
        debug_mode       = check.get("debug", False)

        
        # 1) Detect: yanlış nokta / : / - olsa bile caption olarak yakala
        caption_detect_pat = re.compile(
            r"^(şekil|şek\.?|şekil\.?)\s*\d+(?:\.\d+)*\s*(?:[.:\-–])?\s+.+",
            re.IGNORECASE
        )

        # 2) Strict: beklenen doğru format = "Şekil X.Y Başlık"
        # - X.Y zorunlu
        # - Y’den sonra nokta, ':' veya '-' olmayacak
        caption_strict_pat = re.compile(
            r"^(şekil|şek\.?|şekil\.?)\s*\d+\.\d+\s+.+",
            re.IGNORECASE
        )


        rule_title = (
            f"Şekil Başlıkları (Tez Metni İçinde)\n"
            f"{expected_name}, {expected_size} punto, {expected_spacing} satır aralığı, "
            f"önce {expected_before}, sonra {expected_after}, {expected_align} hizalı"
        )

        errors = []
        figure_idxs = set()
        figure_captions = []
        seen_caption_idxs = set()

        def short_text(s: str, n: int = 90) -> str:
            s = " ".join((s or "").strip().split())
            return (s[:n] + "...") if len(s) > n else s

        def norm_tr(s: str) -> str:
            trans = str.maketrans("çğıöşüÇĞİÖŞÜ", "cgiosuCGIOSU")
            return (s or "").translate(trans).lower().strip()

        # Resim/SmartArt/Chart/Shape yakalama
        def para_has_figure_object(p):
            return bool(p._element.xpath(".//w:drawing") or p._element.xpath(".//w:pict"))

        # i’den başlayarak aşağı/yukarı ilk dolu paragrafı bul
        def find_next_nonempty(start_i, step):
            j = start_i
            while 0 <= j < len(paragraphs):
                t = (paragraphs[j].text or "").strip()
                if t:
                    return j
                j += step
            return None

        # Hizalama (stil + XML)
        def effective_alignment_strict(p):
            al = getattr(p.paragraph_format, "alignment", None)
            if al is not None:
                return al
            s = p.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf and pf.alignment is not None:
                    return pf.alignment
                s = getattr(s, "base_style", None)
            jc = p._element.xpath(".//w:jc")
            if jc:
                val = jc[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                if val == "both":
                    return WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                elif val == "left":
                    return WD_PARAGRAPH_ALIGNMENT.LEFT
                elif val == "right":
                    return WD_PARAGRAPH_ALIGNMENT.RIGHT
                elif val == "center":
                    return WD_PARAGRAPH_ALIGNMENT.CENTER
            return WD_PARAGRAPH_ALIGNMENT.LEFT

        # =====================================================
        # Tarama sınırları (memo öncelikli)
        # =====================================================
        start_idx = None
        end_idx = None

        if memo.get("chapter_first_idx") is not None:
            start_idx = int(memo["chapter_first_idx"])
        elif memo.get("first_chapter_title_idx") is not None:
            start_idx = int(memo["first_chapter_title_idx"])
        elif "chapter_blocks" in memo and memo["chapter_blocks"]:
            start_idx = min([b[0] for b in memo["chapter_blocks"] if b[0] is not None])

        if "references_heading_idx" in memo:
            end_idx = memo["references_heading_idx"]
        else:
            for i, p in enumerate(paragraphs):
                if norm_tr(p.text) == "kaynaklar":
                    end_idx = i
                    break

        debug_file = open(dbg_path("figure_debug.txt"), "w", encoding="utf-8") if debug_mode else None
        if debug_file:
            debug_file.write("=== ŞEKİL DEBUG (CAPTION SADECE ALTA) ===\n\n")
            debug_file.write(f"start_idx={start_idx}, end_idx={end_idx}\n\n")

        if start_idx is None or end_idx is None or start_idx >= end_idx:
            if debug_file:
                debug_file.write("⚠️ Sınırlar belirlenemedi.\n")
                debug_file.close()
            results.append((0, False, rule_title, "Tez metni sınırları belirlenemedi (chapter/references eksik)."))
            return results

        # Etkin biçim değerleri
        def dbg_para_props(p):
            fn_eff = effective_font_name(p)
            fs_eff = effective_font_size_pt(p)
            ls_eff = effective_line_spacing(p, default=expected_spacing)
            sb_eff = effective_space_pt(p, "before")
            sa_eff = effective_space_pt(p, "after")
            al_eff = effective_alignment_strict(p)
            style_name = p.style.name if p.style else ""
            return fn_eff, fs_eff, ls_eff, sb_eff, sa_eff, al_eff, style_name

        def err(head, msg):
            errors.append(f"'{head}': {msg}")

        # =====================================================
        # 1) Nesneyi bul → 2) SADECE ALT caption ara → 3) Üstte caption varsa ihlal yaz
        # =====================================================
        for i in range(start_idx, end_idx):
            p_obj = paragraphs[i]
            if not para_has_figure_object(p_obj):
                continue

            # Altındaki ilk dolu paragraf caption adayıdır
            cap_idx = find_next_nonempty(i + 1, step=+1)

            # Üstte caption var mı? (ihlal açıklaması için bakıyoruz)
            up_idx = find_next_nonempty(i - 1, step=-1)

            up_text = (paragraphs[up_idx].text or "").strip() if up_idx is not None else ""
            down_text = (paragraphs[cap_idx].text or "").strip() if cap_idx is not None else ""

            # 1) Alt caption yoksa
            if cap_idx is None or cap_idx >= end_idx:
                # Üstte caption yazılmışsa özellikle belirt
                if up_idx is not None and up_idx >= start_idx and caption_detect_pat.match(up_text):
                    err(short_text(up_text), "Şekil başlığı şeklin ÜSTÜNDE yazılmış. Kılavuza göre başlık şeklin ALTINDA olmalı.")
                else:
                    err("Şekil nesnesi", "Şekil bulundu ancak altında 'Şekil X.Y ...' başlığı bulunamadı.")
                continue

            # 2) Alt paragraf caption değilse
            if not caption_detect_pat.match(down_text):
                # Üstte caption varsa özellikle belirt
                if up_idx is not None and up_idx >= start_idx and caption_detect_pat.match(up_text):
                    err(short_text(up_text), "Şekil başlığı şeklin ÜSTÜNDE yazılmış. Kılavuza göre başlık şeklin ALTINDA olmalı.")
                else:
                    err(short_text(down_text), "Şekil nesnesinin hemen altında geçerli 'Şekil X.Y ...' başlığı yok.")
                continue

            # 3) Caption doğru yerde → biçim kontrolü
            if cap_idx in seen_caption_idxs:
                continue
            seen_caption_idxs.add(cap_idx)

            figure_idxs.add(cap_idx)
            figure_captions.append((cap_idx, down_text))

            fn, fs, ls, sb, sa, al, st = dbg_para_props(paragraphs[cap_idx])
            head = short_text(down_text)
            
            # Caption yakalandı ama beklenen "Şekil X.Y Başlık" formatında değilse ihlal yaz
            if not caption_strict_pat.match(down_text):
                err(head, "Şekil başlığı numaralandırma formatı hatalı. Beklenen: 'Şekil X.Y Başlık' "
                        "(Y’den sonra nokta, ':' veya '-' kullanılmamalı).")


            if debug_file:
                debug_file.write(f"Nesne idx={i+1} → caption idx={cap_idx+1}\n")
                debug_file.write(f"  Caption: {down_text}\n")
                debug_file.write(f"  Style: {st}\n")
                debug_file.write(f"  Font={fn}, Size={fs}, LS={ls}, Before={sb}, After={sa}, Align={al}\n\n")

            if fn and fn != expected_name:
                err(head, f"Şekil başlığı fontu {fn} yerine {expected_name} olmalı")
            if fs and abs(fs - expected_size) > 0.1:
                err(head, f"Şekil başlığı punto {fs} yerine {expected_size} olmalı")
            if abs(ls - expected_spacing) > 0.1:
                err(head, f"Şekil başlığı satır aralığı {ls} yerine {expected_spacing} olmalı")
            if abs(sb - expected_before) > 1.0:
                err(head, f"Şekil başlığı öncesi {sb} yerine {expected_before} olmalı")
            if abs(sa - expected_after) > 1.0:
                err(head, f"Şekil başlığı sonrası {sa} yerine {expected_after} olmalı")

            if expected_align == "left" and al != WD_PARAGRAPH_ALIGNMENT.LEFT:
                err(head, "Şekil başlığı sola hizalı olmalı")
            elif expected_align == "center" and al != WD_PARAGRAPH_ALIGNMENT.CENTER:
                err(head, "Şekil başlığı ortalı olmalı")
            elif expected_align == "justify" and al != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                err(head, "Şekil başlığı iki yana yaslı olmalı")
            elif expected_align == "right" and al != WD_PARAGRAPH_ALIGNMENT.RIGHT:
                err(head, "Şekil başlığı sağa hizalı olmalı")

        if debug_file:
            debug_file.write("=== KONTROL TAMAMLANDI ===\n")
            debug_file.close()

        memo["figure_forbidden_idxs"] = figure_idxs
        memo["figure_captions"] = figure_captions

        if errors:
            results.append((0, False, rule_title, "; ".join(errors)))
        else:
            results.append((0, True, rule_title, f"{len(figure_idxs)} şekil başlığı bulundu."))


    # ======================================================
    # ŞEKİL NUMARALANDIRMA KONTROLÜ (mükerrer veya sıra hatası)
    # ======================================================
    elif check["check"] == "figure_numbering_check" and check.get("enabled", True):
        import re
        from collections import defaultdict

        rule_title = "Şekil Numaralandırma Kontrolü"
        errors = []

        # Önceki aşamada bulunan şekiller alınır
        figure_captions = memo.get("figure_captions", [])
        if not figure_captions:
            results.append((0, False, rule_title, "Hiç şekil başlığı bulunamadı veya önceki kontrol devre dışıydı."))
            return results

        pattern = re.compile(r"şek(il)?\.?\s*(\d+)(?:\.(\d+))?", re.IGNORECASE)

        # parsed_figs: [(satır, text, main_no, sub_no, full_key)]
        # NOT: float kullanımı (3.10 -> 3.1) hatasına yol açtığı için full_key tuple tutulur.
        parsed_figs = []

        for i, text in figure_captions:
            m = pattern.search(text)
            if not m:
                continue

            main_no = int(m.group(2))                    # Bölüm numarası (ör. 3)
            sub_no  = int(m.group(3)) if m.group(3) else 0  # Alt numara (ör. 10); yoksa 0

            # >>> DÜZELTME: float YOK! (3.10 -> 3.1 hatası burada oluşuyordu)
            # full_key, "tam numarayı" kayıpsız temsil eder.
            full_key = (main_no, sub_no)

            parsed_figs.append((i + 1, text, main_no, sub_no, full_key))

        section_figs = defaultdict(list)
        for idx, text, main, sub, full_key in parsed_figs:
            # main sabit bölüm/grup anahtarı; listede sub ve full_key saklanır
            section_figs[main].append((idx, text, sub, full_key))

        # --- Mükerrer ve sıra atlama kontrolü ---
        for main_no, figs in section_figs.items():
            # >>> DÜZELTME: sıralamayı tuple/float ile değil, alt numara (sub) ile yap
            figs_sorted = sorted(figs, key=lambda x: x[2])  # x[2] = sub
            seen = set()  # >>> DÜZELTME: full_key (tuple) veya sub ile mükerrer kontrolü

            for j, (idx, text, sub, full_key) in enumerate(figs_sorted):
                label = f"Şekil {main_no}.{sub}"
                preview = text[:20].replace("\n", " ") + ("..." if len(text) > 20 else "")

                # Mükerrer kontrolü (kayıpsız anahtar ile)
                if full_key in seen:
                    errors.append(f"{idx}. satır ({preview}): {label} numarası mükerrer.")
                seen.add(full_key)

                # Sıra kontrolü (bir önceki sub + 1 olmalı)
                if j > 0:
                    prev_sub = figs_sorted[j - 1][2]
                    if sub != prev_sub + 1:
                        prev_idx = figs_sorted[j - 1][0]
                        errors.append(
                            f"{idx}. satır ({preview}): {label} numarası sıralı değil "
                            f"(önceki Şekil {main_no}.{prev_sub}, satır {prev_idx})."
                        )

        if errors:
            results.append((0, False, rule_title, "; ".join(errors)))
        else:
            results.append((0, True, rule_title, "Tüm şekil numaraları sıralı ve benzersiz."))



    # ======================================================
    # ÇİZELGE BAŞLIKLARI TESPİTİ ve BİÇİMSEL KONTROLÜ (Tez metni içinde) - GÜNCEL + ATIF FİLTRELİ
    # ======================================================
    elif check["check"] == "table_caption_detector" and check.get("enabled", True):
        import re

        # --- Beklenen biçimsel parametreler ---
        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 10))
        expected_spacing = float(check.get("line_spacing", 1.0))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 12))
        expected_align   = (check.get("alignment", "left") or "left").lower().strip()
        debug_mode       = check.get("debug", False)

        # --- Atıf/cümle filtresi parametreleri (YAML opsiyonel) ---
        max_words = int(check.get("max_words", 15))
        forbidden_contains = check.get("forbidden_contains", None)
        if not isinstance(forbidden_contains, list) or not forbidden_contains:
            # Varsayılanlar: metin içi atıfta sık geçen kalıplar (norm_tr ile karşılaştırılacak)
            forbidden_contains = [
                " ile ",
                " olarak ",
                " goster",         # göster, gösteril, gösterilmektedir...
                " gorul",          # görüldüğü
                " degerlendir",    # değerlendirildi
                " sunul",          # sunulmuştur
                " veril",          # verilmiştir
                " elde edilen",
                " kullan",         # kullanılmıştır
                " aciklan",        # açıklanmıştır
                " incelen",        # incelenmiştir
            ]

        rule_title = (
            f"Çizelge Başlıkları\n"
            f"{expected_name}, {expected_size} punto, {expected_spacing} satır aralığı, "
            f"önce {expected_before}, sonra {expected_after}, {expected_align} hizalı"
        )

        def short_text(s: str, n: int = 90) -> str:
            s = " ".join((s or "").strip().split())
            return (s[:n] + "...") if len(s) > n else s

        # --- Türkçe karakter normalizasyonu ---
        def norm_tr(s):
            trans = str.maketrans({
                "ı": "i", "İ": "i", "ç": "c", "Ç": "c", "ğ": "g", "Ğ": "g",
                "ö": "o", "Ö": "o", "ş": "s", "Ş": "s", "ü": "u", "Ü": "u"
            })
            return (s or "").lower().translate(trans).strip()

        # Daha güvenli: numaradan sonra metin olsun

        
        # 1) Detect: yanlış nokta / : / - olsa bile caption olarak yakala
        table_caption_detect_pat = re.compile(
            r"^cizelge\s*\d+(?:\.\d+)*\s*(?:[.:\-–])?\s+.+",
            re.IGNORECASE
        )

        # 2) Strict: beklenen doğru format = "Çizelge X.Y Başlık"
        # - X.Y zorunlu
        # - Y’den sonra nokta, ':' veya '-' olmayacak
        table_caption_strict_pat = re.compile(
            r"^cizelge\s*\d+\.\d+\s+.+",
            re.IGNORECASE
)


        # =====================================================
        # 🔹 Tarama sınırları (memo öncelikli)
        # =====================================================
        start_idx = None
        end_idx = None

        if memo.get("chapter_first_idx") is not None:
            start_idx = int(memo["chapter_first_idx"])
        elif memo.get("first_chapter_title_idx") is not None:
            start_idx = int(memo["first_chapter_title_idx"])
        elif "chapter_blocks" in memo and memo["chapter_blocks"]:
            start_idx = min([b[0] for b in memo["chapter_blocks"] if b[0] is not None])

        if "references_heading_idx" in memo:
            end_idx = memo["references_heading_idx"]
        else:
            for i, p in enumerate(paragraphs):
                if norm_tr(p.text) == "kaynaklar":
                    end_idx = i
                    break

        # --- Debug dosyası ---
        debug_file = open(dbg_path("table_debug.txt"), "w", encoding="utf-8") if debug_mode else None
        if debug_file:
            debug_file.write("=== ÇİZELGE BAŞLIKLARI DEBUG (ATIF FİLTRELİ) ===\n\n")
            debug_file.write(f"start_idx={start_idx}, end_idx={end_idx}\n")
            debug_file.write(f"max_words={max_words}\n")
            debug_file.write(f"forbidden_contains={forbidden_contains}\n\n")

        if start_idx is None or end_idx is None or start_idx >= end_idx:
            if debug_file:
                debug_file.write("⚠️ Başlangıç veya bitiş bulunamadı — tarama durduruldu.\n")
                debug_file.close()
            results.append((0, False, rule_title, "Tez metni sınırları belirlenemedi (chapter/references eksik)."))
            return results

        # --- Stil zincirinden etkin biçim çözümleme ---
        def resolve_from_styles_parfmt(para, attr):
            val = getattr(para.paragraph_format, attr, None)
            if val is not None:
                return val
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf:
                    v = getattr(pf, attr, None)
                    if v is not None:
                        return v
                s = getattr(s, "base_style", None)
            return None

        def effective_line_spacing_local(para, default=1.0):
            val = resolve_from_styles_parfmt(para, "line_spacing")
            if val is None:
                return float(default)
            if hasattr(val, "pt"):
                return round(val.pt / 12.0, 2)
            try:
                return round(float(val), 2)
            except Exception:
                return float(default)

        def effective_space_pt_local(para, which):
            attr = "space_before" if which == "before" else "space_after"
            val = resolve_from_styles_parfmt(para, attr)
            if val is None:
                return 0.0
            return round(val.pt, 1)

        def effective_alignment_strict(para):
            al = getattr(para.paragraph_format, "alignment", None)
            if al is not None:
                return al
            al = resolve_from_styles_parfmt(para, "alignment")
            if al is not None:
                return al
            jc = para._element.xpath(".//w:jc")
            if jc:
                val = jc[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                if val == "both":
                    return WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                elif val == "left":
                    return WD_PARAGRAPH_ALIGNMENT.LEFT
                elif val == "right":
                    return WD_PARAGRAPH_ALIGNMENT.RIGHT
                elif val == "center":
                    return WD_PARAGRAPH_ALIGNMENT.CENTER
            return WD_PARAGRAPH_ALIGNMENT.LEFT

        def align_to_str(al):
            return (
                "left" if al == WD_PARAGRAPH_ALIGNMENT.LEFT else
                "right" if al == WD_PARAGRAPH_ALIGNMENT.RIGHT else
                "center" if al == WD_PARAGRAPH_ALIGNMENT.CENTER else
                "justify" if al == WD_PARAGRAPH_ALIGNMENT.JUSTIFY else
                "left"
            )

        # --- Metin içi atıf filtresi ---
        def looks_like_intext_reference(original_text: str) -> (bool, str):
            """
            Çizelge başlığı gibi başlayan ama aslında metin içi atıf olan cümleleri elemek için.
            """
            t = (original_text or "").strip()
            if not t:
                return False, ""

            # 1) Nokta ile bitiyorsa genellikle cümledir
            if t.rstrip().endswith("."):
                return True, "Nokta ile bitiyor (cümle/atıf olasılığı yüksek)"

            # 2) Çok uzunsa cümle olma olasılığı artar
            wc = len(t.split())
            if wc > max_words:
                return True, f"Kelime sayısı {wc} > {max_words} (cümle/atıf olasılığı yüksek)"

            # 3) Yasaklı kalıp/fiil/bağlaç içeriyorsa
            low = f" {norm_tr(t)} "
            for frag in forbidden_contains:
                f = (frag or "").strip()
                if not f:
                    continue
                # frag’ı da norm_tr uyumlu kullan (Türkçe karakter vs.)
                f2 = f" {norm_tr(f)} "
                if f2.strip() and f2 in low:
                    return True, f"Metin içi atıf kalıbı içeriyor: '{f.strip()}'"

            return False, ""

        # --- Ana tarama ---
        errors = []
        table_forbidden_idxs = set()
        memo.setdefault("table_captions", [])  # debug bağımsız

        for i in range(start_idx, end_idx):
            p = paragraphs[i]
            text = (p.text or "").strip()
            if not text:
                continue

            norm = norm_tr(text)

            # Liste işaretleri vs. atla
            if text.startswith(("•", "-", "·")):
                continue

            # "Çizelgeler Dizini" vb. atla
            if "cizelgeler" in norm and "dizin" in norm:
                continue

            # Caption formuna uyuyor mu?
            if not table_caption_detect_pat.match(norm):
                continue

            # ✅ Metin içi atıf mı? (ise caption sayma)
            is_ref, why = looks_like_intext_reference(text)
            if is_ref:
                if debug_file:
                    debug_file.write(f"[{i+1}] ATIF OLARAK ELENDİ: {short_text(text,120)} | neden: {why}\n")
                continue

            # Buraya geldiyse: caption olarak kabul
            table_forbidden_idxs.add(i)

            fn = para_font_name(p)
            fs = para_font_size_pt(p)
            ls = effective_line_spacing_local(p, default=expected_spacing)
            sb = effective_space_pt_local(p, "before")
            sa = effective_space_pt_local(p, "after")
            al = effective_alignment_strict(p)

            memo["table_captions"].append({
                "index": i,  # 0-based
                "text": text,
                "font": fn,
                "size": fs,
                "line_spacing": ls,
                "space_before": sb,
                "space_after": sa,
                "alignment": align_to_str(al),
            })

            if debug_file:
                debug_file.write(
                    f"[{i+1}] ÇİZELGE CAPTION: {short_text(text,120)}\n"
                    f"  Font={fn}, Size={fs}, LS={ls}, Before={sb}, After={sa}, Align={al}\n\n"
                )

            head = short_text(text)


            def err(msg):
                errors.append(f"'{head}': {msg}")

            # Caption yakalandı ama beklenen "Çizelge X.Y Başlık" formatında değilse ihlal yaz
            if not table_caption_strict_pat.match(norm):
                err("Çizelge başlığı numaralandırma formatı hatalı. Beklenen: 'Çizelge X.Y Başlık' "
                    "(Y’den sonra nokta, ':' veya '-' kullanılmamalı).")


            # Biçimsel denetim
            if fn and fn != expected_name:
                err(f"Yazı tipi {fn} yerine {expected_name} olmalı")
            if fs and abs(fs - expected_size) > 0.1:
                err(f"Punto {fs} yerine {expected_size} olmalı")
            if abs(ls - expected_spacing) > 0.1:
                err(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
            if abs(sb - expected_before) > 1.0:
                err(f"Öncesi {sb} yerine {expected_before} olmalı")
            if abs(sa - expected_after) > 1.0:
                err(f"Sonrası {sa} yerine {expected_after} olmalı")

            cur = align_to_str(al)
            if expected_align and cur != expected_align:
                err(f"Hizalama {cur} yerine {expected_align} olmalı")

        if debug_file:
            debug_file.write(f"=== TARAMA TAMAMLANDI (caption={len(table_forbidden_idxs)}) ===\n")
            debug_file.close()

        memo["table_forbidden_idxs"] = table_forbidden_idxs


        if errors:
            results.append((0, False, rule_title, "; ".join(errors)))
        else:
            if len(table_forbidden_idxs) == 0:
                results.append((0, False, rule_title, "Tez metni içinde hiç çizelge başlığı bulunamadı."))
            else:
                results.append((0, True, rule_title, f"Tüm {len(table_forbidden_idxs)} çizelge başlığı biçimsel olarak uygun."))




        return results

    # ======================================================
    # ÇİZELGE NUMARALANDIRMA KONTROLÜ (X.Y zorunlu)
    # ======================================================
    elif check["check"] == "table_numbering_check":
        """
        table_caption_detector tarafından toplanan çizelge başlıklarının
        numaralandırmasını kontrol eder:
        - Biçim: Çizelge X.Y (zorunlu)
        - Numara sırası (1.1 → 1.2 → 2.1 ...)
        - Tekrarlanan numaralar
        - Bölüm numarası atlamaları / geriye gitme
        """

        import re

        rule_title = "Çizelge Numaralandırma Kontrolü"

        
        tables = memo.get("table_captions", [])
        if not tables:
            results.append((0, True, rule_title,
                            "Tez metni içinde çizelge başlığı bulunamadı; numaralandırma kontrolü uygulanmadı."))
            return results


        def short_text(s: str, n: int = 90) -> str:
            s = " ".join((s or "").strip().split())
            return (s[:n] + "...") if len(s) > n else s

        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı": "i", "İ": "i", "ç": "c", "Ç": "c", "ğ": "g", "Ğ": "g",
                "ö": "o", "Ö": "o", "ş": "s", "Ş": "s", "ü": "u", "Ü": "u"
            })
            return (s or "").lower().translate(trans).strip()

        # X.Y zorunlu: "cizelge 2.1 ..."
        num_pat_xy = re.compile(r"^cizelge\s+(\d+)\.(\d+)\b", re.IGNORECASE)
        # Tek seviyeli yakalamak için (ihlal mesajını netleştirmek)
        num_pat_x  = re.compile(r"^cizelge\s+(\d+)\b", re.IGNORECASE)

        seen = set()
        prev_main = None
        prev_sub  = None
        errors = []

        for item in tables:
            text = item.get("text", "") or ""
            head = short_text(text)
            tnorm = norm_tr(text)

            m = num_pat_xy.search(tnorm)
            if not m:
                # "Çizelge 2 ..." gibi tek seviye mi?
                m1 = num_pat_x.search(tnorm)
                if m1:
                    errors.append(f"'{head}': Numara biçimi 'Çizelge X.Y' olmalı (örn. Çizelge {m1.group(1)}.1).")
                else:
                    errors.append(f"'{head}': Geçerli numara biçimi bulunamadı (beklenen: Çizelge X.Y).")
                continue

            main_no = int(m.group(1))
            sub_no  = int(m.group(2))
            num_str = f"{main_no}.{sub_no}"

            # Tekrar kontrolü
            if num_str in seen:
                errors.append(f"'{head}': {num_str} numarası tekrarlanmış.")
            else:
                seen.add(num_str)

            # Sıra kontrolü
            if prev_main is not None:
                if main_no == prev_main:
                    # aynı bölüm: sub +1
                    expected = f"{prev_main}.{prev_sub + 1}"
                    if sub_no != prev_sub + 1:
                        errors.append(f"'{head}': {num_str} beklenen {expected} olmalı.")
                elif main_no == prev_main + 1:
                    # yeni bölüm: sub = 1
                    expected = f"{main_no}.1"
                    if sub_no != 1:
                        errors.append(f"'{head}': {num_str} beklenen {expected} olmalı (yeni bölüm).")
                elif main_no > prev_main + 1:
                    errors.append(f"'{head}': {num_str} bölüm atlaması var (önceki bölüm {prev_main}).")
                else:
                    # geriye gitme
                    errors.append(f"'{head}': {num_str} numarası geriye gitmiş (önceki {prev_main}.{prev_sub}).")

            prev_main = main_no
            prev_sub  = sub_no

        if errors:
            results.append((0, False, rule_title, "; ".join(errors)))
        else:
            results.append((0, True, rule_title, f"Tüm {len(tables)} çizelge başlığı X.Y biçiminde, sıralı ve benzersiz."))

        return results

    # ======================================================
    # TEZ METNİ GÖVDE PARAGRAF BİÇİM KONTROLÜ (GİRİŞ → KAYNAKLAR)
    # ======================================================
    elif check["check"] == "body_paragraph_format":
        import re

        # --- Beklenen biçimsel parametreler ---
        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_spacing = float(check.get("line_spacing", 1.5))
        expected_before  = float(check.get("space_before", 6))
        expected_after   = float(check.get("space_after", 6))
        expected_indent  = float(check.get("first_line_indent_cm", 1.25))
        expected_align   = (check.get("alignment", "justify") or "justify").lower().strip()

        start_marker = check.get("start_marker", "^GİRİŞ$")
        end_marker   = check.get("end_marker", "^KAYNAKLAR$")
        debug_mode   = bool(check.get("debug", False))
        skip_words   = [str(w).lower() for w in check.get("skip_contains", [])]

        # --- Denklem / Cambria Math istisnaları (YAML opsiyonel) ---
        skip_fonts = check.get("skip_fonts", ["Cambria Math"])
        if isinstance(skip_fonts, str):
            skip_fonts = [skip_fonts]
        skip_fonts_norm = [str(x).strip().lower() for x in (skip_fonts or []) if str(x).strip()]

        # Sembol taraması opsiyonel (varsayılan açık: güvenli)
        skip_if_contains_math_symbols = check.get("skip_if_contains_math_symbols", True)

        # 🔹 Bölüm, alt başlık ve şekil/çizelge başlıklarını hariç tut (memo)
        forbidden_idxs = set(memo.get("chapter_forbidden_idxs", []))
        forbidden_idxs.update(set(memo.get("subheading_forbidden_idxs", [])))
        forbidden_idxs.update(set(memo.get("figure_forbidden_idxs", [])))
        forbidden_idxs.update(set(memo.get("table_forbidden_idxs", [])))

        # --- Şekil/Tablo başlığı desenleri (ek savunma) ---
        fig_pat   = re.compile(r"^(şek(il)?\.?)\s*\d+(\.\d+)*", re.IGNORECASE)
        table_pat = re.compile(r"^(tablo|çizelge)\s*\d+(\.\d+)*", re.IGNORECASE)

        rule_title = (
            f"TEZ METNİ Gövdesi\n"
            f"{expected_name}, {expected_size} punto, {expected_spacing} satır aralığı, "
            f"önce {expected_before}, sonra {expected_after}, "
            f"ilk satır girintisi {expected_indent} cm, iki yana yaslı"
        )

        # ------------------------------------------------------
        # Yardımcı fonksiyonlar: Stil zincirinden etkin biçim al
        # ------------------------------------------------------
        def resolve_from_styles_parfmt(para, attr_name):
            val = getattr(para.paragraph_format, attr_name, None)
            if val is not None:
                return val
            s = para.style
            while s is not None:
                pf = getattr(s, "paragraph_format", None)
                if pf is not None:
                    v = getattr(pf, attr_name, None)
                    if v is not None:
                        return v
                s = getattr(s, "base_style", None)
            return None

        def effective_alignment(para):
            al = resolve_from_styles_parfmt(para, "alignment")
            return al if al is not None else WD_PARAGRAPH_ALIGNMENT.LEFT

        def effective_line_spacing(para, default=None):
            if default is None:
                default = expected_spacing
            val = resolve_from_styles_parfmt(para, "line_spacing")
            if val is None:
                return float(default)
            if hasattr(val, "pt"):
                return round(val.pt / 12.0, 2)
            try:
                return round(float(val), 2)
            except Exception:
                return float(default)

        def effective_space_pt(para, which="before"):
            attr = "space_before" if which == "before" else "space_after"
            val = resolve_from_styles_parfmt(para, attr)
            if val is None:
                return 0.0
            return round(val.pt, 1)

        def effective_first_line_indent_cm(para):
            val = resolve_from_styles_parfmt(para, "first_line_indent")
            if val is None:
                return 0.0
            return round(val.cm, 2)

        # --- Türkçe karakter normalizasyonu ---
        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı": "i", "İ": "i", "ç": "c", "Ç": "c", "ğ": "g", "Ğ": "g",
                "ö": "o", "Ö": "o", "ş": "s", "Ş": "s", "ü": "u", "Ü": "u"
            })
            return (s or "").lower().translate(trans).strip()

        def preview_text(s: str, n: int = 35) -> str:
            s = " ".join((s or "").strip().split())
            return s if len(s) <= n else (s[:n] + "...")

        # --- Denklem tespiti: Cambria Math run var mı? ---
        def para_has_skip_font(para) -> bool:
            try:
                for r in para.runs:
                    # r.font.name bazen None olabilir
                    fn = (r.font.name or "").strip().lower()
                    if fn and fn in skip_fonts_norm:
                        return True
            except Exception:
                pass
            return False

        # --- Denklem tespiti: sembol taraması (opsiyonel) ---
        math_symbols_pat = re.compile(r"[=<>±×÷∑∫√∞≈≠≤≥∂∆∇→←↔·•^_{}()\[\]]")
        def looks_like_equation_text(t: str) -> bool:
            if not t:
                return False
            # çok kısa metinlerde gereksiz skip olmasın diye bir eşik
            if len(t) < 4:
                return False
            return bool(math_symbols_pat.search(t))

        # ======================================================
        # 🔹 Başlangıç ve Bitiş Paragrafları (memo öncelikli)
        # ======================================================
        start_idx = None
        end_idx = None

        # Start öncelik:
        if memo.get("chapter_first_idx") is not None:
            start_idx = int(memo["chapter_first_idx"])
        elif memo.get("first_chapter_title_idx") is not None:
            start_idx = int(memo["first_chapter_title_idx"]) + 1
        elif "chapter_blocks" in memo and memo["chapter_blocks"]:
            first_block_i = min([b[0] for b in memo["chapter_blocks"] if b[0] is not None])
            start_idx = int(first_block_i) + 2

        # End öncelik:
        if memo.get("references_heading_idx") is not None:
            end_idx = int(memo["references_heading_idx"])
        elif memo.get("references_heading_block_idx") is not None:
            end_idx = int(memo["references_heading_block_idx"])

        # Fallback marker arama
        if start_idx is None or end_idx is None:
            for i, p in enumerate(paragraphs):
                t = norm_tr(p.text)
                if start_idx is None and re.match(norm_tr(start_marker), t, re.IGNORECASE):
                    start_idx = i + 1
                if end_idx is None and re.match(norm_tr(end_marker), t, re.IGNORECASE):
                    end_idx = i
                    break

        if start_idx is None or end_idx is None or start_idx >= end_idx:
            results.append((0, False, rule_title,
                            f"Başlangıç/bitiş bulunamadı veya aralık hatalı. start={start_idx}, end={end_idx}"))
            return results

        # --- Debug ---
        debug_file = None
        if debug_mode:
            debug_file = open(dbg_path("body_debug.txt"), "w", encoding="utf-8")
            debug_file.write("=== TEZ METNİ GÖVDE PARAGRAF DEBUG KAYDI ===\n\n")
            debug_file.write(f"Tarama Aralığı: start={start_idx}, end={end_idx}\n")
            debug_file.write(f"skip_fonts={skip_fonts}\n")
            debug_file.write(f"skip_if_contains_math_symbols={skip_if_contains_math_symbols}\n\n")

        errors = []

        for i in range(start_idx, end_idx):
            p = paragraphs[i]
            text = (p.text or "").strip()
            if not text:
                continue

            # 1) Başlık/alt başlık/şekil/çizelge başlığı atla (memo)
            if i in forbidden_idxs:
                if debug_mode and debug_file:
                    debug_file.write(f"[{i}] (ATLANDI-memo) {preview_text(text, 80)}\n")
                continue

            # 2) Şekil/çizelge başlığı atla (pattern yedek)
            if fig_pat.match(text) or table_pat.match(text):
                if debug_mode and debug_file:
                    debug_file.write(f"[{i}] (ATLANDI-pattern) {preview_text(text, 80)}\n")
                continue

            # 3) Denklem/Cambria Math atla
            if para_has_skip_font(p):
                if debug_mode and debug_file:
                    debug_file.write(f"[{i}] (ATLANDI-skip_font) {preview_text(text, 80)}\n")
                continue

            # 4) Denklem sembol taraması (opsiyonel)
            if skip_if_contains_math_symbols and looks_like_equation_text(text):
                if debug_mode and debug_file:
                    debug_file.write(f"[{i}] (ATLANDI-math_symbols) {preview_text(text, 80)}\n")
                continue

            # 5) Liste/numaralı öğe atla
            if p._element.xpath(".//w:numPr"):
                if debug_mode and debug_file:
                    debug_file.write(f"[{i}] (ATLANDI-liste) {preview_text(text, 80)}\n")
                continue

            # 6) skip_contains
            nt = norm_tr(text)
            if any(word in nt for word in skip_words):
                if debug_mode and debug_file:
                    debug_file.write(f"[{i}] (ATLANDI-skip_contains) {preview_text(text, 80)}\n")
                continue

            # 7) Negatif girinti atla
            pf = p.paragraph_format
            if pf.first_line_indent and pf.first_line_indent.cm < 0:
                if debug_mode and debug_file:
                    debug_file.write(f"[{i}] (ATLANDI-negatif_girinti) {preview_text(text, 80)}\n")
                continue

            # --- Etkin biçim değerleri ---
            fn  = para_font_name(p)
            fs  = para_font_size_pt(p)
            ls  = effective_line_spacing(p, default=expected_spacing)
            sb  = effective_space_pt(p, "before")
            sa  = effective_space_pt(p, "after")
            ind = effective_first_line_indent_cm(p)
            al  = effective_alignment(p)

            if debug_mode and debug_file:
                debug_file.write(
                    f"[{i}] {preview_text(text, 80)}\n"
                    f"  Font={fn}, Size={fs}, LS={ls}, SB={sb}, SA={sa}, IND={ind}, AL={al}\n\n"
                )

            head = preview_text(text, 35)

            def err(msg):
                errors.append(f"'{head}': {msg}")

            if fn and fn != expected_name:
                err(f"Yazı tipi {fn} yerine {expected_name} olmalı")
            if fs and abs(fs - expected_size) > 0.1:
                err(f"Punto {fs} yerine {expected_size} olmalı")
            if abs(ls - expected_spacing) > 0.1:
                err(f"Satır aralığı {ls} yerine {expected_spacing} olmalı")
            if abs(sb - expected_before) > 1.0:
                err(f"Öncesi {sb} yerine {expected_before} olmalı")
            if abs(sa - expected_after) > 1.0:
                err(f"Sonrası {sa} yerine {expected_after} olmalı")
            if abs(ind - expected_indent) > 0.1:
                err(f"İlk satır girintisi {ind} yerine {expected_indent} olmalı")

            if expected_align == "justify" and al != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                err("İki yana yaslı değil")
            elif expected_align == "left" and al != WD_PARAGRAPH_ALIGNMENT.LEFT:
                err("Sola hizalı değil")
            elif expected_align == "center" and al != WD_PARAGRAPH_ALIGNMENT.CENTER:
                err("Ortalanmış değil")
            elif expected_align == "right" and al != WD_PARAGRAPH_ALIGNMENT.RIGHT:
                err("Sağa hizalı değil")

        if debug_file:
            debug_file.write("\n=== KONTROL TAMAMLANDI ===\n")
            debug_file.close()

        if errors:
            results.append((start_idx, False, rule_title, "; ".join(errors)))
        else:
            results.append((start_idx, True, rule_title, "Tüm gövde paragrafları biçimsel olarak uygun."))


    # ============================================================
    # LİSTE  BİÇİM KONTROLÜ (madde/numara)
    # - Sadece liste paragrafları (w:numPr) kontrol edilir
    # - Tez metni sınırları memo'dan alınır (GİRİŞ→KAYNAKLAR)
    # - Rapor: "satır/loc" yok, sadece metin + ölçü
    # ============================================================
    elif check["check"] == "list_paragraph_format":
        import re
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as _ALIGN

        debug_mode = bool(check.get("debug", False))
        debug_file = open(dbg_path("list_debug.txt"), "w", encoding="utf-8") if debug_mode else None

        tol = float(check.get("tolerance_cm", 0.05))
        expected_left = float(check.get("text_indent_cm", 1.27))     # metin başlangıcı
        expected_hang = float(check.get("bullet_indent_cm", 0.63))   # numara/bullet konumu (asılı girinti)

        # --- Tez metni sınırları (memo öncelikli) ---
        start_idx = None
        end_idx = None

        # Start öncelik:
        if memo.get("chapter_first_idx") is not None:
            start_idx = int(memo["chapter_first_idx"])
        elif memo.get("first_chapter_title_idx") is not None:
            start_idx = int(memo["first_chapter_title_idx"]) + 1
        elif "chapter_blocks" in memo and memo["chapter_blocks"]:
            first_block_i = min([b[0] for b in memo["chapter_blocks"] if b[0] is not None])
            start_idx = int(first_block_i) + 2

        # End öncelik:
        if memo.get("references_heading_idx") is not None:
            end_idx = int(memo["references_heading_idx"])
        elif memo.get("references_heading_block_idx") is not None:
            end_idx = int(memo["references_heading_block_idx"])

        # Fallback yoksa: tüm dokümanı taramak yerine güvenli davran
        if start_idx is None or end_idx is None or start_idx >= end_idx:
            # İstersen burada "tüm dokümanı tara" da diyebilirdik
            # ama yanlış pozitifleri azaltmak için kontrollü davranıyoruz.
            if debug_file:
                debug_file.write(
                    f"⚠️ Tez metni sınırları bulunamadı. start={start_idx}, end={end_idx}\n"
                    "Liste kontrolü güvenlik için durduruldu.\n"
                )
                debug_file.close()
            results.append((0, False, "Liste biçimi kontrolü", "Tez metni sınırları belirlenemedi (memo eksik)."))
            return results

        def preview_text(s: str, n: int = 30) -> str:
            s = " ".join((s or "").strip().split())
            return s if len(s) <= n else (s[:n] + "...")

        # --- Etkin indent değerleri: varsa senin fonksiyonlarını kullan, yoksa fallback ---
        def safe_left_indent_cm(p):
            try:
                return float(effective_left_indent_cm(p))
            except Exception:
                # fallback: paragraph_format.left_indent
                li = getattr(p.paragraph_format, "left_indent", None)
                return round(li.cm, 2) if li is not None else 0.0

        def safe_hanging_indent_cm(p):
            try:
                return float(effective_hanging_indent_cm(p))
            except Exception:
                hi = getattr(p.paragraph_format, "first_line_indent", None)
                # hanging indent genelde first_line_indent negatif olur
                if hi is None:
                    return 0.0
                try:
                    v = float(hi.cm)
                    return abs(v)  # negatifse mutlak al
                except Exception:
                    return 0.0

        errors = []

        # ✅ Sadece tez metni aralığında tara
        for idx in range(start_idx, end_idx):
            p = paragraphs[idx]

            # sadece liste paragraflarını hedefle
            if not p._element.xpath(".//w:numPr"):
                continue

            text = (p.text or "").strip()

            # ölçümler
            left_cm = safe_left_indent_cm(p)
            hang_cm = safe_hanging_indent_cm(p)

            # 🔹 python-docx bazı numaralı listelerde değerleri ters/0 verebiliyor — düzelt
            if left_cm == 0 and hang_cm > 0.8:
                left_cm, hang_cm = hang_cm, expected_hang
            if left_cm == 0 and hang_cm == 0:
                left_cm, hang_cm = expected_left, expected_hang
            if hang_cm < 0:
                hang_cm = abs(hang_cm)

            # --- Word otomatik numaralı liste istisnası ---
            try:
                numPr = p._p.pPr.numPr
            except Exception:
                numPr = None

            if hang_cm == 0 and numPr is not None:
                if debug_file:
                    debug_file.write(
                        f"[idx={idx}] (otomatik liste) {preview_text(text, 80) or '(boş)'}\n"
                        f"  Hanging indent 0 cm ama numPr algılandı → hata bastırıldı.\n\n"
                    )
                continue

            # --- Debug kaydı ---
            if debug_file:
                debug_file.write(
                    f"[idx={idx}] (liste)\n"
                    f"  Önizleme: {preview_text(text, 80) or '(boş)'}\n"
                    f"  Left indent: {left_cm:.2f} cm (beklenen ~{expected_left:.2f})\n"
                    f"  Hanging indent: {hang_cm:.2f} cm (beklenen ~{expected_hang:.2f})\n\n"
                )

            # --- Kural ihlali kontrolü ---
            if abs(left_cm - expected_left) > tol or abs(hang_cm - expected_hang) > tol:
                head = preview_text(text, 35) or "(boş paragraf)"
                errors.append(
                    f"'{head}': Metin girintisi {left_cm:.2f} cm (beklenen {expected_left:.2f}), "
                    f"Numara konumu {hang_cm:.2f} cm (beklenen {expected_hang:.2f})"
                )

        if debug_file:
            debug_file.write("\n=== LİSTE KONTROLÜ TAMAMLANDI ===\n")
            debug_file.close()

        rule_title = "Liste biçimi: numara 0.63 cm, metin 1.27 cm girintide olmalı"
        if errors:
            results.append((0, False, rule_title, "; ".join(errors)))
        else:
            results.append((0, True, rule_title, "Tüm liste paragrafları biçimsel olarak uygun."))

        return results


    # ======================================================
    # ÖZGEÇMİŞ BAŞLIĞI TESPİTİ ve BİÇİMSEL KONTROLÜ (SONDAN BAŞA + VARYANT KABULLÜ)
    # - Tez içinde ÖZGEÇMİŞ sayfası bulunmalı
    # - Arama sondan başa doğru yapılır (en sonda olmalı)
    # - "ÖZ GEÇMİŞ", "Ozgecmis", "Özgeçmiş" vb. varyantları yakalar
    # - Bulunsa bile tam olarak "ÖZGEÇMİŞ" yazılması gerektiğini ihlal olarak belirtir
    # ======================================================
    elif check["check"] == "cv_heading_block" and check.get("enabled", True):
        import re

        # 1) YAML parametrelerini al
        expected_name    = check.get("font_name", "Times New Roman")
        expected_size    = float(check.get("font_size_pt", 12))
        expected_bold    = check.get("bold", True)
        expected_caps    = check.get("all_caps", True)
        expected_align   = check.get("alignment", "center").lower()
        expected_spacing = float(check.get("line_spacing", 1.5))
        expected_before  = float(check.get("space_before", 0))
        expected_after   = float(check.get("space_after", 24))
        debug_mode       = check.get("debug", False)

        # Opsiyonel: stil adı kontrolü (verilmezse kontrol etmez)
        expected_style   = check.get("expected_style", None)

        # 2) Varyant listesi (tespit için) – bulunursa “ÖZGEÇMİŞ olmalı” diye uyaracağız
        accepted_variants = check.get("accepted_variants", None)
        if not accepted_variants:
            accepted_variants = [
                "OZGEÇMİŞ", "OZGEÇMIS", "OZGECMIS",
                "ÖZGEÇMIS", "ÖZGEÇMİŞ.",
                "ÖZ GEÇMİŞ", "ÖZ GECMIS", "OZ GECMIS",
                "ÖZGEÇMİŞ:", "ÖZGEÇMİŞ -", "ÖZGEÇMİŞ—",
                "ÖZGEÇMİŞ (CV)", "ÖZGEÇMİŞ (ÖZET)",
            ]

        # 3) “En sonda olmalı” toleransı (fallback):
        #    KAYNAKLAR memo'da YOKSA, docx'te sayfa bilgisi net olmadığı için pratik yaklaşım olarak
        #    belgenin sonundan itibaren son N dolu paragraf penceresi kullanılacak.
        tail_nonempty_limit = int(check.get("tail_nonempty_limit", 80))

        rule_title = (
            f"ÖZGEÇMİŞ Başlığı Biçimsel Kontrolü (En Sonda Olmalı)\n"
            f"{expected_name}, {expected_size} punto, kalın={expected_bold}, "
            f"BÜYÜK HARF={expected_caps}, {expected_align}, "
            f"{expected_spacing} satır aralığı, önce {expected_before}, sonra {expected_after}"
            + (f", stil={expected_style}" if expected_style else "")
        )

        # 4) TR normalizasyon: harfleri sadeleştir, küçült, boşluk/punkt. temizle
        def norm_tr(s: str) -> str:
            trans = str.maketrans({
                "ı": "i", "İ": "i", "ç": "c", "Ç": "c", "ğ": "g", "Ğ": "g",
                "ö": "o", "Ö": "o", "ş": "s", "Ş": "s", "ü": "u", "Ü": "u"
            })
            s = (s or "").lower().translate(trans)
            # boşluk ve noktalama farklarını da yok say
            s = re.sub(r"[\s\.\:\-\–\—\(\)\[\]\{\}]+", "", s)
            return s.strip()

        # 5) Hedef ve varyantları normalize et
        target_raw  = "ÖZGEÇMİŞ"
        target_norm = norm_tr(target_raw)
        variants_norm = [norm_tr(v) for v in (accepted_variants or [])]

        found_idx = None
        found_raw_text = None
        found_is_variant = False

        # ----------------------------------------------------------
        # 5.5) ARAMA PENCERESİNİ BELİRLE (GÜNCELLEME)
        # Amaç:
        #   - KAYNAKLAR başlığını memo'ya kaydettiğimiz için (references_heading_idx),
        #     son sayfaya yakın bölüm aramalarında "sondan -> KAYNAKLAR'a kadar" tarama yapabiliriz.
        #
        # Çalışma:
        #   - Eğer memo["references_heading_idx"] varsa:
        #       ÖZGEÇMİŞ'i sadece KAYNAKLAR'dan SONRAKİ kısımda ararız:
        #       i = len(paragraphs)-1 ... (references_heading_idx+1)
        #   - Eğer KAYNAKLAR yoksa:
        #       fallback olarak eski mantık (tail_nonempty_limit) ile son N dolu paragraf penceresinde ararız.
        # ----------------------------------------------------------
        ref_idx = memo.get("references_heading_idx", None)

        search_start = len(paragraphs) - 1  # her durumda sondan başlayacağız
        search_end   = 0                    # default: en başa kadar

        if ref_idx is not None:
            # KAYNAKLAR başlığı bulunduysa, aramayı KAYNAKLAR'dan sonraki kısım ile sınırla
            try:
                ref_idx_int = int(ref_idx)
            except Exception:
                ref_idx_int = None

            if ref_idx_int is not None and 0 <= ref_idx_int < len(paragraphs):
                # KAYNAKLAR paragrafının bir sonrasına kadar (inclusive döngü mantığı için search_end = ref_idx+1)
                search_end = ref_idx_int + 1
            else:
                # ref_idx bozuksa güvenli fallback: en başa kadar ara
                search_end = 0
                ref_idx_int = None
        else:
            ref_idx_int = None

        # Eğer KAYNAKLAR yoksa, fallback penceresini hesapla (son N dolu paragraf)
        # ve aramayı sadece bu pencereye sıkıştır.
        if ref_idx_int is None:
            nonempty_idxs_from_end = []
            for j in range(len(paragraphs) - 1, -1, -1):
                if (paragraphs[j].text or "").strip():
                    nonempty_idxs_from_end.append(j)
                if len(nonempty_idxs_from_end) >= tail_nonempty_limit:
                    break

            if nonempty_idxs_from_end:
                # pencerenin en eski (küçük) index'i: search_end olarak kullan
                search_end = min(nonempty_idxs_from_end)
            else:
                search_end = 0

        # 6) Sondan başa ara: önce tam hedef, sonra varyant
        #    (GÜNCELLEME: artık tüm doküman yerine search_start -> search_end aralığında arıyoruz)
        for i in range(search_start, search_end - 1, -1):
            raw = (paragraphs[i].text or "").strip()
            if not raw:
                continue
            if norm_tr(raw) == target_norm:
                found_idx = i
                found_raw_text = raw
                found_is_variant = False
                break

        if found_idx is None:
            for i in range(search_start, search_end - 1, -1):
                raw = (paragraphs[i].text or "").strip()
                if not raw:
                    continue
                if norm_tr(raw) in variants_norm:
                    found_idx = i
                    found_raw_text = raw
                    found_is_variant = True
                    break

        # 7) Bulunamadıysa: net hata
        if found_idx is None:
            msg = "⚠️ 'ÖZGEÇMİŞ' başlığı bulunamadı. (ÖZ GEÇMİŞ / Özgeçmiş / Ozgecmis vb. varyantlar da dahil aranmıştır.)"
            results.append((0, False, rule_title, msg))
            return results

        # 8) “En sonda olmalı” kontrolü (GÜNCELLEME):
        #    - Eğer KAYNAKLAR memo'da varsa: arama zaten KAYNAKLAR'dan sonraki kısımda yapıldığı için
        #      bu, pratikte "tez sonu" bölgesinde bulunduğunu garanti eder → tail_window_ok = True
        #    - Eğer KAYNAKLAR yoksa: eski yöntem (son N dolu paragraf penceresi içinde mi?) ile kontrol edilir
        if ref_idx_int is not None:
            tail_window_ok = True
        else:
            # Fallback: eski mantık (son N dolu paragraf listesinde başlık çok geride kalmasın)
            nonempty_idxs_from_end = []
            for j in range(len(paragraphs) - 1, -1, -1):
                if (paragraphs[j].text or "").strip():
                    nonempty_idxs_from_end.append(j)
                if len(nonempty_idxs_from_end) >= tail_nonempty_limit:
                    break

            # başlık, bu “son N dolu paragraf” penceresinin içinde değilse -> en sonda değil
            tail_window_ok = (found_idx in set(nonempty_idxs_from_end))

        # 9) Başlık paragrafının effective özellikleri
        p = paragraphs[found_idx]
        fn = effective_font_name(p)
        fs = effective_font_size_pt(p)
        bb = effective_bold(p)
        ls = effective_line_spacing(p, default=expected_spacing)
        sb = effective_space_pt(p, "before")
        sa = effective_space_pt(p, "after")
        al = effective_alignment(p)
        txt = (p.text or "").strip()
        style_name = p.style.name if getattr(p, "style", None) else ""

        # 10) Hataları topla
        errors = []

        # 10A) Metin tam eşleşme: bulundu ama tam “ÖZGEÇMİŞ” değilse ihlal yaz
        if txt != target_raw:
            # varyant yakalanmışsa daha özellikle vurgula
            if found_is_variant:
                errors.append(f"Başlık bulundu ancak '{txt}' yazılmış. Doğru yazım tam olarak 'ÖZGEÇMİŞ' olmalı")
            else:
                errors.append(f"Başlık metni tam olarak 'ÖZGEÇMİŞ' olmalı (bulunan: '{txt}')")

        # 10B) En sonda olmalı
        if not tail_window_ok:
            errors.append(f"ÖZGEÇMİŞ sayfası tezde en sonda olmalı (başlık çok erken konumda görünüyor)")

        # 10C) Biçim kontrolleri
        if fn and expected_name and fn != expected_name:
            errors.append(f"Yazı tipi {expected_name} olmalı (bulunan: {fn})")

        if fs and abs(fs - expected_size) > 0.1:
            errors.append(f"Punto {expected_size} olmalı (bulunan: {fs})")

        if expected_bold is True and bb is not True:
            errors.append("Kalın (bold) olmalı")
        if expected_bold is False and bb is True:
            errors.append("Kalın (bold) olmamalı")

        if expected_caps:
            # Türkçe büyük harf kontrolü: en güvenlisi ham metni kontrol etmek
            if txt != txt.upper():
                errors.append("Tamamı BÜYÜK HARF olmalı")

        # hizalama
        if expected_align == "center" and al != WD_PARAGRAPH_ALIGNMENT.CENTER:
            errors.append("Ortalanmış olmalı")
        elif expected_align == "left" and al != WD_PARAGRAPH_ALIGNMENT.LEFT:
            errors.append("Sola hizalı olmalı")
        elif expected_align == "right" and al != WD_PARAGRAPH_ALIGNMENT.RIGHT:
            errors.append("Sağa hizalı olmalı")
        elif expected_align == "justify" and al != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            errors.append("İki yana yaslı olmalı")

        # satır aralığı / önce-sonra
        if ls is not None and abs(ls - expected_spacing) > 0.05:
            errors.append(f"Satır aralığı {expected_spacing} olmalı (bulunan: {ls})")
        if sb is not None and abs(sb - expected_before) > 0.5:
            errors.append(f"Paragraf öncesi {expected_before} pt olmalı (bulunan: {sb} pt)")
        if sa is not None and abs(sa - expected_after) > 0.5:
            errors.append(f"Paragraf sonrası {expected_after} pt olmalı (bulunan: {sa} pt)")

        # opsiyonel stil adı
        if expected_style:
            if (style_name or "").strip() != expected_style.strip():
                errors.append(f"Stil adı '{expected_style}' olmalı (bulunan: '{style_name}')")

        # 11) memo’ya kaydet (sonraki kurallar gerekirse kullanır)
        memo["cv_heading_idx"] = found_idx
        memo["cv_heading_block_idx"] = found_idx

        # 12) Sonuç
        if errors:
            results.append((found_idx, False, rule_title, "; ".join(errors)))
        else:
            results.append((found_idx, True, rule_title, "ÖZGEÇMİŞ başlığı bulundu ve biçimsel olarak uygun."))


# KONTROL SONUÇLARINI DÖNDÜR
# ======================================================

    return results # SON DÖNÜŞ KODU SATIRI



# ========================================================================================================================
# ========================================================================================================================
# ========================================================================================================================
# 2. add_metadata_block
# ============================================================
def add_metadata_block(doc, report_config, docx_filename, rules_filename):
    report = report_config.get("report", {})
    info_text = report.get("info_block")
    if info_text:
        para = doc.add_paragraph(info_text)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para.space_after = Pt(12)
    if report.get("show_datetime", False):
        now = datetime.now().strftime("%d.%m.%Y %H:%M")
        para = doc.add_paragraph(f"Tarih: {now}")
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para.space_after = Pt(6)
    if report.get("show_filename", False):
        para = doc.add_paragraph(f"Tez Dosyası: {docx_filename}")
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para.space_after = Pt(12)


# ========================================================================================================================
# ========================================================================================================================
# ========================================================================================================================
    """
    The `create_report` function generates a PDF report with Turkish font support, structured tables,
    and visual indicators for success/failure statuses.
    :return: The `create_report` function returns the file path of the generated PDF report.
    """
# 3. create_report
# ============================================================
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
import re
from datetime import datetime

def init_turkish_pdf_fonts():
    """
    Türkçe karakter desteği için TTF fontları PDF'ye gömer.
    Öncelik:
      1) Proje içi ./fonts/DejaVuSans*.ttf
      2) Linux sistem fontları (Render)
      3) Windows Arial (lokal)
    Olmazsa Helvetica'ya düşer (Türkçe sorun çıkarabilir).
    """
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    normal_font_name = "Helvetica"
    bold_font_name = "Helvetica-Bold"

    base_dir = os.path.dirname(os.path.abspath(__file__))

    candidates = [
        # 1) Repo içi (önerilen)
        (os.path.join(base_dir, "fonts", "DejaVuSans.ttf"),
         os.path.join(base_dir, "fonts", "DejaVuSans-Bold.ttf"),
         "TR_DJV", "TR_DJV_BOLD"),

        # 2) Render / Linux'ta sık görülen sistem font yolları
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
         "TR_DJV", "TR_DJV_BOLD"),

        # 3) Windows (lokal)
        (r"C:\Windows\Fonts\arial.ttf",
         r"C:\Windows\Fonts\arialbd.ttf",
         "TR_ARIAL", "TR_ARIAL_BOLD"),
    ]

    for normal_path, bold_path, n_name, b_name in candidates:
        try:
            if os.path.exists(normal_path):
                pdfmetrics.registerFont(TTFont(n_name, normal_path))
                normal_font_name = n_name

                if os.path.exists(bold_path):
                    pdfmetrics.registerFont(TTFont(b_name, bold_path))
                    bold_font_name = b_name
                else:
                    bold_font_name = normal_font_name

                # ✅ Başarılı yükledik → çık
                return normal_font_name, bold_font_name
        except Exception as e:
            # sıradaki adaya geç
            print(f"Font yükleme denemesi başarısız: {normal_path} -> {e}")

    print("Türkçe font yüklenemedi, Helvetica kullanılacak. (Türkçe karakterler bozulabilir)")
    return normal_font_name, bold_font_name



def wrap_text(text, font_name, font_size, max_width):
    """
    Verilen metni, font ve maksimum genişliğe göre satırlara böler.
    max_width: point cinsinden genişlik (canvas koordinat sistemi).
    """
    from reportlab.pdfbase import pdfmetrics

    text = str(text or "")
    if not text.strip():
        return [""]

    words = text.split()
    lines = []
    current = ""

    for w in words:
        test = (current + " " + w).strip()
        if pdfmetrics.stringWidth(test, font_name, font_size) <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = w

    if current:
        lines.append(current)

    return lines or [""]



def create_report(report_config, results_by_section, docx_filename, rules_filename, student_name=None, app_version_text=None):
    """
    🎨 MODERN 2026 PDF RAPOR OLUŞTURUCU
    =====================================
    Kontrol sonuçlarını profesyonel, modern ve estetik bir PDF raporu olarak üretir.
    
    Özellikler:
    - Gradient başlıklı kapak sayfası
    - Öğrenci adı ve tez başlığı tezden alınır
    - İstatistik özet sayfası
    - Modern tablo tasarımı (zebra striping, yuvarlak köşeli uyarılar)
    - Kurumsal renk paleti (2026 estetiki)
    """
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    
    global memo
    if 'memo' not in globals() or memo is None:
        memo = {}
    if not student_name:
        student_name = memo.get("student_name", "ÖĞRENCİ")
    
    student_name_display = (memo.get("student_name") or student_name or "Bilinmiyor").strip()
    student_name_file = student_name.strip().upper().replace(" ", "_")
    pdf_filename = f"RAPOR_{student_name_file}_{timestamp}.pdf"
    
    # Türkçe fontları başlat
    normal_font, bold_font = init_turkish_pdf_fonts()
    
    # Raporlar klasörü
    base_dir = os.path.dirname(os.path.abspath(__file__))
    reports_dir = os.path.join(base_dir, "Raporlar")
    os.makedirs(reports_dir, exist_ok=True)
    pdf_path = os.path.join(reports_dir, pdf_filename)
    
    c = canvas.Canvas(pdf_path, pagesize=A4)
    width, height = A4
    
    # ════════════════════════════════════════════════════════════════
    # 🎨 MODERN RENK PALETİ (2026 ESTETİĞİ)
    # ════════════════════════════════════════════════════════════════
    COLORS = {
        'primary_dark': (0.05, 0.15, 0.30),      # Koyu Lacivert
        'primary': (0.08, 0.24, 0.45),            # Lacivert
        'primary_light': (0.12, 0.35, 0.55),      # Açık Lacivert
        'accent': (0.70, 0.12, 0.15),             # Kurumsal Bordo
        'accent_light': (0.85, 0.20, 0.22),       # Açık Bordo
        'success': (0.10, 0.58, 0.38),            # Yeşil
        'success_bg': (0.92, 0.98, 0.94),         # Açık Yeşil BG
        'warning': (0.92, 0.60, 0.15),            # Turuncu
        'warning_bg': (1.0, 0.96, 0.88),          # Açık Turuncu BG
        'danger': (0.82, 0.15, 0.18),             # Kırmızı
        'danger_bg': (1.0, 0.94, 0.94),           # Açık Kırmızı BG
        'text_dark': (0.12, 0.14, 0.18),          # Koyu Metin
        'text_medium': (0.35, 0.38, 0.42),        # Orta Metin
        'text_light': (0.55, 0.58, 0.62),         # Açık Metin
        'bg_light': (0.97, 0.98, 0.99),           # Açık Arka Plan
        'bg_card': (1.0, 1.0, 1.0),               # Kart Arka Plan
        'border_light': (0.88, 0.90, 0.92),       # Açık Kenarlık
        'gold': (0.85, 0.68, 0.22),               # Altın
        'divider' : (0.55, 0.57, 0.60),
        # OKÜ kurumsal kırmızı
        'oku_red_dark': (0.40, 0.08, 0.08),   # #661414 (daha koyu, üst kısım için)
        'oku_red':      (0.545, 0.118, 0.118),# #8B1E1E (ana kırmızı)
        'oku_red_light':(0.639, 0.149, 0.149) # #A32626 (alt kısım için daha açık)

    }
    
    # ════════════════════════════════════════════════════════════════
    # 🛠️ YARDIMCI FONKSİYONLAR
    # ════════════════════════════════════════════════════════════════
    
    def draw_gradient_rect(c, x, y, w, h, color1, color2, steps=50):
        """Dikey gradient dikdörtgen çizer"""
        step_h = h / steps
        for i in range(steps):
            r = color1[0] + (color2[0] - color1[0]) * i / steps
            g = color1[1] + (color2[1] - color1[1]) * i / steps
            b = color1[2] + (color2[2] - color1[2]) * i / steps
            c.setFillColorRGB(r, g, b)
            c.rect(x, y + h - (i + 1) * step_h, w, step_h + 0.5, fill=1, stroke=0)
    
    def draw_rounded_card(c, x, y, w, h, radius=8, fill_color=None, stroke_color=None, shadow=True):
        """Gölgeli yuvarlak köşeli kart çizer"""
        c.saveState()
        if shadow:
            # Gölge
            c.setFillColorRGB(0.85, 0.85, 0.88)
            c.roundRect(x + 2, y - 2, w, h, radius, fill=1, stroke=0)
        if fill_color:
            c.setFillColorRGB(*fill_color)
        if stroke_color:
            c.setStrokeColorRGB(*stroke_color)
            c.setLineWidth(0.5)
        c.roundRect(x, y, w, h, radius, fill=1 if fill_color else 0, stroke=1 if stroke_color else 0)
        c.restoreState()
    
    def draw_stat_card(c, x, y, w, h, value, label, icon, color):
        """Modern istatistik kartı çizer - düzeltilmiş layout"""
        c.saveState()
        # Kart arka planı
        draw_rounded_card(c, x, y, w, h, radius=10, fill_color=COLORS['bg_card'], stroke_color=COLORS['border_light'], shadow=True)
        
        # Sol renkli şerit
        c.setFillColorRGB(*color)
        c.roundRect(x, y, 5, h, 2, fill=1, stroke=0)
        
        # Değer (büyük, üstte)
        c.setFont(bold_font, 24)
        c.setFillColorRGB(*COLORS['text_dark'])
        c.drawString(x + 15, y + h - 30, str(value))
        
        # Etiket (küçük, altta)
        c.setFont(normal_font, 9)
        c.setFillColorRGB(*COLORS['text_medium'])
        c.drawString(x + 15, y + 8, label)
        c.restoreState()
    
    def draw_modern_doughnut(c, cx, cy, radius, percentage):
        """Modern doughnut chart çizer"""
        c.saveState()
        
        # Renk seçimi
        if percentage >= 90:
            main_color = COLORS['success']
        elif percentage >= 70:
            main_color = COLORS['warning']
        else:
            main_color = COLORS['danger']
        
        # Arka plan dairesi
        c.setLineWidth(14)
        c.setStrokeColorRGB(0.92, 0.93, 0.95)
        c.setLineCap(1)
        c.arc(cx - radius, cy - radius, cx + radius, cy + radius, 0, 360)
        
        # Değer dairesi
        c.setStrokeColorRGB(*main_color)
        angle = (percentage / 100.0) * 360
        c.arc(cx - radius, cy - radius, cx + radius, cy + radius, 90, -angle)
        
        # Merkez daire (beyaz)
        c.setFillColorRGB(1, 1, 1)
        c.circle(cx, cy, radius - 12, fill=1, stroke=0)
        
        # Merkez yazı
        c.setFont(bold_font, 32)
        c.setFillColorRGB(*COLORS['text_dark'])
        c.drawCentredString(cx, cy + 5, f"%{percentage:.1f}")
        
        c.setFont(normal_font, 11)
        c.setFillColorRGB(*COLORS['text_medium'])
        c.drawCentredString(cx, cy - 15, "UYUM ORANI")
        c.restoreState()
    
    def draw_progress_bar(c, x, y, w, h, percentage, color):
        """Modern ilerleme çubuğu çizer"""
        c.saveState()
        # Arka plan
        c.setFillColorRGB(0.92, 0.93, 0.95)
        c.roundRect(x, y, w, h, h/2, fill=1, stroke=0)
        # Değer
        fill_w = w * (percentage / 100.0)
        if fill_w > h:
            c.setFillColorRGB(*color)
            c.roundRect(x, y, fill_w, h, h/2, fill=1, stroke=0)
        c.restoreState()
    
    def draw_tick(c, cx, cy, size_pt=8, rgb=(0, 0.6, 0)):
        """Yeşil tik işareti çizer"""
        c.saveState()
        c.setFillColorRGB(0.92, 0.98, 0.94)
        c.setStrokeColorRGB(*rgb)
        c.setLineWidth(0.5)
        c.circle(cx, cy, size_pt * 1.1, fill=1, stroke=1)
        c.setStrokeColorRGB(*rgb)
        c.setLineWidth(1.5)
        s = size_pt
        c.line(cx - 0.5*s, cy, cx - 0.1*s, cy - 0.35*s)
        c.line(cx - 0.1*s, cy - 0.35*s, cx + 0.55*s, cy + 0.45*s)
        c.restoreState()
    
    def draw_cross(c, cx, cy, size_pt=8, rgb=(0.82, 0.15, 0.18)):
        """Kırmızı çarpı işareti çizer"""
        c.saveState()
        c.setFillColorRGB(1.0, 0.94, 0.94)
        c.setStrokeColorRGB(*rgb)
        c.setLineWidth(0.5)
        c.circle(cx, cy, size_pt * 1.1, fill=1, stroke=1)
        c.setStrokeColorRGB(*rgb)
        c.setLineWidth(1.5)
        s = size_pt
        c.line(cx - 0.4*s, cy - 0.4*s, cx + 0.4*s, cy + 0.4*s)
        c.line(cx - 0.4*s, cy + 0.4*s, cx + 0.4*s, cy - 0.4*s)
        c.restoreState()
    
    def draw_watermark(c):
        """Filigran çizer"""
        logo_path = os.path.join(base_dir, "static", "logo.png")
        if os.path.exists(logo_path):
            c.saveState()
            c.setFillAlpha(0.06)
            w_w = 16 * cm
            w_h = 16 * cm
            c.drawImage(logo_path, (width - w_w) / 2, (height - w_h) / 2,
                        width=w_w, height=w_h, mask='auto', preserveAspectRatio=True)
            c.restoreState()
    
    def draw_page_footer(c, page_num):
        """Sayfa altbilgisi çizer"""
        c.saveState()
        # Alt çizgi
        c.setStrokeColorRGB(*COLORS['border_light'])
        c.setLineWidth(0.5)
        c.line(2*cm, 1.8*cm, width - 2*cm, 1.8*cm)
        
        # Sol: Tarih
        c.setFont(normal_font, 8)
        c.setFillColorRGB(*COLORS['text_light'])
        c.drawString(2*cm, 1.2*cm, f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        
        # Orta: Kurum
        c.drawCentredString(width/2 + 2*cm, 1.2*cm, "Osmaniye Korkut Ata Üniversitesi - Lisansüstü Eğitim Enstitüsü")
        
        # Sağ: Sayfa no
        c.drawRightString(width - 2*cm, 1.2*cm, f"Sayfa {page_num}")
        c.restoreState()
    
    # ============================================================
    # 📊 İSTATİSTİK HESAPLAMALARI
    # ════════════════════════════════════════════════════════════════
    
    total_checked = 0
    total_ok = 0
    total_fail = 0
    
    table_cols_global = report_config["report"].get("table_columns", [])
    
    def _find_col_idx(cols, candidates):
        cols_lower = [str(c).strip().lower() for c in cols]
        for cand in candidates:
            if cand.lower() in cols_lower:
                return cols_lower.index(cand.lower())
        return None
    
    e_idx_global = _find_col_idx(table_cols_global, ["Evet", "E", "Yes"])
    h_idx_global = _find_col_idx(table_cols_global, ["Hayır", "Hayir", "H", "No"])
    
    # Tüm kurallar üzerinden unweighted toplamlar (Stat Kartları için)
    for sk in report_config["report"].get("section_order", report_config["report"].get("order", [])):
        sec = results_by_section.get(sk, [])
        for res in sec:
            row = list(res)
            total_checked += 1
            if e_idx_global is not None and e_idx_global < len(row):
                v = row[e_idx_global]
                s = str(v).strip()
                if s in ("✔", "✓", "E", "EVET", "True", "1") or (isinstance(v, bool) and v is True):
                    total_ok += 1
            if h_idx_global is not None and h_idx_global < len(row):
                v = row[h_idx_global]
                s = str(v).strip()
                if s in ("✘", "✗", "H", "HAYIR", "False", "0") or (isinstance(v, bool) and v is False):
                    total_fail += 1
    
    # ✅ AĞIRLIKLI GENEL BAŞARI HESABI (Yenikontrol.py logic)
    # 1) Rapor sırası: report.yaml'deki section_order (yoksa order)
    order = report_config["report"].get("section_order", report_config["report"].get("order", []))
    
    # 2) Ön/arka grup tanımları
    front_keys = {
        "general", "inner_cover", "approval", "ethics", 
        "abstract_tr", "abstract_en", "acknowledgements", 
        "toc", "list_of_tables", "list_of_figures", "symbols_abbreviations"
    }
    back_keys = {"references", "appendices", "cv"}
    
    # 3) Tez metni (body)
    body_keys = [k for k in order if (k not in front_keys and k not in back_keys)]
    
    # 4) Bir grubun yüzde hesabı: o gruptaki tüm kurallar içinde OK oranı
    def _group_pct(keys):
        group_total = 0
        group_ok = 0
        for sk in keys:
            sec = results_by_section.get(sk, []) or []
            for res in sec:
                row = list(res)
                group_total += 1
                if e_idx_global is not None and e_idx_global < len(row):
                    v = row[e_idx_global]
                    s = str(v).strip()
                    if s in ("✔", "✓", "E", "EVET", "True", "1") or (isinstance(v, bool) and v is True):
                        group_ok += 1
        return (group_ok / group_total * 100.0) if group_total > 0 else 0.0
    
    # 5) Grup yüzdelerini hesapla
    front_pct = _group_pct(front_keys)
    body_pct = _group_pct(body_keys)
    back_pct = _group_pct(back_keys)
    
    # 6) AĞIRLIKLI genel yüzde
    success_pct = 0.15 * front_pct + 0.80 * body_pct + 0.05 * back_pct

    # ════════════════════════════════════════════════════════════════
    # 📄 SAYFA 1: MODERN KAPAK SAYFASI (2026 ESTETİĞİ)
    # ════════════════════════════════════════════════════════════════
    
    # Filigran
    draw_watermark(c)
    
    # --- Gradient Header (Üst) ---
    header_h = 3.2 * cm
    draw_gradient_rect(c, 0, height - header_h, width, header_h, 
                       COLORS['oku_red_dark'], COLORS['oku_red'])
    
    # Altın şerit
    c.saveState()
    c.setFillColorRGB(*COLORS['divider'])
    c.rect(0, height - header_h - 0.15*cm, width, 0.15*cm, fill=1, stroke=0)
    c.restoreState()
    
    # Logo (Header içinde sol)
    logo_path = os.path.join(base_dir, "static", "logo.png")
    if os.path.exists(logo_path):
        c.drawImage(logo_path, 1.8*cm, height - 2.8*cm, width=2.2*cm, height=2.2*cm, 
                    mask='auto', preserveAspectRatio=True)
    
    # Kurum adı (Header içinde sağda)
    c.saveState()
    c.setFont(bold_font, 18)
    c.setFillColorRGB(1, 1, 1)
    c.drawString(4.5*cm, height - 1.8*cm, "OSMANİYE KORKUT ATA ÜNİVERSİTESİ")
    c.setFont(normal_font, 12)
    c.setFillColorRGB(0.85, 0.88, 0.92)
    c.drawString(4.5*cm, height - 2.5*cm, "Lisansüstü Eğitim Enstitüsü")
    c.restoreState()
    
    y = height - header_h - 1.5*cm
    
    # --- Ana Başlık Kartı (Merkezi, Gölgeli) ---
    card_w = 15 * cm
    card_h = 4.5 * cm
    card_x = (width - card_w) / 2
    card_y = y - card_h - 0.5*cm
    
    draw_rounded_card(c, card_x, card_y, card_w, card_h, radius=12, 
                      fill_color=COLORS['bg_card'], stroke_color=COLORS['border_light'], shadow=True)
    
    # Üst dekoratif çizgi
    c.saveState()
    c.setFillColorRGB(*COLORS['accent'])
    c.roundRect(card_x + 0.3*cm, card_y + card_h - 0.4*cm, card_w - 0.6*cm, 0.25*cm, 2, fill=1, stroke=0)
    c.restoreState()
    
    # Ana Başlık
    c.saveState()
    c.setFont(bold_font, 26)
    c.setFillColorRGB(*COLORS['primary'])
    c.drawCentredString(width/2, card_y + card_h - 1.5*cm, "TEZ YAZIM KURALLARI")
    
    c.setFont(bold_font, 22)
    c.setFillColorRGB(*COLORS['accent'])
    c.drawCentredString(width/2, card_y + card_h - 2.3*cm, "KONTROL RAPORU")
    c.restoreState()
    
    y = card_y - 1.2*cm
    
    # --- Öğrenci Bilgi Kartı ---
    info_card_w = 15 * cm
    info_card_h = 4.6 * cm
    info_card_x = (width - info_card_w) / 2
    info_card_y = y - info_card_h
    
    draw_rounded_card(c, info_card_x, info_card_y, info_card_w, info_card_h, radius=10,
                      fill_color=COLORS['bg_light'], stroke_color=COLORS['border_light'], shadow=False)
    
    # Sol: Etiketler ve değerler
    c.saveState()
    c.setFont(bold_font, 11)
    c.setFillColorRGB(*COLORS['primary'])
    c.drawString(info_card_x + 0.8*cm, info_card_y + info_card_h - 1.0*cm, "ÖĞRENCİ:")
    
    c.setFont(bold_font, 14)
    c.setFillColorRGB(*COLORS['text_dark'])
    c.drawString(info_card_x + 3.5*cm, info_card_y + info_card_h - 1.0*cm, student_name_display)
    
    # TEZ BAŞLIĞI etiketi
    thesis_title = (memo.get("thesis_title") or "Tez Başlığı Belirlenemedi").strip()
    
    
    
        # --- TEZ BAŞLIĞI: Etiket + dinamik genişlik + dinamik satır sayısı ---

    # Etiket konumu
    label_x = info_card_x + 0.8*cm
    label_y = info_card_y + info_card_h - 2.0*cm

    # Etiket
    c.setFont(bold_font, 11)
    c.setFillColorRGB(*COLORS['primary'])
    label_txt = "TEZ BAŞLIĞI:"
    c.drawString(label_x, label_y, label_txt)

    # Etiket genişliğine göre metnin başlayacağı x (dinamik)
    label_w = pdfmetrics.stringWidth(label_txt, bold_font, 11)
    gap = 0.35 * cm
    value_x = label_x + label_w + gap

    # Başlık metni için kullanılabilir genişlik (sağdan 0.8cm pay bırak)
    max_w = (info_card_x + info_card_w - 0.8*cm) - value_x
    if max_w < 2*cm:   # emniyet (çok dar olursa)
        max_w = 2*cm

    # Metni wrap et
    c.setFont(normal_font, 10)
    c.setFillColorRGB(*COLORS['text_dark'])
    title_lines = wrap_text(thesis_title, normal_font, 10, max_w)

    # Kaç satır sığar? (kartın altından güvenli pay bırakarak)
    line_height = 0.45 * cm
    title_top_y = label_y
    title_bottom_y = info_card_y + 0.6*cm  # alt boşluk
    max_title_height = title_top_y - title_bottom_y
    max_lines = max(1, int(max_title_height // line_height))

    # Eğer çok satır varsa, karta sığdığı kadar yaz + sonuna … koy
    visible_lines = title_lines[:max_lines]
    if len(title_lines) > max_lines and visible_lines:
        visible_lines[-1] = (visible_lines[-1].rstrip() + "…")

    # Çiz
    ty = label_y
    for i, ln in enumerate(visible_lines):
        c.drawString(value_x, ty - i*line_height, ln)

    
    

        
    c.restoreState()
    
    y = info_card_y - 1.5*cm
    
    # --- Dashboard Alanı (Doughnut + İstatistikler) ---
    
    # Doughnut Chart (Sol)
    doughnut_cx = 5.5 * cm
    doughnut_cy = y - 3 * cm
    draw_modern_doughnut(c, doughnut_cx, doughnut_cy, 2.5*cm, success_pct)
    

    # Sağ: Bölüm performans çubukları (stat kartlar kalktı → alanı dolduralım)
    stat_y = y - 0.5*cm
    bar_x = 9.2 * cm          # daha sola al
    bar_w = 9.0 * cm          # genişlet (sağa kadar güzel doldurur)
    bar_h = 0.5 * cm

    c.saveState()
    c.setFont(normal_font, 9)
    c.setFillColorRGB(*COLORS['text_medium'])

    # Ön Bölüm
    c.drawString(bar_x, stat_y - 0.8*cm, f"Ön Bölüm: %{front_pct:.0f}")
    draw_progress_bar(c, bar_x, stat_y - 1.45*cm, bar_w, bar_h, front_pct,
                    COLORS['success'] if front_pct >= 70 else COLORS['danger'])

    # Ana Metin
    c.drawString(bar_x, stat_y - 2.3*cm, f"Ana Metin: %{body_pct:.0f}")
    draw_progress_bar(c, bar_x, stat_y - 2.95*cm, bar_w, bar_h, body_pct,
                    COLORS['success'] if body_pct >= 70 else COLORS['danger'])

    # Arka Bölüm
    c.drawString(bar_x, stat_y - 3.8*cm, f"Arka Bölüm: %{back_pct:.0f}")
    draw_progress_bar(c, bar_x, stat_y - 4.45*cm, bar_w, bar_h, back_pct,
                    COLORS['success'] if back_pct >= 70 else COLORS['danger'])

    c.restoreState()
    
    # Küçük özet (dikkat çekmeyen)
    c.saveState()
    c.setFont(normal_font, 8)
    c.setFillColorRGB(*COLORS['text_light'])
    c.drawString(bar_x, stat_y - 5.1*cm,
                f"Toplam: {total_checked}  •  Uygun: {total_ok}  •  Uygunsuz: {total_fail}")
    c.restoreState()


    
    # --- Alt bilgi kutusu ---
    footer_box_y = 2.5 * cm
    c.saveState()
    c.setStrokeColorRGB(*COLORS['border_light'])
    c.setLineWidth(0.5)
    c.line(2*cm, footer_box_y + 1*cm, width - 2*cm, footer_box_y + 1*cm)
    
    c.setFont(normal_font, 9)
    c.setFillColorRGB(*COLORS['text_light'])
    footer_text = f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    if app_version_text:
        footer_text += f"  •  Sürüm: {app_version_text}"
    c.drawCentredString(width/2, footer_box_y + 0.3*cm, footer_text)
    c.drawRightString(width - 2*cm, footer_box_y + 0.3*cm, "Sayfa 1")
    c.restoreState()
    
    # ════════════════════════════════════════════════════════════════
    # KAPAK SAYFASI BİTTİ
    # ════════════════════════════════════════════════════════════════
    # Not: Her bölüm kendi sayfasında başladığı için burada showPage yapmıyoruz
    
    page_num = 1  # Kapak sayfası = 1

    # ------------------------------------------------------------
    # 3. DETAYLI SONUÇLAR (Bölüm tabloları)
    # ------------------------------------------------------------
    for section_key in report_config["report"].get(
        "section_order",
        report_config["report"].get("order", [])
    ):
        # ════════════════════════════════════════════════════════════════
        # ✅ Her bölüm yeni sayfada başlasın
        # ════════════════════════════════════════════════════════════════
        c.showPage()
        page_num += 1
        draw_watermark(c)
        draw_page_footer(c, page_num)
        y = height - 2 * cm
        # Bölüm etiketi
        label = report_config["report"].get(
            "section_labels",
            report_config["report"].get("section_titles", {})
        ).get(section_key, section_key.upper())

        # Bölüm başlığı (Modern stil)
        c.saveState()
        # Sol accent bar
        c.setFillColorRGB(*COLORS['accent'])
        c.roundRect(1.8*cm, y - 0.15*cm, 0.3*cm, 0.55*cm, 2, fill=1, stroke=0)
        
        c.setFont(bold_font, 12)
        c.setFillColorRGB(*COLORS['primary'])
        c.drawString(2.3*cm, y, str(label))
        c.restoreState()
        y -= 0.7 * cm

        section_results = results_by_section.get(section_key, [])

        # Tablo kolon isimleri ve E/H indexleri
        table_cols = report_config["report"].get("table_columns", [])

        def _find_col_idx(cols, candidates):
            cols_lower = [str(c).strip().lower() for c in cols]
            for cand in candidates:
                cand_l = cand.lower()
                if cand_l in cols_lower:
                    return cols_lower.index(cand_l)
            return None

        # Hem (E/H) hem de (Evet/Hayır) destekli
        e_idx = _find_col_idx(table_cols, ["Evet", "E", "Yes"])
        h_idx = _find_col_idx(table_cols, ["Hayır", "Hayir", "H", "No"])

        # E/H normalize ederken sayıları hesapla
        normalized_rows = []
        ok_count = 0
        fail_count = 0

        for res in section_results:
            row = list(res)

            # E sütunu
            if e_idx is not None and e_idx < len(row):
                val = row[e_idx]
                s = str(val).strip()
                if s in ("✔", "✓", "E", "EVET", "True", "1") or (isinstance(val, bool) and val is True):
                    row[e_idx] = "E"
                    ok_count += 1
                else:
                    row[e_idx] = ""

            # H sütunu
            if h_idx is not None and h_idx < len(row):
                val = row[h_idx]
                s = str(val).strip()
                if s in ("✘", "✗", "H", "HAYIR", "False", "0") or (isinstance(val, bool) and val is False):
                    row[h_idx] = "H"
                    fail_count += 1
                else:
                    row[h_idx] = ""

            normalized_rows.append(row)

        total_rules = len(section_results)

        # Bölüm uyum yüzdesi (Toplam kural üzerinden)
        section_pct = (ok_count / total_rules * 100.0) if total_rules > 0 else 0.0
        
        # Renk seçimi
        if section_pct >= 90:
            pct_color = COLORS['success']
        elif section_pct >= 70:
            pct_color = COLORS['warning']
        else:
            pct_color = COLORS['danger']

        # Özet satırı (modern stil)
        c.saveState()
        c.setFont(normal_font, 9)
        c.setFillColorRGB(*COLORS['text_medium'])
        c.drawString(2.3*cm, y, f"Toplam: {total_rules} kural  •  ")
        
        c.setFillColorRGB(*COLORS['success'])
        c.drawString(5.8*cm, y, f"✓ {ok_count} uygun")
        
        c.setFillColorRGB(*COLORS['danger'])
        c.drawString(8.3*cm, y, f"✗ {fail_count} uygunsuz")
        
        c.setFillColorRGB(*pct_color)
        c.setFont(bold_font, 9)
        c.drawString(11.3*cm, y, f"[%{section_pct:.0f} Uyum]")
        c.restoreState()
        y -= 0.45 * cm

        # ------------------------------------------------------------
        # --- TABLO ÇİZİMİ ---
        # ------------------------------------------------------------

        # Sayfa sonu kontrolü (tabloya başlamadan önce)
        min_rows_space = 3  # başlık + en az 2 satır için
        base_line_height = 0.5 * cm
        if y < (2 * cm + min_rows_space * base_line_height):
            c.showPage()
            page_num += 1
            draw_watermark(c)
            draw_page_footer(c, page_num)
            width, height = A4
            y = height - 2 * cm
            c.setFont(normal_font, 10)

        # Tablo geometrisi
        left = 2 * cm
        right = width - 2 * cm
        available_width = right - left

        # ------------------------------------------------------------
        # GÖRÜNÜM GÜNCELLEMESİ
        # - "Açıklama" sütununu başlıktan kaldır
        # - Evet/Hayır sütunları en sağda kalsın
        # - Açıklamalar (sadece Hayır olanlarda) ilgili kuralın ALTINDA,
        #   tek hücre/tek satır olarak (tüm tablo genişliği) yazılsın
        # ------------------------------------------------------------
        def _is_aciklama_col(name) -> bool:
            s = str(name or "").strip().lower()
            return s in ("açıklama", "aciklama", "açiklama")

        table_cols_all = list(table_cols)  # orijinali sakla
        display_col_indices = [i for i, nm in enumerate(table_cols_all) if not _is_aciklama_col(nm)]
        table_cols = [table_cols_all[i] for i in display_col_indices]  # ekranda görünen başlıklar

        # "Açıklama" kolonunun index'i (varsa)
        aciklama_idx = None
        for i, nm in enumerate(table_cols_all):
            if _is_aciklama_col(nm):
                aciklama_idx = i
                break

        col_count = max(1, len(table_cols))

        # Kolon başlıklarına göre ağırlık ver (Evet/Hayır sabit genişlik, Kural geniş)
        weights = []
        for name in table_cols:
            name_lower = str(name or "").strip().lower()
            if name_lower in ("evet", "hayır", "hayir", "e", "h"):
                weights.append(1.2)  # ✅ Evet/Hayır için biraz daha geniş, sabit genişlik
            elif str(name) in ("No#",) or name_lower.startswith("no"):
                weights.append(0.8)
            elif "kural" in name_lower:
                weights.append(7.0)  # Kural sütununu daha da genişlet
            else:
                weights.append(2.0)

        total_w = sum(weights) if weights else col_count
        col_widths = [available_width * w / total_w for w in weights]
        col_x = [left]
        for w_ in col_widths:
            col_x.append(col_x[-1] + w_)

        # 1) BAŞLIK SATIRI (Modern Gradient)
        header_height = base_line_height + 0.35 * cm
        header_top_y = y
        header_bottom_y = y - header_height

        # Gradient header arka planı
        draw_gradient_rect(c, left, header_bottom_y, right - left, header_height,
                          COLORS['primary'], COLORS['primary_light'], steps=30)
        
        c.saveState()
        c.setFont(bold_font, 10)
        c.setFillColorRGB(1, 1, 1)  # Beyaz yazı
        header_text_y = header_bottom_y + 0.28 * cm
        for i, col_name in enumerate(table_cols):
            text_x = col_x[i] + 4
            c.drawString(text_x, header_text_y, str(col_name))
        c.restoreState()

        y = header_bottom_y

        # 2) VERİ SATIRLARI
        data_font_size = 9
        c.setFont(normal_font, data_font_size)
        
        for row_idx, row_all in enumerate(normalized_rows):
            # Ekranda gösterilecek satır (Açıklama sütunu çıkarılmış hali)
            row = [row_all[i] if i < len(row_all) else "" for i in display_col_indices]

            # Hayır olduğunda yazılacak açıklama metni (varsa)
            aciklama_text = ""
            if aciklama_idx is not None and aciklama_idx < len(row_all):
                aciklama_text = str(row_all[aciklama_idx] or "").strip()

            cell_lines_list = []
            max_lines = 1

            for i in range(col_count):
                cell_text = row[i] if i < len(row) else ""
                col_name = table_cols[i] if i < len(table_cols) else ""
                col_lower = str(col_name).lower()

                # Kural metni satır kaydırmalı (wrap)
                if ("kural" in col_lower):
                    raw = str(cell_text or "")
                    chunks = re.split(r'[\n\r]+', raw)

                    parts = []
                    for ch in chunks:
                        ch = ch.strip()
                        if not ch:
                            continue
                        subparts = re.split(r'(?<=[\;])\s+', ch)
                        for sp in subparts:
                            sp = sp.strip()
                            if sp:
                                parts.append(sp)

                    usable_width = col_widths[i] - 4
                    lines = []
                    for p in parts:
                        wrapped = wrap_text(p, normal_font, data_font_size, usable_width)
                        lines.extend(wrapped)

                    if not lines:
                        lines = [""]

                else:
                    lines = [str(cell_text)]

                cell_lines_list.append(lines)
                if len(lines) > max_lines:
                    max_lines = len(lines)

            row_height = max_lines * base_line_height + 0.3 * cm

            # Sayfa sonu kontrolü
            if y - row_height < 2 * cm:
                c.showPage()
                page_num += 1
                draw_watermark(c)
                draw_page_footer(c, page_num)
                width, height = A4
                y = height - 2 * cm

                # ✅ Gradient header ile tutarlı başlık çiz (sayfa bölünmesinde)
                header_top_y = y
                header_bottom_y = y - header_height
                draw_gradient_rect(c, left, header_bottom_y, right - left, header_height,
                                  COLORS['primary'], COLORS['primary_light'], steps=30)
                
                c.saveState()
                c.setFont(bold_font, 10)
                c.setFillColorRGB(1, 1, 1)  # Beyaz yazı
                header_text_y = header_bottom_y + 0.28 * cm
                for i, col_name in enumerate(table_cols):
                    text_x = col_x[i] + 4
                    c.drawString(text_x, header_text_y, str(col_name))
                c.restoreState()

                y = header_bottom_y
                c.setFont(normal_font, data_font_size)

            row_top_y = y
            row_bottom_y = y - row_height

            # ------------------------------------------------------------
            # ✅ Bu satır HAYIR mı? (ve açıklama var mı?)
            # ------------------------------------------------------------
            has_fail = False
            for i in range(col_count):
                col_name = table_cols[i] if i < len(table_cols) else ""
                col_lower = str(col_name).strip().lower()
                if col_lower in ["hayır", "hayir", "h"]:
                    v = row[i] if i < len(row) else ""
                    if str(v).strip() == "H":
                        has_fail = True
                    break

            has_exp = bool(has_fail and aciklama_text)

            # Açıklama satırı varsa yüksekliğini şimdiden hesapla
            exp_font_size = 8

            exp_lines = []
            exp_row_height = 0.0

            if has_exp:

                tx = col_x[1] + 6                      # açıklamanın başlayacağı x
                usable_width = (right - 4) - tx        # sağ sınırdan biraz pay bırak

                # ------------------------------------------------------------
                # ✅ Açıklamayı madde madde ayır ( ; veya satır sonuna göre )
                # ------------------------------------------------------------
                raw_text = str(aciklama_text or "").strip()

                # Önce satır sonlarını ; gibi düşünelim (bazı mesajlar \n ile gelebilir)
                raw_text = re.sub(r"[\r\n]+", "; ", raw_text)

                # ; ile böl -> boşları temizle
                items = [it.strip() for it in re.split(r"[;；]", raw_text) if it.strip()]

                bullet = "- "   # istersen "• " yapabilirsin

                # Hiç bölünemediyse tek madde gibi davran
                if not items:
                    items = [raw_text] if raw_text else []

                exp_lines = []

                for idx, item in enumerate(items):
                    # İlk satırda "Açıklama:" başlığı görünsün
                    if idx == 0:
                        # Açıklama: satırını tek başına ekleyelim (daha okunaklı)
                        exp_lines.extend(wrap_text("Açıklama:", normal_font, exp_font_size, usable_width))

                    # Her maddeyi "- " ile başlat
                    text = f"{bullet}{item}"

                    wrapped = wrap_text(text, normal_font, exp_font_size, usable_width)

                    # wrap_text bazen boş dönerse fallback
                    if not wrapped:
                        wrapped = [text]

                    exp_lines.extend(wrapped)

                # Yükseklik hesapla                
                if not exp_lines:
                    exp_lines = [f"Açıklama: {aciklama_text}"]
                exp_row_height = len(exp_lines) * base_line_height + 0.25 * cm

            # ------------------------------------------------------------
            # ✅ Sayfa sonu kontrolü (ana satır + varsa açıklama satırı birlikte)
            # ------------------------------------------------------------
            total_block_height = row_height + (exp_row_height if has_exp else 0.0)

            if y - total_block_height < 2 * cm:
                c.showPage()
                page_num += 1
                draw_watermark(c)
                draw_page_footer(c, page_num)
                width, height = A4
                y = height - 2 * cm

                # ✅ Gradient header ile tutarlı başlık çiz (sayfa bölünmesinde - açıklama bloğu)
                header_top_y = y
                header_bottom_y = y - header_height
                draw_gradient_rect(c, left, header_bottom_y, right - left, header_height,
                                  COLORS['primary'], COLORS['primary_light'], steps=30)
                
                c.saveState()
                c.setFont(bold_font, 10)
                c.setFillColorRGB(1, 1, 1)  # Beyaz yazı
                header_text_y = header_bottom_y + 0.28 * cm
                for i, col_name in enumerate(table_cols):
                    text_x = col_x[i] + 4
                    c.drawString(text_x, header_text_y, str(col_name))
                c.restoreState()

                y = header_bottom_y
                c.setFont(normal_font, data_font_size)

                # yeni pozisyonla tekrar hesapla
                row_top_y = y
                row_bottom_y = y - row_height

            # ------------------------------------------------------------
            # ✅ SATIR ARKA PLANI (Zebra + Hata Boyama)
            # ------------------------------------------------------------
            block_top_y = row_top_y
            block_bottom_y = row_top_y - total_block_height

            c.saveState()
            if has_fail: 
                c.setFillColorRGB(1.0, 0.94, 0.94) # Çok açık kırmızı (hata vurgusu)
            elif row_idx % 2 == 1:
                c.setFillColorRGB(0.98, 0.98, 0.98) # Çok açık gri (zebra)
            else:
                c.setFillColorRGB(1, 1, 1)
            
            c.rect(left, block_bottom_y, right - left, total_block_height, fill=1, stroke=0)
            c.restoreState()

            c.setLineWidth(0.5)
            c.setStrokeColorRGB(0.8, 0.8, 0.8) # Hafif gri çizgiler

            # Dış çerçeve (tüm blok)
            c.line(left, block_top_y, right, block_top_y)
            c.line(left, block_bottom_y, right, block_bottom_y)
            c.line(left, block_top_y, left, block_bottom_y)
            c.line(right, block_top_y, right, block_bottom_y)

            # No kolonu dikey çizgisi (No | diğerleri ayrımı) -> tüm blok boyunca
            c.line(col_x[1], block_top_y, col_x[1], block_bottom_y)

            # Ana satır iç kolon çizgileri (Kural | Evet | Hayır) -> sadece ANA satır yüksekliği kadar
            # (açıklama satırında bu kolonlar birleşik olacak)
            if col_count >= 3:
                for x in col_x[2:-1]:
                    c.line(x, block_top_y, x, row_bottom_y)

            # Ana satır alt çizgisi
            # - açıklama varsa: No kolonunu kesmeden, sadece col_x[1] -> right arası çiz
            # - açıklama yoksa: tüm genişlikte çiz (normal satır)
            if has_exp:
                c.line(col_x[1], row_bottom_y, right, row_bottom_y)
            else:
                c.line(left, row_bottom_y, right, row_bottom_y)

            # ------------------------------------------------------------
            # ✅ HÜCRE İÇERİKLERİ
            # ------------------------------------------------------------

            # 1) NO yazısı: açıklama satırı da varsa, iki satırın ortasına ortala
            no_text = row[0] if len(row) > 0 else ""
            no_cx = (col_x[0] + col_x[1]) / 2.0
            no_cy = (block_top_y + block_bottom_y) / 2.0  # iki satırın ortası

            # NO'yu çiz (bold/normal tercih senin)
            c.setFont(bold_font, data_font_size)
            c.drawCentredString(no_cx, no_cy - 0.1*cm, str(no_text))
            c.setFont(normal_font, data_font_size)

            # 2) Diğer hücreler (No dışındakiler)
            for i in range(1, col_count):
                col_name = table_cols[i] if i < len(table_cols) else ""
                col_lower = str(col_name).strip().lower()
                cell_val = row[i] if i < len(row) else ""
                s = str(cell_val).strip()

                # Hücre merkezi (tik/cross)
                x_left = col_x[i]
                x_right = col_x[i + 1]
                cx = (x_left + x_right) / 2.0
                cy = (row_top_y + row_bottom_y) / 2.0

                # EVET -> yeşil tik
                if col_lower in ["evet", "e"]:
                    if s == "E":
                        draw_tick(c, cx, cy, size_pt=9, rgb=(0, 0.65, 0))
                    continue

                # HAYIR -> kırmızı çarpı
                if col_lower in ["hayır", "hayir", "h"]:
                    if s == "H":
                        draw_cross(c, cx, cy, size_pt=9, rgb=(0.85, 0, 0))
                    continue

                # Diğer kolonlar: metin bas - dikey ortalı
                lines = cell_lines_list[i]
                text_x = col_x[i] + 4
                # Dikey ortalama: toplam metin yüksekliğini hesapla ve ortala
                total_text_height = len(lines) * base_line_height
                text_start_y = ((row_top_y + row_bottom_y) / 2) + (total_text_height / 2) - 0.1 * cm
                for line_idx, line_text in enumerate(lines):
                    line_y = text_start_y - (line_idx * base_line_height)
                    c.drawString(text_x, line_y, line_text)

            # ------------------------------------------------------------
            # ✅ AÇIKLAMA SATIRI (sadece HAYIR ise) - UYARI KUTUSU STİLİ
            # ------------------------------------------------------------
            if has_exp:
                # Önce eklenecek kutunun genişliğini belirleyelim
                eb_x = col_x[1] + 0.2 * cm
                eb_w = (right - 0.2 * cm) - eb_x
                
                # Metni bu genişliğe göre wrap et
                wrapped_exp = []
                # Çok uzun teknik ifadeleri kısaltabiliriz veya olduğu gibi wrap edebiliriz
                # Örn: "144. satır (belge:144)..."
                for el in exp_lines:
                    # wrap_text(text, font_name, font_size, max_width)
                    # max_width padding düşünülerek verilmeli (örn -15 pt)
                    w_lines = wrap_text(el, normal_font, exp_font_size, eb_w - 15)
                    wrapped_exp.extend(w_lines)
                
                exp_lines = wrapped_exp  # Güncellenmiş satırlar

                # Kutu yüksekliğini metne göre tekrar netleştirelim
                box_h = (len(exp_lines) * base_line_height) + 0.6 * cm
                eb_y = row_bottom_y - box_h + 0.1 * cm
                # eb_x ve eb_w yukarıda wrap işlemi için hesaplanmıştı
                
                c.saveState()
                c.setLineWidth(0.6)
                c.setStrokeColorRGB(1.0, 0.7, 0.7)
                c.setFillColorRGB(1.0, 1.0, 0.98)
                c.roundRect(eb_x, eb_y, eb_w, box_h, 3, fill=1, stroke=1)
                
                c.setFont(bold_font, 8)
                c.setFillColorRGB(0.8, 0, 0)
                c.drawString(eb_x + 0.3*cm, eb_y + box_h - 0.45*cm, "⚠ UYARI / AÇIKLAMA:")
                
                c.setFont(normal_font, exp_font_size)
                c.setFillColorRGB(0.2, 0.2, 0.2)
                tx = eb_x + 0.6 * cm
                ty = eb_y + box_h - 0.9*cm
                for li, tline in enumerate(exp_lines):
                    # "Açıklama:" tekrarını önle
                    if "Açıklama:" in tline: continue
                    c.drawString(tx, ty, tline)
                    ty -= base_line_height
                c.restoreState()

            # Blok bitti -> y’yi en alta al
            y = block_bottom_y



        # Bölümler arası boşluk
        y -= 1 * cm

    c.save()
    return pdf_path



# ============================================================
# 3.5 run_check (API / web entegrasyonu için tek giriş noktası)
# ============================================================
def run_thesis_check(docx_path, rules_path=None, report_path=None, app_version_text=None):
    """
    Dışarıdan (örn. FastAPI) çağırmak için tek fonksiyon.
    - docx_path: kontrol edilecek .docx dosya yolu
    - rules_path / report_path verilmezse, .py dosyasının yanındaki yaml’leri kullanır
    Dönenler:
      pdf_path, results_by_section, student_name
    """
    from pathlib import Path
    import yaml

    base_dir = Path(__file__).parent

    docx_path = Path(docx_path)

    if rules_path is None:
        rules_path = base_dir / "rules.yaml"
    else:
        rules_path = Path(rules_path)

    if report_path is None:
        report_path = base_dir / "report.yaml"
    else:
        report_path = Path(report_path)

    # YAML yükle
    with open(rules_path, "r", encoding="utf-8") as f:
        rules_data = yaml.safe_load(f)

    with open(report_path, "r", encoding="utf-8") as f:
        report_data = yaml.safe_load(f)

    # Belgeyi kontrol et (V50 zaten student_name döndürüyor) :contentReference[oaicite:3]{index=3}
    results_by_section, student_name = process_document(docx_path, rules_data, report_data)

    # PDF raporu üret (create_report pdf_path döndürüyor) :contentReference[oaicite:4]{index=4}
    pdf_path = create_report(report_data, results_by_section, docx_path.name, rules_path.name, student_name, app_version_text=app_version_text)

    return pdf_path, results_by_section, student_name



# ============================================================
# 4. process_document
# ============================================================

def process_document(docx_path, rules_data, report_data):
    """
    Word belgesini açar, rules.yaml kurallarına göre kontrol eder
    ve rapor için tabloya eklenecek sonuçları hazırlar.
    Ayrıca iç kapaktan öğrenci adını tespit edip döndürür.
    """
    
    global memo  # 🔹 memo’yu global tanıtır (run_check fonksiyonuyla paylaşmak için)
    if "memo" not in globals():  # 🔹 Eğer tanımlı değilse oluştur
        memo = {}
    
    print(f"[DOSYA] Word belgesi aciliyor: {docx_path}")
    doc = Document(docx_path)  # Word belgesini yükle
    all_paragraphs = doc.paragraphs  # Tüm paragrafları oku

    # --- Sayfa / satır haritası (Word'ün lastRenderedPageBreak etiketine göre) ---
    # Not: Word belgesi en az bir kez Word ile açılıp kaydedilmiş olmalı ki
    #      lastRenderedPageBreak etiketleri oluşsun.
    page_num_map = {}
    current_page = 1
    for i, p in enumerate(all_paragraphs):
        # Bu paragraftan önce Word yeni sayfa oluşturmuş mu?
        if p._element.xpath(".//w:lastRenderedPageBreak"):
            current_page += 1
        page_num_map[i] = current_page

    # Aynı sayfadaki “satır”ı paragraf sırasına göre hesaba dök
    page_line_map = {}
    last_page = None
    line_on_page = 0
    for idx in range(len(all_paragraphs)):
        page = page_num_map.get(idx, 1)
        if page != last_page:
            # Yeni sayfa → satır sayacını sıfırla
            line_on_page = 1
            last_page = page
        else:
            line_on_page += 1
        page_line_map[idx] = (page, line_on_page)

    # Tüm kurallarda kullanabilmek için memo’ya koy
    memo.clear()             # Her çalıştırmada önceki değerleri sıfırla
    memo["page_line_map"] = page_line_map
    results_by_section = {}  # Bölüm bazlı sonuçlar burada tutulacak
    student_name = None      # 🔑 Öğrenci adını tutmak için değişken
    

    print("[OK] Belge açıldı, paragraf sayısı:", len(all_paragraphs))

    # rules.yaml içindeki tüm bölümleri sırayla işle
    for section_key, section_data in rules_data.get("pages", {}).items():
        if not section_data.get("enabled", False):
            print(f">> {section_key} bölümü atlandı (enabled: false)")
            continue

        print(f"\n[ANALIZ] {section_key.upper()} bölümü kontrol ediliyor...")
        section_results = []  # Bu bölümdeki kuralların sonuçları

        # O bölümdeki tüm kurallar
        for check in section_data.get("rules", []):
            print(f"   ➡ Kural: {check.get('check')}")
            result = run_check(doc, all_paragraphs, check, rules_data) or []  # Her kuralı kontrol et

            for res in result:
                rule_no = len(section_results) + 1
                durum = res[1]
                rule_title = res[2]
                explanation = res[3] if not durum else ""

                yes = "✔" if durum else ""
                no = "✘" if not durum else ""

                row = [rule_no, rule_title, yes, no, explanation]
                section_results.append(row)

                # 🔑 Öğrenci adını yakala (1. yöntem: doğrudan memo'dan)
                if memo.get("student_name"):
                    student_name = memo["student_name"]

                # 🔑 2. yöntem (eski) – “ÖĞRENCİ ADI-SOYADI” başlığına göre
                if student_name is None and "ÖĞRENCİ ADI-SOYADI" in rule_title and durum:
                    student_name = all_paragraphs[res[0]].text.strip()
                    print(f"🆔 Öğrenci adı bulundu (rule_title yöntemi): {student_name}")

                print(f"      ➕ {rule_title.splitlines()[0]}: {'UYGUN' if durum else 'UYGUNSUZ'}")

        results_by_section[section_key] = section_results
        print(f"📊 {section_key} tamamlandı, {len(section_results)} kural işlendi.")

    # --------------------------------------------------------
    # PRELIMINARIES bilgisini oku → run_check içinde kullanılacak
    # --------------------------------------------------------
    if "preliminaries" in rules_data and rules_data["preliminaries"].get("enabled", False):
        prelim_pages = rules_data["preliminaries"].get("pages", [])
        report_data["preliminaries_pages"] = prelim_pages
        print(f"ℹ️ Preliminaries sayfaları: {prelim_pages}")

    # --------------------------------------------------------
    # 🔁 Öğrenci adı son kez kontrol et (GÜNCELLENDİ)
    # --------------------------------------------------------
    # 1️⃣ memo'dan kontrol
    student_name = memo.get("student_name", None)


    # 2️⃣ Sonuç
    if student_name:
        print(f"🎓 Son öğrenci adı değeri (memo’dan): {student_name}")
        
    else:
        student_name = "OGRENCI_ADI"
        print("⚠️ Öğrenci adı memo’da bulunamadı, varsayılan ad kullanılacak.")

    # Artık öğrenci adı da döndürülüyor
    return results_by_section, student_name


# ============================================================
# 5. main
# ============================================================
def main():
    start_time = time.time()

    # --------------------------------------------------------
    # 1. Komut satırı kontrolü
    # --------------------------------------------------------
    # Artık sadece TEZ.docx parametresi bekliyoruz.
    # (rules.yaml ve report.yaml dosyaları .py dosyasının bulunduğu klasörden otomatik yüklenecek)
    if len(sys.argv) != 2:
        logger.warning("⚠️ Komut satırı argümanı verilmedi, varsayılan 'TEZ.docx' kullanılacak.")
        sys.argv.append("TEZ.docx")

    # TEZ.docx dosya yolunu al
    docx_file = Path(sys.argv[1])

    # (İsteğe bağlı ama faydalı) Dosya var mı kontrolü
    if not docx_file.exists():
        logger.error(f"❌ Dosya bulunamadı: {docx_file}")
        sys.exit(1)

    # (İsteğe bağlı) Uzantı kontrolü
    if docx_file.suffix.lower() != ".docx":
        logger.error("❌ Lütfen .docx uzantılı bir dosya verin.")
        sys.exit(1)

    # .py dosyasının bulunduğu klasörden yaml dosyalarını bul
    rules_file  = Path(__file__).parent / "rules.yaml"
    report_file = Path(__file__).parent / "report.yaml"

    # --------------------------------------------------------
    # 2. Kurallar ve rapor şablonunu yükle + kontrol + rapor
    # --------------------------------------------------------
    logger.info("[BILGI] Kurallar yükleniyor...")
    logger.info("[BILGI] Rapor hazırlanıyor...")

    # run_check: yaml oku + belgeyi kontrol et + pdf raporu üret
    pdf_path, results_by_section, student_name = run_thesis_check(docx_file, rules_file, report_file)

    logger.info("[OK] Kurallar yüklendi")

    # --------------------------------------------------------
    # 3. Süre / çıktı
    # --------------------------------------------------------
    duration = time.time() - start_time
    logger.info(f"[BILGI] İşlem süresi: {duration:.2f} saniye")
    logger.info(f"[OK] Rapor oluşturuldu: {pdf_path}")


if __name__ == "__main__":
    main()
# Main Fonksiyonu Sonu ---------------------------------------
